"""
Smart HR — AI-Powered Candidate Search & Analysis Platform
===========================================================

Main FastAPI application entry point. This module contains all API endpoints,
business logic, and integrations for the SmartHR recruitment platform.

Modules:
    - Authentication & Sessions (bcrypt, cookie-based tokens)
    - Document Processing (PDF/Docx extraction via pdfplumber, PyPDF2, python-docx)
    - Google Cloud Storage (per-tenant resume upload & retrieval)
    - Vertex AI / Gemini (resume analysis, HR scorecard generation, JD creation)
    - Discovery Engine (semantic vector search across indexed resumes)
    - Cloud Tasks (async background processing for batch scorecard jobs)
    - Multi-Tenant Management (company isolation, subscription limits)
    - Candidate Pipeline (shortlist, reject, interview scheduling, hiring)

Tech Stack:
    - Backend: FastAPI + Uvicorn (Python 3.11)
    - Database: PostgreSQL 15 via psycopg2 connection pool (see database.py)
    - AI: Google Vertex AI (Gemini 2.5 Flash) — prompts in llm_prompts.py
    - Search: Google Cloud Discovery Engine (vector similarity)
    - Storage: Google Cloud Storage
    - Queue: Google Cloud Tasks
    - Hosting: GCP Cloud Run (Docker container)

Configuration:
    - config.json — GCS bucket, Vertex AI project, Cloud SQL, Cloud Tasks settings
    - service-account.json — GCP credentials (not committed to source control)
    - Environment variables override config.json values in production

Author: Smart HR Team
Version: 1.0.0
"""
# VPS local-mode shim layer. When SMARTHR_LOCAL_MODE=1, replaces google.cloud.*
# and vertexai with local equivalents (filesystem + pgvector + sentence-transformers).
# Must be imported BEFORE any google.cloud / vertexai imports below.
# Always installable now that vps_local is vendored in the repo; the env var
# inside vps_local/__init__.py decides whether shims actually activate.
import vps_local  # noqa: F401  (auto-installs sys.modules shims if SMARTHR_LOCAL_MODE=1)

from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException, Depends, Cookie
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse, RedirectResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request as StarletteRequest
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from typing import Optional
from google.cloud import storage
from google.cloud import discoveryengine_v1
from google import genai
from google.genai import types
import os
import json
import base64
import io
import asyncio
import copy
from typing import List, Dict, Any, AsyncGenerator
import PyPDF2
import pdfplumber
from docx import Document
import warnings
import logging
import re
from database import get_db_manager
import builtins
# Email scraping imports
import imaplib
import email
from email.header import decode_header
import msal
import tempfile
import hashlib
import sqlite3
import bcrypt
import uuid
import secrets
from datetime import datetime, timedelta
import threading
import traceback
import requests
from urllib.parse import urlparse
import mimetypes
import markdown
import sys
import time
from contextlib import contextmanager
from pathlib import Path
import pandas as pd
import numpy as np
try:
    from google.cloud import tasks_v2
except ImportError:
    tasks_v2 = None
from google.protobuf import timestamp_pb2
from search_criteria import parse_search_criteria, apply_criteria_to_results



# LLM Prompts imports
from llm_prompts import (
    RESUME_ANALYSIS_PROMPT,
    RESUME_CONTENT_ANALYSIS_PROMPT,
    ENTITY_EXTRACTION_PROMPT,
    SCORECARD_PROMPT,
    HR_SCORECARD_PROMPT,
    KEYWORD_PROMPT,
    KEYWORD_EXTRACTION_PROMPT,
    # KEYWORD_ANALYSIS_PROMPT,  # Unused - custom inline prompt used instead
    KEYWORD_MATCHING_PROMPT,
    RANKING_PROMPT,
    # COMPANY_OVERVIEW_PROMPT,  # Unused
    JD_PROMPT,
    ENHANCE_PROMPT,
    JD_KEYWORD_PROMPT
)

# -------------------------------------------
# Centralized print filtering for cleaner logs
# -------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s: %(message)s"
)

# Module-level logger used in error paths (previously referenced as `logger`
# without being defined, causing NameError-on-error in 7 admin endpoints).
logger = logging.getLogger(__name__)

_original_print = builtins.print  # Preserve original print for fallback if ever needed

def _filtered_print(*args, **kwargs):
    """
    Filters excessive debug print statements while retaining
    essential summary and error messages.

    Allowed messages:
    1. Final results , errors, and critical status updates
    2. Company/resource setup confirmations
    3. High-level completion messages
    All other verbose debug output is suppressed.
    """
    if not args:
        return

    message = " ".join(str(a) for a in args)

    # Always allow error messages and final results
    CRITICAL_KEYWORDS = ("ERROR", "Error", "FAILED", "Failed", "❌", "⚠️")
    FINAL_RESULTS = ("✅ Successfully", "🎉", "🏆", "💾 Saved", "📊 Results:", "🏢 Company:", "📚 Datastore ID:")

    # Block verbose debug messages
    VERBOSE_PATTERNS = (
        "📝 Raw response preview:", "📝 Content preview:", "📝 Analysis preview:",
        "📝 Query:", "📝 Job requirements:", "📝 Job description:", "📝 Original query:",
        "📝 Response preview:", "🔧 Cleaned JSON text:", "📝 First 200 chars",
        "🧹 Cleaned text for JSON parsing:", "📝 Problematic text:", "📝 Query preview:",
        "🔍 FINAL DEBUG:", "📝 optimized_query:", "🔧 Reconstructed using pattern",
        "🔧 Entity JSON:", "🔧 Scorecard JSON:", "🔧 Last resort reconstruction:",
        "📄 Processing result #", "📄 Processing file", "📄 Candidate", "📄 Scorecard candidate",
        "Part", "text length:", "Got text from candidate", "Got scorecard text",
        "🛡️", "🏁", "finish reason:", "safety ratings:", "📝 Snippet", "Found snippet.",
        "📁 Extracted:", "📁 Using custom folder path:", "☁️ Uploading", "🔗 File path:",
        "📝 Using resolved blob name:", "📄 File:", "Extension:", "MIME:", "📊 File size:",
        "📦 Total batch size:", "🔄 Processing", "files in batch", "📄 Extracting text from",
        "🔍 Trying", "possible path patterns", "pattern:", "Not found with", "Found file using",
        "📊 Query length:", "characters,", "words", "🎯 Optimized", "query:", "🔑 Keywords list:",
        "🎯 Keywords:", "🔑 Sample keywords:", "🔑 Matched", "❌ Missing", "✅ Matched:",
        "🔍 Keyword Analysis:", "🔍 Missing:", "📋", "Company Details:", "📊 Candidate:",
        "📝 Company Details:", "🔄 extract_standardized_keywords", "📊 analyze_candidate_keywords",
        "🔄 Extracted standardized_keywords:", "📝 Extracting from candidates"
    )

    # Allow critical messages
    if any(keyword in message for keyword in CRITICAL_KEYWORDS):
        logging.error(message)
        return

    # Allow final results and important status
    if any(pattern in message for pattern in FINAL_RESULTS):
        logging.info(message)
        return

    # Block verbose debug patterns
    if any(pattern in message for pattern in VERBOSE_PATTERNS):
        logging.debug(message)
        return

    # Allow remaining messages (likely important status updates)
    logging.info(message)

# Override built-in print to use the filtered version
builtins.print = _filtered_print

# Configure maximum file size limits.
# 10 MB is plenty for any resume / JD; the previous 100 MB limit allowed
# storage abuse and memory exhaustion.
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB in bytes
MAX_TOTAL_UPLOAD_SIZE = 100 * 1024 * 1024  # 100MB total for multiple files
MAX_SEARCH_RESULTS = int(os.getenv("MAX_SEARCH_RESULTS", "25"))
MAX_HR_SCORECARD_RESULTS = int(os.getenv("MAX_HR_SCORECARD_RESULTS", str(MAX_SEARCH_RESULTS)))
MAX_SEARCH_FETCH_RESULTS = int(os.getenv("MAX_SEARCH_FETCH_RESULTS", str(max(MAX_SEARCH_RESULTS, 50))))
MAX_HR_SCORECARD_FETCH_RESULTS = int(os.getenv("MAX_HR_SCORECARD_FETCH_RESULTS", str(max(MAX_HR_SCORECARD_RESULTS, 50))))
SEARCH_CRITERIA_OVERFETCH_MULTIPLIER = int(os.getenv("SEARCH_CRITERIA_OVERFETCH_MULTIPLIER", "5"))
SEARCH_CRITERIA_OVERFETCH_EXTRA = int(os.getenv("SEARCH_CRITERIA_OVERFETCH_EXTRA", "15"))
MAX_EMAIL_SCRAPE_LIMIT = int(os.getenv("MAX_EMAIL_SCRAPE_LIMIT", "50"))
STATIC_CACHE_SECONDS = int(os.getenv("STATIC_CACHE_SECONDS", "3600"))
ENABLE_DISCOVERY_SUMMARIES = os.getenv("ENABLE_DISCOVERY_SUMMARIES", "false").lower() in {"1", "true", "yes"}
SEARCH_VECTOR_CACHE_TTL_SECONDS = int(os.getenv("SEARCH_VECTOR_CACHE_TTL_SECONDS", "120"))
SEARCH_VECTOR_CACHE_MAX_ENTRIES = int(os.getenv("SEARCH_VECTOR_CACHE_MAX_ENTRIES", "256"))

SAFE_HTTP_METHODS = {"GET", "HEAD", "OPTIONS"}
CSRF_COOKIE_NAME = "csrf_token"
CSRF_HEADER_NAME = "x-csrf-token"
CSRF_EXEMPT_PATHS = {
    "/api/login",
    "/api/process-hr-scorecard-task",
}

def _is_secure_request(request: StarletteRequest) -> bool:
    return (
        request.headers.get("x-forwarded-proto", "").lower() == "https"
        or request.url.scheme == "https"
    )

def _same_origin_request(request: StarletteRequest) -> bool:
    """Validate Origin/Referer when present for CSRF defense-in-depth."""
    host = request.headers.get("x-forwarded-host") or request.headers.get("host")
    if not host:
        return False

    origin = request.headers.get("origin")
    referer = request.headers.get("referer")
    source = origin or referer
    if not source:
        # Token validation still protects modern same-origin XHR/fetch. Some
        # legacy clients omit Origin/Referer, so do not reject solely on this.
        return True

    try:
        parsed = urlparse(source)
        return parsed.netloc.lower() == host.lower()
    except Exception:
        return False

def clamp_result_count(value: Any, default: int = 10, max_value: int = MAX_SEARCH_RESULTS) -> int:
    """Normalize user-controlled search sizes before calling Discovery/LLM."""
    try:
        count = int(value)
    except (TypeError, ValueError):
        count = default
    if count < 1:
        raise HTTPException(status_code=400, detail="result_count must be at least 1")
    if count > max_value:
        raise HTTPException(status_code=400, detail=f"result_count cannot exceed {max_value}")
    return count

def criteria_fetch_count(requested_count: int, criteria: Any, max_value: int = MAX_SEARCH_FETCH_RESULTS) -> int:
    """Search a deeper candidate pool when exact hard criteria must be met."""
    try:
        requested = int(requested_count)
    except (TypeError, ValueError):
        requested = 10
    requested = max(1, min(requested, max_value))
    if not criteria or criteria.is_empty():
        return requested

    multiplier = max(1, SEARCH_CRITERIA_OVERFETCH_MULTIPLIER)
    extra = max(0, SEARCH_CRITERIA_OVERFETCH_EXTRA)
    return min(max_value, max(requested, requested * multiplier, requested + extra))

def clamp_positive_int(value: Any, default: int, max_value: int, field_name: str) -> int:
    try:
        count = int(value)
    except (TypeError, ValueError):
        count = default
    if count < 1:
        raise HTTPException(status_code=400, detail=f"{field_name} must be at least 1")
    if count > max_value:
        raise HTTPException(status_code=400, detail=f"{field_name} cannot exceed {max_value}")
    return count

def validate_resume_payload(filename: str, content: bytes) -> tuple[bool, str]:
    """Lightweight magic-byte validation for supported resume formats."""
    ext = "." + (filename or "").rsplit(".", 1)[-1].lower() if "." in (filename or "") else ""
    header = (content or b"")[:8]
    if ext == ".pdf":
        return header.startswith(b"%PDF-"), "PDF files must start with a valid PDF signature."
    if ext == ".docx":
        return header.startswith(b"PK\x03\x04") or header.startswith(b"PK\x05\x06") or header.startswith(b"PK\x07\x08"), "DOCX files must be valid ZIP-based Office documents."
    if ext == ".doc":
        return header.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"), "DOC files must be valid OLE compound documents."
    return False, "Unsupported file type."


_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _sanitize_user_text(text: str, max_len: int = 4000) -> str:
    """Strip control characters, collapse excessive whitespace, and clip length.

    Used to defang user-supplied text that flows into LLM prompts. This is not
    a substitute for prompt-engineering safety, but it removes the trivial
    injection vectors (NULs, control chars, runaway lengths)."""
    if not text:
        return ""
    cleaned = _CONTROL_CHARS_RE.sub(" ", str(text))
    cleaned = re.sub(r"\r\n?", "\n", cleaned)
    cleaned = re.sub(r"\n{4,}", "\n\n\n", cleaned)
    cleaned = cleaned.strip()
    if max_len and len(cleaned) > max_len:
        cleaned = cleaned[:max_len]
    return cleaned


def clamp_score(value, default: int = 0) -> int:
    """Coerce an LLM-produced score into the [0, 100] integer range."""
    try:
        score = int(round(float(value)))
    except (TypeError, ValueError):
        return default
    if score < 0:
        return 0
    if score > 100:
        return 100
    return score


_search_vector_cache: dict = {}
_search_vector_cache_lock = threading.Lock()


def _normalize_cache_query(query: str) -> str:
    return re.sub(r"\s+", " ", (query or "").strip().lower())


def _search_cache_get(key: tuple) -> Optional[dict]:
    if SEARCH_VECTOR_CACHE_TTL_SECONDS <= 0:
        return None
    now = time.time()
    with _search_vector_cache_lock:
        cached = _search_vector_cache.get(key)
        if not cached:
            return None
        expires_at, payload = cached
        if expires_at <= now:
            _search_vector_cache.pop(key, None)
            return None
        return copy.deepcopy(payload)


def _search_cache_set(key: tuple, payload: dict) -> None:
    if SEARCH_VECTOR_CACHE_TTL_SECONDS <= 0:
        return
    now = time.time()
    with _search_vector_cache_lock:
        if len(_search_vector_cache) >= SEARCH_VECTOR_CACHE_MAX_ENTRIES:
            oldest_key = min(_search_vector_cache, key=lambda item: _search_vector_cache[item][0])
            _search_vector_cache.pop(oldest_key, None)
        _search_vector_cache[key] = (now + SEARCH_VECTOR_CACHE_TTL_SECONDS, copy.deepcopy(payload))


def _search_cache_clear() -> None:
    with _search_vector_cache_lock:
        _search_vector_cache.clear()

# Alternative Vertex AI implementation for Cloud Run compatibility
def get_vertex_ai_model(model_name="gemini-2.5-flash"):
    """Get Vertex AI model with proper Cloud Run configuration"""
    import os

    # Check if running on Cloud Run
    if os.environ.get('K_SERVICE'):
        print(f"🏃 Running on Cloud Run - using vertexai SDK")
        try:
            import vertexai
            from vertexai.generative_models import GenerativeModel, GenerationConfig

            # Initialize Vertex AI with the project
            vertexai.init(project=PROJECT_ID, location="europe-west4")

            # Create and return the model
            model = GenerativeModel(model_name)
            return model, "vertexai"
        except Exception as e:
            print(f"⚠️ Failed to initialize vertexai SDK: {e}")
            # Fall back to genai

    # For local development or if vertexai fails
    print(f"💻 Using google-genai SDK")
    return None, "genai"

class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    """Adds standard security response headers to every response."""
    async def dispatch(self, request: StarletteRequest, call_next):
        response = await call_next(request)
        # Always-on hardening headers
        response.headers.setdefault("X-Content-Type-Options", "nosniff")
        response.headers.setdefault("X-Frame-Options", "DENY")
        response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
        response.headers.setdefault("Permissions-Policy", "camera=(), microphone=(), geolocation=()")
        # HSTS only when actually served via HTTPS (Cloud Run sets X-Forwarded-Proto)
        if request.headers.get("x-forwarded-proto", "").lower() == "https" or request.url.scheme == "https":
            response.headers.setdefault("Strict-Transport-Security", "max-age=31536000; includeSubDomains")
        # Conservative CSP that allows the CDNs the app already uses (FontAwesome, Chart.js, Google Fonts)
        # Tailwind is now built into /static/css/tailwind.css so we can drop the runtime CDN.
        # Skip for static files & API JSON to avoid breaking downloads / image previews.
        path = request.url.path
        if not path.startswith("/static") and not path.startswith("/api"):
            response.headers.setdefault(
                "Content-Security-Policy",
                "default-src 'self'; "
                "script-src 'self' 'unsafe-inline' https://cdnjs.cloudflare.com https://cdn.jsdelivr.net; "
                "style-src 'self' 'unsafe-inline' https://cdnjs.cloudflare.com https://fonts.googleapis.com https://cdn.jsdelivr.net; "
                "font-src 'self' data: https://fonts.gstatic.com https://cdnjs.cloudflare.com; "
                "img-src 'self' data: https: blob:; "
                "connect-src 'self' https:; "
                "object-src 'none'; "
                "frame-src 'none'; "
                "frame-ancestors 'none'; "
                "form-action 'self'; "
                "base-uri 'self'; "
                "upgrade-insecure-requests"
            )
        return response


class CSRFProtectionMiddleware(BaseHTTPMiddleware):
    """Protect cookie-authenticated unsafe requests with a double-submit token."""

    async def dispatch(self, request: StarletteRequest, call_next):
        path = request.url.path
        method = request.method.upper()
        has_session = bool(request.cookies.get("session_token"))

        if method not in SAFE_HTTP_METHODS and path not in CSRF_EXEMPT_PATHS and has_session:
            if not _same_origin_request(request):
                return JSONResponse(
                    {"detail": "Cross-site request rejected."},
                    status_code=403,
                )

            csrf_cookie = request.cookies.get(CSRF_COOKIE_NAME)
            csrf_header = request.headers.get(CSRF_HEADER_NAME)
            if not csrf_cookie or not csrf_header or not secrets.compare_digest(csrf_cookie, csrf_header):
                return JSONResponse(
                    {"detail": "Invalid or missing CSRF token."},
                    status_code=403,
                )

        response = await call_next(request)

        if method in SAFE_HTTP_METHODS and not request.cookies.get(CSRF_COOKIE_NAME):
            response.set_cookie(
                key=CSRF_COOKIE_NAME,
                value=secrets.token_urlsafe(32),
                httponly=False,
                secure=_is_secure_request(request),
                samesite="strict",
                max_age=7 * 24 * 60 * 60,
            )

        return response


class StaticCacheControlMiddleware(BaseHTTPMiddleware):
    """Add conservative browser caching for static assets."""

    async def dispatch(self, request: StarletteRequest, call_next):
        response = await call_next(request)
        if request.url.path.startswith("/static/") and STATIC_CACHE_SECONDS > 0:
            response.headers.setdefault(
                "Cache-Control",
                f"public, max-age={STATIC_CACHE_SECONDS}",
            )
        return response


class LimitUploadSizeMiddleware(BaseHTTPMiddleware):
    """Middleware to limit upload size and return proper 413 responses for oversized requests"""
    def __init__(self, app, max_upload_size: int):
        super().__init__(app)
        self.max_upload_size = max_upload_size

    async def dispatch(self, request: StarletteRequest, call_next):
        # Check for upload endpoints that should have size limits
        upload_endpoints = ['/api/upload-files', '/api/configure-gcs']

        if request.url.path in upload_endpoints and "content-length" in request.headers:
            try:
                content_length = int(request.headers["content-length"])
                if content_length > self.max_upload_size:
                    print(f"🚫 Request blocked: Content length {content_length:,} bytes exceeds limit {self.max_upload_size:,} bytes")
                    return JSONResponse(
                        content={
                            "detail": f"Request too large. Maximum size allowed: {self.max_upload_size // (1024*1024)}MB, received: {content_length // (1024*1024)}MB",
                            "error_code": "CONTENT_TOO_LARGE",
                            "max_size_mb": self.max_upload_size // (1024*1024),
                            "received_size_mb": content_length // (1024*1024)
                        },
                        status_code=413
                    )
            except (ValueError, TypeError):
                # Reject malformed Content-Length on size-limited endpoints
                # rather than silently letting the request through (would
                # otherwise be an easy size-limit bypass).
                print(f"🚫 Rejecting upload with malformed Content-Length header: {request.headers.get('content-length')!r}")
                return JSONResponse(
                    content={
                        "detail": "Invalid Content-Length header.",
                        "error_code": "INVALID_CONTENT_LENGTH",
                    },
                    status_code=400,
                )

        return await call_next(request)

app = FastAPI(
    title="HR Agent API",
    description="HR Resume Search and Analysis Platform",
    version="1.0.0"
)

# Per-IP rate limiter (used on auth endpoints)
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Add request-size limiting middleware. Individual files are still capped at
# MAX_FILE_SIZE in the upload handler; batches may contain several resumes.
app.add_middleware(LimitUploadSizeMiddleware, max_upload_size=MAX_TOTAL_UPLOAD_SIZE)
# CSRF guard for cookie-authenticated unsafe methods
app.add_middleware(CSRFProtectionMiddleware)
# Security response headers
app.add_middleware(SecurityHeadersMiddleware)
# Browser caching for static files (JS/CSS/images). Kept conservative because
# current asset filenames are not fingerprinted.
app.add_middleware(StaticCacheControlMiddleware)
# Compress responses >= 1 KB. Cuts page-load bytes ~75% for static JS/CSS
# and JSON API payloads. Streaming responses (SSE) are passed through
# uncompressed by FastAPI so /api/smart-search-stream is unaffected.
from fastapi.middleware.gzip import GZipMiddleware
app.add_middleware(GZipMiddleware, minimum_size=1024)

# CORS — allow same-origin by default; widen via CORS_ORIGINS env var (comma-separated).
from fastapi.middleware.cors import CORSMiddleware
_cors_origins_env = os.environ.get("CORS_ORIGINS", "").strip()
_cors_origins = [o.strip() for o in _cors_origins_env.split(",") if o.strip()] if _cors_origins_env else []
if _cors_origins:
    app.add_middleware(
        CORSMiddleware,
        allow_origins=_cors_origins,
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")


def _compute_asset_versions() -> dict[str, str]:
    """Compute short MD5 hashes for cache-busting static assets.

    Hashes are computed once at import time. Templates can use
    `asset_v.main_js` / `asset_v.style_css` to append a `?v=...` query string
    so browsers invalidate cached copies whenever we deploy a change.
    """
    import hashlib as _hashlib
    out: dict[str, str] = {"main_js": "0", "style_css": "0"}
    for key, path in (("main_js", "static/js/main.js"),
                      ("style_css", "static/css/style.css")):
        try:
            with open(path, "rb") as fh:
                out[key] = _hashlib.md5(fh.read()).hexdigest()[:8]
        except Exception:
            pass
    return out


_ASSET_VERSIONS = _compute_asset_versions()
templates.env.globals["asset_v"] = _ASSET_VERSIONS

# Progress tracking for real-time updates
class TokenTracker:
    """Track LLM API usage and token consumption"""
    def __init__(self):
        self.calls = []
        self.total_input_tokens = 0
        self.total_output_tokens = 0
        self.total_calls = 0

    def add_call(self, operation: str, model: str, input_tokens: int = 0, output_tokens: int = 0, success: bool = True):
        """Record an LLM API call"""
        call_data = {
            "operation": operation,
            "model": model,
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "total_tokens": input_tokens + output_tokens,
            "success": success,
            "timestamp": time.time()
        }
        self.calls.append(call_data)

        if success:
            self.total_input_tokens += input_tokens
            self.total_output_tokens += output_tokens
            self.total_calls += 1

    def get_summary(self) -> Dict[str, Any]:
        """Get a summary of all LLM usage"""
        successful_calls = [call for call in self.calls if call["success"]]
        failed_calls = [call for call in self.calls if not call["success"]]

        operations_summary = {}
        for call in successful_calls:
            op = call["operation"]
            if op not in operations_summary:
                operations_summary[op] = {
                    "count": 0,
                    "input_tokens": 0,
                    "output_tokens": 0,
                    "total_tokens": 0
                }
            operations_summary[op]["count"] += 1
            operations_summary[op]["input_tokens"] += call["input_tokens"]
            operations_summary[op]["output_tokens"] += call["output_tokens"]
            operations_summary[op]["total_tokens"] += call["total_tokens"]

        return {
            "total_calls": len(successful_calls),
            "failed_calls": len(failed_calls),
            "total_input_tokens": self.total_input_tokens,
            "total_output_tokens": self.total_output_tokens,
            "total_tokens": self.total_input_tokens + self.total_output_tokens,
            "operations_breakdown": operations_summary,
            "calls_detail": self.calls
        }

class ProgressTracker:
    def __init__(self):
        self.messages = []
        self.current_step = 0
        self.total_steps = 0
        self.token_tracker = TokenTracker()

    def emit(self, message: str, step_type: str = "info"):
        """Emit a progress message"""
        progress_data = {
            "step": self.current_step,
            "total": self.total_steps,
            "message": message,
            "type": step_type,
            "timestamp": __import__('time').time()
        }
        self.messages.append(progress_data)
        print(f"📢 Progress: {message}")
        return progress_data

    def set_total_steps(self, total: int):
        """Set total number of steps"""
        self.total_steps = total

    def next_step(self, message: str = None):
        """Move to next step with optional message"""
        self.current_step += 1
        if message:
            return self.emit(message)

    def error(self, message: str):
        """Emit error message"""
        return self.emit(message, "error")

    def success(self, message: str):
        """Emit success message"""
        return self.emit(message, "success")

    def warning(self, message: str):
        """Emit warning message"""
        return self.emit(message, "warning")

def extract_token_usage(response, operation: str = "unknown", model: str = "unknown") -> Dict[str, int]:
    """Extract token usage information from Gemini API response"""
    try:
        input_tokens = 0
        output_tokens = 0

        # Try to extract usage metadata from response
        if hasattr(response, 'usage_metadata') and response.usage_metadata:
            usage = response.usage_metadata
            input_tokens = getattr(usage, 'prompt_token_count', 0) or 0
            output_tokens = getattr(usage, 'candidates_token_count', 0) or 0
        elif hasattr(response, '_metadata') and response._metadata and hasattr(response._metadata, 'usage') and response._metadata.usage:
            usage = response._metadata.usage
            input_tokens = getattr(usage, 'prompt_token_count', 0) or 0
            output_tokens = getattr(usage, 'completion_token_count', 0) or 0

        # Fallback: estimate tokens based on text length (rough approximation)
        if input_tokens == 0 or output_tokens == 0:
            if hasattr(response, 'text') and response.text:
                # Rough estimation: ~4 characters per token
                output_tokens = max(output_tokens, len(response.text) // 4)
            elif hasattr(response, 'candidates') and response.candidates:
                text_length = 0
                try:
                    for candidate in response.candidates:
                        if candidate and hasattr(candidate, 'content') and candidate.content:
                            if hasattr(candidate.content, 'parts') and candidate.content.parts:
                                for part in candidate.content.parts:
                                    if part and hasattr(part, 'text') and part.text:
                                        text_length += len(part.text)
                except (TypeError, AttributeError) as e:
                    print(f"⚠️ Error iterating candidates for token estimation: {e}")
                    text_length = 0
                output_tokens = max(output_tokens, text_length // 4)

        return {
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "total_tokens": input_tokens + output_tokens
        }
    except Exception as e:
        print(f"⚠️ Error extracting token usage: {e}")
        return {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0}

# Load configuration
def load_config():
    try:
        with open('config.json', 'r') as f:
            config = json.load(f)
        print("✅ Configuration loaded successfully from config.json")
        return config
    except FileNotFoundError:
        print("⚠️ config.json not found, using environment variables")
        # Fallback to environment variables for Cloud Run
        return {
            "gcs": {
                "bucket_name": os.getenv('GCS_BUCKET_NAME', 'your_resume_storage_bucket'),
                "default_folder": os.getenv('GCS_DEFAULT_FOLDER', 'resume/'),
                "credentials_path": os.getenv('GOOGLE_APPLICATION_CREDENTIALS', 'service-account.json')
            },
            "vector_search": {
                "project_id": os.getenv('GOOGLE_CLOUD_PROJECT', 'your-gcp-project-id'),
                "location": os.getenv('VECTOR_SEARCH_LOCATION', 'global'),
                "datastore_id": os.getenv('VECTOR_SEARCH_DATASTORE_ID', 'your_resume_datastore_id')
            },
            "postgresql": {
                "host": os.getenv('DB_HOST', 'your_db_host_ip'),
                "cloud_sql_connection_name": os.getenv('CLOUD_SQL_CONNECTION_NAME', 'your-gcp-project-id:me-central1:your-cloudsql-instance'),
                "port": int(os.getenv('DB_PORT', 5432)),
                "database": os.getenv('DB_NAME', 'postgres'),
                "user": os.getenv('DB_USER', 'postgres'),
                "password": os.getenv('DB_PASSWORD', 'your_db_password_here'),
                "ssl_mode": os.getenv('DB_SSL_MODE', 'require')
            }
        }
    except json.JSONDecodeError:
        print("❌ Invalid configuration file format")
        raise HTTPException(status_code=500, detail="Invalid configuration file format")
    except Exception as e:
        print(f"❌ Error loading configuration: {e}")
        raise HTTPException(status_code=500, detail=f"Configuration error: {str(e)}")

# Load configuration with error handling
try:
    config = load_config()
except Exception as e:
    print(f"❌ Critical error loading configuration: {e}")
    # For development/testing, create minimal config
    config = {
        "gcs": {
            "bucket_name": "your_resume_storage_bucket",
            "default_folder": "resume/",
            "credentials_path": "service-account.json"
        },
        "vector_search": {
            "project_id": "your-gcp-project-id",
            "location": "global",
            "datastore_id": "your_resume_datastore_id"
        }
    }

# Google Cloud Storage configuration from config.json
CREDENTIALS_PATH = config['gcs']['credentials_path']
GCS_BUCKET_NAME = config['gcs']['bucket_name']
DEFAULT_FOLDER = config['gcs']['default_folder']

# Vector Search configuration from config.json
PROJECT_ID = config['vector_search']['project_id']
LOCATION = config['vector_search']['location']
DATASTORE_ID = config['vector_search']['datastore_id']


def gemini_client():
    """Construct a resilient LLM client (Gemini → OpenAI fallback).

    Returns a ``_ResilientLLMClient`` whose ``.models.generate_content(...)`` API
    matches ``google.genai`` exactly, so all 13+ existing call sites work
    unchanged. Behavior:

      * Gemini path used first (Vertex if ADC creds present, else AI Studio with
        ``GEMINI_API_KEY``).
      * On auth/quota errors (``PERMISSION_DENIED``, ``UNAUTHENTICATED``,
        ``RESOURCE_EXHAUSTED``, suspended-key, 401/403/429), trips a 5-minute
        circuit and routes subsequent calls to OpenAI (``OPENAI_API_KEY`` /
        ``OPENAI_MODEL``).
      * If both providers are unavailable, raises a clear error.

    This protects production from a single-provider outage (e.g. a suspended
    API key) without requiring ``SMARTHR_LOCAL_MODE=1``.
    """
    has_adc = bool(
        os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")
        or os.environ.get("GOOGLE_CLOUD_PROJECT")
    )
    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_GEMINI_API_KEY")
    real = None
    try:
        if not has_adc and api_key:
            real = genai.Client(api_key=api_key)
        else:
            real = genai.Client(
                vertexai=True,
                project=PROJECT_ID,
                location="europe-west4",
            )
    except Exception as e:
        # Don't let a Gemini construction error prevent startup if OpenAI works.
        print(f"⚠️ Gemini client construction failed ({e}); will use OpenAI fallback only.")
        real = None
    return _ResilientLLMClient(real)


# === Resilient LLM wrapper (Gemini → OpenAI fallback) ====================
# Module-level circuit-breaker state. When tripped, we skip Gemini entirely
# and go straight to OpenAI for ``trip_seconds`` to avoid wasting RTTs on a
# known-bad key.
_gemini_circuit_state = {"trip_until": 0.0, "fail_count": 0, "last_error": ""}
_GEMINI_PERMANENT_PATTERNS = (
    "permission_denied", "unauthenticated", "api_key_invalid",
    "suspended", "billing", "401 ", "403 ", "429 ",
    "resource_exhausted", "quota",
)


def _is_gemini_auth_or_quota_error(exc: Exception) -> bool:
    msg = str(exc).lower()
    return any(p in msg for p in _GEMINI_PERMANENT_PATTERNS)


def _trip_gemini_circuit(seconds: float = 300.0, err: str = "") -> None:
    _gemini_circuit_state["trip_until"] = time.time() + seconds
    _gemini_circuit_state["fail_count"] += 1
    _gemini_circuit_state["last_error"] = (err or "")[:300]


def _gemini_circuit_open() -> bool:
    return time.time() < _gemini_circuit_state["trip_until"]


_openai_client_singleton = None


def _get_openai_client():
    global _openai_client_singleton
    if _openai_client_singleton is not None:
        return _openai_client_singleton
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return None
    try:
        from openai import OpenAI
        _openai_client_singleton = OpenAI(
            api_key=api_key,
            base_url=os.environ.get("OPENAI_BASE_URL") or None,
        )
        return _openai_client_singleton
    except Exception as e:
        print(f"⚠️ OpenAI client init failed: {e}")
        return None


def _flatten_gemini_contents(contents) -> str:
    """Best-effort flatten of Gemini ``contents`` (str | list[Content]) → prompt str."""
    if isinstance(contents, str):
        return contents
    out = []
    try:
        for c in contents:
            if hasattr(c, "parts"):
                for p in c.parts:
                    t = getattr(p, "text", None)
                    if t:
                        out.append(t)
            elif isinstance(c, str):
                out.append(c)
    except Exception:
        return str(contents)
    return "\n".join(out)


class _UsageMetadataAdapter:
    __slots__ = ("prompt_token_count", "candidates_token_count", "total_token_count")

    def __init__(self, p=0, c=0, t=0):
        self.prompt_token_count = p
        self.candidates_token_count = c
        self.total_token_count = t


class _PartAdapter:
    __slots__ = ("text",)

    def __init__(self, text=""):
        self.text = text


class _ContentAdapter:
    __slots__ = ("role", "parts")

    def __init__(self, role="model", parts=None):
        self.role = role
        self.parts = parts or []


class _CandidateAdapter:
    __slots__ = ("content", "finish_reason", "index")

    def __init__(self, content=None, finish_reason="STOP", index=0):
        self.content = content or _ContentAdapter()
        self.finish_reason = finish_reason
        self.index = index


class _OpenAIResponseAdapter:
    """Wraps an OpenAI ChatCompletion to look like a ``genai`` response."""

    def __init__(self, openai_resp, fallback_provider="openai"):
        try:
            self.text = openai_resp.choices[0].message.content or ""
        except Exception:
            self.text = ""
        try:
            u = openai_resp.usage
            self.usage_metadata = _UsageMetadataAdapter(
                p=getattr(u, "prompt_tokens", 0) or 0,
                c=getattr(u, "completion_tokens", 0) or 0,
                t=getattr(u, "total_tokens", 0) or 0,
            )
        except Exception:
            self.usage_metadata = _UsageMetadataAdapter()
        self.candidates = [
            _CandidateAdapter(
                content=_ContentAdapter(role="model", parts=[_PartAdapter(text=self.text)]),
                finish_reason="STOP",
                index=0,
            )
        ]
        self.prompt_feedback = None
        self._provider = fallback_provider


def _openai_generate(model_name, contents, config):
    client = _get_openai_client()
    if client is None:
        raise RuntimeError(
            "Gemini unavailable and OPENAI_API_KEY not set — no LLM provider available."
        )
    prompt = _flatten_gemini_contents(contents)
    openai_model = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")
    kw = {"model": openai_model}
    if config is not None:
        if getattr(config, "temperature", None) is not None:
            try:
                kw["temperature"] = float(config.temperature)
            except Exception:
                pass
        if getattr(config, "top_p", None) is not None:
            try:
                kw["top_p"] = float(config.top_p)
            except Exception:
                pass
        if getattr(config, "max_output_tokens", None):
            try:
                kw["max_tokens"] = min(int(config.max_output_tokens), 16384)
            except Exception:
                pass
        if getattr(config, "response_mime_type", None) == "application/json":
            kw["response_format"] = {"type": "json_object"}
            if "json" not in prompt.lower():
                prompt = prompt + "\n\nRespond ONLY with valid JSON."
    kw["messages"] = [{"role": "user", "content": prompt}]
    resp = client.chat.completions.create(**kw)
    return _OpenAIResponseAdapter(resp)


class _ResilientModels:
    def __init__(self, real_models):
        self._gemini = real_models  # may be None

    def generate_content(self, model=None, contents=None, config=None, **kwargs):
        # Try Gemini first if available and circuit closed.
        if self._gemini is not None and not _gemini_circuit_open():
            try:
                return self._gemini.generate_content(
                    model=model, contents=contents, config=config, **kwargs
                )
            except Exception as e:
                if _is_gemini_auth_or_quota_error(e):
                    _trip_gemini_circuit(300.0, str(e))
                    print(
                        f"⚠️ Gemini auth/quota error → tripping circuit 5min, "
                        f"using OpenAI fallback. err={str(e)[:160]}"
                    )
                else:
                    print(
                        f"⚠️ Gemini transient error → using OpenAI for this call. "
                        f"err={str(e)[:160]}"
                    )
                # fall through to OpenAI
        elif _gemini_circuit_open():
            # Quiet on the hot path; the trip already logged.
            pass
        return _openai_generate(model, contents, config)

    def list(self, *args, **kwargs):
        """Used by ``_gemini_ping_cached``; report which provider is live."""
        if self._gemini is not None and not _gemini_circuit_open():
            try:
                return self._gemini.list(*args, **kwargs)
            except Exception as e:
                if _is_gemini_auth_or_quota_error(e):
                    _trip_gemini_circuit(300.0, str(e))
                # fall through to OpenAI synthetic ping
        # Synthetic non-empty iterable so callers see "alive" if OpenAI works.
        if _get_openai_client() is not None:
            return iter([{"name": "openai-fallback", "provider": "openai"}])
        # Neither provider available — re-raise something useful.
        raise RuntimeError(
            f"No LLM provider available (gemini circuit open={_gemini_circuit_open()}, "
            f"openai configured={bool(os.environ.get('OPENAI_API_KEY'))})"
        )


class _ResilientLLMClient:
    """Drop-in replacement for ``genai.Client`` with provider fallback."""

    def __init__(self, real_client):
        self._real = real_client
        self.models = _ResilientModels(real_client.models if real_client is not None else None)

    # Some call sites may inspect attributes; pass through to real client when present.
    def __getattr__(self, name):
        if self._real is not None and hasattr(self._real, name):
            return getattr(self._real, name)
        raise AttributeError(name)
# === End resilient LLM wrapper ===========================================



# Set up GCS client with credentials handling
def setup_gcs_credentials():
    """Setup Google Cloud credentials with fallback options"""
    try:
        # VPS local-mode: storage/discovery/vertex are shimmed to local backends,
        # so no GCP credentials are needed. Report success silently.
        if os.environ.get("SMARTHR_LOCAL_MODE", "").lower() in ("1", "true", "yes", "on"):
            print("🏠 SMARTHR_LOCAL_MODE=1 - skipping GCP credential setup (shims active)")
            return True

        # First try the credentials file path
        if os.path.exists(CREDENTIALS_PATH):
            os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = CREDENTIALS_PATH
            print(f"✅ Using credentials file: {CREDENTIALS_PATH}")
            return True

        # If running on Cloud Run, default credentials should work
        if os.getenv('K_SERVICE') or os.getenv('GOOGLE_CLOUD_PROJECT'):
            print("🌩️ Running on Cloud Run - using default service account credentials")
            return True

        # Check if credentials are already set in environment
        if os.getenv('GOOGLE_APPLICATION_CREDENTIALS'):
            print(f"✅ Using credentials from environment variable")
            return True

        print(f"⚠️ Credentials file not found at {CREDENTIALS_PATH} - some features may not work")
        return False
    except Exception as e:
        print(f"⚠️ Error setting up credentials: {e}")
        return False

# Setup credentials
credentials_available = setup_gcs_credentials()

# ---------------------------------------------------------------------------
# Cached GCS storage client.
# Constructing google.cloud.storage.Client() is expensive (~0.5-2s for the
# auth handshake + HTTP pool init). The HR scorecard pipeline used to call
# storage.Client() ~10 times per request (once per candidate during path
# resolution + once per download), serialising on the auth code path and
# adding 10-30s of latency on warm requests. The Client object itself is
# documented as thread-safe for read operations (download, list, exists),
# so a process-wide singleton is the right shape.
# ---------------------------------------------------------------------------
_storage_client_lock = threading.Lock()
_storage_client_singleton: "storage.Client | None" = None


def _get_storage_client() -> "storage.Client":
    """Return a process-wide cached google.cloud.storage.Client.

    Falls back to creating a fresh client on the rare path where the
    singleton can't be initialised (e.g. import-time credential errors).
    """
    global _storage_client_singleton
    if _storage_client_singleton is not None:
        return _storage_client_singleton
    with _storage_client_lock:
        if _storage_client_singleton is None:
            _storage_client_singleton = storage.Client()
    return _storage_client_singleton


def ensure_bucket_folder_exists():
    """Ensure the resume folder exists in the bucket"""
    if not credentials_available:
        print("⚠️ Skipping bucket folder creation - no credentials available")
        return False

    try:
        client = _get_storage_client()
        bucket = client.bucket(GCS_BUCKET_NAME)

        # Check if bucket exists
        if not bucket.exists():
            print(f"⚠️ Bucket {GCS_BUCKET_NAME} does not exist")
            return False

        # Create a placeholder file in the resume folder to ensure it exists
        folder_placeholder = f"{DEFAULT_FOLDER}.gitkeep"
        blob = bucket.blob(folder_placeholder)

        # Only create if it doesn't exist
        if not blob.exists():
            blob.upload_from_string("", content_type="text/plain")
            print(f"✅ Created folder: {DEFAULT_FOLDER}")
        else:
            print(f"✅ Folder {DEFAULT_FOLDER} already exists")

        return True
    except Exception as e:
        print(f"⚠️ Error ensuring folder exists: {str(e)}")
        return False

def create_company_gcs_bucket(company_code: str) -> str:
    """
    Create a company-specific GCS bucket for data isolation.
    Returns the bucket name if successful, None if failed.
    """
    try:
        from google.cloud import storage
        from google.cloud.exceptions import Conflict

        client = _get_storage_client()
        # GCS bucket names must be lowercase and follow DNS naming conventions
        bucket_name = f"{company_code.lower()}-resume-storage"

        print(f"\n🪣 CREATING COMPANY GCS BUCKET")
        print(f"🏢 Company: {company_code}")
        print(f"📦 Bucket: {bucket_name}")

        # Create bucket with proper configuration
        bucket = client.bucket(bucket_name)
        bucket.storage_class = "STANDARD"

        # Create the bucket with location parameter
        bucket = client.create_bucket(bucket, location="US")

        # Create initial folder structure
        resume_folder = f"resume/.gitkeep"
        blob = bucket.blob(resume_folder)
        blob.upload_from_string("# Company resume folder")

        print(f"✅ Successfully created GCS bucket: {bucket_name}")
        return bucket_name

    except Conflict:
        # Bucket already exists
        print(f"⚠️  Bucket {bucket_name} already exists, using existing bucket")
        return bucket_name
    except Exception as e:
        print(f"❌ Failed to create GCS bucket for {company_code}: {e}")
        return None

def create_company_datastore(company_code: str, project_id: str, location: str = "europe-west4") -> str:
    """
    Create a company-specific AI Search datastore for search isolation.
    Returns the datastore ID if successful, None if failed.
    """
    try:
        from google.cloud import discoveryengine_v1
        from google.api_core.client_options import ClientOptions
        from google.api_core.exceptions import AlreadyExists
        import time

        print(f"\n🔍 CREATING COMPANY AI SEARCH DATASTORE")
        print(f"🏢 Company: {company_code}")
        print(f"🎯 Project: {project_id}")
        print(f"🌍 Location: {location}")

        # Initialize the client
        client_options = (
            ClientOptions(api_endpoint=f"{location}-discoveryengine.googleapis.com")
            if location != "global"
            else None
        )

        client = discoveryengine_v1.DataStoreServiceClient(client_options=client_options)

        # Generate unique datastore ID (must match pattern [a-z0-9][a-z0-9-_]*)
        timestamp = int(time.time())
        # Ensure company_code is lowercase and replace any invalid characters
        clean_company_code = company_code.lower().replace('_', '-')
        datastore_id = f"{clean_company_code}-resume-datastore-{timestamp}"

        print(f"📚 Datastore ID: {datastore_id}")

        # Create datastore configuration
        data_store = discoveryengine_v1.DataStore(
            display_name=f"{company_code.title()} Resume Datastore",
            industry_vertical=discoveryengine_v1.IndustryVertical.GENERIC,
            content_config=discoveryengine_v1.DataStore.ContentConfig.CONTENT_REQUIRED,
            solution_types=[discoveryengine_v1.SolutionType.SOLUTION_TYPE_SEARCH],
        )

        # Parent path for the datastore
        parent = f"projects/{project_id}/locations/{location}/collections/default_collection"

        # Create the datastore
        operation = client.create_data_store(
            parent=parent,
            data_store=data_store,
            data_store_id=datastore_id
        )

        print(f"⏳ Creating datastore... This may take a few minutes.")

        # Wait for operation to complete (timeout after 10 minutes)
        result = operation.result(timeout=600)

        print(f"✅ Successfully created AI Search datastore: {datastore_id}")

        # Note: Serving config (default_config) should be created automatically
        print(f"🔧 Datastore created - default serving config should be available automatically")

        return datastore_id

    except AlreadyExists:
        print(f"⚠️  Datastore {datastore_id} already exists, using existing datastore")
        return datastore_id
    except Exception as e:
        print(f"❌ Failed to create AI Search datastore for {company_code}: {e}")
        return None

def get_company_gcs_bucket(company_code: str) -> str:
    """Get the GCS bucket name for a company"""
    return f"{company_code.lower()}-resume-storage"

def get_company_datastore_id(company_code: str, db_manager=None) -> str:
    """Get the datastore ID for a company from the database"""
    if not db_manager:
        db_manager = get_db_manager()

    company = db_manager.get_company_by_code(company_code)
    if company and company.get('datastore_id'):
        return company['datastore_id']

    # Fallback to default datastore if not set
    return DATASTORE_ID


def build_discovery_content_search_spec(discoveryengine_module, summary_result_count: int = 3):
    """Build Discovery content search config.

    Snippets are required by the UI and downstream analysis. Discovery
    summaries are intentionally disabled by default because none of the app's
    core flows consume response.summary, and generated summaries add seconds
    of latency to every search request. Set ENABLE_DISCOVERY_SUMMARIES=true to
    restore the old behavior for experiments.
    """
    content_kwargs = {
        "snippet_spec": discoveryengine_module.SearchRequest.ContentSearchSpec.SnippetSpec(
            return_snippet=True
        )
    }
    if ENABLE_DISCOVERY_SUMMARIES:
        content_kwargs["summary_spec"] = discoveryengine_module.SearchRequest.ContentSearchSpec.SummarySpec(
            summary_result_count=summary_result_count,
            include_citations=True,
        )
    return discoveryengine_module.SearchRequest.ContentSearchSpec(**content_kwargs)


def get_company_resources(user_context: dict = None) -> tuple:
    """
    Get company-specific GCS bucket and datastore ID based on user context.
    Returns (bucket_name, datastore_id)
    """
    if not user_context or not user_context.get('company'):
        if user_context and user_context.get('user_type') != 'super_admin':
            raise HTTPException(status_code=403, detail="Company context required")
        # Super admin or system access - use default resources
        return GCS_BUCKET_NAME, DATASTORE_ID

    # Safe access to company information
    company_info = user_context.get('company', {})
    company_code = None
    company_id = None

    if isinstance(company_info, dict):
        company_code = company_info.get('company_code')
        company_id = company_info.get('company_id') or company_info.get('id')
    else:
        company_code = getattr(company_info, 'company_code', None)
        company_id = getattr(company_info, 'id', None) or getattr(company_info, 'company_id', None)

    print(f"🔍 Company resources lookup: code={company_code}, id={company_id}")

    # Tenant-scoped calls must fail closed when company context is malformed.
    if not company_code and not company_id:
        print(f"⚠️ No company information found for tenant-scoped request")
        raise HTTPException(status_code=403, detail="Company context required")

    # Get company details from database if we have company_id
    if company_id:
        db = get_db_manager()
        company = db.get_company_by_id(company_id)

        if company:
            bucket_name = company.get('gcs_bucket_name') or get_company_gcs_bucket(company_code)
            datastore_id = company.get('datastore_id') or get_company_datastore_id(company_code, db)
            print(f"✅ Found company resources: bucket={bucket_name}, datastore={datastore_id}")
            return bucket_name, datastore_id

    # If we only have company_code, use it to generate resources
    if company_code:
        bucket_name = get_company_gcs_bucket(company_code)
        datastore_id = get_company_datastore_id(company_code)
        print(f"✅ Generated company resources from code: bucket={bucket_name}, datastore={datastore_id}")
        return bucket_name, datastore_id

    print(f"⚠️ Unable to resolve company resources")
    raise HTTPException(status_code=403, detail="Company resources unavailable")

# Cloud Task Management Functions
# =================================

def get_cloud_task_client():
    """Get Google Cloud Tasks client"""
    if tasks_v2 is None:
        raise RuntimeError(
            "google-cloud-tasks is not installed; set USE_CLOUD_TASKS=false "
            "for local/VPS dispatch or add google-cloud-tasks to requirements."
        )

    config = load_config()
    if not config or 'cloud_tasks' not in config:
        raise Exception("Cloud Tasks configuration not found in config.json")

    return tasks_v2.CloudTasksClient()

def create_hr_scorecard_task(
    task_id: str,
    query: str,
    job_title: str,
    result_count: int,
    user_data: dict
) -> Dict[str, Any]:
    """
    Create a cloud task for HR scorecard processing

    Args:
        task_id: Unique identifier for the task
        query: Job description query
        job_title: Job title
        result_count: Number of results to process
        user_data: User information (excluding sensitive data)

    Returns:
        Dict with task information
    """
    # VPS / self-hosted mode: dispatch to local process instead of Cloud Tasks.
    # Enabled by setting USE_CLOUD_TASKS=false in the environment.
    if os.environ.get("USE_CLOUD_TASKS", "true").strip().lower() in ("false", "0", "no"):
        try:
            local_url = os.environ.get(
                "LOCAL_TASK_DISPATCH_URL",
                "http://127.0.0.1:8081/api/process-hr-scorecard-task",
            )
            task_payload = {
                'task_id': task_id,
                'query': query,
                'job_title': job_title,
                'result_count': result_count,
                'user_data': {
                    'id': user_data.get('id'),
                    'email': user_data.get('email'),
                    'company': user_data.get('company'),
                    'user_type': user_data.get('user_type')
                },
                'created_at': datetime.utcnow().isoformat(),
            }

            def _fire_local_task():
                try:
                    import requests as _rq
                    _rq.post(
                        local_url,
                        json=task_payload,
                        headers={'X-Internal-Task': '1'},
                        timeout=(5, 600),
                    )
                except Exception as _e:
                    print(f"⚠️  Local task dispatch failed: {_e}")
                    try:
                        update_task_status(task_id, 'failed', error=str(_e))
                    except Exception:
                        pass

            import threading as _th
            _th.Thread(target=_fire_local_task, daemon=True).start()
            print(f"✅ Local task dispatched (no Cloud Tasks): {task_id}")
            return {
                'success': True,
                'task_name': f"local/{task_id}",
                'task_id': task_id,
                'queue_name': 'local-inproc',
                'scheduled_time': None,
            }
        except Exception as e:
            print(f"❌ Local task dispatch error: {str(e)}")
            return {'success': False, 'error': str(e), 'task_id': task_id}

    try:
        config = load_config()
        cloud_tasks_config = config['cloud_tasks']

        # Create task client
        client = get_cloud_task_client()

        # Construct the fully qualified queue name
        parent = client.queue_path(
            cloud_tasks_config['project_id'],
            cloud_tasks_config['location'],
            cloud_tasks_config['queue_name']
        )

        # Create task payload
        task_payload = {
            'task_id': task_id,
            'query': query,
            'job_title': job_title,
            'result_count': result_count,
            'user_data': {
                'id': user_data.get('id'),
                'email': user_data.get('email'),
                'company': user_data.get('company'),
                'user_type': user_data.get('user_type')
            },
            'created_at': datetime.utcnow().isoformat()
        }

        # Create HTTP request for the task
        target_url = f"{cloud_tasks_config['target_uri']}/api/process-hr-scorecard-task"
        http_request = {
            'http_method': tasks_v2.HttpMethod.POST,
            'url': target_url,
            'headers': {
                'Content-Type': 'application/json',
                'User-Agent': 'HRAgent-CloudTasks/1.0'
            },
            'body': json.dumps(task_payload).encode('utf-8')
        }
        # Attach an OIDC token so the worker can verify the call really came
        # from Cloud Tasks running as our runtime service account.
        sa_email = cloud_tasks_config.get('service_account_email')
        if sa_email:
            http_request['oidc_token'] = {
                'service_account_email': sa_email,
                'audience': target_url,
            }
        task = {
            'name': f"{parent}/tasks/{task_id}",
            'http_request': http_request,
        }

        # Note: previously this scheduled the task 10s in the future, which
        # added a flat 10s of latency to every search. Cloud Tasks dispatches
        # immediately by default, so we omit schedule_time entirely.

        # Create the task
        response = client.create_task(parent=parent, task=task)

        print(f"✅ Cloud task created: {response.name}")

        return {
            'success': True,
            'task_name': response.name,
            'task_id': task_id,
            'queue_name': cloud_tasks_config['queue_name'],
            'scheduled_time': None,
        }

    except Exception as e:
        print(f"❌ Failed to create cloud task: {str(e)}")
        return {
            'success': False,
            'error': str(e),
            'task_id': task_id
        }

def update_task_status(task_id: str, status: str, progress: dict = None, error: str = None) -> bool:
    """
    Update task status in the database

    Args:
        task_id: Task identifier
        status: Task status (pending, running, completed, failed)
        progress: Progress information
        error: Error message if failed

    Returns:
        bool: Success status
    """
    try:
        db = get_db_manager()

        # Update task status
        update_data = {
            'status': status,
            'updated_at': datetime.utcnow().isoformat()
        }

        if progress:
            update_data['progress'] = json.dumps(progress)

        if error:
            update_data['error_message'] = error

        if status == 'completed':
            update_data['completed_at'] = datetime.utcnow().isoformat()

        # Update task status in database
        # Note: This assumes the database manager has a method to update task status
        success = db.update_task_status(task_id, update_data)

        if success:
            print(f"✅ Task status updated: {task_id} -> {status}")
        else:
            print(f"❌ Failed to update task status: {task_id}")

        return success

    except Exception as e:
        print(f"❌ Error updating task status: {str(e)}")
        return False

def save_task_to_database(
    task_id: str,
    query: str,
    job_title: str,
    result_count: int,
    user_id: str,
    company_id: int = None
) -> bool:
    """
    Save task information to database

    Args:
        task_id: Task identifier
        query: Job description query
        job_title: Job title
        result_count: Number of results requested
        user_id: User ID who initiated the task
        company_id: Company ID (optional)

    Returns:
        bool: Success status
    """
    try:
        db = get_db_manager()

        # Save task to database
        success = db.save_hr_scorecard_task(
            task_id=task_id,
            query=query,
            job_title=job_title,
            result_count=result_count,
            user_id=user_id,
            company_id=company_id,
            status='pending',
            created_at=datetime.utcnow().isoformat()
        )

        if success:
            print(f"✅ Task saved to database: {task_id}")
        else:
            print(f"❌ Failed to save task to database: {task_id}")

        return success

    except Exception as e:
        print(f"❌ Error saving task to database: {str(e)}")
        return False

def extract_text_from_file(file_content: bytes, file_path: str) -> str:
    """Extract text from PDF, DOCX, or text files with robust error handling"""
    if not file_content:
        print(f"❌ No file content provided for extraction")
        return None

    file_extension = file_path.lower().split('.')[-1] if file_path else 'unknown'
    file_size = len(file_content)

    print(f"📄 Extracting text from {file_extension.upper()} file (size: {file_size:,} bytes)...")

    try:
        if file_extension == 'pdf':
            extracted_text = extract_text_from_pdf(file_content)
            if extracted_text:
                print(f"✅ PDF text extraction successful (extracted {len(extracted_text):,} characters)")
                return extracted_text
            else:
                print(f"❌ PDF text extraction failed -  no readable text found")
                return None

        elif file_extension in ['docx', 'doc']:
            extracted_text = extract_text_from_docx(file_content)
            if extracted_text:
                print(f"✅ DOCX text extraction successful (extracted {len(extracted_text):,} characters)")
                return extracted_text
            else:
                print(f"❌ DOCX text extraction failed - no readable text found")
                return None

        elif file_extension == 'txt':
            try:
                # Try UTF-8 first
                extracted_text = file_content.decode('utf-8')
                print(f"✅ TXT file decoded as UTF-8 (length: {len(extracted_text):,} characters)")
                return extracted_text
            except UnicodeDecodeError:
                try:
                    # Fallback to latin-1
                    extracted_text = file_content.decode('latin-1')
                    print(f"✅ TXT file decoded as latin-1 (length: {len(extracted_text):,} characters)")
                    return extracted_text
                except UnicodeDecodeError:
                    # Final fallback with error handling
                    extracted_text = file_content.decode('utf-8', errors='ignore')
                    print(f"✅ TXT file decoded with error handling (length: {len(extracted_text):,} characters)")
                    return extracted_text

        else:
            # Try to decode as text for unknown file types
            try:
                extracted_text = file_content.decode('utf-8')
                print(f"✅ Unknown file type decoded as UTF-8 (length: {len(extracted_text):,} characters)")
                return extracted_text
            except UnicodeDecodeError:
                try:
                    extracted_text = file_content.decode('utf-8', errors='ignore')
                    print(f"⚠️ Unknown file type decoded with errors ignored (length: {len(extracted_text):,} characters)")
                    return extracted_text if len(extracted_text.strip()) > 10 else None
                except Exception as e:
                    print(f"❌ Could not decode unknown file type: {str(e)}")
                    return None

    except Exception as e:
        print(f"❌ Unexpected error during text extraction: {str(e)}")
        print(f"📊 File details: extension={file_extension}, size={file_size:,} bytes, path={file_path}")
        return None

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF using multiple robust methods"""
    # Suppress PyPDF2 warnings for cleaner output
    warnings.filterwarnings("ignore", category=UserWarning, module="PyPDF2")
    logging.getLogger("PyPDF2").setLevel(logging.ERROR)

    text = ""

    # Method 1: Try pdfplumber first (better for complex PDFs)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as page_error:
                    print(f"⚠️ pdfplumber page {page_num + 1} extraction failed: {str(page_error)}")
                    continue

        if text.strip():
            print(f"✅ Extracted text using pdfplumber (length: {len(text)}, pages: {len(pdf.pages)})")
            return text.strip()
        else:
            print(f"⚠️ pdfplumber extracted empty text from {len(pdf.pages)} pages")
    except Exception as e:
        print(f"⚠️ pdfplumber failed: {str(e)}")

    # Method 2: Try PyPDF2 with improved error handling
    try:
        import PyPDF2
        from PyPDF2.errors import PdfReadError, PdfReadWarning

        # Suppress PyPDF2 warnings
        warnings.filterwarnings("ignore", category=PdfReadWarning)

        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))

        print(f"📊 PDF has {len(pdf_reader.pages)} pages")

        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as page_error:
                print(f"⚠️ PyPDF2 page {page_num + 1} extraction failed: {str(page_error)}")
                continue

        if text.strip():
            print(f"✅ Extracted text using PyPDF2 (length: {len(text)}, pages: {len(pdf_reader.pages)})")
            return text.strip()
        else:
            print(f"⚠️ PyPDF2 extracted empty text from {len(pdf_reader.pages)} pages")

    except PdfReadError as e:
        print(f"❌ PyPDF2 PDF read error: {str(e)}")
    except Exception as e:
        print(f"❌ PyPDF2 general error: {str(e)}")

    # Method 3: Try alternative PyPDF2 approach (different constructor)
    try:
        import PyPDF2

        file_stream = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(file_stream, strict=False)  # Non-strict mode

        print(f"📊 Alternative PDF reader found {len(pdf_reader.pages)} pages")

        alt_text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    alt_text += page_text + "\n"
            except Exception:
                continue

        if alt_text.strip():
            print(f"✅ Extracted text using PyPDF2 (non-strict) (length: {len(alt_text)})")
            return alt_text.strip()

    except Exception as e:
        print(f"⚠️ PyPDF2 non-strict mode also failed: {str(e)}")

    # Method 4: Try basic text extraction from PDF streams
    try:
        # Sometimes PDFs have text that can be extracted with basic methods
        content_str = file_content.decode('utf-8', errors='ignore')
        if len(content_str) > 100 and any(char.isalnum() for char in content_str):
            # Extract readable text patterns
            import re
            readable_text = re.findall(r'[A-Za-z0-9\s\.,;:!?\-]{10,}', content_str)
            if readable_text:
                fallback_text = ' '.join(readable_text)
                if len(fallback_text) > 50:
                    print(f"✅ Extracted text using fallback method (length: {len(fallback_text)})")
                    return fallback_text
    except Exception as e:
        print(f"⚠️ Fallback text extraction failed: {str(e)}")

    print(f"❌ All PDF extraction methods failed")
    return None

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX files with robust error handling"""
    try:
        doc = Document(io.BytesIO(file_content))

        # Extract text from paragraphs
        paragraphs_text = []
        for i, paragraph in enumerate(doc.paragraphs):
            try:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    paragraphs_text.append(paragraph.text.strip())
            except Exception as para_error:
                print(f"⚠️ DOCX paragraph {i+1} extraction failed: {str(para_error)}")
                continue

        # Extract text from tables if any
        tables_text = []
        try:
            for table_num, table in enumerate(doc.tables):
                try:
                    table_rows = []
                    for row in table.rows:
                        row_cells = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_cells.append(cell.text.strip())
                        if row_cells:
                            table_rows.append(" | ".join(row_cells))
                    if table_rows:
                        tables_text.extend(table_rows)
                except Exception as table_error:
                    print(f"⚠️ DOCX table {table_num+1} extraction failed: {str(table_error)}")
                    continue
        except Exception as tables_error:
            print(f"⚠️ DOCX tables extraction failed: {str(tables_error)}")

        # Combine all text
        all_text = paragraphs_text + tables_text

        if all_text:
            final_text = "\n".join(all_text)
            print(f"✅ Extracted text from DOCX: {len(paragraphs_text)} paragraphs, {len(tables_text)} table rows (total length: {len(final_text)})")
            return final_text
        else:
            print(f"⚠️ DOCX document appears to be empty or contains no readable text")
            return None

    except Exception as e:
        print(f"❌ DOCX extraction failed: {str(e)}")
        # Try alternative approach
        try:
            print(f"🔄 Trying alternative DOCX extraction method...")
            from zipfile import ZipFile

            with ZipFile(io.BytesIO(file_content)) as docx_zip:
                # Try to extract from document.xml
                try:
                    with docx_zip.open('word/document.xml') as xml_file:
                        import xml.etree.ElementTree as ET
                        tree = ET.parse(xml_file)

                        # Extract text from XML
                        text_elements = []
                        for elem in tree.iter():
                            if elem.text:
                                text_elements.append(elem.text)

                        if text_elements:
                            fallback_text = " ".join(text_elements)
                            print(f"✅ Alternative DOCX extraction successful (length: {len(fallback_text)})")
                            return fallback_text

                except Exception as xml_error:
                    print(f"⚠️ Alternative DOCX extraction also failed: {str(xml_error)}")

        except Exception as alt_error:
            print(f"❌ All DOCX extraction methods failed: {str(alt_error)}")

    return None

def send_file_to_gemini_directly(file_content: bytes, file_path: str, query: str) -> Dict[str, Any]:
    """Send file directly to Gemini for analysis with robust error handling"""
    print(f"🤖 Attempting direct file upload to Gemini...")
    print(f"📊 File size: {len(file_content):,} bytes, Path: {file_path}")

    try:
        client = gemini_client()

        # Create a file part with the binary content
        file_extension = file_path.lower().split('.')[-1] if file_path else 'unknown'
        mime_type = {
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'doc': 'application/msword',
            'txt': 'text/plain'
        }.get(file_extension, 'application/octet-stream')

        print(f"📄 File type: {file_extension.upper()}, MIME type: {mime_type}")

        # Check file size limits (Gemini has limits)
        max_size = 10 * 1024 * 1024  # 10MB limit
        if len(file_content) > max_size:
            print(f"⚠️ File size ({len(file_content):,} bytes) exceeds limit ({max_size:,} bytes)")
            return {
                "analysis": f"File too large for direct analysis ({len(file_content):,} bytes > {max_size:,} bytes)",
                "match_score": 0,
                "analyzed_by": "Gemini 2.5 Flash (Direct File)",
                "success": False,
                "error": "File size exceeds limit"
            }

        # Encode file content as base64
        try:
            file_data = base64.b64encode(file_content).decode('utf-8')
            print(f"✅ File encoded to base64 (size: {len(file_data):,} chars)")
        except Exception as encode_error:
            print(f"❌ Base64 encoding failed: {str(encode_error)}")
            return {
                "analysis": f"File encoding failed: {str(encode_error)}",
                "match_score": 0,
                "analyzed_by": "Gemini 2.5 Flash (Direct File)",
                "success": False,
                "error": "Base64 encoding failed"
            }

        analysis_prompt = RESUME_ANALYSIS_PROMPT.format(
            job_posting=query,
            resume_content="[Resume file attached]"
        )

        print(f"🔄 Creating Gemini request with file attachment...")

        contents = [
            types.Content(
                role="user",
                parts=[
                    types.Part(text=analysis_prompt),
                    types.Part(
                        inline_data=types.Blob(
                            mime_type=mime_type,
                            data=file_data
                        )
                    )
                ]
            )
        ]

        generate_content_config = types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
            temperature=0.3,
            top_p=0.8,
            max_output_tokens=5000,
            safety_settings=[
                types.SafetySetting(
                    category="HARM_CATEGORY_HATE_SPEECH",
                    threshold="OFF"
                ),
                types.SafetySetting(
                    category="HARM_CATEGORY_DANGEROUS_CONTENT",
                    threshold="OFF"
                ),
                types.SafetySetting(
                    category="HARM_CATEGORY_SEXUALLY_EXPLICIT",
                    threshold="OFF"
                ),
                types.SafetySetting(
                    category="HARM_CATEGORY_HARASSMENT",
                    threshold="OFF"
                )
            ],
        )

        print(f"🚀 Sending request to Gemini...")
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=contents,
            config=generate_content_config,
        )

        analysis_text = response.text if hasattr(response, 'text') else str(response)

        print(f"✅ Received response from Gemini (length: {len(analysis_text)} chars)")

        # Try to parse JSON response
        try:
            # Clean the response text - remove any markdown formatting
            clean_text = analysis_text.strip()
            if clean_text.startswith('```json'):
                clean_text = clean_text[7:]
            if clean_text.endswith('```'):
                clean_text = clean_text[:-3]
            clean_text = clean_text.strip()

            # Parse JSON
            analysis_json = json.loads(clean_text)
            raw_score = analysis_json.get('match_score') or analysis_json.get('overall_match_score') or analysis_json.get('candidate_overview', {}).get('overall_match_score') or 0
            score = clamp_score(raw_score)
            if score != raw_score:
                print(f"⚠️ Clamped out-of-range LLM score {raw_score!r} → {score}")

            print(f"✅ Direct file analysis completed successfully")
            print(f"👤 Candidate: {analysis_json.get('candidate_name', 'Unknown')}")
            print(f"🎯 Match score: {score}")

            return {
                "analysis": analysis_text,
                "analysis_json": analysis_json,
                "match_score": score,
                "analyzed_by": "Gemini 2.5 Flash (Direct File)",
                "success": True,
                "file_processed": True
            }

        except json.JSONDecodeError as e:
            print(f"⚠️ Failed to parse JSON response: {str(e)}")
            print(f"📝 Raw response length: {len(analysis_text or '')} chars (preview redacted to avoid leaking PII)")

            # Fallback to old score extraction method
            score = clamp_score(extract_score_from_analysis(analysis_text))

            return {
                "analysis": analysis_text,
                "analysis_json": None,
                "match_score": score,
                "analyzed_by": "Gemini 2.5 Flash (Direct File)",
                "success": True,
                "json_parse_error": str(e),
                "file_processed": True
            }

    except Exception as e:
        print(f"❌ Direct file upload failed: {str(e)}")
        print(f"📊 Error details: File size: {len(file_content)}, Extension: {file_path.split('.')[-1] if '.' in file_path else 'unknown'}")

        return {
            "analysis": f"Direct file analysis failed: {str(e)}",
            "match_score": 0,
            "analyzed_by": "Gemini 2.5 Flash (Direct File)",
            "success": False,
            "error": str(e),
            "file_processed": False
        }

_companies_cache = {"data": None, "ts": 0.0}
_COMPANIES_CACHE_TTL = 60.0  # seconds


def _get_companies_cached():
    """60s in-process cache around db.get_all_companies(). Hot path: every
    GCS file fetch calls this for cross-company path resolution."""
    import time as _t
    now = _t.time()
    if _companies_cache["data"] is not None and (now - _companies_cache["ts"]) < _COMPANIES_CACHE_TTL:
        return _companies_cache["data"]
    try:
        rows = get_db_manager().get_all_companies() or []
    except Exception:
        rows = _companies_cache["data"] or []
    _companies_cache["data"] = rows
    _companies_cache["ts"] = now
    return rows


def _sanitize_object_path(p: str) -> str:
    """Sanitize a user-supplied path component for GCS object names.
    Rejects path traversal, absolute paths, control characters."""
    if not p:
        return ""
    # Strip windows drive letter
    if len(p) >= 2 and p[1] == ':':
        p = p[2:]
    # Normalize separators
    p = p.replace('\\', '/').strip()
    # Drop leading slashes
    p = p.lstrip('/')
    # Reject control chars and null bytes
    if any(ord(c) < 32 for c in p):
        return ""
    # Split, drop traversal segments and empty segments, keep alnum/dash/underscore/dot/space
    safe_parts = []
    for seg in p.split('/'):
        seg = seg.strip()
        if not seg or seg in ('.', '..'):
            continue
        # Whitelist conservative chars
        cleaned = ''.join(ch for ch in seg if ch.isalnum() or ch in ' .-_()[]')
        cleaned = cleaned.strip(' .')
        if cleaned:
            safe_parts.append(cleaned)
    return '/'.join(safe_parts)


def resolve_gcs_file_path(file_path: str, user_context: dict = None, bucket_name: str = None) -> str:
    """
    Smart file path resolution that tries multiple path patterns.
    Handles legacy paths (resume/file.pdf) and new company-organized paths.
    """
    print(f"\n🔍 SMART FILE PATH RESOLUTION")
    print(f"📁 Original path: {file_path}")

    try:
        client = _get_storage_client()

        # Use the provided bucket_name or get it from user context
        if not bucket_name and user_context:
            bucket_name, _ = get_company_resources(user_context)
        elif not bucket_name:
            bucket_name = GCS_BUCKET_NAME

        bucket = client.bucket(bucket_name)
        print(f"🪣 Using bucket: {bucket_name}")

        # List of path patterns to try, in order of preference
        paths_to_try = []

        # 1. Try the original path as-is
        original_path = file_path
        if original_path.startswith(f"gs://{GCS_BUCKET_NAME}/"):
            original_path = original_path.replace(f"gs://{GCS_BUCKET_NAME}/", "")
        paths_to_try.append(("original", original_path))

        # FAST PATH: try the original path immediately before doing any expensive
        # work (DB scans, company expansion). The vast majority of file fetches
        # use the canonical path and succeed on the first try; doing the
        # expensive cross-company expansion up-front for every file caused
        # multi-minute latency under parallel load (slow get_all_companies query
        # blocking the DB pool while all worker threads waited).
        try:
            fast_blob = bucket.blob(original_path)
            if fast_blob.exists():
                print(f"✅ Fast-path hit: original path exists -> {original_path}")
                return original_path
        except Exception as _fast_e:
            print(f"⚠️ Fast-path check error (will fall back to full resolution): {_fast_e}")

        # 2. If we have user context, try company-specific paths
        if user_context and user_context.get('company'):
            company_code = user_context['company'].get('company_code')
            if company_code:
                # Extract filename from the original path
                if '/' in original_path:
                    filename = original_path.split('/')[-1]
                    # Try company-organized path
                    company_path = f"{company_code}/resume/{filename}"
                    paths_to_try.append(("company_specific", company_path))
                else:
                    # If original path is just a filename, try company structure
                    company_path = f"{company_code}/resume/{original_path}"
                    paths_to_try.append(("company_filename", company_path))

        # 3. Get all companies and try each one (for cross-company search scenarios)
        try:
            companies = _get_companies_cached()
            for company in companies[:10]:  # Limit to first 10 companies to avoid too many tries
                company_code = company.get('company_code')
                if company_code:
                    if '/' in original_path:
                        filename = original_path.split('/')[-1]
                        company_path = f"{company_code}/resume/{filename}"
                    else:
                        company_path = f"{company_code}/resume/{original_path}"

                    # Only add if not already in the list
                    if not any(path == company_path for _, path in paths_to_try):
                        paths_to_try.append((f"company_{company_code}", company_path))
        except Exception as e:
            print(f"⚠️ Could not load companies for path resolution: {e}")

        # 4. Try system path (for super admin uploads)
        if '/' in original_path:
            filename = original_path.split('/')[-1]
            system_path = f"system/resume/{filename}"
        else:
            system_path = f"system/resume/{original_path}"
        paths_to_try.append(("system", system_path))

        # Now try each path until we find one that exists
        print(f"🔍 Trying {len(paths_to_try)} possible path patterns...")

        for i, (pattern_name, path_to_try) in enumerate(paths_to_try, 1):
            print(f"   {i}. {pattern_name}: {path_to_try}")
            blob = bucket.blob(path_to_try)

            if blob.exists():
                print(f"✅ Found file using {pattern_name} pattern: {path_to_try}")
                return path_to_try
            else:
                print(f"   ❌ Not found with {pattern_name} pattern")

        print(f"❌ File not found with any path pattern")
        print(f"🔧 Suggestion: Check if the file was uploaded with the correct company organization")
        return None

    except Exception as e:
        print(f"❌ Error during file path resolution: {str(e)}")
        return None

def get_file_content_from_gcs(file_path: str, user_context: dict = None) -> str:
    """Get file content from Google Cloud Storage with smart path resolution"""
    # Get company-specific bucket
    bucket_name, datastore_id = get_company_resources(user_context)

    print(f"\n📁 GCS FILE RETRIEVAL")
    print(f"🔗 File path: {file_path}")
    print(f"🪣 Bucket: {bucket_name}")

    try:
        client = _get_storage_client()
        bucket = client.bucket(bucket_name)
        print(f"✅ GCS client and bucket initialized")

        # First try to resolve the correct file path
        resolved_path = resolve_gcs_file_path(file_path, user_context, bucket_name)

        if not resolved_path:
            print(f"❌ Could not resolve file path: {file_path}")
            return None

        blob_name = resolved_path
        print(f"📝 Using resolved blob name: {blob_name}")

        blob = bucket.blob(blob_name)
        print(f"🔍 Checking if blob exists...")

        if not blob.exists():
            print(f"❌ Blob does not exist: {blob_name}")
            return None

        print(f"✅ Blob exists, downloading content...")
        # Download content as bytes first
        file_content = blob.download_as_bytes()
        print(f"✅ File downloaded successfully (size: {len(file_content)} bytes)")

        # Extract text based on file type
        extracted_text = extract_text_from_file(file_content, file_path)
        if extracted_text:
            return extracted_text
        else:
            # If text extraction fails, we'll return the file content for direct Gemini upload
            print(f"⚠️ Text extraction failed, will try direct file upload to Gemini")
            return None

    except Exception as e:
        print(f"❌ Error getting file content: {str(e)}")
        return None

def get_file_content_and_bytes_from_gcs(file_path: str, user_context: dict = None) -> tuple:
    """Get both text content and raw bytes from GCS file with smart path resolution"""
    # Get company-specific bucket
    bucket_name, datastore_id = get_company_resources(user_context)

    print(f"\n📁 GCS FILE RETRIEVAL (WITH BYTES)")
    print(f"🔗 File path: {file_path}")
    print(f"🪣 Bucket: {bucket_name}")

    try:
        client = _get_storage_client()
        bucket = client.bucket(bucket_name)
        print(f"✅ GCS client and bucket initialized")

        # First try to resolve the correct file path
        resolved_path = resolve_gcs_file_path(file_path, user_context, bucket_name)

        if not resolved_path:
            print(f"❌ Could not resolve file path: {file_path}")
            return None, None

        blob_name = resolved_path
        print(f"📝 Using resolved blob name: {blob_name}")

        blob = bucket.blob(blob_name)
        print(f"🔍 Checking if blob exists...")

        if not blob.exists():
            print(f"❌ Blob does not exist: {blob_name}")
            return None, None

        print(f"✅ Blob exists, downloading content...")
        # Download content as bytes
        file_content = blob.download_as_bytes()
        print(f"✅ File downloaded successfully (size: {len(file_content)} bytes)")

        # Extract text based on file type
        extracted_text = extract_text_from_file(file_content, file_path)

        return extracted_text, file_content

    except Exception as e:
        print(f"❌ Error getting file content: {str(e)}")
        return None, None

def get_mime_type_for_datastore(file_path: str) -> str:
    """Get the correct MIME type for AI Search datastore based on file extension"""
    file_extension = file_path.lower().split('.')[-1] if '.' in file_path else ''

    mime_types = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'txt': 'text/plain',
        'html': 'text/html',
        'htm': 'text/html',
        'xml': 'application/xml',
        'json': 'application/json',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'xls': 'application/vnd.ms-excel.sheet.macroenabled.12',
        'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'png': 'image/png',
        'gif': 'image/gif',
        'bmp': 'image/bmp',
        'tiff': 'image/tiff'
    }

    mime_type = mime_types.get(file_extension, 'application/octet-stream')
    print(f"📄 File: {file_path} -> Extension: {file_extension} -> MIME: {mime_type}")

    return mime_type

def upload_to_vector_datastore(file_path: str, company_id: int = None, company_code: str = None, user_id: str = None) -> bool:
    """
    Upload a document to the company-specific vector datastore with proper metadata.
    This enables incremental indexing during file upload.
    """
    print(f"\n📊 VECTOR DATASTORE UPLOAD")
    print(f"📁 File path: {file_path}")
    print(f"🏢 Company: {company_code} (ID: {company_id})")
    print(f"👤 User: {user_id}")

    try:
        from google.cloud import discoveryengine_v1
        from google.api_core.client_options import ClientOptions
        from datetime import datetime

        # Get company-specific resources
        user_context = {'company': {'company_code': company_code, 'company_id': company_id}} if company_code else None
        bucket_name, datastore_id = get_company_resources(user_context)

        print(f"🪣 Using bucket: {bucket_name}")
        print(f"📚 Using datastore: {datastore_id}")

        # Initialize the client
        client_options = (
            ClientOptions(api_endpoint=f"{LOCATION}-discoveryengine.googleapis.com")
            if LOCATION != "global"
            else None
        )

        client = discoveryengine_v1.DocumentServiceClient(client_options=client_options)

        # Prepare the document with company metadata
        # Sanitize document_id to only contain [a-zA-Z0-9-_]
        document_id = file_path.replace('/', '_').replace('.', '_')
        # Remove or replace all invalid characters with underscores
        import re
        document_id = re.sub(r'[^a-zA-Z0-9_-]', '_', document_id)
        # Remove multiple consecutive underscores
        document_id = re.sub(r'_+', '_', document_id)
        # Remove leading/trailing underscores
        document_id = document_id.strip('_')

        # Create structured data with company metadata
        structured_data = {
            "file_path": file_path,
            "company_id": str(company_id) if company_id else "system",
            "company_code": company_code or "system",
            "uploaded_by": user_id or "system",
            "upload_timestamp": datetime.now().isoformat(),
            "tenant_isolation": company_code or "system",
            "gcs_uri": f"gs://{bucket_name}/{file_path}"
        }

        # Extract filename for better searchability
        filename = file_path.split('/')[-1]
        structured_data["filename"] = filename
        structured_data["file_extension"] = filename.split('.')[-1] if '.' in filename else "unknown"

        # Create the document
        document = discoveryengine_v1.Document(
            id=document_id,
            struct_data=structured_data,
            # Use the GCS URI as the content URI
            content=discoveryengine_v1.Document.Content(
                uri=f"gs://{bucket_name}/{file_path}",
                mime_type=get_mime_type_for_datastore(file_path)
            )
        )

        # Prepare the request
        parent = f"projects/{PROJECT_ID}/locations/{LOCATION}/collections/default_collection/dataStores/{datastore_id}/branches/default_branch"

        request = discoveryengine_v1.CreateDocumentRequest(
            parent=parent,
            document=document,
            document_id=document_id
        )

        print(f"📤 Uploading to datastore: {datastore_id}")
        print(f"🆔 Document ID: {document_id}")
        print(f"🏷️ Metadata: {structured_data}")

        # Create the document
        operation = client.create_document(request=request)
        print(f"✅ Document upload initiated")
        print(f"🔄 Operation: {operation.name}")

        return True

    except Exception as e:
        print(f"❌ Vector datastore upload failed: {str(e)}")
        return False

def create_company_search_request(
    query: str,
    company_code: str = None,
    result_count: int = 10,
    max_value: int = MAX_SEARCH_RESULTS,
) -> 'discoveryengine_v1.SearchRequest':
    """
    Create a search request using company-specific datastore for automatic tenant isolation.
    """
    from google.cloud import discoveryengine_v1
    result_count = min(max(int(result_count or 10), 1), max_value)

    print(f"\n🔍 CREATING COMPANY-SPECIFIC SEARCH")
    print(f"📝 Query: {query}")

    # Get company-specific datastore (automatic isolation)
    user_context = {'company': {'company_code': company_code}} if company_code else None
    bucket_name, datastore_id = get_company_resources(user_context)

    print(f"📚 Using datastore: {datastore_id}")
    print(f"🏢 Company: {company_code or 'Default (Super Admin)'}")

    # Base serving config
    serving_config = f"projects/{PROJECT_ID}/locations/{LOCATION}/collections/default_collection/dataStores/{datastore_id}/servingConfigs/default_config"

    # Create the search request
    request = discoveryengine_v1.SearchRequest(
        serving_config=serving_config,
        query=query,
        page_size=result_count,
        content_search_spec=build_discovery_content_search_spec(discoveryengine_v1),
    )

    print(f"✅ Search will automatically be isolated to company datastore")

    return request

def get_search_client(location: str = None):
    """Get a properly configured search client"""
    from google.cloud import discoveryengine_v1
    from google.api_core.client_options import ClientOptions

    if not location:
        location = LOCATION

    client_options = (
        ClientOptions(api_endpoint=f"{location}-discoveryengine.googleapis.com")
        if location != "global"
        else None
    )

    return discoveryengine_v1.SearchServiceClient(client_options=client_options)

def create_universal_search_request(
    query: str,
    user_context: dict = None,
    result_count: int = 10,
    max_value: int = MAX_SEARCH_RESULTS,
) -> 'discoveryengine_v1.SearchRequest':
    """
    Universal search request creator that automatically uses the right datastore.
    Works for both company users and super admins.
    """
    from google.cloud import discoveryengine_v1
    result_count = min(max(int(result_count or 10), 1), max_value)

    print(f"\n🔍 CREATING UNIVERSAL SEARCH REQUEST")
    print(f"📝 Query: {query[:100]}{'...' if len(query) > 100 else ''}")

    # Get appropriate datastore based on user context
    bucket_name, datastore_id = get_company_resources(user_context)

    # Extract company info for logging
    company_code = "Super Admin (Default)"
    if user_context and user_context.get('company'):
        company_info = user_context['company']
        if isinstance(company_info, dict):
            company_code = company_info.get('company_code', 'Unknown')
        else:
            company_code = getattr(company_info, 'company_code', 'Unknown')

    print(f"🏢 Company: {company_code}")
    print(f"📚 Using datastore: {datastore_id}")

    # Create serving config path
    serving_config = f"projects/{PROJECT_ID}/locations/{LOCATION}/collections/default_collection/dataStores/{datastore_id}/servingConfigs/default_config"

    # Create the search request
    request = discoveryengine_v1.SearchRequest(
        serving_config=serving_config,
        query=query,
        page_size=result_count,
        content_search_spec=build_discovery_content_search_spec(discoveryengine_v1),
    )

    print(f"✅ Search request created with automatic datastore isolation")

    return request

def extract_score_from_analysis(analysis_text: str) -> int:
    """Extract a 0-100 numerical score from LLM analysis text.

    Strategy: try anchored "OVERALL MATCH SCORE …" patterns first, then fall
    back to a single `score: NN` / `**NN/100**` / `NN/100` capture. Returns 0
    when nothing parseable in [0, 100] is found.

    Previously this function spammed dozens of `print` calls on every call AND
    biased its fallback to 80-100 only (so any 50-79 score returned 0). Now it
    is silent on the happy path, has no biased fallback, and uses `\b` word
    boundaries to avoid grabbing digits inside identifiers/dates.
    """
    if not analysis_text:
        return 0

    # Most specific → most generic. All capture a 1-3 digit number; final
    # validator clamps to [0,100]. `[\s\S]{0,80}?` is a lazy any-char (incl.
    # newlines) gap so we can span "(0-100):\n\n" between header and value
    # without greedy runaway.
    patterns = (
        r'OVERALL\s+MATCH\s+SCORE[\s\S]{0,80}?\*\*\s*(\d{1,3})\s*(?:/\s*100|%)?\s*\*\*',
        r'OVERALL\s+MATCH\s+SCORE[\s\S]{0,80}?(\d{1,3})\s*/\s*100\b',
        r'OVERALL\s+MATCH\s+SCORE[\s\S]{0,80}?(\d{1,3})\s*%',
        r'OVERALL\s+MATCH\s+SCORE\s*(?:\([^)]*\))?\s*[:\-]?\s*(\d{1,3})\b',
        r'(?:final|match)\s+score[\s\S]{0,40}?(\d{1,3})\s*/\s*100\b',
        r'\*\*\s*(\d{1,3})\s*/\s*100\s*\*\*',
        r'\b(\d{1,3})\s*/\s*100\b',
        r'\bscore\s*[:=]\s*(\d{1,3})\b',
    )
    for pat in patterns:
        m = re.search(pat, analysis_text, re.IGNORECASE)
        if m:
            try:
                score = int(m.group(1))
            except (ValueError, IndexError):
                continue
            if 0 <= score <= 100:
                return score
    return 0

def create_basic_entities_fallback(resume_content: str) -> Dict[str, Any]:
    """Create basic entity structure from resume text using pattern matching"""
    import re

    print(f"🔄 Creating basic entities fallback from resume text")

    # Basic patterns for extraction
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'[\+]?[1-9]?[0-9]{7,15}'
    name_pattern = r'^([A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)'

    # Extract basic info
    emails = re.findall(email_pattern, resume_content)
    phones = re.findall(phone_pattern, resume_content)

    # Try to find name (usually at the beginning)
    lines = resume_content.split('\n')
    name = "Not available"
    for line in lines[:5]:  # Check first 5 lines
        line = line.strip()
        if len(line) > 5 and len(line) < 50 and ' ' in line:
            # Simple heuristic for name detection
            words = line.split()
            if len(words) >= 2 and all(word.isalpha() and word[0].isupper() for word in words[:2]):
                name = line
                break

    # Extract skills (common technical terms)
    skill_patterns = [
        r'\bPython\b', r'\bJava\b', r'\bSQL\b', r'\bMachine Learning\b', r'\bML\b',
        r'\bTensorFlow\b', r'\bPyTorch\b', r'\bPandas\b', r'\bNumPy\b', r'\bScikit-learn\b',
        r'\bAWS\b', r'\bGCP\b', r'\bAzure\b', r'\bDocker\b', r'\bKubernetes\b',
        r'\bReact\b', r'\bNode\.js\b', r'\bJavaScript\b', r'\bHTML\b', r'\bCSS\b'
    ]

    skills = []
    for pattern in skill_patterns:
        matches = re.findall(pattern, resume_content, re.IGNORECASE)
        if matches:
            skills.extend([match for match in matches if match not in skills])

    # Extract education info
    education_patterns = [
        r'\b(Bachelor|Master|PhD|B\.S\.|M\.S\.|B\.A\.|M\.A\.)\b.*?(?:in|of)\s+([A-Za-z\s]+)',
        r'\b(University|College|Institute)\s+([A-Za-z\s]+)'
    ]

    education_info = {"degree": "Not available", "university": "Not available", "graduation_year": "Not available"}
    for pattern in education_patterns:
        matches = re.findall(pattern, resume_content, re.IGNORECASE)
        if matches:
            education_info["degree"] = matches[0][0] if matches[0][0] else "Not available"
            education_info["university"] = matches[0][1] if len(matches[0]) > 1 else "Not available"
            break

    # Extract years of experience
    exp_pattern = r'(\d+)\s*\+?\s*years?\s+(?:of\s+)?experience'
    exp_matches = re.findall(exp_pattern, resume_content, re.IGNORECASE)
    experience_years = f"{exp_matches[0]} Years" if exp_matches else "Not specified"

    basic_entities = {
        "candidate_overview": {
            "name": name,
            "location": "Not available",
            "email": emails[0] if emails else "Not available",
            "phone": phones[0] if phones else "Not available",
            "current_role": "Not available",
            "experience_years": experience_years
        },
        "technical_skills": skills[:5],  # Top 5 skills
        "tools_technologies": skills[5:10] if len(skills) > 5 else [],
        "education": education_info,
        "work_experience": [
            {
                "year_range": "Not available",
                "role": "Not available",
                "company": "Not available",
                "key_responsibilities": ["Not available"]
            }
        ]
    }

    print(f"✅ Basic entities created: name={name}, skills={len(skills)}, email={emails[0] if emails else 'None'}")
    return basic_entities

async def extract_candidate_entities(resume_content: str, token_tracker: TokenTracker = None) -> Dict[str, Any]:
    """First LLM call: Extract candidate information and entities from resume"""
    print(f"\n👤 STEP 1: CANDIDATE ENTITY EXTRACTION")
    print(f"📄 Resume content length: {len(resume_content)} characters")

    try:
        client = gemini_client()

        entity_extraction_prompt = ENTITY_EXTRACTION_PROMPT.format(
            resume_content=resume_content
        )

        contents = [
            types.Content(
                role="user",
                parts=[types.Part(text=entity_extraction_prompt)]
            )
        ]

        generate_content_config = types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
            temperature=0.1,
            top_p=0.8,
            max_output_tokens=5000,  # Increased for better responses
            safety_settings=[
                types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="OFF"),
                types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="OFF"),
                types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="OFF"),
                types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="OFF")
            ],
        )

        print(f"🚀 Extracting candidate entities...")

        # Retry mechanism for entity extraction
        max_retries = 2
        response = None
        for attempt in range(max_retries + 1):
            try:
                response = await asyncio.to_thread(
                    client.models.generate_content,
                    model="gemini-2.5-flash",
                    contents=contents,
                    config=generate_content_config,
                )

                # Track token usage
                if token_tracker:
                    token_usage = extract_token_usage(response, "entity_extraction", "gemini-2.5-flash")
                    token_tracker.add_call(
                        operation="entity_extraction",
                        model="gemini-2.5-flash",
                        input_tokens=token_usage["input_tokens"],
                        output_tokens=token_usage["output_tokens"],
                        success=True
                    )
                    print(f"📊 Token usage - Input: {token_usage['input_tokens']}, Output: {token_usage['output_tokens']}, Total: {token_usage['total_tokens']}")

                break
            except Exception as e:
                print(f"⚠️ Entity extraction attempt {attempt + 1} failed: {str(e)}")
                if attempt == max_retries:
                    print(f"❌ All entity extraction attempts failed")
                    return {"success": False, "error": f"All attempts failed: {str(e)}"}
                import asyncio
                # Exponential backoff: 1s, 2s, 4s, 8s ... capped at 32s
                await asyncio.sleep(min(32, 2 ** attempt))

        # Extract response text with comprehensive null handling and debugging
        analysis_text = ""
        try:
            print(f"🔍 Analyzing response structure...")
            print(f"📊 Response type: {type(response)}")
            print(f"📊 Has text attr: {hasattr(response, 'text')}")
            print(f"📊 Has candidates attr: {hasattr(response, 'candidates')}")

            # First try the direct text attribute
            if hasattr(response, 'text') and response.text is not None:
                text_content = str(response.text).strip()
                print(f"📝 Direct text length: {len(text_content)}")
                if len(text_content) > 0:
                    analysis_text = text_content
                    print(f"✅ Got response via .text attribute")
                else:
                    print(f"⚠️ Direct text is empty")

            # Then try candidates approach
            if not analysis_text and hasattr(response, 'candidates') and response.candidates is not None:
                candidates_count = len(response.candidates) if response.candidates else 0
                print(f"🔍 Trying candidates approach ({candidates_count} candidates)")

                for i, candidate in enumerate(response.candidates):
                    print(f"  📄 Candidate {i}: {type(candidate)}")
                    if candidate is not None:
                        print(f"    Has content: {hasattr(candidate, 'content')}")
                        print(f"    Has finish_reason: {hasattr(candidate, 'finish_reason')}")

                        if hasattr(candidate, 'finish_reason'):
                            print(f"    Finish reason: {candidate.finish_reason}")

                        if hasattr(candidate, 'content') and candidate.content is not None:
                            print(f"    Content type: {type(candidate.content)}")
                            print(f"    Has parts: {hasattr(candidate.content, 'parts')}")

                            if hasattr(candidate.content, 'parts') and candidate.content.parts is not None:
                                parts_count = len(candidate.content.parts) if candidate.content.parts else 0
                                print(f"    Parts count: {parts_count}")

                                for j, part in enumerate(candidate.content.parts):
                                    print(f"      Part {j}: {type(part)}")
                                    if part is not None and hasattr(part, 'text') and part.text is not None:
                                        part_text = str(part.text).strip()
                                        print(f"      Part {j} text length: {len(part_text)}")
                                        if len(part_text) > 0:
                                            analysis_text += part_text
                                            print(f"✅ Got text from candidate {i}, part {j}")

            # Check for safety ratings or blocked content
            if not analysis_text and hasattr(response, 'candidates') and response.candidates:
                for i, candidate in enumerate(response.candidates):
                    if hasattr(candidate, 'safety_ratings'):
                        print(f"🛡️ Candidate {i} safety ratings: {candidate.safety_ratings}")
                    if hasattr(candidate, 'finish_reason'):
                        finish_reason = str(candidate.finish_reason)
                        print(f"🏁 Candidate {i} finish reason: {finish_reason}")
                        if 'SAFETY' in finish_reason or 'BLOCKED' in finish_reason:
                            print(f"⚠️ Content may have been blocked by safety filters")

            # Final fallback - try to extract any available text
            if not analysis_text:
                print(f"⚠️ Using string conversion fallback")
                full_response_str = str(response) if response is not None else ""
                print(f"📝 Full response string length: {len(full_response_str)}")
                if len(full_response_str) > 100:  # If we have substantial content
                    analysis_text = full_response_str
                    print(f"✅ Using full response string as fallback")

        except Exception as extract_error:
            print(f"❌ Error extracting entity response text: {str(extract_error)}")
            print(f"📊 Response type: {type(response)}")
            analysis_text = str(response) if response is not None else ""

        if not analysis_text:
            print(f"❌ Empty entity extraction response - creating basic fallback")
            # Create a basic entity structure from the resume text
            basic_entities = create_basic_entities_fallback(resume_content)
            return {
                "success": True,
                "entities": basic_entities,
                "raw_response": "Fallback entity extraction",
                "fallback_used": True
            }

        print(f"📄 Entity extraction response length: {len(analysis_text)} characters")

        # Parse JSON response
        try:
            clean_text = analysis_text.strip()
            if clean_text.startswith('```json'):
                clean_text = clean_text[7:]
            if clean_text.endswith('```'):
                clean_text = clean_text[:-3]
            clean_text = clean_text.strip()

            # Fix common JSON issues
            import re

            # Handle unterminated strings at end of lines
            clean_text = re.sub(r':\s*"([^"]*?)$', r': "\1"', clean_text, flags=re.MULTILINE)

            # Handle unterminated strings in the middle
            clean_text = re.sub(r':\s*"([^"]*?)\n\s*[,}]', r': "\1",', clean_text)

            # Remove trailing commas
            clean_text = re.sub(r',(\s*[}\]])', r'\1', clean_text)

            # Ensure proper closing
            if not clean_text.endswith('}'):
                open_braces = clean_text.count('{') - clean_text.count('}')
                open_brackets = clean_text.count('[') - clean_text.count(']')
                clean_text += ']' * open_brackets + '}' * open_braces
                print(f"🔧 Entity JSON: Added {open_brackets} ] and {open_braces} }}")

            entities = json.loads(clean_text)

            print(f"✅ Entity extraction successful")
            print(f"👤 Candidate: {entities.get('candidate_overview', {}).get('name', 'Unknown')}")

            return {
                "success": True,
                "entities": entities,
                "raw_response": analysis_text
            }

        except json.JSONDecodeError as e:
            print(f"⚠️ Failed to parse entity JSON: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "raw_response": analysis_text
            }

    except Exception as e:
        print(f"❌ Entity extraction failed: {str(e)}")
        return {"success": False, "error": str(e)}

def determine_dynamic_score_categories(job_title: str, query: str) -> Dict[str, str]:
    """Determine appropriate score categories based on job title and query"""
    job_lower = job_title.lower() if job_title else ""
    query_lower = query.lower() if query else ""
    combined_text = f"{job_lower} {query_lower}"

    # Define category mappings for different role types
    category_mappings = {
        "technical": {
            "primary_competency": "technical_skills",
            "secondary_competency": "problem_solving",
            "experience_relevance": "system_design",
            "cultural_behavioral_fit": "collaboration"
        },
        "management": {
            "primary_competency": "leadership_experience",
            "secondary_competency": "strategic_thinking",
            "experience_relevance": "team_management",
            "cultural_behavioral_fit": "business_acumen"
        },
        "sales": {
            "primary_competency": "relationship_building",
            "secondary_competency": "negotiation_skills",
            "experience_relevance": "target_achievement",
            "cultural_behavioral_fit": "market_knowledge"
        },
        "marketing": {
            "primary_competency": "creative_skills",
            "secondary_competency": "analytics_knowledge",
            "experience_relevance": "brand_management",
            "cultural_behavioral_fit": "digital_marketing"
        },
        "hr": {
            "primary_competency": "people_management",
            "secondary_competency": "policy_knowledge",
            "experience_relevance": "organizational_skills",
            "cultural_behavioral_fit": "employee_relations"
        },
        "finance": {
            "primary_competency": "financial_analysis",
            "secondary_competency": "regulatory_knowledge",
            "experience_relevance": "attention_to_detail",
            "cultural_behavioral_fit": "reporting_skills"
        },
        "design": {
            "primary_competency": "design_skills",
            "secondary_competency": "creative_thinking",
            "experience_relevance": "user_experience",
            "cultural_behavioral_fit": "technical_proficiency"
        },
        "operations": {
            "primary_competency": "process_improvement",
            "secondary_competency": "project_management",
            "experience_relevance": "analytical_thinking",
            "cultural_behavioral_fit": "stakeholder_management"
        },
        "consulting": {
            "primary_competency": "analytical_skills",
            "secondary_competency": "client_management",
            "experience_relevance": "problem_solving",
            "cultural_behavioral_fit": "presentation_skills"
        },
        "customer_service": {
            "primary_competency": "communication_skills",
            "secondary_competency": "problem_resolution",
            "experience_relevance": "empathy",
            "cultural_behavioral_fit": "product_knowledge"
        }
    }

    # Keywords to identify role types
    role_keywords = {
        "technical": ["developer", "engineer", "data scientist", "programmer", "software", "backend", "frontend", "fullstack", "devops", "ai", "ml", "python", "java", "react", "node"],
        "management": ["manager", "director", "vp", "head", "chief", "team lead", "supervisor", "coordinator"],
        "sales": ["sales", "account manager", "business development", "client acquisition", "revenue"],
        "marketing": ["marketing", "brand", "digital marketing", "content", "social media", "campaign", "advertising"],
        "hr": ["hr", "human resources", "recruiter", "talent", "people operations", "employee relations"],
        "finance": ["finance", "accounting", "financial analyst", "controller", "auditor", "budget", "financial"],
        "design": ["designer", "ui", "ux", "graphic", "visual", "creative", "art director", "product design"],
        "operations": ["operations", "project manager", "process", "logistics", "supply chain", "analyst"],
        "consulting": ["consultant", "advisor", "strategy", "business analyst", "transformation"],
        "customer_service": ["customer service", "support", "help desk", "client relations", "customer success"]
    }

    # Determine role type based on keywords
    role_type = "technical"  # default
    max_matches = 0

    for role, keywords in role_keywords.items():
        matches = sum(1 for keyword in keywords if keyword in combined_text)
        if matches > max_matches:
            max_matches = matches
            role_type = role

    return category_mappings.get(role_type, category_mappings["technical"])

def create_basic_scorecard_fallback(entities: Dict[str, Any], query: str, job_title: str) -> Dict[str, Any]:
    """Create a basic HR scorecard when LLM fails"""
    print(f"🔄 Creating basic scorecard fallback")

    # Get candidate info
    candidate_info = entities.get('candidate_overview', {})
    candidate_name = candidate_info.get('name', 'Unknown')
    candidate_skills = entities.get('technical_skills', [])
    candidate_tools = entities.get('tools_technologies', [])

    # Simple scoring based on keyword matching
    query_lower = query.lower()
    all_candidate_skills = candidate_skills + candidate_tools

    # Count skill matches
    skill_matches = 0
    matched_skills = []
    for skill in all_candidate_skills:
        if skill.lower() in query_lower:
            skill_matches += 1
            matched_skills.append(skill)

    # Calculate basic scores
    total_skills = len(all_candidate_skills)
    technical_score = min(90, (skill_matches / max(1, total_skills)) * 100 + 50) if total_skills > 0 else 50
    overall_score = int(technical_score * 0.8 + 60 * 0.2)  # Weighted average

    # Determine match status
    if overall_score >= 80:
        match_status = "Strong Fit"
    elif overall_score >= 60:
        match_status = "Medium Fit"
    else:
        match_status = "Weak Fit"

    # Generate basic alternative roles based on skills
    alternative_roles = []
    if 'data' in query_lower or 'scientist' in query_lower:
        alternative_roles = [
            {"role_title": "Data Analyst", "match_percentage": 75, "reasoning": "Strong analytical skills", "department": "Analytics"},
            {"role_title": "Business Intelligence Analyst", "match_percentage": 65, "reasoning": "Data interpretation skills", "department": "Business"}
        ]
    elif 'developer' in query_lower or 'engineer' in query_lower:
        alternative_roles = [
            {"role_title": "Software Developer", "match_percentage": 80, "reasoning": "Programming experience", "department": "Engineering"},
            {"role_title": "Technical Analyst", "match_percentage": 70, "reasoning": "Technical background", "department": "Technology"}
        ]
    else:
        alternative_roles = [
            {"role_title": "Technical Specialist", "match_percentage": 70, "reasoning": "Relevant technical skills", "department": "Technology"},
            {"role_title": "Analyst", "match_percentage": 65, "reasoning": "Analytical capabilities", "department": "Operations"}
        ]

    # Basic tenure prediction
    exp_years = candidate_info.get('experience_years', 'Not specified')
    if isinstance(exp_years, str) and any(char.isdigit() for char in exp_years):
        years = int(''.join(filter(str.isdigit, exp_years)))
        if years >= 5:
            tenure_estimate = "4-6 years"
            tenure_score = 85
        elif years >= 2:
            tenure_estimate = "2-4 years"
            tenure_score = 75
        else:
            tenure_estimate = "1-3 years"
            tenure_score = 65
    else:
        tenure_estimate = "2-4 years"
        tenure_score = 70

    # Get dynamic score categories
    dynamic_categories = determine_dynamic_score_categories(job_title, query)

    basic_scorecard = {
        "candidate_overview": {
            "name": candidate_name,
            "location": candidate_info.get('location', 'Not available'),
            "experience_years": candidate_info.get('experience_years', 'Not specified'),
            "position_applied_for": job_title,
            "email": candidate_info.get('email', 'Not available'),
            "phone": candidate_info.get('phone', 'Not available'),
            "overall_match_score": overall_score,
            "match_status": match_status
        },
        "score_breakdown": {
            dynamic_categories["primary_competency"]: {
                "score": int(technical_score),
                "comment": f"Found {skill_matches} matching skills out of {total_skills} total skills"
            },
            dynamic_categories["secondary_competency"]: {
                "score": 65,
                "comment": "Supporting skills assessment"
            },
            dynamic_categories["experience_relevance"]: {
                "score": 60,
                "comment": "Basic experience assessment"
            },
            "keyword_technical_match": {
                "score": 50,
                "comment": "Keyword analysis will be updated with standardized matching"
            }
        },
        "analysis_summary": {
            "ai_analysis": f"Candidate {candidate_name} shows {match_status.lower()} with {skill_matches} relevant technical skills matching the job requirements.",
            "resume_highlights": [
                f"Technical skills: {', '.join(matched_skills[:3]) if matched_skills else 'General technical background'}",
                f"Experience level: {candidate_info.get('experience_years', 'Not specified')}",
                f"Overall match score: {overall_score}%"
            ],

        },

        "tenure_prediction": {
            "estimated_tenure": tenure_estimate,
            "confidence_level": "Medium",
            "tenure_score": tenure_score,
            "factors": {
                "job_stability_history": {
                    "score": 70,
                    "analysis": "Basic assessment based on experience level"
                },
                "career_progression": {
                    "score": 65,
                    "analysis": "Shows progression potential"
                },
                "industry_alignment": {
                    "score": 70,
                    "analysis": "Good fit with industry requirements"
                }
            },
        }
    }

    print(f"✅ Basic scorecard created: {candidate_name} - {match_status} ({overall_score}%)")
    return basic_scorecard

async def generate_hr_scorecard_from_entities(entities: Dict[str, Any], query: str, job_title: str, token_tracker: TokenTracker = None) -> Dict[str, Any]:
    """Second LLM call: Generate HR scorecard by comparing entities with job requirements"""
    print(f"\n🏢 STEP 2: HR SCORECARD GENERATION")
    print(f"📝 Job requirements: '{query[:100]}...'")

    try:
        client = gemini_client()

        # Get candidate info safely
        candidate_info = entities.get('candidate_overview', {})
        candidate_name = candidate_info.get('name', 'Unknown')
        candidate_location = candidate_info.get('location', 'Not available')
        candidate_experience = candidate_info.get('experience_years', 'Not specified')
        candidate_email = candidate_info.get('email', 'Not available')
        candidate_phone = candidate_info.get('phone', 'Not available')

        # Prepare resume content for scorecard generation
        resume_content = f"""
CANDIDATE:
Name: {candidate_name}
Experience: {candidate_experience}
Skills: {', '.join(entities.get('technical_skills', [])[:5])}
Tools: {', '.join(entities.get('tools_technologies', [])[:5])}
Education: {entities.get('education', {}).get('degree', 'Not available')}
Work History: {', '.join([f"{exp.get('company', 'Company')} ({exp.get('duration', 'Duration')})" for exp in entities.get('work_experience', [])[:3]])}
"""

        scorecard_prompt = SCORECARD_PROMPT.format(
            job_posting=query,
            resume_content=resume_content
        )

        contents = [
            types.Content(
                role="user",
                parts=[types.Part(text=scorecard_prompt)]
            )
        ]

        generate_content_config = types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
            temperature=0.2,
            top_p=0.8,
            max_output_tokens=5000,
            safety_settings=[
                types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="OFF"),
                types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="OFF"),
                types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="OFF"),
                types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="OFF")
            ],
        )

        print(f"🚀 Generating HR scorecard...")

        # Retry mechanism for scorecard generation
        max_retries = 2
        response = None
        for attempt in range(max_retries + 1):
            try:
                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=contents,
                    config=generate_content_config,
                )

                # Track token usage
                if token_tracker:
                    token_usage = extract_token_usage(response, "scorecard_generation", "gemini-2.5-flash")
                    token_tracker.add_call(
                        operation="scorecard_generation",
                        model="gemini-2.5-flash",
                        input_tokens=token_usage["input_tokens"],
                        output_tokens=token_usage["output_tokens"],
                        success=True
                    )
                    print(f"📊 Token usage - Input: {token_usage['input_tokens']}, Output: {token_usage['output_tokens']}, Total: {token_usage['total_tokens']}")

                break
            except Exception as e:
                print(f"⚠️ Scorecard generation attempt {attempt + 1} failed: {str(e)}")
                if attempt == max_retries:
                    print(f"❌ All scorecard generation attempts failed")
                    return {"success": False, "error": f"All attempts failed: {str(e)}"}
                import asyncio
                # Exponential backoff: 1s, 2s, 4s, 8s ... capped at 32s
                await asyncio.sleep(min(32, 2 ** attempt))

        # Extract response text with comprehensive null handling and debugging
        analysis_text = ""
        try:
            print(f"🔍 Analyzing scorecard response structure...")
            print(f"📊 Response type: {type(response)}")
            print(f"📊 Has text attr: {hasattr(response, 'text')}")
            print(f"📊 Has candidates attr: {hasattr(response, 'candidates')}")

            # First try the direct text attribute
            if hasattr(response, 'text') and response.text is not None:
                text_content = str(response.text).strip()
                print(f"📝 Direct scorecard text length: {len(text_content)}")
                if len(text_content) > 0:
                    analysis_text = text_content
                    print(f"✅ Got scorecard response via .text attribute")
                else:
                    print(f"⚠️ Direct scorecard text is empty")

            # Then try candidates approach
            if not analysis_text and hasattr(response, 'candidates') and response.candidates is not None:
                candidates_count = len(response.candidates) if response.candidates else 0
                print(f"🔍 Trying scorecard candidates approach ({candidates_count} candidates)")

                for i, candidate in enumerate(response.candidates):
                    print(f"  📄 Scorecard candidate {i}: {type(candidate)}")
                    if candidate is not None:
                        print(f"    Has content: {hasattr(candidate, 'content')}")
                        print(f"    Has finish_reason: {hasattr(candidate, 'finish_reason')}")

                        if hasattr(candidate, 'finish_reason'):
                            print(f"    Finish reason: {candidate.finish_reason}")

                        if hasattr(candidate, 'content') and candidate.content is not None:
                            print(f"    Content type: {type(candidate.content)}")
                            print(f"    Has parts: {hasattr(candidate.content, 'parts')}")

                            if hasattr(candidate.content, 'parts') and candidate.content.parts is not None:
                                parts_count = len(candidate.content.parts) if candidate.content.parts else 0
                                print(f"    Parts count: {parts_count}")

                                for j, part in enumerate(candidate.content.parts):
                                    print(f"      Part {j}: {type(part)}")
                                    if part is not None and hasattr(part, 'text') and part.text is not None:
                                        part_text = str(part.text).strip()
                                        print(f"      Part {j} text length: {len(part_text)}")
                                        if len(part_text) > 0:
                                            analysis_text += part_text
                                            print(f"✅ Got scorecard text from candidate {i}, part {j}")

            # Check for safety ratings or blocked content
            if not analysis_text and hasattr(response, 'candidates') and response.candidates:
                for i, candidate in enumerate(response.candidates):
                    if hasattr(candidate, 'safety_ratings'):
                        print(f"🛡️ Scorecard candidate {i} safety ratings: {candidate.safety_ratings}")
                    if hasattr(candidate, 'finish_reason'):
                        finish_reason = str(candidate.finish_reason)
                        print(f"🏁 Scorecard candidate {i} finish reason: {finish_reason}")
                        if 'SAFETY' in finish_reason or 'BLOCKED' in finish_reason:
                            print(f"⚠️ Scorecard content may have been blocked by safety filters")

            # Final fallback
            if not analysis_text:
                print(f"⚠️ Using scorecard string conversion fallback")
                full_response_str = str(response) if response is not None else ""
                print(f"📝 Full scorecard response string length: {len(full_response_str)}")
                if len(full_response_str) > 100:
                    analysis_text = full_response_str
                    print(f"✅ Using full scorecard response string as fallback")

        except Exception as extract_error:
            print(f"❌ Error extracting scorecard response text: {str(extract_error)}")
            print(f"📊 Scorecard response type: {type(response)}")
            analysis_text = str(response) if response is not None else ""

        if not analysis_text:
            print(f"❌ Empty scorecard generation response - creating basic fallback")
            # Create a basic scorecard from entities
            basic_scorecard = create_basic_scorecard_fallback(entities, query, job_title)
            return {
                "success": True,
                "scorecard": basic_scorecard,
                "raw_response": "Fallback scorecard generation",
                "fallback_used": True
            }

        print(f"📄 Scorecard response length: {len(analysis_text)} characters")

        # Parse JSON response
        try:
            clean_text = analysis_text.strip()
            if clean_text.startswith('```json'):
                clean_text = clean_text[7:]
            if clean_text.endswith('```'):
                clean_text = clean_text[:-3]
            clean_text = clean_text.strip()

            # Fix common JSON issues
            import re

            # Handle unterminated strings at end of lines
            clean_text = re.sub(r':\s*"([^"]*?)$', r': "\1"', clean_text, flags=re.MULTILINE)

            # Handle unterminated strings in the middle
            clean_text = re.sub(r':\s*"([^"]*?)\n\s*[,}]', r': "\1",', clean_text)

            # Remove trailing commas
            clean_text = re.sub(r',(\s*[}\]])', r'\1', clean_text)

            # Ensure proper closing
            if not clean_text.endswith('}'):
                open_braces = clean_text.count('{') - clean_text.count('}')
                open_brackets = clean_text.count('[') - clean_text.count(']')
                clean_text += ']' * open_brackets + '}' * open_braces
                print(f"🔧 Scorecard JSON: Added {open_brackets} ] and {open_braces} }}")

            scorecard = json.loads(clean_text)
            overall_score = scorecard.get('candidate_overview', {}).get('overall_match_score', 0)

            print(f"✅ HR scorecard generated successfully")
            print(f"🎯 Overall match score: {overall_score}%")

            return {
                "success": True,
                "scorecard": scorecard,
                "raw_response": analysis_text
            }

        except json.JSONDecodeError as e:
            print(f"⚠️ Failed to parse scorecard JSON: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "raw_response": analysis_text
            }

    except Exception as e:
        print(f"❌ Scorecard generation failed: {str(e)}")
        return {"success": False, "error": str(e)}

async def analyze_resume_with_split_approach(resume_content: str, query: str, job_title: str = "Position", standardized_keywords: List[str] = None, token_tracker: TokenTracker = None) -> Dict[str, Any]:
    """Split HR scorecard analysis into two lighter LLM calls"""
    print(f"\n🔄 USING SPLIT APPROACH FOR HR SCORECARD")

    # Step 1: Extract candidate entities
    entity_result = await extract_candidate_entities(resume_content)

    if not entity_result.get('success'):
        print(f"❌ Entity extraction failed, falling back to basic analysis")
        return analyze_resume_with_gemini(resume_content, query)

    entities = entity_result['entities']

    # Step 2: Generate HR scorecard from entities
    scorecard_result = await generate_hr_scorecard_from_entities(entities, query, job_title)

    if not scorecard_result.get('success'):
        print(f"❌ Scorecard generation failed, falling back to basic analysis")
        return analyze_resume_with_gemini(resume_content, query)

    scorecard = scorecard_result['scorecard']
    overall_score = scorecard.get('candidate_overview', {}).get('overall_match_score', 0)

    # Apply standardized keywords if available
    if standardized_keywords and len(standardized_keywords) > 0:
        print(f"🔍 Applying STANDARDIZED keyword analysis to split approach result with {len(standardized_keywords)} keywords")
        keyword_analysis = await asyncio.to_thread(analyze_candidate_keywords, resume_content, standardized_keywords, token_tracker)
        scorecard['keyword_coverage'] = keyword_analysis
        print(f"✅ Applied standardized keywords to split approach: {keyword_analysis['jd_keywords_matched']}/{keyword_analysis['total_jd_keywords']} matched")
    else:
        print(f"⚠️ No standardized keywords provided to split approach")

    print(f"✅ Split approach successful!")
    print(f"👤 Candidate: {scorecard.get('candidate_overview', {}).get('name', 'Unknown')}")
    print(f"🎯 Match score: {overall_score}%")

    return {
        "hr_scorecard": scorecard,
        "analysis": f"Split approach HR scorecard analysis completed successfully",
        "match_score": overall_score,
        "analyzed_by": "Gemini 2.5 Flash Split Approach",
        "success": True,
        "scorecard_type": "split_approach",
        "entity_extraction": entity_result,
        "scorecard_generation": scorecard_result
    }

def adjust_scorecard_with_keyword_analysis(scorecard_json: Dict[str, Any], keyword_analysis: Dict[str, Any]) -> Dict[str, Any]:
    """
    Adjusts the HR scorecard overall score to properly integrate keyword matching results.
    This ensures that keyword coverage has appropriate weight in the final score calculation.
    """
    try:
        # Get keyword matching metrics
        matched_keywords = keyword_analysis.get('jd_keywords_matched', 0)
        total_keywords = keyword_analysis.get('total_jd_keywords', 1)  # Avoid division by zero
        keyword_match_percentage = (matched_keywords / total_keywords) * 100 if total_keywords > 0 else 0

        # Calculate keyword technical match score based on actual matching
        if keyword_match_percentage >= 80:
            keyword_technical_score = 90 + (keyword_match_percentage - 80) * 0.5  # 90-95 range
        elif keyword_match_percentage >= 60:
            keyword_technical_score = 75 + (keyword_match_percentage - 60) * 0.75  # 75-90 range
        elif keyword_match_percentage >= 40:
            keyword_technical_score = 60 + (keyword_match_percentage - 40) * 0.75  # 60-75 range
        elif keyword_match_percentage >= 20:
            keyword_technical_score = 40 + (keyword_match_percentage - 20) * 1.0   # 40-60 range
        else:
            keyword_technical_score = keyword_match_percentage * 2  # 0-40 range

        # Cap the score at 95
        keyword_technical_score = min(95, max(0, keyword_technical_score))

        # Update the keyword_technical_match score in breakdown
        score_breakdown = scorecard_json.get('score_breakdown', {})

        # Find the keyword technical match category (it might have different names)
        keyword_category_key = None
        for key in score_breakdown.keys():
            if 'keyword' in key.lower() or 'technical_match' in key.lower():
                keyword_category_key = key
                break

        if keyword_category_key:
            # Update the keyword technical score and comment
            score_breakdown[keyword_category_key]['score'] = int(keyword_technical_score)
            score_breakdown[keyword_category_key]['comment'] = f"Matched {matched_keywords}/{total_keywords} required keywords ({keyword_match_percentage:.1f}% coverage). " + \
                ("Strong keyword alignment with job requirements." if keyword_match_percentage >= 70 else
                 "Moderate keyword coverage - some key terms missing." if keyword_match_percentage >= 40 else
                 "Limited keyword coverage - significant gaps in required technical terms.")

        # Recalculate overall score with proper weighting
        scores = []
        weights = []

        for category, details in score_breakdown.items():
            if isinstance(details, dict) and 'score' in details:
                score = details['score']
                scores.append(score)

                # Give higher weight to keyword matching to ensure it significantly impacts the score
                if 'keyword' in category.lower() or 'technical_match' in category.lower():
                    weights.append(0.35)  # 35% weight for keyword matching
                else:
                    weights.append(0.65 / (len(score_breakdown) - 1))  # Distribute remaining 65% among other categories

        # Calculate weighted average
        if scores and weights:
            weighted_score = sum(score * weight for score, weight in zip(scores, weights))
            adjusted_overall_score = int(weighted_score)

            # Update overall score
            candidate_overview = scorecard_json.get('candidate_overview', {})
            original_score = candidate_overview.get('overall_match_score', 0)
            candidate_overview['overall_match_score'] = adjusted_overall_score

            # Update match status based on new score
            if adjusted_overall_score >= 80:
                candidate_overview['match_status'] = "Strong Fit"
            elif adjusted_overall_score >= 60:
                candidate_overview['match_status'] = "Medium Fit"
            else:
                candidate_overview['match_status'] = "Weak Fit"

            print(f"🔄 Score adjustment: {original_score}% → {adjusted_overall_score}% (keyword coverage: {keyword_match_percentage:.1f}%)")
            print(f"🎯 Updated match status: {candidate_overview['match_status']}")

        return scorecard_json

    except Exception as e:
        print(f"❌ Error adjusting scorecard with keyword analysis: {str(e)}")
        return scorecard_json


_MISSING_NAME_TOKENS = {
    '', 'n/a', 'na', 'unknown', 'not provided', 'not available',
    'not specified', 'none', 'null', 'tbd', 'candidate', 'full name',
    'full candidate name', 'candidate name', 'candidate name (not provided)',
    'candidate not provided', 'candidate (not provided)',
}


def _derive_name_from_filename(file_path: Optional[str]) -> Optional[str]:
    """Derive a human-friendly candidate name from a resume filename.

    Examples:
        Resume_Mohd_Maahir.pdf            -> "Mohd Maahir"
        resumes/Resume_A.H._Muskan.txt    -> "A.H. Muskan"
        cv-john-doe-2025.docx             -> "John Doe 2025"
    """
    if not file_path:
        return None
    try:
        base = file_path.rsplit('/', 1)[-1].rsplit('\\', 1)[-1]
        # strip extension
        if '.' in base:
            base = base.rsplit('.', 1)[0]
        # strip common prefixes
        for prefix in ('Resume_', 'resume_', 'CV_', 'cv_', 'Resume-', 'resume-', 'CV-', 'cv-'):
            if base.startswith(prefix):
                base = base[len(prefix):]
                break
        # underscores / hyphens -> spaces
        cleaned = re.sub(r'[_\-]+', ' ', base).strip()
        # collapse whitespace
        cleaned = re.sub(r'\s+', ' ', cleaned)
        if not cleaned:
            return None
        # Title-case but preserve dotted initials like "A.H."
        parts = []
        for tok in cleaned.split(' '):
            if not tok:
                continue
            if '.' in tok and len(tok) <= 5:
                parts.append(tok.upper())
            else:
                parts.append(tok[:1].upper() + tok[1:])
        return ' '.join(parts) or None
    except Exception:
        return None


def _is_missing_value(value: Any) -> bool:
    if value is None:
        return True
    if not isinstance(value, str):
        return False
    norm = value.strip().lower()
    if not norm:
        return True
    if norm in _MISSING_NAME_TOKENS:
        return True
    # things like "(not provided)", "n / a"
    stripped = re.sub(r'[\s()\[\]\.\-_/]', '', norm)
    return stripped in {'notprovided', 'notavailable', 'notspecified', 'na', 'none', 'null', 'unknown'}


def _sanitize_scorecard_identity(scorecard: Dict[str, Any], file_path: Optional[str]) -> Dict[str, Any]:
    """Replace placeholder identity fields in candidate_overview with real fallbacks.

    - name: derived from the resume filename when LLM returned "Not Provided" / similar.
    - email/phone/location: normalized to "Not Provided" when missing (no UI-jarring strings).
    """
    if not isinstance(scorecard, dict):
        return scorecard
    overview = scorecard.get('candidate_overview')
    if not isinstance(overview, dict):
        return scorecard

    if _is_missing_value(overview.get('name')):
        fallback = _derive_name_from_filename(file_path)
        if fallback:
            overview['name'] = fallback
            print(f"🪪 Filled missing candidate name from filename: {fallback}")

    for field in ('email', 'phone', 'location'):
        if _is_missing_value(overview.get(field)):
            overview[field] = 'Not Provided'

    return scorecard


_EVIDENCE_STOP_WORDS = {
    "and", "the", "for", "with", "from", "that", "this", "have", "has",
    "are", "was", "were", "will", "shall", "must", "should", "years",
    "year", "experience", "candidate", "developer", "engineer", "role",
    "position", "required", "requirements", "looking", "ideal",
}


def _clean_resume_line(value: Any) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    text = re.sub(r"^[•\-\*\u2022\s]+", "", text).strip()
    return text


def _extract_query_terms(query: str, scorecard: Dict[str, Any]) -> List[str]:
    terms = set()
    for word in re.findall(r"[A-Za-z][A-Za-z0-9+#.]{2,}", query or ""):
        lower = word.lower()
        if lower not in _EVIDENCE_STOP_WORDS:
            terms.add(lower)

    keyword_coverage = scorecard.get("keyword_coverage") if isinstance(scorecard, dict) else {}
    if isinstance(keyword_coverage, dict):
        for key in ("matched_keywords", "missing_keywords"):
            for keyword in keyword_coverage.get(key) or []:
                for word in re.findall(r"[A-Za-z][A-Za-z0-9+#.]{2,}", str(keyword)):
                    lower = word.lower()
                    if lower not in _EVIDENCE_STOP_WORDS:
                        terms.add(lower)

    return sorted(terms, key=len, reverse=True)


def _extract_resume_highlights(resume_content: str, query: str, scorecard: Dict[str, Any], limit: int = 4) -> List[str]:
    if not resume_content:
        return []

    terms = _extract_query_terms(query, scorecard)
    if not terms:
        return []

    raw_parts = re.split(r"[\r\n]+|(?<=[.!?])\s+", resume_content)
    scored = []
    seen = set()

    for part in raw_parts:
        line = _clean_resume_line(part)
        if len(line) < 28 or len(line) > 280:
            continue
        lower = line.lower()
        if "@" in line and len(line.split()) <= 6:
            continue

        score = sum(2 for term in terms if term in lower)
        if re.search(r"\b\d+\+?\s*(years?|yrs?)\b", lower):
            score += 2
        if re.search(r"\b\d+%|\b\d+x\b|\b\d+\+?\s*(projects?|clients?|users?|members?)\b", lower):
            score += 1
        if score <= 0:
            continue

        normalized = re.sub(r"\W+", "", lower)[:120]
        if normalized in seen:
            continue
        seen.add(normalized)
        scored.append((score, len(line), line))

    scored.sort(key=lambda item: (-item[0], item[1]))
    return [line for _, _, line in scored[:limit]]


def _extract_career_timeline(resume_content: str, scorecard: Dict[str, Any], limit: int = 4) -> List[Dict[str, Any]]:
    if not resume_content:
        return []

    month = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*"
    date_token = rf"(?:{month}\s+)?(?:19|20)\d{{2}}|Present|Current"
    date_range = re.compile(rf"({date_token})\s*(?:-|–|—|to)\s*({date_token})", re.I)
    role_words = re.compile(r"\b(engineer|developer|manager|analyst|consultant|lead|architect|specialist|scientist|recruiter|executive|intern)\b", re.I)
    terms = _extract_query_terms("", scorecard)

    timeline = []
    seen = set()
    for raw in re.split(r"[\r\n]+", resume_content):
        line = _clean_resume_line(raw)
        if len(line) < 18 or len(line) > 220:
            continue
        match = date_range.search(line)
        if not match:
            continue

        before = _clean_resume_line(line[:match.start()])
        after = _clean_resume_line(line[match.end():])
        context = before or after
        if not context or not role_words.search(context):
            continue

        pieces = [p.strip(" ,") for p in re.split(r"\s+(?:at|@)\s+|\s+[|]\s+|\s+[–—-]\s+", context, maxsplit=1) if p.strip(" ,")]
        role = pieces[0][:80] if pieces else "Position"
        company = pieces[1][:80] if len(pieces) > 1 else "Company not specified"
        year_range = f"{match.group(1)} - {match.group(2)}"
        unique_key = (role.lower(), company.lower(), year_range.lower())
        if unique_key in seen:
            continue
        seen.add(unique_key)

        lower_line = line.lower()
        key_skills = [term for term in terms if term in lower_line][:6]
        timeline.append({
            "role": role,
            "company": company,
            "year_range": year_range,
            "key_skills": key_skills,
        })
        if len(timeline) >= limit:
            break

    return timeline


def _enrich_scorecard_evidence(
    scorecard: Dict[str, Any],
    resume_content: Optional[str],
    query: str,
) -> Dict[str, Any]:
    """Fill missing evidence fields from resume text/snippets without changing scoring."""
    if not isinstance(scorecard, dict) or not resume_content:
        return scorecard

    analysis_summary = scorecard.get("analysis_summary")
    if not isinstance(analysis_summary, dict):
        analysis_summary = {}
        scorecard["analysis_summary"] = analysis_summary

    highlights = analysis_summary.get("resume_highlights")
    if not isinstance(highlights, list):
        highlights = []

    if len([h for h in highlights if h]) < 2:
        extracted = _extract_resume_highlights(resume_content, query, scorecard)
        if extracted:
            existing = {_clean_resume_line(h).lower() for h in highlights if h}
            for item in extracted:
                if _clean_resume_line(item).lower() not in existing:
                    highlights.append(item)
                if len(highlights) >= 4:
                    break
            analysis_summary["resume_highlights"] = highlights
            scorecard["resume_snippets"] = highlights

    timeline = scorecard.get("career_timeline")
    if not isinstance(timeline, list) or not timeline:
        extracted_timeline = _extract_career_timeline(resume_content, scorecard)
        if extracted_timeline:
            scorecard["career_timeline"] = extracted_timeline

    return scorecard


async def analyze_resume_with_hr_scorecard(resume_content: str, query: str, job_title: str = "Position", standardized_keywords: List[str] = None, token_tracker: TokenTracker = None, file_path: Optional[str] = None) -> Dict[str, Any]:
    """Generate comprehensive HR scorecard analysis for resume screening"""
    print(f"\n🏢 HR SCORECARD ANALYSIS STARTING")
    print(f"📝 Job requirements: '{query[:100]}...'")
    print(f"📄 Resume content length: {len(resume_content)} characters")

    # Retry configuration
    max_retries = 2
    retry_delay = 3  # seconds

    for attempt in range(max_retries + 1):
        print(f"\n🔄 LLM ATTEMPT {attempt + 1}/{max_retries + 1}")

        try:
            print(f"🔧 Initializing Gemini client...")
            client = gemini_client()
            print(f"✅ Gemini client initialized successfully")

            model = "gemini-2.5-flash"
            print(f"🧠 Using model: {model} for HR scorecard generation")

            # Create comprehensive HR scorecard prompt
            hr_scorecard_prompt = HR_SCORECARD_PROMPT.format(
                job_posting=query,
                resume_content=resume_content
            )

            contents = [
                types.Content(
                    role="user",
                    parts=[types.Part(text=hr_scorecard_prompt)]
                )
            ]

            generate_content_config = types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
                temperature=0.2,  # Lower temperature for consistent HR analysis
                top_p=0.8,
                # 10000 = comfortable headroom for Gemini 2.5 Flash (which
                # spends ~2.4-3.4k tokens on internal "thinking" before the
                # actual JSON response) so the scorecard always emits a full
                # schema. An earlier 8192 attempt occasionally truncated
                # mid-JSON, dropping fields like career_timeline /
                # detailed_analysis on some candidates.
                max_output_tokens=4096,
                safety_settings=[
                    types.SafetySetting(
                        category="HARM_CATEGORY_HATE_SPEECH",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_DANGEROUS_CONTENT",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_SEXUALLY_EXPLICIT",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_HARASSMENT",
                        threshold="OFF"
                    )
                ],
            )

            print(f"\n🚀 GENERATING HR SCORECARD WITH LLM...")
            print(f"📝 PROMPT BEING SENT TO GEMINI:")
            print(hr_scorecard_prompt[:10] + "... [TRUNCATED FOR READABILITY]")
            print(f"📊 Prompt length: {len(hr_scorecard_prompt)} characters")
            print(f"🧠 Model: {model}")
            print(f"🎯 Temperature: {generate_content_config.temperature}")
            print(f"⏰ Sending request to Gemini...")

            # Run blocking SDK call in a thread so event loop can process
            # other candidates concurrently (this is the dominant latency).
            response = await asyncio.to_thread(
                client.models.generate_content,
                model=model,
                contents=contents,
                config=generate_content_config,
            )

            # Track token usage
            if token_tracker:
                token_usage = extract_token_usage(response, "hr_scorecard_direct", model)
                token_tracker.add_call(
                    operation="hr_scorecard_direct",
                    model=model,
                    input_tokens=token_usage["input_tokens"],
                    output_tokens=token_usage["output_tokens"],
                    success=True
                )
                print(f"📊 Token usage - Input: {token_usage['input_tokens']}, Output: {token_usage['output_tokens']}, Total: {token_usage['total_tokens']}")

            print(f"✅ Response received from Gemini on attempt {attempt + 1}")

            # Extract response text with comprehensive null handling
            analysis_text = ""
            try:
                if hasattr(response, 'text') and response.text is not None:
                    analysis_text = response.text
                elif hasattr(response, 'candidates') and response.candidates is not None:
                    for candidate in response.candidates:
                        if candidate is not None and hasattr(candidate, 'content') and candidate.content is not None:
                            if hasattr(candidate.content, 'parts') and candidate.content.parts is not None:
                                for part in candidate.content.parts:
                                    if part is not None and hasattr(part, 'text') and part.text is not None:
                                        analysis_text += part.text
                else:
                    analysis_text = str(response) if response is not None else ""
            except Exception as extract_error:
                print(f"❌ Error extracting HR scorecard response text: {str(extract_error)}")
                analysis_text = str(response) if response is not None else ""

            if not analysis_text:
                print(f"❌ Empty HR scorecard response from Gemini")
                print(f"🔄 Attempting split approach fallback...")
                # Fallback to split approach but still apply standardized keywords
                split_result = await analyze_resume_with_split_approach(resume_content, query, job_title, standardized_keywords)

                # Keywords are now applied within the split approach function

                # If split approach also fails, create a basic analysis
                if not split_result.get('success'):
                    print(f"⚠️ Split approach also failed, creating basic analysis...")
                    basic_entities = create_basic_entities_fallback(resume_content)
                    basic_scorecard = create_basic_scorecard_fallback(basic_entities, query, job_title)

                    # Apply standardized keywords to basic scorecard if available
                    if standardized_keywords and len(standardized_keywords) > 0:
                        print(f"🔍 Applying standardized keywords to basic scorecard")
                        keyword_analysis = await asyncio.to_thread(analyze_candidate_keywords, resume_content, standardized_keywords, token_tracker)
                        basic_scorecard['keyword_coverage'] = keyword_analysis
                        basic_scorecard = adjust_scorecard_with_keyword_analysis(basic_scorecard, keyword_analysis)
                        print(f"✅ Applied standardized keywords to basic scorecard")

                    # Fill in missing identity fields from filename
                    basic_scorecard = _sanitize_scorecard_identity(basic_scorecard, file_path)
                    basic_scorecard = _enrich_scorecard_evidence(basic_scorecard, resume_content, query)

                    return {
                        "hr_scorecard": basic_scorecard,
                        "analysis": "Basic analysis created due to LLM response issues",
                        "match_score": basic_scorecard.get('candidate_overview', {}).get('overall_match_score', 50),
                        "analyzed_by": "Fallback Basic Analysis",
                        "success": True,
                        "scorecard_type": "basic_fallback",
                        "fallback_reason": "Empty LLM responses"
                    }

                # Sanitize identity fields on split-approach result too
                try:
                    if isinstance(split_result, dict) and isinstance(split_result.get('hr_scorecard'), dict):
                        split_result['hr_scorecard'] = _sanitize_scorecard_identity(
                            split_result['hr_scorecard'], file_path
                        )
                        split_result['hr_scorecard'] = _enrich_scorecard_evidence(
                            split_result['hr_scorecard'], resume_content, query
                        )
                except Exception as _se:
                    print(f"⚠️ split_result sanitize error: {_se}")
                return split_result

            print(f"📄 HR scorecard response length: {len(analysis_text)} characters")
            print(f"📄 FULL LLM RESPONSE:")
            print(analysis_text)

            # Parse JSON response with enhanced error handling
            try:
                # Clean the response text
                clean_text = analysis_text.strip()
                if clean_text.startswith('```json'):
                    clean_text = clean_text[7:]
                if clean_text.endswith('```'):
                    clean_text = clean_text[:-3]
                clean_text = clean_text.strip()

                # Handle potential truncation and malformed JSON
                if not clean_text.endswith('}'):
                    print(f"🔧 JSON appears truncated, attempting reconstruction...")

                    # Try to find the last complete section - more comprehensive patterns
                    import re
                    patterns = [
                        # Try to end at complete career timeline
                        r'.*"career_timeline":\s*\[[^\]]*\]',
                        # Try to end at resume snippets
                        r'.*"resume_snippets":\s*\[[^\]]*\]',
                        # Try to end at ai_summary
                        r'.*"ai_summary":\s*"[^"]*"',
                        # Try to end at keyword coverage
                        r'.*"matched_keywords":\s*\[[^\]]*\]',
                        # Try to end at score breakdown
                        r'.*"soft_skills":\s*\{[^}]*\}',
                        # Fallback patterns
                        r'.*"availability":\s*"[^"]*"',
                        r'.*"salary_expectation":\s*"[^"]*"',
                        r'.*"industry_experience":\s*"[^"]*"',
                        r'.*"strengths":\s*\[[^\]]*\]'
                    ]

                    reconstructed = False
                    for i, pattern in enumerate(patterns):
                        match = re.search(pattern, clean_text, re.DOTALL)
                        if match:
                            clean_text = match.group(0)
                            # Close any open objects/arrays
                            open_braces = clean_text.count('{') - clean_text.count('}')
                            open_brackets = clean_text.count('[') - clean_text.count(']')
                            clean_text += ']' * open_brackets + '}' * open_braces
                            print(f"🔧 Reconstructed using pattern {i+1}: {pattern[:50]}...")
                            reconstructed = True
                            break

                    if not reconstructed:
                        # Last resort: try to remove the incomplete last line and close properly
                        lines = clean_text.split('\n')
                        if len(lines) > 1:
                            # Remove the last line if it looks incomplete
                            last_line = lines[-1].strip()
                            if last_line and not last_line.endswith(('"', '}', ']', ',')):
                                print(f"🔧 Removing incomplete last line: {last_line}")
                                clean_text = '\n'.join(lines[:-1])

                        # Close any open structures
                        open_braces = clean_text.count('{') - clean_text.count('}')
                        open_brackets = clean_text.count('[') - clean_text.count(']')

                        # Add missing closing characters
                        clean_text += ']' * open_brackets + '}' * open_braces
                        print(f"🔧 Last resort reconstruction: added {open_brackets} brackets, {open_braces} braces")

                # Fix common JSON issues before parsing
                import re

                # Handle unterminated strings at end of lines
                clean_text = re.sub(r':\s*"([^"]*?)$', r': "\1"', clean_text, flags=re.MULTILINE)

                # Handle unterminated strings in the middle that might be missing closing quote
                clean_text = re.sub(r':\s*"([^"]*?)\n\s*[,}]', r': "\1",', clean_text)

                # Remove any trailing commas before closing braces/brackets
                clean_text = re.sub(r',(\s*[}\]])', r'\1', clean_text)

                # If JSON still incomplete, try to close it properly
                if not clean_text.endswith('}'):
                    # Count unclosed braces and brackets
                    open_braces = clean_text.count('{') - clean_text.count('}')
                    open_brackets = clean_text.count('[') - clean_text.count(']')

                    # Add missing closing characters
                    clean_text += ']' * open_brackets + '}' * open_braces
                    print("🔧 Added {} closing brackets and {} closing braces".format(open_brackets, open_braces))

                try:
                    scorecard_json = json.loads(clean_text)
                    overall_score = scorecard_json.get('candidate_overview', {}).get('overall_match_score', 0)

                    # ALWAYS apply standardized keyword analysis (replace any LLM-generated keywords)
                    print(f"\n🔍 KEYWORD APPLICATION CHECKPOINT:")
                    print(f"   🔑 standardized_keywords: {standardized_keywords}")
                    print(f"   🔑 standardized_keywords is None: {standardized_keywords is None}")
                    print(f"   🔑 standardized_keywords length: {len(standardized_keywords) if standardized_keywords else 0}")
                    print(f"   📋 Current scorecard keyword_coverage before replacement: {scorecard_json.get('keyword_coverage', 'NOT_FOUND')}")

                    if standardized_keywords and len(standardized_keywords) > 0:
                        print(f"\n✅ APPLYING STANDARDIZED KEYWORD ANALYSIS")
                        print(f"🔑 Using {len(standardized_keywords)} keywords for analysis")
                        print(f"🎯 Keywords: {standardized_keywords[:10]}{'...' if len(standardized_keywords) > 10 else ''}")
                        print(f"🔄 Calling analyze_candidate_keywords() function...")

                        keyword_analysis = await asyncio.to_thread(analyze_candidate_keywords, resume_content, standardized_keywords, token_tracker)

                        print(f"📊 analyze_candidate_keywords() returned: {keyword_analysis}")

                        # FORCE replace the keyword_coverage with standardized analysis
                        print(f"🔄 Replacing scorecard keyword_coverage with standardized analysis...")
                        scorecard_json['keyword_coverage'] = keyword_analysis
                        print(f"✅ Replacement complete. New keyword_coverage: {scorecard_json['keyword_coverage']}")

                        print(f"✅ Standardized keyword analysis: {keyword_analysis['jd_keywords_matched']}/{keyword_analysis['total_jd_keywords']} matched")
                        print(f"🔑 Matched ({len(keyword_analysis['matched_keywords'])}): {keyword_analysis['matched_keywords'][:5]}{'...' if len(keyword_analysis['matched_keywords']) > 5 else ''}")
                        print(f"❌ Missing ({len(keyword_analysis['missing_keywords'])}): {keyword_analysis['missing_keywords'][:5]}{'...' if len(keyword_analysis['missing_keywords']) > 5 else ''}")
                        print(f"🎯 TOTAL KEYWORDS USED: {keyword_analysis['total_jd_keywords']} (SAME for ALL candidates)")
                    else:
                        print(f"\n❌ NO STANDARDIZED KEYWORDS - THIS IS THE PROBLEM!")
                        print(f"⚠️ standardized_keywords is: {standardized_keywords}")
                        print(f"⚠️ Length is: {len(standardized_keywords) if standardized_keywords else 'None/Empty'}")
                        print(f"🔄 Keeping existing LLM-generated keyword analysis as fallback")
                        # Keep the existing keyword_coverage from LLM if we have no standardized keywords
                        existing_keywords = scorecard_json.get('keyword_coverage', {})
                        if not existing_keywords.get('matched_keywords') and not existing_keywords.get('missing_keywords'):
                            print(f"🔄 Creating empty keyword analysis to maintain consistency")
                            scorecard_json['keyword_coverage'] = {
                                "jd_keywords_matched": 0,
                                "total_jd_keywords": 0,
                                "missing_keywords": [],
                                "matched_keywords": []
                            }

                    print(f"✅ HR scorecard generated successfully")
                    print(f"👤 Candidate: {scorecard_json.get('candidate_overview', {}).get('name', 'Unknown')}")
                    print(f"🎯 Overall match score: {overall_score}%")

                    # Add completion separator to distinguish between resume processes
                    print("=" * 50)

                    # Adjust scorecard with keyword analysis
                    adjusted_scorecard = adjust_scorecard_with_keyword_analysis(scorecard_json, scorecard_json['keyword_coverage'])

                    # Fill in missing identity fields (name/email/phone) with sane fallbacks
                    adjusted_scorecard = _sanitize_scorecard_identity(adjusted_scorecard, file_path)
                    adjusted_scorecard = _enrich_scorecard_evidence(adjusted_scorecard, resume_content, query)

                    return {
                        "hr_scorecard": adjusted_scorecard,
                        "analysis": analysis_text,
                        "match_score": adjusted_scorecard.get('candidate_overview', {}).get('overall_match_score', 0),
                        "analyzed_by": "Gemini 2.5 Flash HR Scorecard",
                        "success": True,
                        "scorecard_type": "comprehensive_hr"
                    }

                except json.JSONDecodeError as e:
                    print(f"⚠️ Failed to parse HR scorecard JSON: {str(e)}")
                    print(f"🔧 Problematic JSON snippet: {clean_text[-200:]}")
                    print(f"🔍 DETAILED HR SCORECARD JSON ERROR:")
                    print(f"   - Error message: {str(e)}")
                    print(f"   - Error position: {e.pos if hasattr(e, 'pos') else 'Unknown'}")
                    print(f"   - Error line: {e.lineno if hasattr(e, 'lineno') else 'Unknown'}")
                    print(f"   - Error column: {e.colno if hasattr(e, 'colno') else 'Unknown'}")
                    print(f"   - Attempt number: {attempt + 1}")
                    print(f"   - Full clean text:\n{clean_text}")
                    print(f"   - Original analysis text:\n{analysis_text}")
                    raise Exception(f"DETAILED ERROR: HR scorecard JSON parsing failed on attempt {attempt + 1} - {str(e)}. Full response: {analysis_text}")

            except Exception as json_error:
                print(f"❌ JSON parsing failed: {str(json_error)}")
                print(f"🔍 DETAILED HR SCORECARD PROCESSING ERROR:")
                print(f"   - Error message: {str(json_error)}")
                print(f"   - Error type: {type(json_error).__name__}")
                print(f"   - Error args: {json_error.args}")
                print(f"   - Attempt number: {attempt + 1}")
                print(f"   - Analysis text length: {len(analysis_text) if analysis_text else 0}")
                print(f"   - Analysis text preview: {analysis_text[:300] if analysis_text else 'None'}")
                raise Exception(f"DETAILED ERROR: HR scorecard processing failed on attempt {attempt + 1} - {str(json_error)}. Analysis text: {analysis_text}")

        except Exception as attempt_error:
            print(f"❌ Attempt {attempt + 1} failed: {str(attempt_error)}")
            if attempt < max_retries:
                print(f"🔄 Retrying in {retry_delay} seconds...")
                await asyncio.sleep(retry_delay)
                continue
            else:
                print(f"❌ All attempts failed, falling back to split approach...")
                break

    # If we reach here, all attempts failed
    print(f"❌ HR SCORECARD ANALYSIS FAILED: All attempts exhausted")
    print(f"🔍 DETAILED HR SCORECARD FAILURE ANALYSIS:")
    print(f"   - Max retries: {max_retries}")
    print(f"   - Retry delay: {retry_delay} seconds")
    print(f"   - Resume content length: {len(resume_content)} characters")
    print(f"   - Query length: {len(query)} characters")
    print(f"   - Job title: {job_title}")
    print(f"   - Standardized keywords: {standardized_keywords}")
    # Resume content preview redacted to avoid leaking PII into logs.

    # Collect all the error details and raise a comprehensive error
    error_details = []
    error_details.append(f"All {max_retries} attempts failed for HR scorecard analysis")
    error_details.append(f"Resume length: {len(resume_content)} characters")
    error_details.append(f"Query: {query}")
    error_details.append(f"Job title: {job_title}")

    comprehensive_error = "\n".join(error_details)
    raise Exception(f"DETAILED ERROR: HR SCORECARD ANALYSIS COMPLETELY FAILED\n{comprehensive_error}")

def analyze_resume_with_gemini(resume_content: str, query: str) -> Dict[str, Any]:
    """Analyze resume content with Gemini LLM and provide matching score"""
    print(f"\n🤖 GEMINI ANALYSIS STARTING")
    print(f"📝 Query for analysis: '{query}'")
    print(f"📄 Resume content length: {len(resume_content)} characters")

    try:
        print(f"🔧 Initializing Gemini client...")
        client = gemini_client()
        print(f"✅ Gemini client initialized successfully")

        model = "gemini-2.5-flash"
        print(f"🧠 Using model: {model}")

        # Create analysis prompt
        print(f"📝 Creating analysis prompt...")
        analysis_prompt = RESUME_CONTENT_ANALYSIS_PROMPT.format(
            job_posting=query,
            resume_content=resume_content
        )

        contents = [
            types.Content(
                role="user",
                parts=[
                    types.Part(text=analysis_prompt)
                ]
            )
        ]

        generate_content_config = types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
            temperature=0.3,  # Lower temperature for more consistent analysis
            top_p=0.8,
            max_output_tokens=5000,
            safety_settings=[
                types.SafetySetting(
                    category="HARM_CATEGORY_HATE_SPEECH",
                    threshold="OFF"
                ),
                types.SafetySetting(
                    category="HARM_CATEGORY_DANGEROUS_CONTENT",
                    threshold="OFF"
                ),
                types.SafetySetting(
                    category="HARM_CATEGORY_SEXUALLY_EXPLICIT",
                    threshold="OFF"
                ),
                types.SafetySetting(
                    category="HARM_CATEGORY_HARASSMENT",
                    threshold="OFF"
                )
            ],
        )

        # Generate content
        print(f"🚀 Sending request to Gemini...")
        response = client.models.generate_content(
            model=model,
            contents=contents,
            config=generate_content_config,
        )
        print(f"✅ Received response from Gemini")

        # Extract response text with comprehensive null handling
        analysis_text = ""
        try:
            if hasattr(response, 'text') and response.text is not None:
                analysis_text = response.text
            elif hasattr(response, 'candidates') and response.candidates is not None:
                for candidate in response.candidates:
                    if candidate is not None and hasattr(candidate, 'content') and candidate.content is not None:
                        if hasattr(candidate.content, 'parts') and candidate.content.parts is not None:
                            for part in candidate.content.parts:
                                if part is not None and hasattr(part, 'text') and part.text is not None:
                                    analysis_text += part.text
            else:
                analysis_text = str(response) if response is not None else ""
        except Exception as extract_error:
            print(f"❌ Error extracting basic analysis response text: {str(extract_error)}")
            analysis_text = str(response) if response is not None else ""

        if not analysis_text:
            print(f"❌ Empty response from Gemini")
            return {
                "analysis": "Empty response from Gemini",
                "match_score": 0,
                "analyzed_by": "Gemini 2.5 Flash",
                "success": False,
                "error": "Empty response"
            }

        print(f"📄 Analysis text length: {len(analysis_text)} characters")
        print(f"📝 Analysis preview: {analysis_text[:300]}..." if len(analysis_text) > 300 else f"📝 Full analysis: {analysis_text}")

        # Try to parse JSON response
        print(f"🔍 Parsing JSON response from Gemini...")
        try:
            # Clean the response text - remove any markdown formatting
            clean_text = analysis_text.strip()
            if clean_text.startswith('```json'):
                clean_text = clean_text[7:]
            if clean_text.endswith('```'):
                clean_text = clean_text[:-3]
            clean_text = clean_text.strip()

            # Parse JSON
            analysis_json = json.loads(clean_text)
            raw_score = analysis_json.get('match_score') or analysis_json.get('overall_match_score') or analysis_json.get('candidate_overview', {}).get('overall_match_score') or 0
            score = clamp_score(raw_score)
            if score != raw_score:
                print(f"⚠️ Clamped out-of-range LLM score {raw_score!r} → {score}")

            print(f"✅ Successfully parsed JSON response")
            print(f"👤 Candidate: {analysis_json.get('candidate_name', 'Unknown')}")
            print(f"🎯 Match score: {score}")

            final_result = {
                "analysis": analysis_text,
                "analysis_json": analysis_json,
                "match_score": score,
                "analyzed_by": "Gemini 2.5 Flash",
                "success": True
            }
            print(f"✅ Gemini analysis completed successfully")
            return final_result

        except json.JSONDecodeError as e:
            print(f"⚠️ Failed to parse JSON, falling back to text analysis: {str(e)}")
            # Fallback to old score extraction method
            score = clamp_score(extract_score_from_analysis(analysis_text))

            final_result = {
                "analysis": analysis_text,
                "analysis_json": None,
                "match_score": score,
                "analyzed_by": "Gemini 2.5 Flash",
                "success": True,
                "json_parse_error": str(e)
            }
            print(f"✅ Gemini analysis completed with fallback")
            return final_result

    except Exception as e:
        print(f"❌ GEMINI ANALYSIS FAILED: {str(e)}")
        error_result = {
            "analysis": f"Analysis failed: {str(e)}",
            "match_score": 0,
            "analyzed_by": "Gemini 2.5 Flash",
            "success": False,
            "error": str(e)
        }
        print(f"📄 Returning error result: {error_result}")
        return error_result

# Initialize folder on startup
@app.on_event("startup")
async def startup_event():
    """Startup event with robust error handling"""
    print("🚀 Starting HR Agent API...")

    # Increase asyncio default executor pool so blocking GCS/PDF extraction
    # and Gemini SDK calls (wrapped in asyncio.to_thread) can run more
    # candidates concurrently. Default is min(32, cpu+4) which is too small
    # on Cloud Run with 2 vCPU.
    try:
        import concurrent.futures, asyncio as _asyncio
        loop = _asyncio.get_running_loop()
        loop.set_default_executor(concurrent.futures.ThreadPoolExecutor(max_workers=64))
        print("✅ asyncio default executor scaled to 64 workers")
    except Exception as _exec_err:
        print(f"⚠️ Could not scale executor: {_exec_err}")

    try:
        # Try to ensure bucket folder exists
        bucket_success = ensure_bucket_folder_exists()
        if bucket_success:
            print("✅ GCS bucket folder setup completed")
        else:
            print("⚠️ GCS bucket folder setup skipped - continuing with limited functionality")
    except Exception as e:
        print(f"⚠️ Error during GCS setup: {e} - continuing without GCS functionality")

    try:
        # Test database connection
        from database import get_db_manager
        db = get_db_manager()
        if db:
            print("✅ Database connection initialized successfully")
            try:
                db.ensure_perf_indexes()
                print("✅ Performance indexes ensured")
            except Exception as _idx_err:
                print(f"⚠️ Could not ensure perf indexes: {_idx_err}")
            try:
                db.ensure_integrations_tables()
                # Seed global (company_id IS NULL) default email templates
                # if none yet exist.
                from services.email_service import DEFAULT_TEMPLATES
                existing = db.list_email_templates(company_id=None)
                existing_names = {(r['kind'], r['name']) for r in (existing or [])}
                for tpl in DEFAULT_TEMPLATES:
                    if (tpl['kind'], tpl['name']) in existing_names:
                        continue
                    db.upsert_email_template(
                        company_id=None,
                        kind=tpl['kind'],
                        name=tpl['name'],
                        subject=tpl['subject'],
                        body=tpl['body'],
                        is_default=tpl.get('is_default', False),
                    )
                print("✅ Integration tables ensured and default email templates seeded")
            except Exception as _intg_err:
                print(f"⚠️ Integrations setup skipped: {_intg_err}")
            try:
                cleanup_stats = db.cleanup_stale_data()
                print(f"🧹 Retention cleanup: {cleanup_stats}")
            except Exception as _cleanup_err:
                print(f"⚠️ Retention cleanup failed: {_cleanup_err}")
        else:
            print("⚠️ Database connection failed - some features may not work")
    except Exception as e:
        print(f"⚠️ Database setup error: {e} - continuing with limited functionality")

    print("🎉 HR Agent API startup completed")

    if os.getenv("SMARTHR_LOCAL_MODE", "").lower() in {"1", "true", "yes", "on"}:
        try:
            import asyncio as _asyncio
            from vps_local.discovery_shim import warm_search_runtime

            await _asyncio.wait_for(
                _asyncio.to_thread(warm_search_runtime),
                timeout=float(os.getenv("SMARTHR_SEARCH_WARMUP_TIMEOUT", "30")),
            )
            print("Local search runtime warmed")
        except Exception as _warm_err:
            print(f"Local search warmup skipped: {_warm_err}")

# Health check endpoint for Cloud Run
async def health_check():
    """Lightweight liveness probe for Cloud Run / nginx upstream checks."""
    return {"status": "healthy", "timestamp": time.time()}


# Cache the result of the Gemini ping so we don't hit the API on every probe.
_gemini_ping_cache: dict = {"ts": 0.0, "ok": None, "err": None}


def _gemini_ping_cached(ttl_sec: float = 60.0) -> tuple[bool, str | None]:
    """Verify the Gemini client can be constructed (and, on first call this
    minute, list available models). Result is cached for *ttl_sec* seconds so
    /health/deep is cheap to poll."""
    now = time.time()
    if _gemini_ping_cache["ok"] is not None and (now - _gemini_ping_cache["ts"]) < ttl_sec:
        return bool(_gemini_ping_cache["ok"]), _gemini_ping_cache["err"]
    try:
        c = gemini_client()
        # Cheapest call we know: list models (shim short-circuits to ['stub'])
        next(iter(c.models.list()), None)
        _gemini_ping_cache.update(ts=now, ok=True, err=None)
        return True, None
    except Exception as e:
        msg = str(e)[:200]
        _gemini_ping_cache.update(ts=now, ok=False, err=msg)
        return False, msg


async def health_check_deep():
    """Readiness probe that pings critical dependencies.

    Returns 200 if every dependency that is configured is reachable; 503
    otherwise. Each dependency reports its own pass/fail so a caller can
    triage. Optional dependencies that aren't configured are reported as
    'skipped' and do not fail the overall check.
    """
    from fastapi import Response
    import json as _json

    checks: dict[str, dict] = {}
    overall_ok = True

    # Database + pgvector (combined roundtrip)
    db_start = time.time()
    try:
        from database import get_db_manager
        db = get_db_manager()
        if db is None:
            checks["database"] = {"status": "fail", "error": "no manager"}
            overall_ok = False
        elif hasattr(db, "get_cursor"):
            # Real DatabaseManager API: context-managed cursor.
            with db.get_cursor(dict_cursor=False) as (cur, _conn):
                cur.execute("SELECT 1")
                cur.fetchone()
                checks["database"] = {"status": "ok", "ms": int((time.time() - db_start) * 1000)}
                # pgvector liveness: count embedding rows
                pv_start = time.time()
                try:
                    cur.execute("SELECT COUNT(*) FROM resume_embeddings")
                    (row_count,) = cur.fetchone()
                    checks["pgvector"] = {
                        "status": "ok",
                        "rows": int(row_count),
                        "ms": int((time.time() - pv_start) * 1000),
                    }
                except Exception as e:
                    checks["pgvector"] = {"status": "fail", "error": str(e)[:200]}
                    overall_ok = False
        else:
            conn = db.get_connection() if hasattr(db, "get_connection") else None
            if conn is not None:
                try:
                    cur = conn.cursor()
                    cur.execute("SELECT 1")
                    cur.fetchone()
                    checks["database"] = {"status": "ok", "ms": int((time.time() - db_start) * 1000)}
                    # pgvector liveness: count embedding rows
                    pv_start = time.time()
                    try:
                        cur.execute("SELECT COUNT(*) FROM resume_embeddings")
                        (row_count,) = cur.fetchone()
                        checks["pgvector"] = {
                            "status": "ok",
                            "rows": int(row_count),
                            "ms": int((time.time() - pv_start) * 1000),
                        }
                    except Exception as e:
                        checks["pgvector"] = {"status": "fail", "error": str(e)[:200]}
                        overall_ok = False
                    cur.close()
                finally:
                    try:
                        conn.close()
                    except Exception:
                        pass
            else:
                checks["database"] = {"status": "ok", "ms": int((time.time() - db_start) * 1000), "note": "no usable connection api"}
    except Exception as e:
        checks["database"] = {"status": "fail", "error": str(e)[:200]}
        overall_ok = False

    # Storage roundtrip: write + read + delete a probe file under /app/storage.
    # Under SMARTHR_LOCAL_MODE this exercises the local filesystem mount; on
    # Cloud Run it would still resolve to the configured GCS-shim path.
    try:
        st_start = time.time()
        probe_dir = os.path.join(os.environ.get("SMARTHR_STORAGE_ROOT", "/app/storage"), "_healthcheck")
        os.makedirs(probe_dir, exist_ok=True)
        probe_path = os.path.join(probe_dir, f"probe_{int(time.time()*1000)}.tmp")
        with open(probe_path, "w", encoding="utf-8") as f:
            f.write("ok")
        with open(probe_path, "r", encoding="utf-8") as f:
            ok = f.read() == "ok"
        try:
            os.remove(probe_path)
        except OSError:
            pass
        if ok:
            checks["storage"] = {"status": "ok", "ms": int((time.time() - st_start) * 1000), "path": probe_dir}
        else:
            checks["storage"] = {"status": "fail", "error": "read mismatch"}
            overall_ok = False
    except Exception as e:
        checks["storage"] = {"status": "fail", "error": str(e)[:200]}
        overall_ok = False

    # GCS (only if configured)
    try:
        bucket = os.environ.get("GCS_BUCKET_NAME")
        if not bucket:
            try:
                _cfg = load_config()
                bucket = _cfg.get("gcs_bucket") if isinstance(_cfg, dict) else None
            except Exception:
                bucket = None
        if not bucket:
            checks["gcs"] = {"status": "skipped"}
        else:
            from google.cloud import storage
            gcs_start = time.time()
            client = storage.Client()
            client.bucket(bucket).exists()
            checks["gcs"] = {"status": "ok", "bucket": bucket, "ms": int((time.time() - gcs_start) * 1000)}
    except Exception as e:
        checks["gcs"] = {"status": "fail", "error": str(e)[:200]}
        overall_ok = False

    # LLM provider ping (cached 60s to avoid hitting the API on every probe).
    # With the resilient wrapper, "ok" means *some* provider can serve requests
    # (Gemini preferred, OpenAI fallback). We surface circuit-breaker state so
    # operators can see when we've degraded onto OpenAI.
    try:
        has_gemini_creds = bool(
            os.environ.get("GEMINI_API_KEY")
            or os.environ.get("GOOGLE_GEMINI_API_KEY")
            or os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")
            or os.environ.get("GOOGLE_CLOUD_PROJECT")
        )
        has_openai = bool(os.environ.get("OPENAI_API_KEY"))
        if not has_gemini_creds and not has_openai:
            checks["gemini"] = {"status": "skipped", "note": "no LLM credentials in env"}
        else:
            gem_start = time.time()
            ok, err = _gemini_ping_cached()
            circuit_open = _gemini_circuit_open()
            provider = "gemini"
            if circuit_open:
                provider = "openai" if has_openai else "none"
            entry = {
                "ms": int((time.time() - gem_start) * 1000),
                "cached_ttl_s": 60,
                "provider": provider,
                "gemini_circuit_open": circuit_open,
                "gemini_fail_count": _gemini_circuit_state.get("fail_count", 0),
            }
            if ok:
                # Healthy via *some* provider. If we're on OpenAI fallback,
                # surface that as a warning but don't fail the overall check
                # (the system is serving requests).
                entry["status"] = "ok"
                if circuit_open:
                    entry["note"] = "serving via OpenAI fallback (Gemini circuit open)"
                    entry["last_gemini_error"] = (
                        _gemini_circuit_state.get("last_error", "") or ""
                    )[:200]
                checks["gemini"] = entry
            else:
                entry["status"] = "fail"
                entry["error"] = err or "unknown"
                checks["gemini"] = entry
                overall_ok = False
    except Exception as e:
        checks["gemini"] = {"status": "fail", "error": str(e)[:200]}
        overall_ok = False

    body = {"status": "healthy" if overall_ok else "degraded",
            "timestamp": time.time(),
            "checks": checks}
    return Response(
        content=_json.dumps(body),
        media_type="application/json",
        status_code=200 if overall_ok else 503,
    )


# Register on both /health and /healthz. Stacking @app.get decorators only
# registers one of the routes on some FastAPI versions; using add_api_route
# explicitly is the safe pattern.
app.add_api_route("/health", health_check, methods=["GET"])
app.add_api_route("/healthz", health_check, methods=["GET"])
app.add_api_route("/health/deep", health_check_deep, methods=["GET"])

# Authentication dependency
security = HTTPBearer()

async def get_current_user(
    request: Request,
    session_token: Optional[str] = Cookie(None)
):
    """Get the current authenticated user from session"""
    if not session_token:
        return None

    db = get_db_manager()
    user = db.verify_session(session_token)
    return user

async def require_auth(
    request: Request,
    session_token: Optional[str] = Cookie(None)
):
    """Require authentication for protected routes"""
    user = await get_current_user(request, session_token)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return user

async def require_super_admin(
    request: Request,
    session_token: Optional[str] = Cookie(None)
):
    """Require super admin role"""
    user = await require_auth(request, session_token)
    if user['user_type'] != 'super_admin':
        raise HTTPException(status_code=403, detail="Super admin access required")
    return user

async def require_tenant_admin(
    request: Request,
    session_token: Optional[str] = Cookie(None)
):
    """Require tenant admin role"""
    user = await require_auth(request, session_token)
    if user['user_type'] not in ['super_admin', 'tenant_admin']:
        raise HTTPException(status_code=403, detail="Admin access required")
    return user

@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request, session_token: Optional[str] = Cookie(None)):
    """Root route - redirect to login if not authenticated"""
    user = await get_current_user(request, session_token)
    if not user:
        return RedirectResponse(url="/login", status_code=302)
    return templates.TemplateResponse(request, "index.html", {"user": user})

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    """Login page"""
    return templates.TemplateResponse(request, "login.html")

@app.post("/api/login")
@limiter.limit("5/minute")
async def login(
    request: Request,
    email: str = Form(...),
    password: str = Form(...)
):
    """Handle login"""
    db = get_db_manager()
    user = db.authenticate_user(email, password)

    if not user:
        raise HTTPException(status_code=401, detail="Invalid email or password")

    # Create session
    ip_address = request.client.host
    user_agent = request.headers.get('user-agent')
    session_token = db.create_session(user['id'], ip_address, user_agent)

    # Create response with session cookie
    response = JSONResponse({
        "success": True,
        "user": user,
        "redirect": "/dashboard" if user['user_type'] == 'super_admin' else "/"
    })

    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        # `strict` blocks the cookie on any cross-site request, including
        # top-level GET navigations from external sites (e.g. an email
        # link). That's the safe default for an internal HR tool. If we
        # later need bookmarklets or external SSO redirects we can soften
        # this back to "lax".
        samesite="strict",
        max_age=7 * 24 * 60 * 60  # 7 days
    )

    response.set_cookie(
        key=CSRF_COOKIE_NAME,
        value=secrets.token_urlsafe(32),
        httponly=False,
        secure=True,
        samesite="strict",
        max_age=7 * 24 * 60 * 60
    )

    return response

@app.post("/api/logout")
async def logout(
    request: Request,
    session_token: Optional[str] = Cookie(None)
):
    """Handle logout"""
    if session_token:
        db = get_db_manager()
        db.logout_user(session_token)

    response = JSONResponse({"success": True})
    response.delete_cookie("session_token")
    response.delete_cookie(CSRF_COOKIE_NAME)
    return response

@app.get("/api/me")
async def get_me(user: dict = Depends(require_auth)):
    """Get current user info"""
    return {"user": user}

# Admin routes
@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request, user: dict = Depends(require_auth)):
    """Admin dashboard"""
    if user['user_type'] == 'tenant_user':
        return RedirectResponse(url="/", status_code=302)
    return templates.TemplateResponse(request, "dashboard.html", {"user": user})

@app.get("/api/companies")
async def get_companies(user: dict = Depends(require_super_admin)):
    """Get all companies (super admin only)"""
    try:
        db = get_db_manager()
        companies = db.get_all_companies()
        return {"companies": companies}
    except Exception as e:
        logger.error(f"Error fetching companies: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch companies")

@app.post("/api/companies")
async def create_company(
    company_name: str = Form(...),
    company_code: str = Form(...),
    subscription_plan: str = Form("basic"),
    max_users: int = Form(10),
    max_resumes: int = Form(1000),
    max_searches: int = Form(10000),
    user: dict = Depends(require_super_admin)
):
    """Create a new company (super admin only)"""
    db = get_db_manager()
    try:
        # Create company in database
        company_id = db.create_tenant_company(
            company_name=company_name,
            company_code=company_code,
            created_by=user['id'],
            subscription_plan=subscription_plan,
            max_users=max_users,
            max_resumes=max_resumes,
            max_searches=max_searches
        )

        # Create company-specific GCS bucket
        gcs_bucket_name = create_company_gcs_bucket(company_code)

        # Create company-specific AI Search datastore
        company_datastore_id = create_company_datastore(company_code, PROJECT_ID, LOCATION)

        # Update company with the created resources
        if gcs_bucket_name or company_datastore_id:
            db.update_company_resources(company_id, gcs_bucket_name, company_datastore_id)

        return {"success": True, "company_id": company_id}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/api/companies/{company_id}/set-resources")
async def set_company_resources(
    company_id: int,
    gcs_bucket_name: str = Form(None),
    datastore_id: str = Form(None),
    user: dict = Depends(require_super_admin)
):
    """Explicitly set the GCS bucket and/or Discovery Engine datastore for a tenant.
    Use to point a tenant at existing pre-provisioned resources (e.g. shared global ones)."""
    db = get_db_manager()
    company = db.get_company_by_id(company_id)
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")
    db.update_company_resources(company_id, gcs_bucket_name, datastore_id)
    refreshed = db.get_company_by_id(company_id)
    return {
        "success": True,
        "company_id": company_id,
        "gcs_bucket_name": refreshed.get('gcs_bucket_name'),
        "datastore_id": refreshed.get('datastore_id'),
    }

@app.post("/api/companies/{company_id}/provision-resources")
async def provision_company_resources(
    company_id: int,
    force: bool = Form(False),
    user: dict = Depends(require_super_admin)
):
    """Provision (or re-provision) GCS bucket and Discovery Engine datastore for an existing tenant.
    Idempotent: skips creation if the company already has resources unless force=true."""
    db = get_db_manager()
    company = db.get_company_by_id(company_id)
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")
    company_code = company.get('company_code')
    if not company_code:
        raise HTTPException(status_code=400, detail="Company missing company_code")

    existing_bucket = company.get('gcs_bucket_name')
    existing_datastore = company.get('datastore_id')

    new_bucket = existing_bucket
    new_datastore = existing_datastore

    bucket_action = "kept"
    datastore_action = "kept"
    errors = []

    if not existing_bucket or force:
        try:
            new_bucket = create_company_gcs_bucket(company_code)
            bucket_action = "created"
        except Exception as e:
            errors.append(f"bucket: {e}")
            logger.exception("Failed to provision bucket for company %s", company_code)

    if not existing_datastore or force:
        try:
            new_datastore = create_company_datastore(company_code, PROJECT_ID, LOCATION)
            datastore_action = "created"
        except Exception as e:
            errors.append(f"datastore: {e}")
            logger.exception("Failed to provision datastore for company %s", company_code)

    if (new_bucket and new_bucket != existing_bucket) or (new_datastore and new_datastore != existing_datastore):
        db.update_company_resources(company_id, new_bucket, new_datastore)

    return {
        "success": len(errors) == 0,
        "company_id": company_id,
        "company_code": company_code,
        "gcs_bucket_name": new_bucket,
        "datastore_id": new_datastore,
        "bucket_action": bucket_action,
        "datastore_action": datastore_action,
        "errors": errors,
    }

@app.get("/api/companies/{company_id}/users")
async def get_company_users(
    company_id: int,
    limit: int = 500,
    offset: int = 0,
    user: dict = Depends(require_tenant_admin)
):
    """Get users in a company (paginated; max 1000 per page)."""
    # Check if user has access to this company
    user_company_id = None
    if user.get('company'):
        if isinstance(user['company'], dict):
            user_company_id = user['company'].get('company_id') or user['company'].get('id')
        else:
            user_company_id = getattr(user['company'], 'id', None)

    if user['user_type'] != 'super_admin' and user_company_id != company_id:
        raise HTTPException(status_code=403, detail="Access denied")

    db = get_db_manager()
    users = db.get_company_users(company_id, limit=limit, offset=offset)
    return {"users": users, "limit": limit, "offset": offset, "count": len(users)}

@app.post("/api/users")
async def create_user(
    email: str = Form(...),
    password: str = Form(...),
    full_name: str = Form(...),
    user_type: str = Form(...),
    company_id: Optional[int] = Form(None),
    user: dict = Depends(require_tenant_admin)
):
    """Create a new user"""
    # Validate permissions
    if user['user_type'] == 'tenant_admin':
        # Get user's company ID safely
        user_company_id = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                user_company_id = user['company'].get('company_id') or user['company'].get('id')
            else:
                user_company_id = getattr(user['company'], 'id', None)

        # Tenant admin can only create tenant_user in their company
        if user_type != 'tenant_user' or company_id != user_company_id:
            raise HTTPException(status_code=403, detail="Access denied")

    db = get_db_manager()
    try:
        user_id = db.create_user(
            email=email,
            password=password,
            full_name=full_name,
            user_type=user_type,
            company_id=company_id,
            created_by=user['id']
        )
        return {"success": True, "user_id": user_id}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.put("/api/users/{user_id}")
async def update_user_endpoint(
    user_id: int,
    full_name: Optional[str] = Form(None),
    email: Optional[str] = Form(None),
    user_type: Optional[str] = Form(None),
    is_active: Optional[str] = Form(None),
    new_password: Optional[str] = Form(None),
    user: dict = Depends(require_tenant_admin),
):
    """Update an existing user.

    - super_admin can edit any user (including elevating to tenant_admin / super_admin).
    - tenant_admin can only edit users that belong to their own company and may
      not change a user's user_type to super_admin.
    """
    db = get_db_manager()
    target = db.get_user_by_id(user_id)
    if not target:
        raise HTTPException(status_code=404, detail="User not found")

    # Resolve current user's company id
    actor_company_id = None
    if user.get('company'):
        if isinstance(user['company'], dict):
            actor_company_id = user['company'].get('company_id') or user['company'].get('id')
        else:
            actor_company_id = getattr(user['company'], 'id', None)

    # Authorization: tenant_admin must share a company with the target
    if user['user_type'] == 'tenant_admin':
        if not actor_company_id:
            raise HTTPException(status_code=403, detail="Access denied")
        company_user_ids = {u['id'] for u in (db.get_company_users(actor_company_id) or [])}
        if user_id not in company_user_ids:
            raise HTTPException(status_code=403, detail="Access denied")
        # Tenant admin cannot grant super_admin
        if user_type == 'super_admin':
            raise HTTPException(status_code=403, detail="Cannot promote to super_admin")
        # Tenant admin cannot edit a super_admin
        if (target.get('user_type') if isinstance(target, dict) else target['user_type']) == 'super_admin':
            raise HTTPException(status_code=403, detail="Cannot edit a super_admin")

    # Validate user_type
    if user_type is not None and user_type not in ('super_admin', 'tenant_admin', 'tenant_user'):
        raise HTTPException(status_code=400, detail="Invalid user_type")

    # Build updates dict (only include fields that were provided)
    updates: dict = {}
    if full_name is not None and full_name.strip():
        updates['full_name'] = full_name.strip()
    if email is not None and email.strip():
        updates['email'] = email.strip().lower()
    if user_type is not None:
        updates['user_type'] = user_type
    if is_active is not None:
        updates['is_active'] = str(is_active).lower() in ('true', '1', 'yes', 'on')

    try:
        if updates:
            db.update_user(user_id, updates, user['id'])
        if new_password:
            if len(new_password) < 8:
                raise HTTPException(status_code=400, detail="Password must be at least 8 characters")
            db.change_password(user_id, new_password)
        if not updates and not new_password:
            raise HTTPException(status_code=400, detail="No changes provided")
        return {"success": True, "user_id": user_id}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to update user {user_id}: {e}")
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/api/users/{user_id}")
async def delete_user_endpoint(
    user_id: int,
    user: dict = Depends(require_tenant_admin),
):
    """Delete a user.

    Authorization:
    - super_admin can delete any user except themselves.
    - tenant_admin can delete tenant_user / tenant_admin in their own company,
      but cannot delete a super_admin and cannot delete themselves.
    """
    if user_id == user['id']:
        raise HTTPException(status_code=400, detail="You cannot delete your own account")

    db = get_db_manager()
    target = db.get_user_by_id(user_id)
    if not target:
        raise HTTPException(status_code=404, detail="User not found")

    target_type = target.get('user_type') if isinstance(target, dict) else target['user_type']

    # Tenant-admin scope check
    if user['user_type'] == 'tenant_admin':
        actor_company_id = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                actor_company_id = user['company'].get('company_id') or user['company'].get('id')
            else:
                actor_company_id = getattr(user['company'], 'id', None)
        if not actor_company_id:
            raise HTTPException(status_code=403, detail="Access denied")
        company_user_ids = {u['id'] for u in (db.get_company_users(actor_company_id) or [])}
        if user_id not in company_user_ids:
            raise HTTPException(status_code=403, detail="Access denied")
        if target_type == 'super_admin':
            raise HTTPException(status_code=403, detail="Cannot delete a super_admin")

    try:
        ok = db.delete_user(user_id, user['id'])
        if not ok:
            raise HTTPException(status_code=404, detail="User not found")
        return {"success": True, "user_id": user_id}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete user {user_id}: {e}")
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/users")
async def get_users(user: dict = Depends(require_auth)):
    """Get users based on user type"""
    try:
        db = get_db_manager()

        if user['user_type'] == 'super_admin':
            users = db.get_all_users()
        elif user['user_type'] == 'tenant_admin':
            company_id = None
            if user.get('company'):
                if isinstance(user['company'], dict):
                    company_id = user['company'].get('company_id') or user['company'].get('id')
                else:
                    company_id = getattr(user['company'], 'id', None)
            users = db.get_company_users(company_id) if company_id else []
        else:
            raise HTTPException(status_code=403, detail="Access denied")

        return {"users": users}
    except Exception as e:
        logger.error(f"Error fetching users: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch users")

@app.post("/api/change-password")
async def change_password_api(
    current_password: str = Form(...),
    new_password: str = Form(...),
    user: dict = Depends(require_auth)
):
    """Allow the logged-in user to change their own password."""
    if len(new_password) < 8:
        raise HTTPException(status_code=400, detail="New password must be at least 8 characters")
    db = get_db_manager()
    # Verify current password
    verified = db.authenticate_user(user['email'], current_password)
    if not verified:
        raise HTTPException(status_code=401, detail="Current password is incorrect")
    ok = db.change_password(user['id'], new_password)
    if not ok:
        raise HTTPException(status_code=500, detail="Failed to update password")
    return {"success": True, "message": "Password updated"}


@app.put("/api/companies/{company_id}")
async def update_company(
    company_id: int,
    company_name: Optional[str] = Form(None),
    subscription_plan: Optional[str] = Form(None),
    max_users: Optional[int] = Form(None),
    max_resumes: Optional[int] = Form(None),
    is_active: Optional[bool] = Form(None),
    user: dict = Depends(require_super_admin),
):
    """Update a tenant company (super admin only)."""
    fields = {
        "company_name": company_name,
        "subscription_plan": subscription_plan,
        "max_users": max_users,
        "max_resumes": max_resumes,
        "is_active": is_active,
    }
    fields = {k: v for k, v in fields.items() if v is not None}
    if not fields:
        raise HTTPException(status_code=400, detail="No fields to update")
    set_clause = ", ".join(f"{k} = %s" for k in fields.keys())
    values = list(fields.values()) + [company_id]
    try:
        db = get_db_manager()
        with db.get_cursor() as (cursor, conn):
            cursor.execute(
                f"UPDATE tenant_companies SET {set_clause}, updated_at = CURRENT_TIMESTAMP WHERE id = %s",
                values,
            )
            if cursor.rowcount == 0:
                raise HTTPException(status_code=404, detail="Company not found")
        return {"success": True, "updated": list(fields.keys())}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating company {company_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to update company")


@app.get("/api/system-stats")
async def get_system_stats(user: dict = Depends(require_super_admin)):
    """Get system-wide statistics (super admin only)"""
    try:
        db = get_db_manager()
        stats = db.get_system_stats()
        return {"stats": stats}
    except Exception as e:
        logger.error(f"Error fetching system stats: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch system stats")

@app.get("/api/company-stats")
async def get_company_stats(user: dict = Depends(require_tenant_admin)):
    """Get company statistics (tenant admin)"""
    try:
        db = get_db_manager()
        company_id = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                company_id = user['company'].get('company_id') or user['company'].get('id')
            else:
                company_id = getattr(user['company'], 'id', None)

        if not company_id:
            raise HTTPException(status_code=400, detail="No company associated with user")

        stats = db.get_company_stats(company_id)
        return {"stats": stats}
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error fetching company stats: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch company stats")

@app.get("/api/company-resource-usage")
async def get_company_resource_usage(user: dict = Depends(require_tenant_admin)):
    """Get company resource usage details for tenant admin dashboard"""
    try:
        db = get_db_manager()

        # Get the admin's company ID
        company_id = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                company_id = user['company'].get('company_id') or user['company'].get('id')
            else:
                company_id = getattr(user['company'], 'id', None)

        if not company_id:
            raise HTTPException(status_code=400, detail="No company associated with user")

        # Get comprehensive company stats
        stats = db.get_company_stats(company_id)

        # Calculate usage percentages
        resume_usage_percent = 0
        search_usage_percent = 0

        if stats.get('max_resumes') and stats.get('max_resumes') > 0:
            resume_usage_percent = min(100, (stats.get('total_resumes', 0) / stats.get('max_resumes', 1)) * 100)

        # For searches, use a default limit of 10000 if not set
        max_searches = stats.get('max_searches', 10000)  # Default search limit
        if max_searches > 0:
            search_usage_percent = min(100, (stats.get('total_searches', 0) / max_searches) * 100)

        return {
            "success": True,
            "resource_usage": {
                "resumes": {
                    "current": stats.get('total_resumes', 0),
                    "maximum": stats.get('max_resumes', 1000),
                    "usage_percent": round(resume_usage_percent, 1)
                },
                "searches": {
                    "current": stats.get('total_searches', 0),
                    "maximum": max_searches,
                    "usage_percent": round(search_usage_percent, 1)
                },
                "users": {
                    "current": stats.get('total_users', 0),
                    "maximum": stats.get('max_users', 10)
                },
                "subscription_plan": stats.get('subscription_plan', 'basic')
            }
        }
    except HTTPException:
        # Don't mask 4xx errors (e.g. "no company associated") as 500.
        raise
    except Exception as e:
        print(f"❌ Error fetching company resource usage: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch company resource usage: {str(e)}")

@app.get("/stream-demo", response_class=HTMLResponse)
async def stream_demo(request: Request, user: dict = Depends(require_auth)):
    """Demo page to show real-time progress streaming"""
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Real-time Resume Search</title>
        <style>
            body { font-family: Arial, sans-serif; margin: 20px; }
            #progress {
                background: #f5f5f5;
                padding: 20px;
                border-radius: 5px;
                margin: 20px 0;
                height: 400px;
                overflow-y: auto;
                border: 1px solid #ddd;
            }
            .progress-item {
                margin: 5px 0;
                padding: 8px;
                border-radius: 3px;
            }
            .info { background-color: #e7f3ff; color: #0066cc; }
            .success { background-color: #e7f9e7; color: #006600; }
            .warning { background-color: #fff3cd; color: #856404; }
            .error { background-color: #f8d7da; color: #721c24; }
            #query { width: 100%; padding: 10px; margin: 10px 0; }
            button {
                background: #007bff;
                color: white;
                padding: 10px 20px;
                border: none;
                border-radius: 5px;
                cursor: pointer;
            }
            button:disabled { background: #6c757d; cursor: not-allowed; }
        </style>
    </head>
    <body>
        <h1>🔍 Real-time Resume Search Demo</h1>
        <textarea id="query" placeholder="Enter job description or requirements..." rows="4">
Looking for a Senior Python Developer with 5+ years experience in web development, AWS cloud services, and database management. Must have experience with Django, React, PostgreSQL, and agile methodologies. Leadership experience preferred.
        </textarea>
        <button id="searchBtn" onclick="startSearch()">🚀 Start Smart Search</button>
        <button id="stopBtn" onclick="stopSearch()" disabled>⏹️ Stop</button>

        <div id="progress"></div>

        <script>
            let eventSource = null;

            function startSearch() {
                const query = document.getElementById('query').value;
                if (!query.trim()) {
                    alert('Please enter a job description');
                    return;
                }

                document.getElementById('searchBtn').disabled = true;
                document.getElementById('stopBtn').disabled = false;
                document.getElementById('progress').innerHTML = '';

                // Create Server-Sent Events connection
                const formData = new FormData();
                formData.append('query', query);
                formData.append('result_count', '5');

                // Use fetch to POST the form data and get streaming response
                fetch('/api/smart-search-stream', {
                    method: 'POST',
                    body: formData
                })
                .then(response => {
                    const reader = response.body.getReader();
                    const decoder = new TextDecoder();

                    function readStream() {
                        reader.read().then(({ done, value }) => {
                            if (done) {
                                searchComplete();
                                return;
                            }

                            const chunk = decoder.decode(value);
                            const lines = chunk.split('\\n');

                            lines.forEach(line => {
                                if (line.startsWith('data: ')) {
                                    try {
                                        const data = JSON.parse(line.substring(6));
                                        displayProgress(data);
                                    } catch (e) {
                                        console.log('Parse error:', e, line);
                                    }
                                }
                            });

                            readStream();
                        });
                    }

                    readStream();
                })
                .catch(error => {
                    console.error('Error:', error);
                    searchComplete();
                });
            }

            function displayProgress(data) {
                const progressDiv = document.getElementById('progress');

                if (data.type === 'final_result') {
                    // Display final results
                    const resultDiv = document.createElement('div');
                    resultDiv.className = 'progress-item success';
                    resultDiv.innerHTML = `
                        <strong>🎉 Search Results:</strong><br>
                        📊 Total Documents: ${data.data.total_results}<br>
                        🧠 Successfully Analyzed: ${data.data.analyzed_count}<br>
                        📝 Optimized Query: ${data.data.optimized_query.substring(0, 100)}...
                    `;
                    progressDiv.appendChild(resultDiv);
                } else if (data.type === 'error') {
                    const errorDiv = document.createElement('div');
                    errorDiv.className = 'progress-item error';
                    errorDiv.innerHTML = `❌ Error: ${data.message}`;
                    progressDiv.appendChild(errorDiv);
                } else if (data.message) {
                    // Regular progress message
                    const messageDiv = document.createElement('div');
                    messageDiv.className = `progress-item ${data.type}`;
                    messageDiv.innerHTML = `${data.message} (Step ${data.step}/${data.total})`;
                    progressDiv.appendChild(messageDiv);
                }

                // Auto-scroll to bottom
                progressDiv.scrollTop = progressDiv.scrollHeight;
            }

            function stopSearch() {
                if (eventSource) {
                    eventSource.close();
                }
                searchComplete();
            }

            function searchComplete() {
                document.getElementById('searchBtn').disabled = false;
                document.getElementById('stopBtn').disabled = true;
            }
        </script>
    </body>
    </html>
    """
    return HTMLResponse(content=html_content)

# Email scraping functionality
RESUME_ATTACHMENT_EXTENSIONS = {'.pdf', '.doc', '.docx'}
RESUME_FILENAME_KEYWORDS = {
    'resume', 'cv', 'curriculum', 'profile', 'application',
    'bio-data', 'biodata',
}


def is_resume_related_email(subject: str, body: str) -> bool:
    """Check if email contains resume-related keywords."""
    RESUME_KEYWORDS = [
        'resume', 'cv', 'curriculum vitae', 'interview', 'application',
        'job application', 'hiring', 'candidate', 'profile', 'applicant',
        'cover letter', 'attached my', 'attached is my', 'please find attached',
        'consider my application', 'applying for', 'applied for'
    ]
    text_to_check = f"{subject} {body}".lower()
    return any(keyword.lower() in text_to_check for keyword in RESUME_KEYWORDS)


def is_supported_resume_attachment(filename: str) -> bool:
    """Return True for file types the resume parser/upload path supports."""
    if not filename:
        return False
    return Path(filename).suffix.lower() in RESUME_ATTACHMENT_EXTENSIONS


def is_resume_attachment(filename: str) -> bool:
    """Check if attachment filename itself suggests it's a resume."""
    if not filename:
        return False

    filename_lower = filename.lower()
    if not is_supported_resume_attachment(filename):
        return False
    return any(keyword in filename_lower for keyword in RESUME_FILENAME_KEYWORDS)


def should_collect_email_attachment(filename: str, email_is_resume_related: bool) -> bool:
    """Decide whether to download/upload an email attachment as a resume.

    Many real candidate resumes are named only with the candidate's name
    (e.g. "Muhammad Ali.pdf"). Requiring "resume" or "cv" in the filename
    causes false negatives. Keep the extension allow-list strict, then allow
    candidate-named PDF/DOC/DOCX files when the email context is resume-like.
    Final magic-byte validation still happens before upload.
    """
    if not is_supported_resume_attachment(filename):
        return False
    return email_is_resume_related or is_resume_attachment(filename)

def safe_decode_filename(filename: str) -> str:
    """Safely decode attachment filename."""
    try:
        if filename:
            decoded_name, enc = decode_header(filename)[0]
            if isinstance(decoded_name, bytes):
                return decoded_name.decode(enc or "utf-8", errors="ignore")
            return decoded_name
    except Exception as e:
        print(f"⚠️ Failed to decode filename: {filename}, error: {e}")
    return filename or "unknown_file"


def scrape_imap_resume_attachments(
    email_address: str,
    password: str,
    host: str,
    provider_label: str,
    limit: int = 100,
) -> List[Dict[str, Any]]:
    """Scrape recent IMAP messages for resume attachments."""
    results = []
    print(f"Connecting to {provider_label} IMAP for {email_address}")
    require_resume_context = not (
        provider_label.lower() == "outlook"
        and os.getenv("OUTLOOK_REQUIRE_RESUME_CONTEXT", "false").lower() not in {"1", "true", "yes", "on"}
    )

    with imaplib.IMAP4_SSL(host, 993) as imap:
        status, _ = imap.login(email_address, password)
        if status != "OK":
            raise ConnectionError(f"Failed to authenticate with {provider_label} IMAP")

        imap.select("INBOX")
        status, message_numbers = imap.search(None, "ALL")
        if status != "OK":
            raise RuntimeError(f"Failed to retrieve emails from {provider_label} IMAP")

        email_ids = message_numbers[0].split()
        if not email_ids:
            return []

        for eid in reversed(email_ids[-min(limit, len(email_ids)):]):
            try:
                status, msg_data = imap.fetch(eid, "(RFC822)")
                if status != "OK" or not msg_data or not msg_data[0]:
                    continue

                msg = email.message_from_bytes(msg_data[0][1])
                subject = "No Subject"
                subject_header = msg.get("Subject")
                if subject_header:
                    decoded_subject, encoding = decode_header(subject_header)[0]
                    subject = (
                        decoded_subject.decode(encoding or "utf-8", errors="ignore")
                        if isinstance(decoded_subject, bytes)
                        else decoded_subject
                    )

                body_plain = ""
                body_html = ""
                attachment_parts = []

                for part in msg.walk() if msg.is_multipart() else [msg]:
                    content_type = part.get_content_type()
                    content_disposition = str(part.get("Content-Disposition", ""))
                    if "attachment" in content_disposition:
                        filename = safe_decode_filename(part.get_filename())
                        if filename and is_supported_resume_attachment(filename):
                            attachment_parts.append((filename, content_type, part))
                        continue

                    if content_type in {"text/plain", "text/html"}:
                        payload = part.get_payload(decode=True)
                        if payload:
                            charset = part.get_content_charset() or "utf-8"
                            decoded = payload.decode(charset, errors="ignore")
                            if content_type == "text/plain" and not body_plain:
                                body_plain = decoded
                            elif content_type == "text/html" and not body_html:
                                body_html = decoded

                body_text = body_plain or body_html or ""
                email_related = is_resume_related_email(subject, body_text)
                attachments = []

                for filename, content_type, part in attachment_parts:
                    if require_resume_context and not should_collect_email_attachment(filename, email_related):
                        continue
                    payload = part.get_payload(decode=True)
                    if not payload:
                        continue
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    clean_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
                    unique_filename = f"{timestamp}_{eid.decode()}_{clean_filename}"
                    attachments.append({
                        "filename": unique_filename,
                        "original_filename": filename,
                        "content": payload,
                        "content_type": content_type,
                    })

                if attachments:
                    snippet_source = body_plain or body_html
                    results.append({
                        "id": eid.decode(),
                        "from": msg.get("From", "Unknown"),
                        "subject": subject,
                        "date": msg.get("Date", "Unknown"),
                        "body_plain": body_plain,
                        "body_html": body_html,
                        "snippet": snippet_source.strip()[:500] if snippet_source else "",
                        "attachments": attachments,
                        "has_resume_attachment": True,
                    })
            except Exception as e:
                print(f"Failed to process {provider_label} IMAP email {eid!r}: {e}")
                continue

        imap.close()
        imap.logout()

    print(f"{provider_label} IMAP scraping complete. Found {len(results)} resume emails")
    return results


def scrape_gmail_attachments(email_address: str, password: str, limit: int = 100) -> List[Dict[str, Any]]:
    """Scrape Gmail for resume attachments."""
    results = []

    try:
        print(f"📧 Connecting to Gmail for {email_address}")

        # Connect to Gmail over SSL
        with imaplib.IMAP4_SSL("imap.gmail.com", 993) as imap:
            # Authenticate
            status, _ = imap.login(email_address, password)
            if status != "OK":
                raise ConnectionError("Failed to authenticate with Gmail IMAP")

            print("✅ Successfully authenticated with Gmail")

            # Select the mailbox
            imap.select("INBOX")

            # Search for all emails
            status, message_numbers = imap.search(None, "ALL")
            if status != "OK":
                raise RuntimeError("Failed to retrieve emails from Gmail")

            email_ids = message_numbers[0].split()
            if not email_ids:
                print("📭 No emails found in inbox")
                return []

            total_emails = len(email_ids)
            print(f"📊 Found {total_emails} total emails. Processing up to {limit} emails...")

            # Process emails from newest to oldest
            emails_to_check = min(limit, total_emails)
            latest_ids = email_ids[-emails_to_check:]

            for i, eid in enumerate(reversed(latest_ids), 1):
                try:
                    # Fetch email
                    status, msg_data = imap.fetch(eid, "(RFC822)")
                    if status != "OK":
                        print(f"⚠️ Failed to fetch email ID {eid.decode()}")
                        continue

                    raw_email = msg_data[0][1]
                    msg = email.message_from_bytes(raw_email)

                    # Decode subject
                    subject = "No Subject"
                    try:
                        subject_header = msg.get("Subject")
                        if subject_header:
                            decoded_subject, encoding = decode_header(subject_header)[0]
                            if isinstance(decoded_subject, bytes):
                                subject = decoded_subject.decode(encoding or "utf-8", errors="ignore")
                            else:
                                subject = decoded_subject
                    except Exception as e:
                        print(f"⚠️ Failed to decode subject for email {eid.decode()}: {e}")

                    from_ = msg.get("From", "Unknown")
                    date_ = msg.get("Date", "Unknown")

                    # Extract email body
                    body_plain = ""
                    body_html = ""
                    attachments = []

                    if msg.is_multipart():
                        for part in msg.walk():
                            content_type = part.get_content_type()
                            content_disposition = str(part.get("Content-Disposition", ""))

                            if "attachment" in content_disposition:
                                # Process attachment
                                filename = safe_decode_filename(part.get_filename())

                                if is_resume_attachment(filename):
                                    try:
                                        payload = part.get_payload(decode=True)
                                        if payload:
                                            # Create unique filename
                                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                            clean_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
                                            unique_filename = f"{timestamp}_{eid.decode()}_{clean_filename}"

                                            attachments.append({
                                                "filename": unique_filename,
                                                "original_filename": filename,
                                                "content": payload,
                                                "content_type": content_type
                                            })
                                            print(f"📎 Found resume attachment: {filename}")
                                    except Exception as e:
                                        print(f"❌ Failed to process attachment {filename}: {e}")
                                continue

                            # Extract text content
                            try:
                                payload = part.get_payload(decode=True)
                                if payload:
                                    charset = part.get_content_charset() or "utf-8"
                                    decoded = payload.decode(charset, errors="ignore")

                                    if content_type == "text/plain" and not body_plain:
                                        body_plain = decoded
                                    elif content_type == "text/html" and not body_html:
                                        body_html = decoded
                            except Exception as e:
                                print(f"⚠️ Failed to decode part for email {eid.decode()}: {e}")
                    else:
                        # Single-part message
                        try:
                            payload = msg.get_payload(decode=True)
                            if payload:
                                charset = msg.get_content_charset() or "utf-8"
                                decoded = payload.decode(charset, errors="ignore")

                                if msg.get_content_type() == "text/plain":
                                    body_plain = decoded
                                elif msg.get_content_type() == "text/html":
                                    body_html = decoded
                        except Exception as e:
                            print(f"⚠️ Failed to decode single-part email {eid.decode()}: {e}")

                    # Check if this email is resume-related
                    body_text = body_plain or body_html or ""

                    if attachments or is_resume_related_email(subject, body_text):
                        snippet_source = body_plain or body_html

                        results.append({
                            "id": eid.decode(),
                            "from": from_,
                            "subject": subject,
                            "date": date_,
                            "body_plain": body_plain,
                            "body_html": body_html,
                            "snippet": snippet_source.strip()[:500] if snippet_source else "",
                            "attachments": attachments,
                            "has_resume_attachment": len(attachments) > 0
                        })

                        if attachments:
                            print(f"📧 Found resume email with {len(attachments)} attachment(s): {subject[:50]}...")

                except Exception as e:
                    print(f"❌ Error processing email {eid.decode()}: {e}")
                    continue

            # Close connection
            imap.close()
            imap.logout()

            print(f"✅ Gmail scraping complete! Found {len(results)} resume-related emails")
            return results

    except Exception as e:
        print(f"❌ Gmail scraping failed: {e}")
        raise

def scrape_outlook_attachments(email_address: str, password: str, limit: int = 100) -> List[Dict[str, Any]]:
    """Scrape Outlook 365 for resume attachments using Microsoft Graph API."""
    results = []

    try:
        print(f"📧 Connecting to Outlook 365 for {email_address}")

        # Microsoft Graph API configuration
        authority = "https://login.microsoftonline.com/common"
        client_id = os.getenv("OUTLOOK_CLIENT_ID")
        client_secret = os.getenv("OUTLOOK_CLIENT_SECRET")

        if not client_id or not client_secret:
            print("⚠️ Outlook 365 API credentials not configured. Set OUTLOOK_CLIENT_ID and OUTLOOK_CLIENT_SECRET environment variables.")
            try:
                return scrape_imap_resume_attachments(
                    email_address,
                    password,
                    "outlook.office365.com",
                    "Outlook",
                    limit,
                )
            except Exception as imap_error:
                print(f"Outlook IMAP fallback failed: {imap_error}")
                raise RuntimeError(
                    "Outlook authentication failed. Use the correct mailbox password/app password and make sure IMAP is enabled for this Microsoft 365 mailbox."
                ) from imap_error

        # Create MSAL app
        app = msal.ConfidentialClientApplication(
            client_id,
            authority=authority,
            client_credential=client_secret
        )

        # Get token using Resource Owner Password Credentials (ROPC) flow
        scopes = ["https://graph.microsoft.com/Mail.Read"]

        result = app.acquire_token_by_username_password(
            username=email_address,
            password=password,
            scopes=scopes
        )

        if "access_token" not in result:
            print(f"❌ Failed to authenticate with Outlook 365: {result.get('error_description', 'Unknown error')}")
            try:
                return scrape_imap_resume_attachments(
                    email_address,
                    password,
                    "outlook.office365.com",
                    "Outlook",
                    limit,
                )
            except Exception as imap_error:
                print(f"Outlook IMAP fallback failed: {imap_error}")
                raise RuntimeError(
                    "Outlook authentication failed. Use the correct mailbox password/app password and make sure IMAP is enabled for this Microsoft 365 mailbox."
                ) from imap_error

        access_token = result["access_token"]
        print("✅ Successfully authenticated with Outlook 365")

        # Get messages from inbox
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json'
        }

        # Fetch recent messages that have attachments. Filtering by resume
        # keywords at the Graph query layer misses real candidates whose
        # subject/body does not contain "resume"/"cv" exactly, so classify
        # message context locally after we fetch the body.
        messages_url = f"https://graph.microsoft.com/v1.0/me/messages"
        params = {
            '$filter': 'hasAttachments eq true',
            '$top': limit,
            '$orderby': 'receivedDateTime desc',
            '$select': 'id,subject,from,receivedDateTime,body,hasAttachments'
        }

        response = requests.get(messages_url, headers=headers, params=params, timeout=30)

        if response.status_code != 200:
            print(f"❌ Failed to retrieve messages: {response.status_code} - {response.text}")
            try:
                return scrape_imap_resume_attachments(
                    email_address,
                    password,
                    "outlook.office365.com",
                    "Outlook",
                    limit,
                )
            except Exception as imap_error:
                print(f"Outlook IMAP fallback failed: {imap_error}")
                raise RuntimeError(
                    "Outlook message access failed. Check Microsoft 365 mailbox permissions and IMAP access."
                ) from imap_error

        messages = response.json().get('value', [])
        print(f"📊 Found {len(messages)} messages with attachments")
        outlook_require_context = os.getenv("OUTLOOK_REQUIRE_RESUME_CONTEXT", "false").lower() in {"1", "true", "yes", "on"}

        for message in messages:
            try:
                message_id = message['id']
                subject = message.get('subject', 'No Subject')
                from_email = message.get('from', {}).get('emailAddress', {}).get('address', 'Unknown')
                date_received = message.get('receivedDateTime', 'Unknown')
                body_content = message.get('body', {}).get('content', '')
                email_related = is_resume_related_email(subject, body_content)

                # Get attachments for this message
                attachments_url = f"https://graph.microsoft.com/v1.0/me/messages/{message_id}/attachments"
                attach_response = requests.get(attachments_url, headers=headers, timeout=30)

                if attach_response.status_code != 200:
                    print(f"⚠️ Failed to get attachments for message {message_id}")
                    continue

                attachments_data = attach_response.json().get('value', [])
                resume_attachments = []

                for attachment in attachments_data:
                    filename = attachment.get('name', '')
                    content_type = attachment.get('contentType', '')

                    if (
                        is_supported_resume_attachment(filename)
                        and (
                            not outlook_require_context
                            or should_collect_email_attachment(filename, email_related)
                        )
                    ):
                        # Get attachment content
                        attachment_id = attachment['id']
                        attachment_url = f"https://graph.microsoft.com/v1.0/me/messages/{message_id}/attachments/{attachment_id}"
                        content_response = requests.get(attachment_url, headers=headers, timeout=60)

                        if content_response.status_code == 200:
                            attachment_data = content_response.json()
                            content_bytes = attachment_data.get('contentBytes', '')

                            if content_bytes:
                                try:
                                    # Decode base64 content
                                    decoded_content = base64.b64decode(content_bytes)

                                    # Create unique filename
                                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                    clean_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
                                    unique_filename = f"{timestamp}_{message_id[:8]}_{clean_filename}"

                                    resume_attachments.append({
                                        "filename": unique_filename,
                                        "original_filename": filename,
                                        "content": decoded_content,
                                        "content_type": content_type
                                    })
                                    print(f"📎 Found resume attachment: {filename}")
                                except Exception as e:
                                    print(f"❌ Failed to decode attachment {filename}: {e}")

                if resume_attachments:
                    results.append({
                        "id": message_id,
                        "from": from_email,
                        "subject": subject,
                        "date": date_received,
                        "body_plain": body_content,
                        "body_html": "",
                        "snippet": body_content[:500] if body_content else "",
                        "attachments": resume_attachments,
                        "has_resume_attachment": len(resume_attachments) > 0
                    })

                    print(f"📧 Found resume email with {len(resume_attachments)} attachment(s): {subject[:50]}...")

            except Exception as e:
                print(f"❌ Error processing message {message.get('id', 'unknown')}: {e}")
                continue

        print(f"✅ Outlook 365 scraping complete! Found {len(results)} resume-related emails")
        return results

    except Exception as e:
        print(f"❌ Outlook 365 scraping failed: {e}")
        raise

async def process_scraped_attachments(scraped_emails: List[Dict[str, Any]], folder_path: str, user: dict) -> Dict[str, Any]:
    """Process scraped email attachments and upload them to GCS like regular file uploads."""

    # Get company-specific resources
    bucket_name, datastore_id = get_company_resources(user)

    if not bucket_name:
        raise HTTPException(status_code=400, detail="Google Cloud Storage bucket not configured for your company")

    if not (os.path.exists(CREDENTIALS_PATH) or os.environ.get("K_SERVICE") or os.environ.get("SMARTHR_LOCAL_MODE", "").lower() in ("1", "true", "yes", "on")):
        raise HTTPException(status_code=400, detail="Credentials file not found")

    # Create folder structure
    if folder_path and folder_path.strip():
        safe_folder = _sanitize_object_path(folder_path)
        if not safe_folder:
            raise HTTPException(status_code=400, detail="Invalid folder path")
        normalized_folder = safe_folder.rstrip('/') + '/'
        upload_folder = f"email_attachments/{normalized_folder}"
    else:
        upload_folder = "email_attachments/"

    print(f"📁 Using upload folder: {upload_folder}")
    print(f"🪣 Using company bucket: {bucket_name}")

    db = get_db_manager()

    # Get company_id for limit checking
    company_id = None
    if user.get('company'):
        if isinstance(user['company'], dict):
            company_id = user['company'].get('company_id') or user['company'].get('id')
        else:
            company_id = getattr(user['company'], 'id', None)

    # Check resume upload limit
    if company_id:
        try:
            company_stats = db.get_company_stats(company_id)
            max_resumes = company_stats['max_resumes'] or 1000
            current_resumes = company_stats['total_resumes'] or 0

            # Count total attachments to upload
            total_attachments = sum(len(email_data['attachments']) for email_data in scraped_emails)

            if current_resumes + total_attachments > max_resumes:
                remaining_slots = max_resumes - current_resumes
                if remaining_slots <= 0:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Company has reached maximum resume storage limit ({max_resumes} resumes). Cannot upload any new resumes."
                    )
                else:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Upload limit exceeded! You can only upload {remaining_slots} more resume(s). Current: {current_resumes}/{max_resumes}, Attempting: {total_attachments}"
                    )
        except HTTPException:
            raise
        except Exception as e:
            print(f"❌ Error checking resume limits: {e}")

    client = _get_storage_client()
    bucket = client.bucket(bucket_name)

    upload_results = []
    success_count = 0
    error_count = 0

    print(f"🔄 Processing {len(scraped_emails)} emails with attachments")

    for email_data in scraped_emails:
        email_id = email_data['id']
        email_subject = email_data['subject']

        for attachment in email_data['attachments']:
            try:
                filename = attachment['filename']
                original_filename = attachment['original_filename']
                content = attachment['content']
                content_type = attachment['content_type']
                safe_name = _sanitize_object_path(filename) or _sanitize_object_path(original_filename) or 'attachment'

                print(f"📤 Uploading {original_filename} from email: {email_subject[:50]}...")

                # Check file size
                if len(content) > MAX_FILE_SIZE:
                    print(f"❌ File {original_filename} rejected: too large ({len(content):,} bytes)")
                    upload_results.append({
                        "filename": original_filename,
                        "success": False,
                        "error": f"File too large. Maximum size allowed: {MAX_FILE_SIZE // (1024*1024)}MB",
                        "source": "email_attachment",
                        "email_id": email_id
                    })
                    error_count += 1
                    continue

                is_valid_payload, payload_error = validate_resume_payload(original_filename, content)
                if not is_valid_payload:
                    print(f"❌ File {original_filename} rejected: invalid resume payload")
                    upload_results.append({
                        "filename": original_filename,
                        "success": False,
                        "error": payload_error,
                        "source": "email_attachment",
                        "email_id": email_id
                    })
                    error_count += 1
                    continue

                # Upload to GCS
                blob_name = f"{upload_folder}{safe_name}"
                blob = bucket.blob(blob_name)

                blob.upload_from_string(
                    content,
                    content_type=content_type
                )

                # Track upload in database
                db.track_resume_upload(
                    file_name=original_filename,
                    file_path=blob_name,
                    file_size=len(content),
                    mime_type=content_type,
                    user_id=user['id'],
                    company_id=company_id
                )

                # Upload to vector datastore
                company_code = None
                if user.get('company'):
                    if isinstance(user['company'], dict):
                        company_code = user['company'].get('company_code')
                    else:
                        company_code = getattr(user['company'], 'company_code', None)

                print(f"🔍 Adding {original_filename} to vector datastore")
                vector_upload_success = upload_to_vector_datastore(
                    file_path=blob_name,
                    company_id=company_id,
                    company_code=company_code,
                    user_id=str(user['id'])
                )

                upload_results.append({
                    "filename": original_filename,
                    "success": True,
                    "gcs_path": f"gs://{bucket_name}/{blob_name}",
                    "size": len(content),
                    "folder": upload_folder,
                    "vector_indexed": vector_upload_success,
                    "source": "email_attachment",
                    "email_id": email_id,
                    "email_subject": email_subject
                })

                success_count += 1
                print(f"✅ Successfully processed {original_filename}")

            except Exception as e:
                print(f"❌ Failed to process attachment {attachment.get('original_filename', 'unknown')}: {str(e)}")
                upload_results.append({
                    "filename": attachment.get('original_filename', 'unknown'),
                    "success": False,
                    "error": f"Upload failed: {str(e)}",
                    "source": "email_attachment",
                    "email_id": email_id
                })
                error_count += 1

    print(f"✅ Email attachment processing complete! {success_count} successful, {error_count} failed")

    if success_count:
        _search_cache_clear()

    return {
        "results": upload_results,
        "email_stats": {
            "total_emails": len(scraped_emails),
            "successful_uploads": success_count,
            "failed_uploads": error_count,
            "total_attachments": sum(len(email_data['attachments']) for email_data in scraped_emails)
        }
    }

@app.post("/api/scrape-email-attachments")
@limiter.limit("5/minute")
async def scrape_email_attachments(
    request: Request,
    email_provider: str = Form(...),  # "gmail" or "outlook"
    email_address: str = Form(...),
    password: str = Form(...),
    limit: int = Form(100),
    folder_path: str = Form(None),
    user: dict = Depends(require_auth)
):
    """Scrape email attachments from Gmail or Outlook 365 and upload them to GCS."""
    limit = clamp_positive_int(limit, default=25, max_value=MAX_EMAIL_SCRAPE_LIMIT, field_name="limit")

    print(f"📧 EMAIL SCRAPING STARTED")
    print(f"📧 Provider: {email_provider}")
    print(f"📧 Email: {email_address}")
    print(f"📧 Limit: {limit}")

    if email_provider.lower() not in ["gmail", "outlook"]:
        raise HTTPException(status_code=400, detail="Email provider must be 'gmail' or 'outlook'")

    try:
        # Scrape emails based on provider. Both functions are blocking
        # (imaplib / synchronous `requests`) so off-load to a thread to
        # avoid stalling the event loop.
        if email_provider.lower() == "gmail":
            scraped_emails = await asyncio.to_thread(
                scrape_gmail_attachments, email_address, password, limit
            )
        else:  # outlook
            scraped_emails = await asyncio.to_thread(
                scrape_outlook_attachments, email_address, password, limit
            )

        if not scraped_emails:
            return {
                "message": "No resume-related emails with attachments found",
                "results": [],
                "email_stats": {
                    "total_emails": 0,
                    "successful_uploads": 0,
                    "failed_uploads": 0,
                    "total_attachments": 0
                }
            }

        # Process and upload the scraped attachments
        result = await process_scraped_attachments(scraped_emails, folder_path, user)

        return {
            "message": f"Successfully scraped {email_provider} and processed attachments",
            **result
        }

    except Exception as e:
        print(f"❌ EMAIL SCRAPING FAILED: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Email scraping failed: {str(e)}")

@app.post("/api/configure-gcs")
@limiter.limit("2/minute")
async def configure_gcs(
    request: Request,
    bucket_name: str = Form(...),
    user: dict = Depends(require_super_admin)
):
    """Configure Google Cloud Storage bucket name (for manual override)"""
    global GCS_BUCKET_NAME

    try:
        # Test the connection with local credentials
        client = _get_storage_client()
        bucket = client.bucket(bucket_name)

        # Test if bucket exists and is accessible
        if not bucket.exists():
            raise HTTPException(status_code=400, detail="Bucket does not exist or is not accessible")

        # Update bucket name if different from config
        if bucket_name != config['gcs']['bucket_name']:
            GCS_BUCKET_NAME = bucket_name
            print(f"Bucket name updated to: {bucket_name}")

        return {"success": True, "message": "Google Cloud Storage configured successfully"}

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Configuration failed: {str(e)}")

@app.post("/api/upload-files")
@limiter.limit("30/minute")
async def upload_files(
    request: Request,
    files: List[UploadFile] = File(default=None),
    folder_path: str = Form(None),
    batch_info: str = Form(None),
    # Email scraping options
    scrape_email: bool = Form(False),
    email_provider: str = Form(None),  # "gmail" or "outlook"
    email_address: str = Form(None),
    email_password: str = Form(None),
    email_limit: int = Form(100),
    user: dict = Depends(require_auth)
):
    """Upload multiple files to company-specific Google Cloud Storage bucket with batch support.

    Can also scrape email attachments from Gmail or Outlook 365 if scrape_email is True.
    """

    # Handle email scraping option
    if scrape_email:
        email_limit = clamp_positive_int(
            email_limit,
            default=25,
            max_value=MAX_EMAIL_SCRAPE_LIMIT,
            field_name="email_limit",
        )
        if not email_provider or not email_address or not email_password:
            raise HTTPException(status_code=400, detail="Email provider, address, and password are required for email scraping")

        print(f"📧 EMAIL SCRAPING MODE ACTIVATED")
        print(f"📧 Provider: {email_provider}")
        print(f"📧 Email: {email_address}")
        print(f"📧 Limit: {email_limit}")

        try:
            # Scrape emails based on provider. Off-load blocking IMAP /
            # HTTP calls to a worker thread.
            if email_provider.lower() == "gmail":
                scraped_emails = await asyncio.to_thread(
                    scrape_gmail_attachments, email_address, email_password, email_limit
                )
            elif email_provider.lower() == "outlook":
                scraped_emails = await asyncio.to_thread(
                    scrape_outlook_attachments, email_address, email_password, email_limit
                )
            else:
                raise HTTPException(status_code=400, detail="Email provider must be 'gmail' or 'outlook'")

            if not scraped_emails:
                return {
                    "message": "No resume-related emails with attachments found",
                    "results": [],
                    "batch_stats": {
                        "successful": 0,
                        "failed": 0,
                        "total": 0
                    },
                    "email_stats": {
                        "total_emails": 0,
                        "successful_uploads": 0,
                        "failed_uploads": 0,
                        "total_attachments": 0
                    }
                }

            # Process and upload the scraped attachments
            result = await process_scraped_attachments(scraped_emails, folder_path, user)

            return {
                "message": f"Successfully scraped {email_provider} and processed attachments",
                "results": result["results"],
                "batch_stats": {
                    "successful": result["email_stats"]["successful_uploads"],
                    "failed": result["email_stats"]["failed_uploads"],
                    "total": result["email_stats"]["total_attachments"]
                },
                "email_stats": result["email_stats"]
            }

        except Exception as e:
            print(f"❌ EMAIL SCRAPING FAILED: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Email scraping failed: {str(e)}")

    # Regular file upload mode
    if not files:
        raise HTTPException(status_code=400, detail="No files provided for upload")

    # Log batch information
    batch_log = f" (Batch: {batch_info})" if batch_info else ""
    print(f"\n📤 BATCH UPLOAD STARTED{batch_log}")
    print(f"📁 Files in this batch: {len(files)}")

    # Calculate total batch size
    total_batch_size = sum(file.size if hasattr(file, 'size') else 0 for file in files)
    print(f"📦 Total batch size: {total_batch_size:,} bytes ({total_batch_size / (1024*1024):.1f}MB)")
    if total_batch_size and total_batch_size > MAX_TOTAL_UPLOAD_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"Batch too large. Maximum total upload size is {MAX_TOTAL_UPLOAD_SIZE // (1024*1024)}MB."
        )

    # Get company-specific resources
    bucket_name, datastore_id = get_company_resources(user)

    if not bucket_name:
        raise HTTPException(status_code=400, detail="Google Cloud Storage bucket not configured for your company")

    if not (os.path.exists(CREDENTIALS_PATH) or os.environ.get("K_SERVICE") or os.environ.get("SMARTHR_LOCAL_MODE", "").lower() in ("1", "true", "yes", "on")):
        raise HTTPException(status_code=400, detail="Credentials file not found")

    # Create folder structure - no company prefix needed since each company has its own bucket
    if folder_path and folder_path.strip():
        # Sanitize: strip ../, leading slashes, drive letters, control chars
        safe_folder = _sanitize_object_path(folder_path)
        if not safe_folder:
            raise HTTPException(status_code=400, detail="Invalid folder path")
        normalized_folder = safe_folder.rstrip('/') + '/'
        upload_folder = normalized_folder
        print(f"📁 Using custom folder path: {folder_path} -> {upload_folder}")
    else:
        upload_folder = DEFAULT_FOLDER
        print(f"📁 Using default folder: {upload_folder}")

    print(f"🪣 Using company bucket: {bucket_name}")

    db = get_db_manager()

    try:
        # Get company_id early for limit checking
        company_id = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                company_id = user['company'].get('company_id') or user['company'].get('id')
            else:
                company_id = getattr(user['company'], 'id', None)

        # Check resume upload limit BEFORE starting any uploads
        if company_id:
            try:
                # Get current resume count and limits
                company_stats = db.get_company_stats(company_id)
                max_resumes = company_stats['max_resumes'] or 1000
                current_resumes = company_stats['total_resumes'] or 0

                # Check if uploading these files would exceed the limit
                files_to_upload = len([f for f in files if ('.' + f.filename.split('.')[-1].lower()) in config['upload']['allowed_extensions']])

                if current_resumes + files_to_upload > max_resumes:
                    remaining_slots = max_resumes - current_resumes
                    if remaining_slots <= 0:
                        raise HTTPException(
                            status_code=400,
                            detail=f"Company has reached maximum resume storage limit ({max_resumes} resumes). Cannot upload any new resumes."
                        )
                    else:
                        raise HTTPException(
                            status_code=400,
                            detail=f"Upload limit exceeded! You can only upload {remaining_slots} more resume(s). Current: {current_resumes}/{max_resumes}, Attempting: {files_to_upload}"
                        )

            except HTTPException:
                raise  # Re-raise HTTP exceptions
            except Exception as e:
                print(f"❌ ERROR checking resume limits: {e}")
                import traceback
                traceback.print_exc()
                # Continue with upload if limit check fails (fallback behavior)

        client = _get_storage_client()
        bucket = client.bucket(bucket_name)

        upload_results = []
        batch_success_count = 0
        batch_error_count = 0

        print(f"\n🔄 Processing {len(files)} files in batch{batch_log}")

        for i, file in enumerate(files):
            file_number = i + 1
            print(f"\n📄 Processing file {file_number}/{len(files)}: {file.filename}")

            # Validate file type using config
            allowed_extensions = config['upload']['allowed_extensions']
            file_ext = '.' + file.filename.split('.')[-1].lower()

            if file_ext not in allowed_extensions:
                print(f"❌ File {file.filename} rejected: unsupported type {file_ext}")
                upload_results.append({
                    "filename": file.filename,
                    "success": False,
                    "error": f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
                })
                batch_error_count += 1
                continue

            # Check file size before reading content
            file_size = file.size if hasattr(file, 'size') else None
            if file_size and file_size > MAX_FILE_SIZE:
                print(f"❌ File {file.filename} rejected: too large ({file_size:,} bytes > {MAX_FILE_SIZE:,} bytes)")
                upload_results.append({
                    "filename": file.filename,
                    "success": False,
                    "error": f"File too large. Maximum size allowed: {MAX_FILE_SIZE // (1024*1024)}MB, file size: {file_size // (1024*1024)}MB"
                })
                batch_error_count += 1
                continue

            # Read file content with size validation
            print(f"📖 Reading file content for {file.filename}")
            content = await file.read()
            if len(content) > MAX_FILE_SIZE:
                print(f"❌ File {file.filename} rejected: content too large ({len(content):,} bytes)")
                upload_results.append({
                    "filename": file.filename,
                    "success": False,
                    "error": f"File too large. Maximum size allowed: {MAX_FILE_SIZE // (1024*1024)}MB, actual size: {len(content) // (1024*1024)}MB"
                })
                batch_error_count += 1
                continue

            is_valid_payload, payload_error = validate_resume_payload(file.filename, content)
            if not is_valid_payload:
                print(f"❌ File {file.filename} rejected: invalid payload for {file_ext}")
                upload_results.append({
                    "filename": file.filename,
                    "success": False,
                    "error": payload_error
                })
                batch_error_count += 1
                continue

            try:
                # Sanitize filename to prevent path injection via crafted filenames
                safe_name = _sanitize_object_path(file.filename) or 'unnamed'
                # Upload to GCS
                blob_name = f"{upload_folder}{safe_name}"
                blob = bucket.blob(blob_name)

                print(f"☁️ Uploading {file.filename} to GCS: {blob_name}")
                blob.upload_from_string(
                    content,
                    content_type=file.content_type
                )

                # Track upload in database (limit already checked above)
                db.track_resume_upload(
                    file_name=file.filename,
                    file_path=blob_name,
                    file_size=len(content),
                    mime_type=file.content_type,
                    user_id=user['id'],
                    company_id=company_id
                )

                # Upload to vector datastore for immediate searchability (incremental indexing)
                company_code = None
                if user.get('company'):
                    if isinstance(user['company'], dict):
                        company_code = user['company'].get('company_code')
                    else:
                        company_code = getattr(user['company'], 'company_code', None)

                print(f"🔍 Adding {file.filename} to vector datastore")
                vector_upload_success = upload_to_vector_datastore(
                    file_path=blob_name,
                    company_id=company_id,
                    company_code=company_code,
                    user_id=str(user['id'])
                )

                upload_results.append({
                    "filename": file.filename,
                    "success": True,
                    "gcs_path": f"gs://{bucket_name}/{blob_name}",
                    "size": len(content),
                    "folder": upload_folder,
                    "vector_indexed": vector_upload_success,
                    "batch_info": batch_info
                })

                batch_success_count += 1
                print(f"✅ Successfully processed {file.filename}")

            except Exception as e:
                print(f"❌ Failed to process {file.filename}: {str(e)}")
                upload_results.append({
                    "filename": file.filename,
                    "success": False,
                    "error": f"Upload failed: {str(e)}",
                    "batch_info": batch_info
                })
                batch_error_count += 1

            # Reset file position for next iteration
            await file.seek(0)

        # Log batch completion
        print(f"\n✅ BATCH UPLOAD COMPLETED{batch_log}")
        print(f"📊 Results: {batch_success_count} successful, {batch_error_count} failed")

        if batch_success_count:
            _search_cache_clear()

        return {
            "results": upload_results,
            "batch_info": batch_info,
            "batch_stats": {
                "successful": batch_success_count,
                "failed": batch_error_count,
                "total": len(files)
            }
        }

    except Exception as e:
        print(f"❌ BATCH UPLOAD FAILED{batch_log}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Upload process failed: {str(e)}")

@app.get("/api/gcs-status")
async def get_gcs_status(user: dict = Depends(require_tenant_admin)):
    """Get current Google Cloud Storage configuration status (admin only).

    In SMARTHR_LOCAL_MODE the storage backend is the local filesystem via
    `vps_local.storage_shim`, so we always report as configured and skip the
    legacy GCS credential / bucket prompts.
    """
    local_mode = os.environ.get("SMARTHR_LOCAL_MODE", "").lower() in ("1", "true", "yes", "on")
    if local_mode:
        is_super = user.get("user_type") == "super_admin"
        bucket_name, _ = get_company_resources(user) if not is_super else (GCS_BUCKET_NAME or "smarthr-local", DATASTORE_ID)
        return {
            "configured": True,
            "bucket_name": bucket_name if is_super else None,
            "default_folder": DEFAULT_FOLDER,
            "credentials_found": True,
            "credentials_path": None,
            "local_mode": True,
        }

    credentials_exist = os.path.exists(CREDENTIALS_PATH) or bool(os.environ.get("K_SERVICE"))
    is_super = user.get("user_type") == "super_admin"
    bucket_name, _ = get_company_resources(user) if not is_super else (GCS_BUCKET_NAME, DATASTORE_ID)
    return {
        "configured": bool(bucket_name) and credentials_exist,
        "bucket_name": bucket_name if is_super else None,
        "default_folder": DEFAULT_FOLDER,
        "credentials_found": credentials_exist,
        "credentials_path": CREDENTIALS_PATH if is_super else None
    }

@app.get("/api/config")
async def get_config(user: dict = Depends(require_auth)):
    """Get client-safe configuration (auth required, infrastructure details redacted)"""
    return {
        "upload": config.get("upload", {}),
        "gcs": {
            "default_folder": config.get("gcs", {}).get("default_folder", ""),
        },
    }

@app.post("/api/search-vector")
async def search_vector(
    query: str = Form(...),
    result_count: int = Form(10),
    user: dict = Depends(require_auth),
):
    """Search the vector store with a query and return the requested count."""

    query = _sanitize_user_text(query, max_len=4000)
    result_count = clamp_result_count(result_count, default=10)
    if not query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    if not (os.path.exists(CREDENTIALS_PATH) or os.environ.get("K_SERVICE") or os.environ.get("SMARTHR_LOCAL_MODE", "").lower() in ("1", "true", "yes", "on")):
        raise HTTPException(status_code=400, detail="Credentials file not found")

    try:
        # Resolve resources once for cache key + response payload.
        bucket_name, datastore_id = get_company_resources(user)
        cache_key = ("search-vector", datastore_id, _normalize_cache_query(query), result_count)
        cached_response = _search_cache_get(cache_key)
        if cached_response is not None:
            return cached_response

        search_criteria = parse_search_criteria(query)
        fetch_count = criteria_fetch_count(result_count, search_criteria)
        if fetch_count != result_count:
            print(f"Criteria overfetch enabled: requested={result_count}, pool={fetch_count}")

        # Use universal search helpers for automatic datastore isolation
        client = get_search_client()

        # Create search request with automatic datastore switching
        request = create_universal_search_request(
            query=query,
            user_context=user,
            result_count=fetch_count,
            max_value=MAX_SEARCH_FETCH_RESULTS
        )
        print(f"✅ Search request created with page_size={result_count}, snippets=True")

        # Perform the search
        print(f"\n🔍 STEP 3: EXECUTING VECTOR SEARCH")
        response = client.search(request)
        print(f"✅ Vector search completed successfully")

        results = []
        for result in response.results:
                document_data = {
                    "id": result.id,
                    "document_name": result.document.name if result.document else "Unknown",
                    "relevance_score": getattr(result, 'relevance_score', 0.0),
                }

                # Extract document data if available
                if result.document and hasattr(result.document, 'struct_data'):
                    struct_data = result.document.struct_data
                    if struct_data:
                        document_data["content"] = dict(struct_data)

                # Extract derived struct data if available
                if hasattr(result.document, 'derived_struct_data') and result.document.derived_struct_data:
                    document_data["derived_data"] = dict(result.document.derived_struct_data)

                # Extract snippets if available
                if hasattr(result, 'document') and result.document:
                    if hasattr(result.document, 'derived_struct_data') and result.document.derived_struct_data:
                        snippets = result.document.derived_struct_data.get('snippets', [])
                        if snippets:
                            document_data["snippets"] = [dict(snippet) for snippet in snippets]

                results.append(document_data)

        exact_match_shortfall = 0
        if not search_criteria.is_empty():
            profile_statuses = [
                (r.get("content") or {}).get("criteria_profile_status")
                for r in results
                if "criteria_profile_status" in (r.get("content") or {})
            ]
            if profile_statuses:
                results = [
                    r for r in results
                    if (r.get("content") or {}).get("criteria_profile_status") == "exact"
                ][:result_count]
                exact_match_shortfall = max(0, result_count - len(results))

        response_payload = {
            "query": query,
            "total_results": len(results),
            "requested_result_count": result_count,
            "candidate_pool_size": fetch_count,
            "exact_match_shortfall": exact_match_shortfall,
            "results": results,
            "datastore_id": datastore_id,
            "project_id": PROJECT_ID,
            "search_criteria": search_criteria.to_dict(),
        }
        _search_cache_set(cache_key, response_payload)
        return response_payload

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Vector search failed: {str(e)}")

@app.post("/api/search-and-analyze")
async def search_and_analyze_resume(query: str = Form(...), result_count: int = Form(3), user: dict = Depends(require_auth)):
    """Search vector store, retrieve N results, and analyze ALL with Gemini LLM"""
    query = _sanitize_user_text(query, max_len=4000)
    result_count = clamp_result_count(result_count, default=3)

    # Get user's datastore info
    bucket_name, datastore_id = get_company_resources(user)

    print(f"\n🔍 STARTING ENHANCED SEARCH AND ANALYSIS")
    print(f"📝 Query: '{query}'")
    print(f"📊 Datastore: {datastore_id}")
    print(f"🏢 Project: {PROJECT_ID}")
    print(f"🌍 Location: {LOCATION}")
    print(f"📈 Will retrieve top {result_count} results")

    if not query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    if not (os.path.exists(CREDENTIALS_PATH) or os.environ.get("K_SERVICE") or os.environ.get("SMARTHR_LOCAL_MODE", "").lower() in ("1", "true", "yes", "on")):
        raise HTTPException(status_code=400, detail="Credentials file not found")

    try:
        # First, perform vector search
        print(f"\n🚀 STEP 1: INITIALIZING VECTOR SEARCH")

        # Use universal search helpers for automatic datastore isolation
        client = get_search_client()
        print(f"✅ Search client initialized with automatic datastore isolation")

        # Create the search request with automatic datastore switching
        print(f"\n🔎 STEP 2: CREATING SEARCH REQUEST")
        request = create_universal_search_request(
            query=query,
            user_context=user,
            result_count=result_count
        )

        # Perform the search
        print(f"\n🔍 STEP 3: EXECUTING VECTOR SEARCH")
        response = client.search(request)

        results = []
        all_analyses = []

        print(f"\n📊 STEP 4: PROCESSING SEARCH RESULTS")
        result_list = list(response.results)
        found_count = len(result_list)
        print(f"📈 Found {found_count} results")
        print(f"🧠 Will analyze ALL {found_count} results")

        for i, result in enumerate(result_list):
            print(f"\n📄 Processing result #{i+1}:")
            print(f"   🆔 ID: {result.id}")
            print(f"   📝 Document: {result.document.name if result.document else 'Unknown'}")
            print(f"   ⭐ Relevance: {getattr(result, 'relevance_score', 'N/A')}")

            document_data = {
                "id": result.id,
                "document_name": result.document.name if result.document else "Unknown",
                "relevance_score": getattr(result, 'relevance_score', 0.0),
            }

            # Extract document data if available
            if result.document and hasattr(result.document, 'struct_data'):
                struct_data = result.document.struct_data
                if struct_data:
                    document_data["content"] = dict(struct_data)

            # Extract derived struct data if available
            if hasattr(result.document, 'derived_struct_data') and result.document.derived_struct_data:
                document_data["derived_data"] = dict(result.document.derived_struct_data)

            # Extract snippets if available
            if hasattr(result, 'document') and result.document:
                if hasattr(result.document, 'derived_struct_data') and result.document.derived_struct_data:
                    snippets = result.document.derived_struct_data.get('snippets', [])
                    if snippets:
                        document_data["snippets"] = [dict(snippet) for snippet in snippets]

            results.append(document_data)

            # Analyze ALL results with Gemini
            if result.document:
                print(f"\n🧠 STEP 5: ANALYZING RESULT #{i+1} WITH GEMINI AI")
                print(f"🎯 Analyzing document: {result.document.name if result.document else 'Unknown'}")
                try:
                    # Try to get the file content from the document URI or name
                    file_content = None
                    print(f"📁 Attempting to retrieve file content...")

                    # First, check derived_struct_data for useful content
                    if hasattr(result.document, 'derived_struct_data') and result.document.derived_struct_data:
                        derived_data = dict(result.document.derived_struct_data)
                        print(f"📊 Derived struct data: {derived_data}")

                        # Check if we have a direct GCS link first
                        if 'link' in derived_data and derived_data['link']:
                            gcs_link = derived_data['link']
                            print(f"🔗 Found GCS link: {gcs_link}")
                            if gcs_link.startswith('gs://'):
                                # Extract the file path from gs://bucket/path format
                                gcs_path = gcs_link.replace('gs://', '').split('/', 1)
                                if len(gcs_path) == 2:
                                    bucket_name, file_path = gcs_path
                                    print(f"📁 Extracted: bucket={bucket_name}, path={file_path}")
                                    if bucket_name == GCS_BUCKET_NAME:
                                        # Get both text content and raw bytes
                                        file_content, file_bytes = get_file_content_and_bytes_from_gcs(file_path, None)

                                        if file_content:
                                            print(f"✅ Retrieved text content from GCS (length: {len(file_content)})")
                                        elif file_bytes:
                                            print(f"⚠️ Text extraction failed, but we have file bytes. Trying direct Gemini upload...")
                                            # Try direct file upload to Gemini
                                            direct_analysis = send_file_to_gemini_directly(file_bytes, file_path, query)
                                            if direct_analysis['success']:
                                                print(f"✅ Direct file analysis successful!")
                                                top_result_analysis = direct_analysis
                                                document_data["gemini_analysis"] = top_result_analysis
                                                file_content = "Content analyzed directly by Gemini from file"
                                            else:
                                                print(f"❌ Direct file analysis also failed")
                                        else:
                                            print(f"❌ Failed to retrieve any content from GCS link")

                        # If no content from GCS link, try to extract from snippets
                        if not file_content and 'snippets' in derived_data and derived_data['snippets']:
                            snippets = derived_data['snippets']
                            print(f"🔍 Processing {len(snippets)} snippet objects...")
                            snippet_texts = []
                            for j, snip in enumerate(snippets):
                                print(f"   📝 Snippet {j+1}: {type(snip)}")
                                try:
                                    # Try to convert the MapComposite to dict and extract values
                                    snip_dict = dict(snip)
                                    print(f"   📊 Snippet dict: {snip_dict}")
                                    for key in ['snippet', 'content', 'text', 'extractive_segment']:
                                        if key in snip_dict and snip_dict[key]:
                                            snippet_texts.append(str(snip_dict[key]))
                                            print(f"   ✅ Found snippet.{key}: {str(snip_dict[key])[:100]}...")
                                            break
                                except Exception as e:
                                    print(f"   ❌ Error converting snippet to dict: {e}")

                            if snippet_texts:
                                file_content = " ".join(snippet_texts)
                                print(f"✅ Found content from {len(snippet_texts)} snippets (total length: {len(file_content)})")

                    # If we have content, analyze it
                    if file_content:
                        print(f"🧠 Sending to Gemini AI for analysis...")
                        print(f"📝 Content preview: {file_content[:200]}..." if len(file_content) > 200 else f"📝 Full content: {file_content}")
                        current_analysis = analyze_resume_with_gemini(file_content, query)
                        print(f"✅ Gemini analysis completed")
                        print(f"🎯 Match score: {current_analysis.get('match_score', 'N/A')}")
                        print(f"✅ Analysis success: {current_analysis.get('success', False)}")
                        document_data["gemini_analysis"] = current_analysis
                        all_analyses.append(current_analysis)
                    else:
                        print(f"❌ No content available for analysis")
                        current_analysis = {
                            "analysis": "Could not retrieve file content for analysis",
                            "match_score": 0,
                            "success": False
                        }
                        document_data["gemini_analysis"] = current_analysis
                        all_analyses.append(current_analysis)

                except Exception as e:
                    print(f"❌ Error during analysis: {str(e)}")
                    current_analysis = {
                        "analysis": f"Analysis failed: {str(e)}",
                        "match_score": 0,
                        "success": False,
                        "error": str(e)
                    }
                    document_data["gemini_analysis"] = current_analysis
                    all_analyses.append(current_analysis)

        # Sort analyses by match score (highest first)
        all_analyses.sort(key=lambda x: x.get('match_score', 0), reverse=True)

        final_response = {
            "query": query,
            "total_results": len(results),
            "results": results,
            "all_analyses": all_analyses,
            "top_result_analysis": all_analyses[0] if all_analyses else None,
            "analyzed_count": len(all_analyses),
            "datastore_id": get_company_resources(user)[1],
            "project_id": PROJECT_ID,
            "enhanced_search": True
        }

        print(f"\n🎉 SEARCH AND ANALYSIS COMPLETED")
        print(f"📊 Total results: {len(results)}")
        print(f"🧠 Analyzed results: {len(all_analyses)}")
        print(f"🎯 AI Analysis: {'✅ Success' if all_analyses else '❌ Failed'}")
        if all_analyses:
            best_analysis = all_analyses[0]
            print(f"🏆 Best match score: {best_analysis.get('match_score', 'N/A')}")
            scores = [a.get('match_score', 0) for a in all_analyses if a.get('success')]
            if scores:
                print(f"📊 All scores: {scores}")
        print(f"📝 Query processed: '{query}'")
        print(f"✅ Returning results to frontend")

        return final_response

    except Exception as e:
        print(f"❌ Enhanced search failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Enhanced search failed: {str(e)}")

def extract_keywords_for_single_query(query: str, token_tracker: TokenTracker = None) -> Dict[str, Any]:
    """Extract keywords and create ONE optimized search query with all important points"""
    print(f"\n🧠 STEP 1: LLM KEYWORD EXTRACTION FOR SINGLE OPTIMIZED QUERY")
    print(f"📝 Original query: '{query[:200]}...' ({len(query)} chars)")

    try:
        client = gemini_client()

        keyword_prompt = KEYWORD_PROMPT.format(job_posting=query)

        contents = [
            types.Content(
                role="user",
                parts=[types.Part(text=keyword_prompt)]
            )
        ]

        # Try different model versions with fallbacks
        models_to_try = ["gemini-2.5-flash", "gemini-1.5-pro", "gemini-pro"]
        generate_response = None

        for model_name in models_to_try:
            try:
                print(f"🤖 Trying model: {model_name}")
                generate_response = client.models.generate_content(
                    model=model_name,
                    contents=contents,
                    config=types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
                        temperature=0.1,
                        # Gemini 2.5 Flash burns ~2.4-3.4k tokens on internal "thinking" even with
                        # thinking_budget=0, and those count against max_output_tokens. Pad generously
                        # so the actual JSON response (~300-600 tokens) is never truncated.
                        max_output_tokens=12000,
                        response_mime_type="application/json",
                        safety_settings=[
                            types.SafetySetting(
                                category="HARM_CATEGORY_HATE_SPEECH",
                                threshold="OFF"
                            ),
                            types.SafetySetting(
                                category="HARM_CATEGORY_DANGEROUS_CONTENT",
                                threshold="OFF"
                            ),
                            types.SafetySetting(
                                category="HARM_CATEGORY_SEXUALLY_EXPLICIT",
                                threshold="OFF"
                            ),
                            types.SafetySetting(
                                category="HARM_CATEGORY_HARASSMENT",
                                threshold="OFF"
                            )
                        ],
                    ),
                )
                print(f"✅ Successfully connected to model: {model_name}")

                # Track token usage
                if token_tracker:
                    token_usage = extract_token_usage(generate_response, "keyword_extraction", model_name)
                    token_tracker.add_call(
                        operation="keyword_extraction",
                        model=model_name,
                        input_tokens=token_usage["input_tokens"],
                        output_tokens=token_usage["output_tokens"],
                        success=True
                    )
                    print(f"📊 Token usage - Input: {token_usage['input_tokens']}, Output: {token_usage['output_tokens']}, Total: {token_usage['total_tokens']}")

                break
            except Exception as model_error:
                print(f"❌ Model {model_name} failed: {str(model_error)}")
                continue

        if not generate_response:
            print(f"❌ All models failed for keyword extraction")
            return create_fallback_single_query(query)

        # Extract response text with proper null checking
        response_text = ""
        try:
            if hasattr(generate_response, 'text') and generate_response.text is not None:
                response_text = generate_response.text.strip()
                print(f"✅ LLM response received via .text attribute")
            elif hasattr(generate_response, 'candidates') and generate_response.candidates is not None:
                print(f"📝 Extracting from candidates (count: {len(generate_response.candidates)})")
                for candidate in generate_response.candidates:
                    if candidate is not None and hasattr(candidate, 'content') and candidate.content is not None:
                        if hasattr(candidate.content, 'parts') and candidate.content.parts is not None:
                            for part in candidate.content.parts:
                                if part is not None and hasattr(part, 'text') and part.text is not None:
                                    response_text += part.text
                if response_text:
                    print(f"✅ LLM response extracted from candidates")
                else:
                    print(f"❌ No text found in candidates")
                    print(f"🔍 DETAILED ERROR: LLM response structure analysis:")
                    print(f"   - Response object type: {type(generate_response)}")
                    print(f"   - Has .candidates: {hasattr(generate_response, 'candidates')}")
                    if hasattr(generate_response, 'candidates'):
                        print(f"   - Candidates count: {len(generate_response.candidates) if generate_response.candidates else 'None'}")
                        if generate_response.candidates:
                            for i, candidate in enumerate(generate_response.candidates):
                                print(f"   - Candidate {i}: {type(candidate)}")
                                print(f"     - Has .content: {hasattr(candidate, 'content')}")
                                if hasattr(candidate, 'content'):
                                    print(f"     - Content type: {type(candidate.content)}")
                                    print(f"     - Has .parts: {hasattr(candidate.content, 'parts')}")
                                    if hasattr(candidate.content, 'parts'):
                                        print(f"     - Parts count: {len(candidate.content.parts) if candidate.content.parts else 'None'}")
                    raise Exception(f"DETAILED ERROR: No extractable text found in LLM response. Response structure: {str(generate_response)[:500]}")
            else:
                print(f"⚠️ Trying string conversion of response object")
                response_text = str(generate_response) if generate_response is not None else ""
        except Exception as extract_error:
            print(f"❌ Error extracting response text: {str(extract_error)}")
            print(f"🔍 DETAILED ERROR: {extract_error}")
            print(f"   - Error type: {type(extract_error).__name__}")
            print(f"   - Error args: {extract_error.args}")
            raise Exception(f"DETAILED ERROR: Failed to extract text from LLM response: {str(extract_error)}")

        if response_text and len(response_text.strip()) > 0:
            print(f"📝 Response text length: {len(response_text)} characters")
            print(f"📝 Response preview: {response_text[:100]}...")

            try:
                # Clean JSON response with better truncation handling
                clean_text = response_text.strip()

                # Remove markdown code blocks
                if "```json" in clean_text:
                    json_start = clean_text.find("```json") + 7
                    json_end = clean_text.find("```", json_start)
                    if json_end != -1:
                        clean_text = clean_text[json_start:json_end].strip()
                    else:
                        clean_text = clean_text[json_start:].strip()
                elif "```" in clean_text:
                    # Handle generic code blocks
                    first_triple = clean_text.find("```")
                    if first_triple != -1:
                        json_start = first_triple + 3
                        json_end = clean_text.find("```", json_start)
                        if json_end != -1:
                            clean_text = clean_text[json_start:json_end].strip()
                        else:
                            clean_text = clean_text[json_start:].strip()

                # Try to find JSON object boundaries if still having issues
                if not clean_text.startswith('{'):
                    brace_start = clean_text.find('{')
                    if brace_start != -1:
                        clean_text = clean_text[brace_start:]

                # Handle truncated JSON by finding the last complete field
                if not clean_text.endswith('}'):
                    # Look for the last complete field before truncation
                    last_complete_field_patterns = [
                        r'.*"education":\s*"[^"]*"',  # education field
                        r'.*"domain_part":\s*"[^"]*"',  # domain_part field
                        r'.*"role_part":\s*"[^"]*"',  # role_part field
                        r'.*"skills_part":\s*"[^"]*"',  # skills_part field
                        r'.*"experience_part":\s*"[^"]*"',  # experience_part field
                    ]

                    import re
                    for pattern in last_complete_field_patterns:
                        match = re.search(pattern, clean_text, re.DOTALL)
                        if match:
                            clean_text = match.group(0)
                            # Close any open objects/arrays
                            open_braces = clean_text.count('{') - clean_text.count('}')
                            open_brackets = clean_text.count('[') - clean_text.count(']')

                            # Add closing brackets and braces
                            clean_text += ']' * open_brackets + '}' * open_braces
                            print(f"🔧 Reconstructed truncated JSON")
                            break
                    else:
                        # If no pattern matches, try to close the JSON properly
                        brace_end = clean_text.rfind('}')
                        if brace_end != -1:
                            clean_text = clean_text[:brace_end + 1]
                        else:
                            # Try to add missing closing brace
                            if clean_text.count('{') > clean_text.count('}'):
                                clean_text += '}'

                print(f"🔧 Cleaned JSON text: {clean_text[:200]}...")

                # Salvage: if JSON is truncated mid-string in optimized_query (common
                # when LLM produces a giant boolean expression), try to parse first,
                # then on failure regex-extract the partial optimized_query value.
                try:
                    keywords_data = json.loads(clean_text)
                except json.JSONDecodeError as je:
                    print(f"⚠️ JSON parse failed ({je}); attempting salvage of optimized_query")
                    import re as _re
                    salvage_match = _re.search(
                        r'"optimized_query"\s*:\s*"((?:[^"\\]|\\.)*)',
                        clean_text,
                    )
                    if salvage_match and salvage_match.group(1).strip():
                        salvaged = salvage_match.group(1).strip()
                        # Drop dangling trailing partial-word and odd characters
                        salvaged = _re.sub(r'[\s,;:]+$', '', salvaged)
                        keywords_data = {"optimized_query": salvaged, "query_components": {}, "extracted_info": {}}
                        print(f"🩹 Salvaged optimized_query (length {len(salvaged)})")
                    else:
                        raise
                print(f"🔍 FULL LLM JSON RESPONSE:")
                print(f"   Keys in response: {list(keywords_data.keys())}")
                print(f"   Full response: {keywords_data}")

                # Extract the single optimized query with validation
                optimized_query = str(keywords_data.get('optimized_query', '')).strip()
                print(f"🔍 Extracted optimized_query: '{optimized_query}' (length: {len(optimized_query)})")

                # Ensure we have a valid query
                if not optimized_query or len(optimized_query) < 5:
                    # Fallback: create query from components if available
                    query_components = keywords_data.get('query_components', {})
                    experience_part = query_components.get('experience_part', 'experienced')
                    skills_part = query_components.get('skills_part', 'technical')
                    role_part = query_components.get('role_part', 'specialist')
                    domain_part = query_components.get('domain_part', 'technology')

                    optimized_query = f"{experience_part} {skills_part} {role_part} {domain_part}".strip()
                    print(f"⚠️ Created fallback query from components: '{optimized_query}'")

                # Strip boolean-search noise that occasionally leaks in even though
                # the prompt forbids it. These tokens hurt downstream embeddings.
                _boolean_noise = {"OR", "AND", "NOT"}
                optimized_query = ' '.join(
                    w for w in optimized_query.replace('(', ' ').replace(')', ' ').replace('"', ' ').split()
                    if w.upper() not in _boolean_noise
                )

                # Remove duplicate words and validate word count
                words = optimized_query.split()
                unique_words = []
                seen_words = set()

                for word in words:
                    word_lower = word.lower()
                    if word_lower not in seen_words:
                        unique_words.append(word)
                        seen_words.add(word_lower)

                optimized_query = ' '.join(unique_words)
                word_count = len(unique_words)

                # Hard caps: 40 words OR 600 chars (whichever hits first).
                if word_count > 40:
                    unique_words = unique_words[:40]
                    optimized_query = ' '.join(unique_words)
                    word_count = len(unique_words)
                    print(f"✂️ Capped query to 40 words")
                if len(optimized_query) > 600:
                    optimized_query = optimized_query[:600].rsplit(' ', 1)[0]
                    word_count = len(optimized_query.split())
                    print(f"✂️ Capped query to 600 chars")
                elif word_count < 10:
                    print(f"⚠️ Query quite short ({word_count} words) - may need more comprehensive input")

                print(f"\n📋 EXTRACTED SINGLE OPTIMIZED QUERY:")
                print(f"🎯 Optimized query: '{optimized_query}'")
                print(f"📊 Query length: {len(optimized_query)} characters, {word_count} words")

                if word_count < 50:
                    print(f"⚠️ Warning: Query has only {word_count} words, target is 150-200 words for comprehensive search")

                # Extract detailed info for analysis
                extracted_info = keywords_data.get('extracted_info', {})
                print(f"📝 Key skills: {extracted_info.get('key_skills', [])}")
                print(f"🎖️ Seniority: {extracted_info.get('seniority_level', 'Not specified')}")
                print(f"⏱️ Experience: {extracted_info.get('experience_years', 'Not specified')}")
                print(f"✅ LLM single-query extraction successful!")

                return {
                    "success": True,
                    "keywords_data": keywords_data,
                    "optimized_query": optimized_query,
                    "enhanced_query": optimized_query,  # Use optimized query as primary
                    "original_query": query,
                    "single_query_strategy": True,
                    "extracted_info": extracted_info,
                    "query_components": keywords_data.get('query_components', {})
                }

            except json.JSONDecodeError as e:
                print(f"❌ Failed to parse LLM response as JSON: {e}")
                print(f"📝 Raw response that failed: {response_text[:500]}...")
                print(f"🔍 DETAILED JSON ERROR:")
                print(f"   - Error message: {str(e)}")
                print(f"   - Error position: {e.pos if hasattr(e, 'pos') else 'Unknown'}")
                print(f"   - Error line: {e.lineno if hasattr(e, 'lineno') else 'Unknown'}")
                print(f"   - Error column: {e.colno if hasattr(e, 'colno') else 'Unknown'}")
                print(f"   - Full response text:\n{response_text}")
                raise Exception(f"DETAILED ERROR: JSON parsing failed - {str(e)}. Full response: {response_text}")
        else:
            print(f"❌ Empty or null response text from LLM")
            print(f"🔍 DETAILED ERROR: LLM returned empty response")
            print(f"   - Response text: '{response_text}'")
            print(f"   - Response text length: {len(response_text) if response_text else 0}")
            print(f"   - Response text type: {type(response_text)}")
            print(f"   - Response text is None: {response_text is None}")
            print(f"   - Response text strip length: {len(response_text.strip()) if response_text else 0}")
            raise Exception(f"DETAILED ERROR: LLM returned empty or null response. Response: '{response_text}'")

    except Exception as e:
        print(f"❌ LLM keyword extraction failed: {str(e)}")
        print(f"📊 Error type: {type(e).__name__}")
        print(f"🔍 DETAILED ERROR: {str(e)}")
        print(f"   - Error type: {type(e).__name__}")
        print(f"   - Error args: {e.args}")
        print(f"   - Error traceback:")
        import traceback
        traceback.print_exc()
        raise Exception(f"DETAILED ERROR: LLM keyword extraction failed - {str(e)}")

def create_fallback_single_query(query: str) -> Dict[str, Any]:
    """Create fallback single optimized query using ONLY words from the original query"""
    import re

    print(f"🔄 Using fallback single-query extraction - ONLY extracting existing keywords")

    # Clean and normalize the query
    query_lower = query.lower()

    # Extract all meaningful words (3+ characters, alphanumeric) from the original query
    all_words = re.findall(r'\b[a-zA-Z]{3,}\b', query)
    unique_words = []
    seen_words = set()

    # Keep original case and remove duplicates
    for word in all_words:
        word_lower = word.lower()
        if word_lower not in seen_words and len(word) >= 3:
            unique_words.append(word)
            seen_words.add(word_lower)

    # Enhanced patterns to identify important terms that EXIST in the original query
    experience_patterns = [
        r'(\d+)\s*\+?\s*years?', r'senior', r'junior', r'lead', r'principal',
        r'manager', r'director', r'entry.?level', r'mid.?level', r'experienced'
    ]

    skill_patterns = [
        r'python', r'java', r'javascript', r'react', r'node\.?js', r'angular', r'vue',
        r'aws', r'azure', r'gcp', r'docker', r'kubernetes', r'sql', r'mongodb',
        r'tensorflow', r'pytorch', r'scikit.?learn', r'pandas', r'numpy',
        r'machine\s+learning', r'ai', r'artificial\s+intelligence', r'data\s+science',
        r'html', r'css', r'typescript', r'php', r'c\+\+', r'c#', r'golang', r'rust'
    ]

    role_patterns = [
        r'engineer', r'developer', r'scientist', r'analyst', r'manager',
        r'architect', r'consultant', r'specialist', r'lead', r'director',
        r'full.?stack', r'backend', r'frontend', r'devops', r'data\s+scientist',
        r'software\s+engineer', r'product\s+manager', r'project\s+manager'
    ]

    domain_patterns = [
        r'healthcare', r'finance', r'fintech', r'retail', r'e.?commerce',
        r'ai', r'machine\s+learning', r'data', r'web', r'mobile', r'cloud',
        r'blockchain', r'cybersecurity', r'nlp', r'computer\s+vision',
        r'saas', r'b2b', r'b2c', r'startup', r'enterprise'
    ]

    def extract_matches_from_original(patterns, text, max_matches=10):
        """Extract matches for given patterns - ONLY from original text"""
        matches = []
        for pattern in patterns:
            found = re.findall(pattern, text, re.IGNORECASE)
            if found:
                for match in found:
                    if isinstance(match, tuple):
                        match = ' '.join(match)
                    clean_match = re.sub(r'[^\w\s+]', ' ', str(match)).strip()
                    if clean_match and len(clean_match) > 2:
                        matches.append(clean_match)
                        if len(matches) >= max_matches:
                            break
            if len(matches) >= max_matches:
                break
        return matches

    # Extract terms that actually exist in the original query
    experience_terms = extract_matches_from_original(experience_patterns, query, 5)
    skill_terms = extract_matches_from_original(skill_patterns, query, 15)
    role_terms = extract_matches_from_original(role_patterns, query, 8)
    domain_terms = extract_matches_from_original(domain_patterns, query, 8)

    # Build optimized query using ONLY words from the original
    important_words = []

    # Add found experience terms (from original)
    if experience_terms:
        important_words.extend(experience_terms[:3])

    # Add found skill terms (from original)
    if skill_terms:
        important_words.extend(skill_terms[:10])

    # Add found role terms (from original)
    if role_terms:
        important_words.extend(role_terms[:5])

    # Add found domain terms (from original)
    if domain_terms:
        important_words.extend(domain_terms[:5])

    # Add other important words from original query (filtering out common stop words)
    stop_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'between', 'among', 'within', 'without', 'under', 'over', 'again', 'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'only', 'own', 'same', 'than', 'too', 'very', 'can', 'will', 'just', 'should', 'now'}

    for word in unique_words:
        word_lower = word.lower()
        if word_lower not in stop_words and len(word) >= 3:
            # Check if it's not already included
            if not any(word_lower in term.lower() for term in important_words):
                important_words.append(word)

    # Remove duplicates while preserving order
    final_words = []
    seen = set()
    for word in important_words:
        word_lower = word.lower()
        if word_lower not in seen:
            final_words.append(word)
            seen.add(word_lower)

    # Combine into optimized query (limit to reasonable length)
    optimized_query = ' '.join(final_words[:50])  # Limit to 50 most important words

    # Ensure minimum length by including more original words if needed
    if len(final_words) < 10 and len(unique_words) > len(final_words):
        additional_words = [w for w in unique_words if w.lower() not in seen and w.lower() not in stop_words]
        final_words.extend(additional_words[:10-len(final_words)])
        optimized_query = ' '.join(final_words)

    print(f"📊 Fallback single-query extraction results:")
    print(f"   🎯 Experience terms: {experience_terms}")
    print(f"   🔧 Skill terms: {skill_terms}")
    print(f"   💼 Role terms: {role_terms}")
    print(f"   🏢 Domain terms: {domain_terms}")
    print(f"   ✨ Optimized query: '{optimized_query}'")

    return {
        "success": False,
        "fallback": True,
        "optimized_query": optimized_query,
        "enhanced_query": optimized_query,
        "original_query": query,
        "single_query_strategy": True,
        "extracted_info": {
            "experience_years": experience_terms[0] if experience_terms else "Not specified",
            "seniority_level": "Senior" if "senior" in optimized_query.lower() else "Not specified",
            "key_skills": skill_terms[:3],
            "job_titles": role_terms[:2],
            "domain_areas": domain_terms[:2],
            "must_have_requirements": skill_terms[:2],
            "location": "Not specified",
            "education": "Not specified"
        },
        "query_components": {
            "experience_part": experience_terms[0] if experience_terms else "experienced",
            "skills_part": ' '.join(skill_terms[:3]) if skill_terms else "technical",
            "role_part": ' '.join(role_terms[:2]) if role_terms else "specialist",
            "domain_part": ' '.join(domain_terms[:2]) if domain_terms else ""
        },
        "fallback_stats": {
            "experience_terms_found": len(experience_terms),
            "skill_terms_found": len(skill_terms),
            "role_terms_found": len(role_terms),
            "domain_terms_found": len(domain_terms)
        }
    }


def rank_results_with_llm(results: List[Dict], original_query: str, keywords_data: Dict) -> List[Dict]:
    """Use LLM to intelligently rank and score the search results"""
    print(f"\n🧠 STEP 3: LLM INTELLIGENT RANKING")
    print(f"📊 Ranking {len(results)} candidates")

    if not results:
        return results

    try:
        client = gemini_client()

        # Prepare candidate summaries for ranking
        candidates_for_ranking = []
        for i, result in enumerate(results):
            analysis = result.get('gemini_analysis', {})
            if analysis.get('success') and analysis.get('analysis_json'):
                candidate_info = analysis['analysis_json']
                candidates_for_ranking.append({
                    "index": i,
                    "name": candidate_info.get('candidate_name', f'Candidate {i+1}'),
                    "role": candidate_info.get('current_role', 'Unknown'),
                    "experience": candidate_info.get('experience_years', 'Unknown'),
                    "skills": candidate_info.get('matching_skills', []),
                    "strengths": candidate_info.get('key_strengths', []),
                    "initial_score": analysis.get('match_score', 0)
                })

        if not candidates_for_ranking:
            print("⚠️ No candidates with successful analysis found for ranking")
            return results

        ranking_prompt = RANKING_PROMPT.format(
            job_posting=original_query,
            candidates_data=json.dumps(candidates_for_ranking, indent=2)
        )

        contents = [
            types.Content(
                role="user",
                parts=[types.Part(text=ranking_prompt)]
            )
        ]

        # Try different model versions with fallbacks
        models_to_try = ["gemini-2.5-flash", "gemini-1.5-pro", "gemini-pro", "gemini-1.0-pro-latest"]
        generate_response = None

        for model_name in models_to_try:
            try:
                print(f"🤖 Trying ranking model: {model_name}")
                generate_response = client.models.generate_content(
                    model=model_name,
                    contents=contents,
                    config=types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
                        temperature=0.1,
                        max_output_tokens=4096,  # Maximum allowed for Gemini 2.5 Flash
                    ),
                )
                print(f"✅ Successfully connected to ranking model: {model_name}")
                break
            except Exception as model_error:
                print(f"❌ Ranking model {model_name} failed: {str(model_error)}")
                continue

        if not generate_response:
            print(f"❌ All ranking models failed")
            return results

        if generate_response:
            print(f"✅ LLM ranking response received")

            # Check if response has text attribute
            response_text = ""
            if hasattr(generate_response, 'text') and generate_response.text:
                response_text = generate_response.text.strip()
            elif hasattr(generate_response, 'candidates') and generate_response.candidates:
                # Try to get text from candidates
                for candidate in generate_response.candidates:
                    if hasattr(candidate, 'content') and candidate.content:
                        for part in candidate.content.parts:
                            if hasattr(part, 'text'):
                                response_text += part.text
            else:
                response_text = str(generate_response)

            if response_text:
                print(f"📝 Ranking response text length: {len(response_text)}")

                try:
                    # Clean up the response
                    if "```json" in response_text:
                        response_text = response_text.split("```json")[1].split("```")[0].strip()
                    elif "```" in response_text:
                        response_text = response_text.split("```")[1].split("```")[0].strip()

                    ranking_data = json.loads(response_text)
                    ranked_candidates = ranking_data.get('ranked_candidates', [])

                    print(f"🏆 LLM ranking completed for {len(ranked_candidates)} candidates")

                    # Apply the new rankings to the original results
                    ranked_results = []
                    for ranked_candidate in ranked_candidates:
                        original_index = ranked_candidate['index']
                        if 0 <= original_index < len(results):
                            result = results[original_index].copy()

                            # Update the Gemini analysis with new ranking info
                            if 'gemini_analysis' in result:
                                result['gemini_analysis']['match_score'] = ranked_candidate['final_score']
                                result['gemini_analysis']['ranking_reason'] = ranked_candidate.get('ranking_reason', '')
                                result['gemini_analysis']['llm_ranked'] = True
                                result['gemini_analysis']['ranking_concerns'] = ranked_candidate.get('concerns', [])

                            ranked_results.append(result)

                    # Add any results that weren't ranked
                    ranked_indices = {rc['index'] for rc in ranked_candidates}
                    for i, result in enumerate(results):
                        if i not in ranked_indices:
                            ranked_results.append(result)

                    print(f"✅ Final ranked results: {len(ranked_results)} candidates")
                    return ranked_results

                except json.JSONDecodeError as e:
                    print(f"❌ Failed to parse LLM ranking response: {e}")
                    print(f"📝 Raw response: {response_text}")
                    return results
            else:
                print(f"❌ Empty ranking response from LLM")
                return results
        else:
            print(f"❌ No ranking response from LLM")
            return results

    except Exception as e:
        print(f"❌ LLM ranking failed: {str(e)}")
        return results

@app.post("/api/enhanced-search")
@limiter.limit("15/minute")
async def enhanced_search_with_single_query(
    request: Request,
    query: str = Form(...),
    result_count: int = Form(20),
    use_llm_optimization: bool = Form(True),
    user: dict = Depends(require_auth)
):
    """Enhanced search with single optimized query extraction and analysis"""
    query = _sanitize_user_text(query, max_len=4000)
    result_count = clamp_result_count(result_count, default=20)
    print(f"\n🚀 ENHANCED SINGLE-QUERY SEARCH PIPELINE STARTED")
    print(f"📝 Query: '{query}'")
    print(f"📊 Requested results: {result_count}")

    try:
        # Step 1: LLM Single Query Optimization (blocking SDK call — run in
        # a thread so the event loop stays responsive for other requests).
        keyword_result = await asyncio.to_thread(extract_keywords_for_single_query, query)
        enhanced_query = keyword_result.get('optimized_query', query)
        keywords_data = keyword_result.get('keywords_data', {})

        print(f"\n🔍 STEP 2: VECTOR SEARCH WITH OPTIMIZED SINGLE QUERY")
        print(f"🎯 Using optimized query: '{enhanced_query}'")
        print(f"📊 Query components extracted:")
        if keyword_result.get('extracted_info'):
            info = keyword_result['extracted_info']
            print(f"   🎖️ Seniority: {info.get('seniority_level', 'Not specified')}")
            print(f"   ⏱️ Experience: {info.get('experience_years', 'Not specified')}")
            print(f"   🔧 Key skills: {info.get('key_skills', [])}")
            print(f"   💼 Job titles: {info.get('job_titles', [])}")

        # Step 2: Vector Search with Single Optimized Query. When the user
        # specifies hard criteria, search a larger internal pool so the final
        # response can still contain the requested number of exact matches.
        search_criteria = parse_search_criteria(query)
        fetch_count = criteria_fetch_count(result_count, search_criteria)
        if fetch_count != result_count:
            print(f"🎯 Criteria overfetch enabled: requested={result_count}, pool={fetch_count}")

        # Use universal search client and request
        client = get_search_client()
        request = create_universal_search_request(
            enhanced_query,
            user,
            fetch_count,
            max_value=MAX_SEARCH_FETCH_RESULTS,
        )

        response = client.search(request)

        # Process search results and get Gemini analysis
        print(f"\n📄 PROCESSING VECTOR SEARCH RESULTS")
        result_list = list(response.results)
        found_count = len(result_list)
        print(f"📈 Found {found_count} vector search results")

        async def _process_one(i, result):
            print(f"\n📄 Processing result #{i+1}: {result.document.name if result.document else 'Unknown'}")

            document_data = {
                "id": result.id,
                "document_name": result.document.name if result.document else "Unknown",
                "relevance_score": getattr(result, 'relevance_score', 0.0),
            }
            file_path = None

            # Get Gemini analysis for each result (same logic, just isolated per-candidate)
            if result.document:
                try:
                    file_content = None

                    if hasattr(result.document, 'derived_struct_data') and result.document.derived_struct_data:
                        derived_data = dict(result.document.derived_struct_data)

                        if 'link' in derived_data and derived_data['link']:
                            gcs_link = derived_data['link']
                            if gcs_link.startswith('gs://'):
                                gcs_path = gcs_link.replace('gs://', '').split('/', 1)
                                if len(gcs_path) == 2:
                                    bucket_name, file_path = gcs_path
                                    if bucket_name == GCS_BUCKET_NAME:
                                        file_content, file_bytes = await asyncio.to_thread(
                                            get_file_content_and_bytes_from_gcs, file_path, None
                                        )

                                        if not file_content and file_bytes:
                                            direct_analysis = await asyncio.to_thread(
                                                send_file_to_gemini_directly, file_bytes, file_path, enhanced_query
                                            )
                                            if direct_analysis['success']:
                                                document_data["gemini_analysis"] = direct_analysis
                                                file_content = "Direct analysis"

                        if not file_content and 'snippets' in derived_data and derived_data['snippets']:
                            snippets = derived_data['snippets']
                            snippet_texts = []
                            for snip in snippets:
                                try:
                                    snip_dict = dict(snip)
                                    for key in ['snippet', 'content', 'text', 'extractive_segment']:
                                        if key in snip_dict and snip_dict[key]:
                                            snippet_texts.append(str(snip_dict[key]))
                                            break
                                except Exception:
                                    pass

                            if snippet_texts:
                                file_content = " ".join(snippet_texts)

                    # Analyze with primary query if we have content
                    if file_content and 'gemini_analysis' not in document_data:
                        # Extract job title from query for better scorecard context
                        job_title_match = re.search(r'(?:job title|position|role):\s*([^\n\r]+)', query, re.IGNORECASE)
                        job_title = job_title_match.group(1).strip() if job_title_match else "Position"

                        # Use comprehensive HR scorecard analysis
                        current_analysis = await analyze_resume_with_hr_scorecard(file_content, enhanced_query, job_title, file_path=file_path)
                        document_data["gemini_analysis"] = current_analysis
                    elif 'gemini_analysis' not in document_data:
                        document_data["gemini_analysis"] = {
                            "analysis": "Could not retrieve file content",
                            "match_score": 0,
                            "success": False
                        }

                except Exception as e:
                    document_data["gemini_analysis"] = {
                        "analysis": f"Analysis failed: {str(e)}",
                        "match_score": 0,
                        "success": False
                    }

            return document_data

        # Run candidate analyses in batches. With hard criteria, stop as soon
        # as the requested number of exact matches has been found.
        if search_criteria.is_empty():
            results = await asyncio.gather(*[
                _process_one(i, r) for i, r in enumerate(result_list)
            ])
        else:
            results = []
            batch_size = max(result_count, min(15, max(1, fetch_count)))
            for batch_start in range(0, len(result_list), batch_size):
                batch = result_list[batch_start:batch_start + batch_size]
                batch_results = await asyncio.gather(*[
                    _process_one(batch_start + i, r) for i, r in enumerate(batch)
                ])
                apply_criteria_to_results(batch_results, search_criteria, strict=True, min_kept=0)
                results.extend(batch_results)
                apply_criteria_to_results(results, search_criteria, strict=True, min_kept=0)
                results[:] = results[:result_count]
                if len(results) >= result_count:
                    break

        # Step 3: Optional LLM Ranking (if requested). The ranker uses a
        # blocking Gemini call; offload to a thread so the event loop can
        # service concurrent requests while we wait.
        if use_llm_optimization and results:
            print(f"\n🧠 STEP 3: LLM INTELLIGENT RANKING")
            ranked_results = await asyncio.to_thread(rank_results_with_llm, results, query, keywords_data)
        else:
            print(f"\n📊 STEP 3: SKIPPING LLM RANKING")
            ranked_results = results

        # Prepare final response with single-query approach data
        successful_analyses = [r['gemini_analysis'] for r in ranked_results if r.get('gemini_analysis', {}).get('success')]
        successful_analyses.sort(key=lambda x: x.get('match_score', 0), reverse=True)

        # Apply hard search criteria (min years, languages, location) extracted
        # from the original query. Uses the same parser as multi-query &
        # hr-scorecard endpoints so behaviour is consistent across all search
        # methods. Candidates failing criteria are score-penalised and demoted.
        if not search_criteria.is_empty():
            before_criteria_count = len(ranked_results)
            apply_criteria_to_results(ranked_results, search_criteria, strict=True, min_kept=0)
            ranked_results[:] = ranked_results[:result_count]
            successful_analyses = [
                r['gemini_analysis'] for r in ranked_results
                if r.get('gemini_analysis', {}).get('success')
            ]
            exact_dropped = before_criteria_count - len(ranked_results)
            print(
                f"Exact criteria applied: {search_criteria.to_dict()} "
                f"({exact_dropped} non-matches removed)"
            )

        final_response = {
            "query": query,
            "optimized_query": enhanced_query,
            "query_optimization": keyword_result,
            "search_strategy": "single_optimized_query",
            "total_results": len(ranked_results),
            "requested_result_count": result_count,
            "candidate_pool_size": fetch_count,
            "exact_match_shortfall": max(0, result_count - len(ranked_results)),
            "results": ranked_results,
            "all_analyses": successful_analyses,
            "top_result_analysis": successful_analyses[0] if successful_analyses else None,
            "analyzed_count": len(successful_analyses),
            "llm_optimization_used": use_llm_optimization,
            "single_query_approach": True,
            "datastore_id": get_company_resources(user)[1],
            "project_id": PROJECT_ID,
            "search_criteria": search_criteria.to_dict(),
        }

        print(f"\n🎉 ENHANCED SINGLE-QUERY PIPELINE COMPLETED")
        print(f"📊 Total results: {len(ranked_results)}")
        print(f"🧠 Successfully analyzed: {len(successful_analyses)}")
        print(f"🎯 Optimizations: Query={'✅' if keyword_result.get('success') else '⚠️ Fallback'}, Ranking={'✅' if use_llm_optimization else '⏭️ Skipped'}")
        if successful_analyses:
            print(f"🏆 Best match score: {successful_analyses[0].get('match_score', 'N/A')}")

        return final_response

    except Exception as e:
        print(f"❌ Enhanced single-query search pipeline failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Enhanced single-query search failed: {str(e)}")

@app.post("/api/smart-search-stream")
@limiter.limit("10/minute")
async def smart_search_with_streaming(
    request: Request,
    query: str = Form(...),
    result_count: int = Form(10),
    user: dict = Depends(require_auth)
):
    """Smart search with real-time progress streaming using Server-Sent Events"""
    query = _sanitize_user_text(query, max_len=4000)
    result_count = clamp_result_count(result_count, default=10)

    async def generate_progress():
        tracker = ProgressTracker()
        try:
            # Step 1: Initialize
            tracker.set_total_steps(5)
            progress_msg = tracker.emit('🚀 Starting smart search...', 'info')
            yield f"data: {json.dumps(progress_msg)}\n\n"
            await asyncio.sleep(0.1)

            # Step 2: Extract keywords
            tracker.next_step()
            progress_msg = tracker.emit('🧠 Extracting keywords from job description...', 'info')
            yield f"data: {json.dumps(progress_msg)}\n\n"

            keyword_result = extract_keywords_for_single_query(query)
            optimized_query = keyword_result.get('optimized_query', query)

            if keyword_result.get('success'):
                progress_msg = tracker.success('✅ Keywords extracted successfully')
                yield f"data: {json.dumps(progress_msg)}\n\n"

                extracted_info = keyword_result.get('extracted_info', {})
                if extracted_info.get('key_skills'):
                    skills_text = ', '.join(extracted_info['key_skills'][:5])
                    msg_text = f'🔧 Key skills found: {skills_text}'
                    progress_msg = tracker.emit(msg_text, 'info')
                    yield f"data: {json.dumps(progress_msg)}\n\n"
            else:
                progress_msg = tracker.warning('⚠️ Using fallback keyword extraction')
                yield f"data: {json.dumps(progress_msg)}\n\n"

            word_count = len(optimized_query.split())
            query_preview = optimized_query[:100] + "..."
            msg_text = f'📝 Generated optimized query ({word_count} words): {query_preview}'
            progress_msg = tracker.emit(msg_text, 'success')
            yield f"data: {json.dumps(progress_msg)}\n\n"

            # Step 3: Perform search with detailed progress
            tracker.next_step()
            progress_msg = tracker.emit('🔍 Searching vector database...', 'info')
            yield f"data: {json.dumps(progress_msg)}\n\n"

            # Store initial message count
            initial_msg_count = len(tracker.messages)

            # Execute search with progress tracking
            search_result = await execute_smart_search_with_progress(query, result_count, tracker, user)

            # Yield all new progress messages that were added during search
            for msg in tracker.messages[initial_msg_count:]:
                yield f"data: {json.dumps(msg)}\n\n"
                await asyncio.sleep(0.05)  # Small delay for better streaming

            # Step 4: Complete
            tracker.next_step()
            progress_msg = tracker.success('🎉 Search completed!')
            yield f"data: {json.dumps(progress_msg)}\n\n"

            # Send final result
            yield f"data: {json.dumps({'type': 'final_result', 'data': search_result})}\n\n"

        except Exception as e:
            error_msg = tracker.error(f'❌ Search failed: {str(e)}')
            yield f"data: {json.dumps(error_msg)}\n\n"
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return StreamingResponse(
        generate_progress(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "Content-Encoding": "identity",
        }
    )

async def execute_smart_search_with_progress(query: str, result_count: int, tracker: ProgressTracker, user: dict):
    """Execute smart search with real-time progress updates"""
    try:
        # Extract keywords
        keyword_result = extract_keywords_for_single_query(query)
        optimized_query = keyword_result.get('optimized_query', query)

        # Perform vector search using universal search helper
        client = get_search_client()
        request = create_universal_search_request(optimized_query, user, result_count)

        response = client.search(request)
        result_list = list(response.results)
        found_count = len(result_list)

        tracker.emit(f'📊 Found {found_count} matching documents')

        # Process results with REAL analysis and progress updates
        results = []
        successful_analyses = []

        for i, result in enumerate(result_list):
            current_doc = i + 1
            doc_name = result.document.name if result.document else 'Unknown'

            # Extract filename from document name for cleaner display
            clean_name = doc_name.split('/')[-1] if '/' in doc_name else doc_name

            tracker.emit(f'📄 Processing document {current_doc}/{found_count}: {clean_name}')

            document_data = {
                "id": result.id,
                "document_name": doc_name,
                "relevance_score": getattr(result, 'relevance_score', 0.0),
            }

            # Get and analyze document content
            if result.document:
                try:
                    file_content = None

                    if hasattr(result.document, 'derived_struct_data') and result.document.derived_struct_data:
                        derived_data = dict(result.document.derived_struct_data)

                        if 'link' in derived_data and derived_data['link']:
                            gcs_link = derived_data['link']
                            if gcs_link.startswith('gs://'):
                                tracker.emit(f'📁 Downloading from GCS: {clean_name}')

                                gcs_path = gcs_link.replace('gs://', '').split('/', 1)
                                if len(gcs_path) == 2:
                                    bucket_name, file_path = gcs_path
                                    if bucket_name == GCS_BUCKET_NAME:
                                        file_content, file_bytes = get_file_content_and_bytes_from_gcs(file_path, None)

                                        if file_content:
                                            tracker.emit(f'✅ Text extracted successfully ({len(file_content)} chars)')
                                        elif file_bytes:
                                            tracker.emit('🤖 Sending file directly to Gemini AI...')
                                            direct_analysis = send_file_to_gemini_directly(file_bytes, file_path, optimized_query)
                                            if direct_analysis['success']:
                                                document_data["gemini_analysis"] = direct_analysis
                                                successful_analyses.append(direct_analysis)
                                                score = direct_analysis.get('match_score', 0)
                                                tracker.emit(f'✅ Direct AI analysis completed - Score: {score}%', 'success')
                                                file_content = "Direct analysis"

                        if not file_content and 'snippets' in derived_data and derived_data['snippets']:
                            tracker.emit('📝 Extracting from document snippets...')
                            snippets = derived_data['snippets']
                            snippet_texts = []
                            for snip in snippets:
                                try:
                                    snip_dict = dict(snip)
                                    for key in ['snippet', 'content', 'text', 'extractive_segment']:
                                        if key in snip_dict and snip_dict[key]:
                                            snippet_texts.append(str(snip_dict[key]))
                                            break
                                except Exception:
                                    pass

                            if snippet_texts:
                                file_content = " ".join(snippet_texts)
                                tracker.emit(f'✅ Extracted {len(snippet_texts)} snippets')

                    # Analyze with Gemini if we have content
                    if file_content and 'gemini_analysis' not in document_data:
                        clean_name = doc_name.split('/')[-1] if '/' in doc_name else doc_name
                        tracker.emit(f'🧠 Analyzing with AI: {clean_name}')

                        # Extract job title from query for better scorecard context
                        job_title_match = re.search(r'(?:job title|position|role):\s*([^\n\r]+)', query, re.IGNORECASE)
                        job_title = job_title_match.group(1).strip() if job_title_match else "Position"

                        # Use comprehensive HR scorecard analysis
                        current_analysis = await analyze_resume_with_hr_scorecard(file_content, optimized_query, job_title, file_path=doc_name)
                        document_data["gemini_analysis"] = current_analysis

                        if current_analysis.get('success'):
                            successful_analyses.append(current_analysis)
                            score = current_analysis.get('match_score', 0)
                            tracker.emit(f'✅ HR Scorecard complete - Match score: {score}%', 'success')

                            # Add additional info for HR scorecard
                            if current_analysis.get('hr_scorecard'):
                                scorecard = current_analysis['hr_scorecard']
                                candidate_name = scorecard.get('candidate_overview', {}).get('name', 'Unknown')
                                match_status = scorecard.get('candidate_overview', {}).get('match_status', 'Unknown')
                                tracker.emit(f'📊 {candidate_name} - {match_status}', 'info')
                        else:
                            tracker.emit('⚠️ HR Scorecard completed with issues', 'warning')
                    elif 'gemini_analysis' not in document_data:
                        tracker.emit('⚠️ No content available for analysis', 'warning')
                        document_data["gemini_analysis"] = {
                            "analysis": "Could not retrieve file content",
                            "match_score": 0,
                            "success": False
                        }

                except Exception as e:
                    tracker.emit(f'❌ Error processing {clean_name}: {str(e)}', 'error')
                    document_data["gemini_analysis"] = {
                        "analysis": f"Analysis failed: {str(e)}",
                        "match_score": 0,
                        "success": False
                    }

            results.append(document_data)

            # Small delay to make progress visible
            await asyncio.sleep(0.1)

        # Sort by match score
        tracker.emit('📊 Sorting results by match score...')
        successful_analyses.sort(key=lambda x: x.get('match_score', 0), reverse=True)
        results.sort(key=lambda x: x.get('gemini_analysis', {}).get('match_score', 0), reverse=True)

        if successful_analyses:
            best_score = successful_analyses[0].get('match_score', 0)
            tracker.emit(f'🏆 Best match score: {best_score}%', 'success')

        # Prepare final response
        final_response = {
            "query": query,
            "optimized_query": optimized_query,
            "query_optimization_success": keyword_result.get('success', False),
            "extracted_info": keyword_result.get('extracted_info', {}),
            "search_strategy": "smart_single_query_streaming",
            "total_results": len(results),
            "results": results,
            "successful_analyses": successful_analyses,
            "top_match": successful_analyses[0] if successful_analyses else None,
            "analyzed_count": len(successful_analyses),
            "datastore_id": get_company_resources(user)[1],
            "project_id": PROJECT_ID,
            "completed": True
        }

        return final_response

    except Exception as e:
        tracker.emit(f'❌ Search failed: {str(e)}', 'error')
        raise HTTPException(status_code=500, detail=f"Smart search failed: {str(e)}")

async def execute_smart_search_internal(query: str, result_count: int, tracker: ProgressTracker = None, user: dict = Depends(require_auth)):
    """Internal function to execute smart search with optional progress tracking"""
    try:
        # Extract keywords
        keyword_result = extract_keywords_for_single_query(query)
        optimized_query = keyword_result.get('optimized_query', query)

        # Perform vector search using universal search helper
        client = get_search_client()
        request = create_universal_search_request(optimized_query, user, result_count)

        response = client.search(request)
        result_list = list(response.results)
        found_count = len(result_list)

        # Process results with REAL analysis (not placeholder)
        results = []
        successful_analyses = []

        for i, result in enumerate(result_list):
            current_doc = i + 1
            doc_name = result.document.name if result.document else 'Unknown'

            document_data = {
                "id": result.id,
                "document_name": doc_name,
                "relevance_score": getattr(result, 'relevance_score', 0.0),
            }

            # Get and analyze document content (REAL PROCESSING)
            if result.document:
                try:
                    file_content = None

                    if hasattr(result.document, 'derived_struct_data') and result.document.derived_struct_data:
                        derived_data = dict(result.document.derived_struct_data)

                        if 'link' in derived_data and derived_data['link']:
                            gcs_link = derived_data['link']
                            if gcs_link.startswith('gs://'):
                                gcs_path = gcs_link.replace('gs://', '').split('/', 1)
                                if len(gcs_path) == 2:
                                    bucket_name, file_path = gcs_path
                                    if bucket_name == GCS_BUCKET_NAME:
                                        file_content, file_bytes = get_file_content_and_bytes_from_gcs(file_path, None)

                                        if not file_content and file_bytes:
                                            direct_analysis = send_file_to_gemini_directly(file_bytes, file_path, optimized_query)
                                            if direct_analysis['success']:
                                                document_data["gemini_analysis"] = direct_analysis
                                                successful_analyses.append(direct_analysis)
                                                file_content = "Direct analysis"

                        if not file_content and 'snippets' in derived_data and derived_data['snippets']:
                            snippets = derived_data['snippets']
                            snippet_texts = []
                            for snip in snippets:
                                try:
                                    snip_dict = dict(snip)
                                    for key in ['snippet', 'content', 'text', 'extractive_segment']:
                                        if key in snip_dict and snip_dict[key]:
                                            snippet_texts.append(str(snip_dict[key]))
                                            break
                                except Exception:
                                    pass

                            if snippet_texts:
                                file_content = " ".join(snippet_texts)

                    # Analyze with Gemini if we have content
                    if file_content and 'gemini_analysis' not in document_data:
                        clean_name = doc_name.split('/')[-1] if '/' in doc_name else doc_name
                        tracker.emit(f'🧠 Analyzing with AI: {clean_name}')

                        # Extract job title from query for better scorecard context
                        job_title_match = re.search(r'(?:job title|position|role):\s*([^\n\r]+)', query, re.IGNORECASE)
                        job_title = job_title_match.group(1).strip() if job_title_match else "Position"

                        # Use comprehensive HR scorecard analysis
                        current_analysis = await analyze_resume_with_hr_scorecard(file_content, optimized_query, job_title, file_path=doc_name)
                        document_data["gemini_analysis"] = current_analysis

                        if current_analysis.get('success'):
                            successful_analyses.append(current_analysis)
                            score = current_analysis.get('match_score', 0)
                            tracker.emit(f'✅ HR Scorecard complete - Match score: {score}%', 'success')

                            # Add additional info for HR scorecard
                            if current_analysis.get('hr_scorecard'):
                                scorecard = current_analysis['hr_scorecard']
                                candidate_name = scorecard.get('candidate_overview', {}).get('name', 'Unknown')
                                match_status = scorecard.get('candidate_overview', {}).get('match_status', 'Unknown')
                                tracker.emit(f'📊 {candidate_name} - {match_status}', 'info')
                        else:
                            tracker.emit('⚠️ HR Scorecard completed with issues', 'warning')
                    elif 'gemini_analysis' not in document_data:
                        tracker.emit('⚠️ No content available for analysis', 'warning')
                        document_data["gemini_analysis"] = {
                            "analysis": "Could not retrieve file content",
                            "match_score": 0,
                            "success": False
                        }

                except Exception as e:
                    tracker.emit(f'❌ Error processing {clean_name}: {str(e)}', 'error')
                    document_data["gemini_analysis"] = {
                        "analysis": f"Analysis failed: {str(e)}",
                        "match_score": 0,
                        "success": False
                    }

            results.append(document_data)

        # Sort by match score
        successful_analyses.sort(key=lambda x: x.get('match_score', 0), reverse=True)
        results.sort(key=lambda x: x.get('gemini_analysis', {}).get('match_score', 0), reverse=True)

        # Hard-filter: penalize candidates missing user-specified criteria
        # (min years, languages, location) and attach criteria_match metadata.
        search_criteria = parse_search_criteria(query)
        if not search_criteria.is_empty():
            apply_criteria_to_results(results, search_criteria, strict=True, min_kept=0)
            results[:] = results[:result_count]
            # Re-sort successful_analyses to match adjusted scores
            successful_analyses = [
                r.get('gemini_analysis') for r in results
                if r.get('gemini_analysis', {}).get('success')
            ]

        # Prepare final response
        final_response = {
            "query": query,
            "optimized_query": optimized_query,
            "query_optimization_success": keyword_result.get('success', False),
            "extracted_info": keyword_result.get('extracted_info', {}),
            "search_strategy": "smart_single_query_streaming",
            "total_results": len(results),
            "results": results,
            "successful_analyses": successful_analyses,
            "top_match": successful_analyses[0] if successful_analyses else None,
            "analyzed_count": len(successful_analyses),
            "datastore_id": get_company_resources(user)[1],
            "project_id": PROJECT_ID,
            "search_criteria": search_criteria.to_dict(),
            "completed": True
        }

        return final_response

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Smart search failed: {str(e)}")

@app.post("/api/smart-search")
@limiter.limit("20/minute")
async def smart_search_legacy_redirect(
    request: Request,
    user: dict = Depends(require_auth)
):
    """DEPRECATED: legacy synchronous smart-search.

    Removed because the synchronous variant routinely exceeded the 90-second
    request budget. Clients should use either:
      - POST /api/smart-search-stream  (SSE, incremental results)
      - POST /api/hr-scorecard-search  (async task + polling)
      - POST /api/enhanced-search      (parallelized single-query)
    """
    raise HTTPException(
        status_code=410,
        detail="/api/smart-search is deprecated. Use /api/smart-search-stream, "
               "/api/hr-scorecard-search, or /api/enhanced-search instead."
    )

@app.get("/api/vector-status")
async def get_vector_status(user: dict = Depends(require_auth)):
    """Get current vector search configuration status"""
    credentials_exist = os.path.exists(CREDENTIALS_PATH) or bool(os.environ.get("K_SERVICE"))
    return {
        "configured": credentials_exist,
        "project_id": PROJECT_ID,
        "location": LOCATION,
        "datastore_id": get_company_resources(user)[1],
        "credentials_found": credentials_exist,
        "credentials_path": CREDENTIALS_PATH
    }

@app.post("/api/multi-query-search")
async def multi_query_search_and_analyze(
    query: str = Form(...),
    result_count: int = Form(15),
    use_multi_search: bool = Form(True),
    user: dict = Depends(require_auth)
):
    """Execute exactly 4 structured searches and analyze all results with LLM"""
    query = _sanitize_user_text(query, max_len=4000)
    result_count = clamp_result_count(result_count, default=15)
    print(f"\n🚀 4-QUERY SEARCH PIPELINE STARTED")
    print(f"📝 Original query length: {len(query)} characters")
    print(f"📊 Target results per query: {result_count}")

    try:
        # Check search limits BEFORE performing any expensive operations
        db = get_db_manager()
        company_id = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                company_id = user['company'].get('company_id') or user['company'].get('id')
            else:
                company_id = getattr(user['company'], 'id', None)

        if company_id:
            try:
                db.check_search_limit(company_id)
                print(f"✅ Search limit check passed for company {company_id}")
            except Exception as e:
                print(f"❌ Search limit exceeded: {e}")
                raise HTTPException(status_code=400, detail=str(e))
        # Step 1: Extract 4 structured search queries
        # Use the single query extraction as fallback since extract_keywords_with_llm doesn't exist
        keyword_result = extract_keywords_for_single_query(query)

        if not keyword_result.get('success'):
            print("⚠️ Using fallback 4-query extraction")

        # Prepare exactly 4 search queries from the single optimized query
        optimized_query = keyword_result.get('optimized_query', query)
        extracted_info = keyword_result.get('extracted_info', {})

        # Create 4 focused queries from the extracted information
        experience_query = f"{extracted_info.get('seniority_level', 'senior')} {extracted_info.get('experience_years', 'experienced')} professional"
        skills_query = ' '.join(extracted_info.get('key_skills', ['python', 'sql', 'aws'])[:5])
        role_query = ' '.join(extracted_info.get('job_titles', ['developer', 'engineer'])[:3])
        domain_query = ' '.join(extracted_info.get('domain_areas', ['technology', 'software'])[:3])

        search_queries = [
            ("experience", experience_query if experience_query.strip() else 'senior experience'),
            ("skills", skills_query if skills_query.strip() else 'python sql aws'),
            ("role", role_query if role_query.strip() else 'developer engineer'),
            ("domain", domain_query if domain_query.strip() else 'technology software')
        ]

        print(f"\n🔍 EXECUTING EXACTLY 4 STRUCTURED SEARCHES:")
        for i, (search_type, search_query) in enumerate(search_queries, 1):
            print(f"   {i}. {search_type.upper()}: '{search_query}'")

        # Step 2: Execute all 4 searches in parallel
        all_search_results = {}
        combined_document_ids = set()

        # Use universal search client
        client = get_search_client()

        async def _run_one_search(search_type, search_query):
            try:
                print(f"\n🔎 Executing {search_type.upper()} search: '{search_query}'")
                req = create_universal_search_request(search_query, user, min(result_count, 10))
                response = await asyncio.to_thread(client.search, req)
                results = list(response.results)
                print(f"   📊 Found {len(results)} results for {search_type} search")
                return search_type, search_query, results, None
            except Exception as e:
                print(f"   ❌ {search_type} search failed: {str(e)}")
                return search_type, search_query, [], str(e)

        # gather is order-preserving; results come back in the same order as input
        gathered = await asyncio.gather(*[
            _run_one_search(st, sq) for st, sq in search_queries
        ])

        for search_type, search_query, results, err in gathered:
            entry = {
                'query': search_query,
                'results': results,
                'count': len(results),
            }
            if err is not None:
                entry['error'] = err
            all_search_results[search_type] = entry
            for result in results:
                combined_document_ids.add(result.id)

        # Step 3: Combine and score results from all 4 searches with deduplication
        print(f"\n📊 COMBINING RESULTS FROM 4 SEARCHES WITH DEDUPLICATION")
        print(f"   🔢 Unique documents found: {len(combined_document_ids)}")

        # Enhanced scoring system for 4-query approach with proper deduplication
        document_scores = {}
        search_weights = {
            'experience': 1.0,  # Experience matching is critical
            'skills': 1.2,      # Skills are most important
            'role': 0.9,        # Role titles are important but flexible
            'domain': 0.8       # Domain is good but not critical
        }

        # Track raw counts for debugging
        total_raw_results = 0
        duplicate_encounters = 0

        for search_type, search_data in all_search_results.items():
            for i, result in enumerate(search_data['results']):
                total_raw_results += 1
                doc_id = result.id
                relevance_score = getattr(result, 'relevance_score', 0.0)

                # Position-based scoring (earlier results get higher scores).
                # Clamp at 0 so positions beyond 12 don't subtract from the
                # candidate's own relevance_score (bug fix: previously
                # `1.0 - i*0.08` went negative for i>=13, penalising deep
                # results rather than just neutralising the position bonus).
                position_score = max(0.0, 1.0 - (i * 0.08))

                # Calculate weighted score for this search
                weighted_score = (relevance_score + position_score) * search_weights.get(search_type, 0.5)

                if doc_id not in document_scores:
                    # First time seeing this document
                    document_scores[doc_id] = {
                        'document': result,
                        'searches': [search_type],
                        'total_score': weighted_score,
                        'search_scores': {search_type: weighted_score},
                        'best_search': search_type,
                        'best_position': i,
                        'multi_search_bonus': 0,
                        'relevance_scores': {search_type: relevance_score}
                    }
                else:
                    # Document found in multiple searches - this is a duplicate
                    duplicate_encounters += 1
                    existing_doc = document_scores[doc_id]

                    # Add this search to the list
                    existing_doc['searches'].append(search_type)
                    existing_doc['search_scores'][search_type] = weighted_score
                    existing_doc['relevance_scores'][search_type] = relevance_score

                    # Update total score with multi-search bonus
                    existing_doc['total_score'] += weighted_score * 0.6  # Multi-search bonus

                    # Check if this is a better version of the document (higher relevance or better position)
                    current_best_score = existing_doc['relevance_scores'][existing_doc['best_search']]
                    if relevance_score > current_best_score or (relevance_score == current_best_score and i < existing_doc['best_position']):
                        # This version is better - update the document reference
                        existing_doc['document'] = result
                        existing_doc['best_search'] = search_type
                        existing_doc['best_position'] = i
                        print(f"   🔄 Updated document {doc_id} with better version from {search_type} search")

        print(f"   📈 Raw results processed: {total_raw_results}")
        print(f"   🔍 Duplicate encounters: {duplicate_encounters}")
        print(f"   ✅ Unique documents after deduplication: {len(document_scores)}")

        # Apply additional multi-search bonuses for documents found in multiple searches
        for doc_id, doc_data in document_scores.items():
            search_count = len(doc_data['searches'])
            if search_count > 1:
                # Exponential bonus for documents found in multiple categories
                multi_bonus = (search_count - 1) * 0.5  # 0.5 bonus for each additional search type
                doc_data['total_score'] += multi_bonus
                doc_data['multi_search_bonus'] = multi_bonus

                print(f"   🏆 Document {doc_id} found in {search_count} searches, bonus: +{multi_bonus:.2f}")

        # Step 4: Select top candidates based on combined scoring (already deduplicated)
        sorted_documents = sorted(
            document_scores.values(),
            key=lambda x: x['total_score'],
            reverse=True
        )[:result_count]

        print(f"   🎯 Top {len(sorted_documents)} unique documents selected for LLM analysis")

        # Show deduplication summary
        multi_search_docs = [doc for doc in sorted_documents if len(doc['searches']) > 1]
        if multi_search_docs:
            print(f"   🔗 Documents found in multiple searches: {len(multi_search_docs)}")
            for doc in multi_search_docs[:3]:  # Show top 3 as examples
                doc_name = doc['document'].document.name if doc['document'].document else 'Unknown'
                print(f"      - {doc_name}: found in {doc['searches']} (score: {doc['total_score']:.2f})")

        # Step 5: Send ALL selected unique documents to LLM for analysis
        print(f"\n🧠 ANALYZING ALL UNIQUE DOCUMENTS WITH LLM")
        final_results = []

        # Use the skills query as the primary analysis query (most specific)
        analysis_query = skills_query if skills_query.strip() else optimized_query[:200]
        print(f"🎯 Using analysis query: '{analysis_query}'")

        async def _analyze_one(i, doc_data):
            result = doc_data['document']
            print(f"\n📄 Analyzing unique document #{i+1}: {result.document.name if result.document else 'Unknown'}")
            print(f"   🎯 Found in searches: {doc_data['searches']}")
            print(f"   🏅 Best version from: {doc_data['best_search']} search (position {doc_data['best_position']})")
            print(f"   📊 Combined score: {doc_data['total_score']:.3f}")
            print(f"   🏆 Multi-search bonus: {doc_data['multi_search_bonus']:.3f}")

            document_data = {
                "id": result.id,
                "document_name": result.document.name if result.document else "Unknown",
                "relevance_score": getattr(result, 'relevance_score', 0.0),
                "combined_score": doc_data['total_score'],
                "found_in_searches": doc_data['searches'],
                "search_scores": doc_data['search_scores'],
                "multi_search_bonus": doc_data['multi_search_bonus'],
                "best_search": doc_data['best_search'],
                "best_position": doc_data['best_position'],
                "is_duplicate_resolved": len(doc_data['searches']) > 1
            }

            # Extract and analyze document content with LLM
            if result.document:
                try:
                    file_content = None

                    if hasattr(result.document, 'derived_struct_data') and result.document.derived_struct_data:
                        derived_data = dict(result.document.derived_struct_data)

                        # Try to get content from GCS link
                        if 'link' in derived_data and derived_data['link']:
                            gcs_link = derived_data['link']
                            if gcs_link.startswith('gs://'):
                                gcs_path = gcs_link.replace('gs://', '').split('/', 1)
                                if len(gcs_path) == 2:
                                    bucket_name, file_path = gcs_path
                                    if bucket_name == GCS_BUCKET_NAME:
                                        file_content, file_bytes = await asyncio.to_thread(
                                            get_file_content_and_bytes_from_gcs, file_path, None
                                        )

                                        if not file_content and file_bytes:
                                            # Send directly to Gemini with analysis query
                                            print(f"   🤖 Sending file directly to Gemini for analysis...")
                                            direct_analysis = await asyncio.to_thread(
                                                send_file_to_gemini_directly, file_bytes, file_path, analysis_query
                                            )
                                            if direct_analysis['success']:
                                                document_data["gemini_analysis"] = direct_analysis
                                                print(f"   ✅ Direct file analysis successful (score: {direct_analysis.get('match_score', 'N/A')})")
                                                file_content = "Direct analysis completed"

                        # Fallback to snippets if no direct content
                        if not file_content and 'snippets' in derived_data and derived_data['snippets']:
                            snippets = derived_data['snippets']
                            snippet_texts = []
                            for snip in snippets:
                                try:
                                    snip_dict = dict(snip)
                                    for key in ['snippet', 'content', 'text', 'extractive_segment']:
                                        if key in snip_dict and snip_dict[key]:
                                            snippet_texts.append(str(snip_dict[key]))
                                            break
                                except Exception:
                                    pass

                            if snippet_texts:
                                file_content = " ".join(snippet_texts)
                                print(f"   📝 Using snippet content (length: {len(file_content)})")

                    # Analyze with LLM if we have content and no analysis yet
                    if file_content and 'gemini_analysis' not in document_data:
                        print(f"   🧠 Generating HR scorecard...")

                        # Extract job title from query for better scorecard context
                        job_title_match = re.search(r'(?:job title|position|role):\s*([^\n\r]+)', query, re.IGNORECASE)
                        job_title = job_title_match.group(1).strip() if job_title_match else "Position"

                        # Use comprehensive HR scorecard analysis
                        current_analysis = await analyze_resume_with_hr_scorecard(file_content, analysis_query, job_title, file_path=(result.document.name if result.document else None))
                        document_data["gemini_analysis"] = current_analysis
                        print(f"   ✅ HR Scorecard completed (score: {current_analysis.get('match_score', 'N/A')})")
                    elif 'gemini_analysis' not in document_data:
                        print(f"   ❌ No content available for analysis")
                        document_data["gemini_analysis"] = {
                            "analysis": "Could not retrieve file content",
                            "match_score": 0,
                            "success": False
                        }

                except Exception as e:
                    print(f"   ❌ Analysis error: {str(e)}")
                    document_data["gemini_analysis"] = {
                        "analysis": f"Analysis failed: {str(e)}",
                        "match_score": 0,
                        "success": False
                    }

            return document_data

        # Run all candidate analyses in parallel (preserves order via gather).
        final_results = await asyncio.gather(*[
            _analyze_one(i, doc_data) for i, doc_data in enumerate(sorted_documents)
        ])

        # Prepare final response with comprehensive data including deduplication info
        successful_analyses = [r['gemini_analysis'] for r in final_results if r.get('gemini_analysis', {}).get('success')]
        successful_analyses.sort(key=lambda x: x.get('match_score', 0), reverse=True)

        # Hard-filter: apply user-specified criteria (min years, languages, location)
        search_criteria = parse_search_criteria(query)
        if not search_criteria.is_empty():
            apply_criteria_to_results(final_results, search_criteria, strict=True, min_kept=0)
            final_results[:] = final_results[:result_count]
            successful_analyses = [
                r['gemini_analysis'] for r in final_results
                if r.get('gemini_analysis', {}).get('success')
            ]

        # Calculate deduplication statistics
        total_documents_before_dedup = sum(all_search_results[k]['count'] for k in all_search_results.keys())
        documents_after_dedup = len(final_results)
        duplicates_removed = total_documents_before_dedup - documents_after_dedup

        # Count documents found in multiple searches
        multi_search_documents = len([r for r in final_results if r.get('is_duplicate_resolved', False)])

        final_response = {
            "query": query,
            "search_strategy": "4_query_structured_deduplicated",
            "search_queries": {
                "experience": all_search_results['experience']['query'],
                "skills": all_search_results['skills']['query'],
                "role": all_search_results['role']['query'],
                "domain": all_search_results['domain']['query']
            },
            "optimized_query": optimized_query,
            "extracted_info": extracted_info,
            "search_results_count": {k: v['count'] for k, v in all_search_results.items()},
            "deduplication_stats": {
                "total_raw_results": total_documents_before_dedup,
                "unique_documents": documents_after_dedup,
                "duplicates_removed": duplicates_removed,
                "multi_search_documents": multi_search_documents,
                "deduplication_rate": round((duplicates_removed / total_documents_before_dedup * 100), 2) if total_documents_before_dedup > 0 else 0
            },
            "total_unique_documents": len(combined_document_ids),
            "analyzed_documents": len(final_results),
            "results": final_results,
            "all_analyses": successful_analyses,
            "top_result_analysis": successful_analyses[0] if successful_analyses else None,
            "keyword_extraction": keyword_result,
            "datastore_id": get_company_resources(user)[1],
            "project_id": PROJECT_ID,
            "four_query_approach": True,
            "duplicate_handling": "enabled",
            "search_criteria": search_criteria.to_dict()
        }

        print(f"\n🎉 4-QUERY SEARCH PIPELINE WITH DEDUPLICATION COMPLETED")
        print(f"📊 Executed 4 structured searches:")
        for search_type, count in final_response["search_results_count"].items():
            print(f"   - {search_type.capitalize()}: {count} results")
        print(f"🔢 Raw results: {total_documents_before_dedup}")
        print(f"✅ Unique documents after deduplication: {documents_after_dedup}")
        print(f"🗑️ Duplicates removed: {duplicates_removed}")
        print(f"🔗 Documents found in multiple searches: {multi_search_documents}")
        print(f"📈 Deduplication rate: {final_response['deduplication_stats']['deduplication_rate']}%")
        print(f"🧠 Successfully analyzed: {len(successful_analyses)} documents")
        print(f"🏆 Best match score: {successful_analyses[0].get('match_score', 'N/A') if successful_analyses else 'N/A'}")

        return final_response

    except Exception as e:
        print(f"❌ 4-query search pipeline failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"4-query search failed: {str(e)}")

def deduplicate_search_results(all_search_results: Dict, search_weights: Dict = None) -> List[Dict]:
    """Deduplicate search results and select the best version of each document"""

    if search_weights is None:
        search_weights = {'default': 1.0}

    document_scores = {}
    total_raw_results = 0
    duplicate_encounters = 0

    for search_type, search_data in all_search_results.items():
        for i, result in enumerate(search_data.get('results', [])):
            total_raw_results += 1
            doc_id = result.id
            relevance_score = getattr(result, 'relevance_score', 0.0)

            # Position-based scoring. Clamp at 0 to prevent results past
            # position 12 from being penalised below their own relevance score.
            position_score = max(0.0, 1.0 - (i * 0.08))

            # Apply search type weighting
            weight = search_weights.get(search_type, 1.0)
            weighted_score = (relevance_score + position_score) * weight

            if doc_id not in document_scores:
                # First encounter of this document
                document_scores[doc_id] = {
                    'document': result,
                    'searches': [search_type],
                    'total_score': weighted_score,
                    'search_scores': {search_type: weighted_score},
                    'best_search': search_type,
                    'best_position': i,
                    'multi_search_bonus': 0,
                    'relevance_scores': {search_type: relevance_score}
                }
            else:
                # Duplicate found
                duplicate_encounters += 1
                existing_doc = document_scores[doc_id]

                # Add this search occurrence
                existing_doc['searches'].append(search_type)
                existing_doc['search_scores'][search_type] = weighted_score
                existing_doc['relevance_scores'][search_type] = relevance_score

                # Multi-search bonus
                existing_doc['total_score'] += weighted_score * 0.6

                # Check if this is a better version
                current_best_score = existing_doc['relevance_scores'][existing_doc['best_search']]
                if relevance_score > current_best_score or (relevance_score == current_best_score and i < existing_doc['best_position']):
                    existing_doc['document'] = result
                    existing_doc['best_search'] = search_type
                    existing_doc['best_position'] = i

    # Apply multi-search bonuses
    for doc_data in document_scores.values():
        search_count = len(doc_data['searches'])
        if search_count > 1:
            multi_bonus = (search_count - 1) * 0.5
            doc_data['total_score'] += multi_bonus
            doc_data['multi_search_bonus'] = multi_bonus

    print(f"📊 Deduplication results: {total_raw_results} raw → {len(document_scores)} unique (removed {duplicate_encounters} duplicates)")

    return list(document_scores.values())

@app.post("/api/hr-scorecard-search")
@limiter.limit("15/minute")
async def hr_scorecard_search(
    request: Request,
    query: str = Form(...),
    result_count: int = Form(10),
    job_title: str = Form("Position"),
    user: dict = Depends(require_auth)
):
    """Submit HR scorecard search as a cloud task for asynchronous processing"""
    query = _sanitize_user_text(query, max_len=4000)
    job_title = _sanitize_user_text(job_title, max_len=200) or "Position"
    result_count = clamp_result_count(result_count, default=10, max_value=MAX_HR_SCORECARD_RESULTS)
    print(f"\n🏢 HR SCORECARD SEARCH TASK SUBMISSION STARTED")
    print(f"📝 Job description: '{query[:100]}...'")
    print(f"💼 Position: {job_title}")
    print(f"📊 Requested results: {result_count}")

    try:
        # Check search limits BEFORE creating task
        db = get_db_manager()
        company_id = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                company_id = user['company'].get('company_id') or user['company'].get('id')
            else:
                company_id = getattr(user['company'], 'id', None)

        if company_id:
            try:
                db.check_search_limit(company_id)
                print(f"✅ Search limit check passed for company {company_id}")
            except Exception as e:
                print(f"❌ Search limit exceeded: {e}")
                raise HTTPException(status_code=400, detail=str(e))

        # Generate unique task ID
        task_id = f"hr-scorecard-{uuid.uuid4().hex[:12]}-{int(time.time())}"

        # Prepare user data (excluding sensitive information)
        user_data = {
            'id': user.get('id'),
            'email': user.get('email'),
            'company': user.get('company'),
            'user_type': user.get('user_type')
        }

        # Save task to database first
        success = save_task_to_database(
            task_id=task_id,
            query=query,
            job_title=job_title,
            result_count=result_count,
            user_id=user['id'],
            company_id=company_id
        )

        if not success:
            raise HTTPException(
                status_code=500,
                detail="Failed to save task to database"
            )

        # Create cloud task
        task_result = create_hr_scorecard_task(
            task_id=task_id,
            query=query,
            job_title=job_title,
            result_count=result_count,
            user_data=user_data
        )

        if not task_result['success']:
            # Update task status to failed
            update_task_status(task_id, 'failed', error=task_result['error'])
            raise HTTPException(
                status_code=500,
                detail=f"Failed to create cloud task: {task_result['error']}"
            )

        print(f"✅ HR scorecard search task submitted successfully: {task_id}")

        return {
            "success": True,
            "task_id": task_id,
            "message": "HR scorecard search task submitted successfully",
            "status": "pending",
            "query": query,
            "job_title": job_title,
            "result_count": result_count,
            "task_details": {
                "task_name": task_result['task_name'],
                "queue_name": task_result['queue_name'],
                "scheduled_time": task_result.get('scheduled_time')
            },
            "check_status_url": f"/api/hr-scorecard-task-status/{task_id}",
            "estimated_processing_time": "2-5 minutes"
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error submitting HR scorecard search task: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to submit HR scorecard search task: {str(e)}"
        )

# The old processing logic has been moved to the task processing endpoint
# This endpoint now only submits tasks for asynchronous processing



        # Store keywords in a common variable for ALL candidates to use
        MASTER_KEYWORDS_FOR_ANALYSIS = standardized_keywords.copy()

        print(f"\n✅ MASTER KEYWORDS ESTABLISHED:")
        print(f"🔑 standardized_keywords variable: {standardized_keywords}")
        print(f"🔑 MASTER_KEYWORDS_FOR_ANALYSIS variable: {MASTER_KEYWORDS_FOR_ANALYSIS}")
        print(f"🔑 Total keywords for analysis: {len(MASTER_KEYWORDS_FOR_ANALYSIS)}")
        print(f"🔑 Keywords list: {MASTER_KEYWORDS_FOR_ANALYSIS[:15]}{'...' if len(MASTER_KEYWORDS_FOR_ANALYSIS) > 15 else ''}")
        print(f"🎯 These SAME keywords will be used for ALL candidates to ensure consistency")

        # CRITICAL: Validate that we have keywords before proceeding
        if not MASTER_KEYWORDS_FOR_ANALYSIS or len(MASTER_KEYWORDS_FOR_ANALYSIS) == 0:
            print(f"🚨 CRITICAL FAILURE: MASTER_KEYWORDS_FOR_ANALYSIS is empty after all processing!")

        # Step 2: Get optimized query for search (separate from keywords)
        print(f"\n🔍 STEP 2: OPTIMIZING SEARCH QUERY")
        search_optimization = extract_keywords_for_single_query(query, token_tracker)
        optimized_query = search_optimization.get('optimized_query', query)

        print(f"🔍 Search query optimization success: {search_optimization.get('success', False)}")
        print(f"🎯 Optimized search query: '{optimized_query[:100]}...'")

        print(f"\n🔍 EXECUTING HR SCORECARD SEARCH")
        print(f"🎯 Optimized query: '{optimized_query[:100]}...'")

        # Step 3: Vector search using company-specific datastore
        # Use universal search client
        client = get_search_client()

        # Use company-specific datastore (automatic isolation)
        company_code = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                company_code = user['company'].get('company_code')
            else:
                company_code = getattr(user['company'], 'company_code', None)
        request = create_company_search_request(
            query=optimized_query,
            company_code=company_code,
            result_count=result_count
        )

        response = client.search(request)
        results = []
        hr_scorecards = []

        # Get results from company-specific datastore (no filtering needed - isolation is automatic)
        result_list = list(response.results)
        found_count = len(result_list)
        print(f"📈 Found {found_count} candidates from company-specific datastore")

        if found_count == 0:
            company_name = user.get('company', {}).get('company_name', 'your company')
            print(f"📋 No candidates found in datastore")

            # Calculate final processing time and token summary
            processing_time = time.time() - start_time
            token_summary = token_tracker.get_summary()

            return {
                "query": query,
                "job_title": job_title,
                "optimized_query": optimized_query,
                "query_optimization_success": search_optimization.get('success', False),
                "search_strategy": "hr_scorecard_comprehensive",
                "total_results": 0,
                "results": [],
                "hr_scorecards": [],
                "top_scorecard": None,
                "standardized_keywords": {
                    "keywords_extracted": len(MASTER_KEYWORDS_FOR_ANALYSIS),
                    "keywords_list": MASTER_KEYWORDS_FOR_ANALYSIS,
                    "extraction_success": jd_keyword_result.get('success', False),
                    "source": "standardized_jd_extraction",
                    "fallback_used": jd_keyword_result.get('fallback_used', False)
                },
                "hr_metrics": {
                    "successful_scorecards": 0,
                    "strong_fits": 0,
                    "medium_fits": 0,
                    "weak_fits": 0,
                    "scorecard_success_rate": 0
                },
                "datastore_id": get_company_resources(user)[1],
                "project_id": PROJECT_ID,
                "analysis_timestamp": __import__('time').time(),
                "processing_stats": {
                    "processing_time_seconds": round(processing_time, 2),
                    "token_usage": token_summary
                },
                "message": f"No resumes found for {company_name}. Please upload resumes first.",
                "empty_datastore": True
            }

        # Step 4: Generate HR scorecards for filtered results
        print(f"\n📄 GENERATING HR SCORECARDS FOR COMPANY CANDIDATES")

        for i, result in enumerate(result_list):
            print(f"\n🏢 Processing candidate #{i+1}: {result.document.name if result.document else 'Unknown'}")

            document_data = {
                "id": result.id,
                "document_name": result.document.name if result.document else "Unknown",
                "relevance_score": getattr(result, 'relevance_score', 0.0),
                "file_path": None,  # Will be set below if found
            }

            # Get and analyze document content with HR scorecard
            if result.document:
                try:
                    file_content = None

                    if hasattr(result.document, 'derived_struct_data') and result.document.derived_struct_data:
                        derived_data = dict(result.document.derived_struct_data)

                        if 'link' in derived_data and derived_data['link']:
                            gcs_link = derived_data['link']
                            if gcs_link.startswith('gs://'):
                                gcs_path = gcs_link.replace('gs://', '').split('/', 1)
                                if len(gcs_path) == 2:
                                    bucket_name, file_path = gcs_path
                                    # Get user's company resources to check if this file belongs to them
                                    user_bucket, user_datastore = get_company_resources(user)
                                    if bucket_name == user_bucket:
                                        document_data["file_path"] = file_path  # Store file path for download
                                        file_content, file_bytes = get_file_content_and_bytes_from_gcs(file_path, user)

                                        if not file_content and file_bytes:
                                            # For direct file processing, we'll use basic analysis
                                            # but note it in the scorecard
                                            direct_analysis = send_file_to_gemini_directly(file_bytes, file_path, optimized_query)
                                            if direct_analysis['success']:
                                                document_data["gemini_analysis"] = direct_analysis
                                                document_data["analysis_type"] = "direct_file_basic"
                                                file_content = "Direct analysis"

                        if not file_content and 'snippets' in derived_data and derived_data['snippets']:
                            snippets = derived_data['snippets']
                            snippet_texts = []
                            for snip in snippets:
                                try:
                                    snip_dict = dict(snip)
                                    for key in ['snippet', 'content', 'text', 'extractive_segment']:
                                        if key in snip_dict and snip_dict[key]:
                                            snippet_texts.append(str(snip_dict[key]))
                                            break
                                except Exception:
                                    pass

                            if snippet_texts:
                                file_content = " ".join(snippet_texts)

                    # Generate HR scorecard if we have content
                    if file_content and 'gemini_analysis' not in document_data:
                        print(f"\n🏢 GENERATING COMPREHENSIVE HR SCORECARD FOR CANDIDATE #{i+1}")

                        # CRITICAL DEBUG - Show all keyword variables at this point
                        print(f"\n🔍 KEYWORD VARIABLES AT CANDIDATE PROCESSING:")
                        print(f"   🔑 MASTER_KEYWORDS_FOR_ANALYSIS: {MASTER_KEYWORDS_FOR_ANALYSIS}")
                        print(f"   🔑 MASTER_KEYWORDS_FOR_ANALYSIS type: {type(MASTER_KEYWORDS_FOR_ANALYSIS)}")
                        try:
                            print(f"   🔑 MASTER_KEYWORDS_FOR_ANALYSIS length: {len(MASTER_KEYWORDS_FOR_ANALYSIS)}")
                        except:
                            print(f"   🔑 MASTER_KEYWORDS_FOR_ANALYSIS length: CANNOT_GET_LENGTH")
                        print(f"   🔑 standardized_keywords: {standardized_keywords}")
                        print(f"   🔑 standardized_keywords type: {type(standardized_keywords)}")
                        try:
                            print(f"   🔑 standardized_keywords length: {len(standardized_keywords)}")
                        except:
                            print(f"   🔑 standardized_keywords length: CANNOT_GET_LENGTH")

                        # Use the validated MASTER_KEYWORDS_FOR_ANALYSIS (no fallbacks)
                        keywords_to_use = MASTER_KEYWORDS_FOR_ANALYSIS
                        print(f"✅ Using MASTER_KEYWORDS_FOR_ANALYSIS: {len(keywords_to_use)} keywords")

                        print(f"🔍 FINAL DEBUG: keywords_to_use = {keywords_to_use[:5]}... (len={len(keywords_to_use)})")
                        print(f"🔍 keywords_to_use type = {type(keywords_to_use)}")
                        print(f"🔍 keywords_to_use is None = {keywords_to_use is None}")
                        print(f"🔍 About to call analyze_resume_with_hr_scorecard with {len(keywords_to_use)} keywords")

                        # CRITICAL: Verify keywords are not None before passing
                        if not keywords_to_use or len(keywords_to_use) == 0:
                            print(f"🚨 CRITICAL BUG: keywords_to_use is empty at function call!")
                            print(f"🔧 Forcing emergency keywords for this candidate")
                            keywords_to_use = ['Python', 'Programming', 'Development', 'Software', 'Technical', 'Experience']

                        print(f"\n🔍 FUNCTION CALL DEBUG:")
                        print(f"   📞 About to call: analyze_resume_with_hr_scorecard()")
                        print(f"   📄 file_content length: {len(file_content)}")
                        print(f"   📝 optimized_query: '{optimized_query[:50]}...'")
                        print(f"   💼 job_title: '{job_title}'")
                        print(f"   🔑 keywords_to_use: {keywords_to_use}")
                        print(f"   🔑 keywords_to_use type: {type(keywords_to_use)}")
                        print(f"   🔑 keywords_to_use length: {len(keywords_to_use)}")
                        print(f"   🚀 Making function call now...")

                        hr_analysis = await analyze_resume_with_hr_scorecard(file_content, optimized_query, job_title, keywords_to_use, token_tracker, file_path=document_data.get("file_path") or document_data.get("document_name"))
                        document_data["gemini_analysis"] = hr_analysis
                        document_data["analysis_type"] = "comprehensive_hr_scorecard"

                        if hr_analysis.get('success') and hr_analysis.get('hr_scorecard'):
                            hr_scorecards.append(hr_analysis)
                            scorecard = hr_analysis['hr_scorecard']
                            candidate_name = scorecard.get('candidate_overview', {}).get('name', 'Unknown')
                            match_status = scorecard.get('candidate_overview', {}).get('match_status', 'Unknown')
                            score = hr_analysis.get('match_score', 0)
                            print(f"✅ HR Scorecard: {candidate_name} - {match_status} ({score}%)")
                        else:
                            print(f"⚠️ HR Scorecard generation had issues")
                    elif 'gemini_analysis' not in document_data:
                        document_data["gemini_analysis"] = {
                            "analysis": "Could not retrieve file content for HR scorecard",
                            "match_score": 0,
                            "success": False,
                            "scorecard_type": "no_content"
                        }
                        document_data["analysis_type"] = "no_content"

                except Exception as e:
                    print(f"❌ Error processing candidate: {str(e)}")
                    document_data["gemini_analysis"] = {
                        "analysis": f"HR scorecard generation failed: {str(e)}",
                        "match_score": 0,
                        "success": False,
                        "error": str(e),
                        "scorecard_type": "error"
                    }
                    document_data["analysis_type"] = "error"

            results.append(document_data)

        # Sort by HR scorecard match score
        hr_scorecards.sort(key=lambda x: x.get('match_score', 0), reverse=True)
        results.sort(key=lambda x: x.get('gemini_analysis', {}).get('match_score', 0), reverse=True)

        # Hard-filter: apply user-specified criteria
        hr_search_criteria = parse_search_criteria(query)
        if not hr_search_criteria.is_empty():
            apply_criteria_to_results(results, hr_search_criteria, strict=True, min_kept=0)
            results[:] = results[:result_count]
            hr_scorecards = [
                r['gemini_analysis'] for r in results
                if r.get('gemini_analysis', {}).get('success')
            ]

        # Calculate HR metrics
        successful_scorecards = len([r for r in results if r.get('analysis_type') == 'comprehensive_hr_scorecard'])
        strong_fits = len([r for r in results if r.get('gemini_analysis', {}).get('hr_scorecard', {}).get('candidate_overview', {}).get('match_status') == 'Strong Fit'])
        medium_fits = len([r for r in results if r.get('gemini_analysis', {}).get('hr_scorecard', {}).get('candidate_overview', {}).get('match_status') == 'Medium Fit'])
        weak_fits = len([r for r in results if r.get('gemini_analysis', {}).get('hr_scorecard', {}).get('candidate_overview', {}).get('match_status') == 'Weak Fit'])

        final_response = {
            "query": query,
            "job_title": job_title,
            "optimized_query": optimized_query,
            "query_optimization_success": search_optimization.get('success', False),
            "search_strategy": "hr_scorecard_comprehensive",
            "total_results": len(results),
            "results": results,
            "hr_scorecards": hr_scorecards,
            "top_scorecard": hr_scorecards[0] if hr_scorecards else None,
            "standardized_keywords": {
                "keywords_extracted": len(MASTER_KEYWORDS_FOR_ANALYSIS),
                "keywords_list": MASTER_KEYWORDS_FOR_ANALYSIS,
                "extraction_success": jd_keyword_result.get('success', False),
                "source": "standardized_jd_extraction",
                "fallback_used": jd_keyword_result.get('fallback_used', False)
            },
            "hr_metrics": {
                "successful_scorecards": successful_scorecards,
                "strong_fits": strong_fits,
                "medium_fits": medium_fits,
                "weak_fits": weak_fits,
                "scorecard_success_rate": round((successful_scorecards / len(results) * 100), 2) if results else 0
            },
            "datastore_id": get_company_resources(user)[1],  # Use company-specific datastore
            "project_id": PROJECT_ID,
            "analysis_timestamp": __import__('time').time(),
            "processing_stats": {
                "processing_time_seconds": processing_time,
                "token_usage": token_summary
            }
        }

        print(f"\n🎉 HR SCORECARD SEARCH COMPLETED")
        print(f"📊 Total candidates: {len(results)}")
        print(f"🏢 Successful HR scorecards: {successful_scorecards}")
        print(f"🟢 Strong fits: {strong_fits}")
        print(f"🟡 Medium fits: {medium_fits}")
        print(f"🔴 Weak fits: {weak_fits}")
        print(f"🔑 Standardized keywords used: {len(MASTER_KEYWORDS_FOR_ANALYSIS)} (SAME for ALL candidates)")

        # Calculate processing time and get token usage summary
        end_time = time.time()
        processing_time = round(end_time - start_time, 2)
        token_summary = token_tracker.get_summary()

        print(f"\n📊 TOKEN USAGE SUMMARY:")
        print(f"   🔢 Total LLM calls: {token_summary['total_calls']}")
        print(f"   📈 Total input tokens: {token_summary['total_input_tokens']:,}")
        print(f"   📉 Total output tokens: {token_summary['total_output_tokens']:,}")
        print(f"   💰 Total tokens: {token_summary['total_tokens']:,}")
        print(f"   ⏱️ Processing time: {processing_time}s")
        print(f"   📋 Operations breakdown:")
        for operation, stats in token_summary['operations_breakdown'].items():
            print(f"      • {operation}: {stats['count']} calls, {stats['total_tokens']:,} tokens")

        # Show keyword analysis summary for verification
        if hr_scorecards:
            best_candidate = hr_scorecards[0].get('hr_scorecard', {}).get('candidate_overview', {})
            print(f"🏆 Top candidate: {best_candidate.get('name', 'Unknown')} ({hr_scorecards[0].get('match_score', 0)}%)")

            # Verify keyword consistency by showing counts
            print(f"\n🔍 KEYWORD ANALYSIS VERIFICATION:")
            keyword_totals = set()  # Track unique total counts to verify consistency
            for i, scorecard in enumerate(hr_scorecards[:5]):  # Show top 5 for verification
                candidate_name = scorecard.get('hr_scorecard', {}).get('candidate_overview', {}).get('name', f'Candidate {i+1}')
                keyword_coverage = scorecard.get('hr_scorecard', {}).get('keyword_coverage', {})
                matched_count = keyword_coverage.get('jd_keywords_matched', 0)
                total_count = keyword_coverage.get('total_jd_keywords', 0)
                keyword_totals.add(total_count)
                print(f"  📋 {candidate_name}: {matched_count}/{total_count} keywords matched")

            # Check for consistency
            if len(keyword_totals) == 1:
                print(f"  ✅ CONSISTENCY CHECK PASSED: All candidates use {list(keyword_totals)[0]} total keywords")
            else:
                print(f"  ❌ CONSISTENCY CHECK FAILED: Found different total keyword counts: {keyword_totals}")
                print(f"  🔧 This indicates the standardized keyword system needs debugging")

        # Save to PostgreSQL database
        try:
            db = get_db_manager()

            # Get company ID from user (handle different user object structures)
            company_id = None
            if user.get('company'):
                if isinstance(user['company'], dict):
                    company_id = user['company'].get('company_id') or user['company'].get('id')
                else:
                    company_id = getattr(user['company'], 'id', None)

            print(f"🏢 Company ID for database save: {company_id}")

            # Save search history
            search_id = db.save_search_history(
                query=query,
                job_title=job_title,
                result_count=len(results),
                user_id=user['id'],
                company_id=company_id,
                search_method='hr-scorecard'
            )

            # Save search results with tenant context
            if results:
                db.save_search_results(search_id, results, company_id=company_id, user_id=user['id'])
                print(f"💾 Saved search history and {len(results)} results to database (search_id: {search_id})")

                # Retrieve saved results to get their database IDs and attach to response
                try:
                    saved_rows = db.get_search_results(search_id)
                    file_to_id = {row['file_path']: str(row['id']) for row in saved_rows}

                    for res in results:
                        if 'file_path' in res and res['file_path'] in file_to_id:
                            res['id'] = file_to_id[res['file_path']]
                except Exception as map_err:
                    print(f"⚠️ Unable to map DB IDs back to results: {map_err}")

            # Add search_id to response
            final_response['search_id'] = search_id

        except Exception as db_error:
            print(f"❌ Database save error: {str(db_error)}")
            # Don't fail the entire request if database save fails
            final_response['database_save_error'] = str(db_error)

@app.post("/api/process-hr-scorecard-task")
async def process_hr_scorecard_task(request: Request):
    """Process HR scorecard search task - called by Cloud Tasks"""
    # Verify the caller is Cloud Tasks (OIDC token signed for our runtime SA).
    # Skip verification only when explicitly running locally (no K_SERVICE).
    if os.environ.get("K_SERVICE"):
        try:
            auth_header = request.headers.get("Authorization", "")
            if not auth_header.startswith("Bearer "):
                raise HTTPException(status_code=401, detail="Missing OIDC bearer token")
            token = auth_header.split(" ", 1)[1]
            from google.oauth2 import id_token as _gid
            from google.auth.transport import requests as _greq
            expected_sa = config.get("cloud_tasks", {}).get("service_account_email")
            target_uri = config.get("cloud_tasks", {}).get("target_uri", "")
            audience = f"{target_uri}/api/process-hr-scorecard-task" if target_uri else None
            claims = _gid.verify_oauth2_token(token, _greq.Request(), audience=audience)
            if expected_sa and claims.get("email") != expected_sa:
                raise HTTPException(status_code=403, detail="Untrusted task caller")
        except HTTPException:
            raise
        except Exception as _oidc_err:
            print(f"❌ OIDC verification failed: {_oidc_err}")
            raise HTTPException(status_code=401, detail="Invalid OIDC token")
    try:
        # Get task payload from request
        task_payload = await request.json()

        task_id = task_payload.get('task_id')
        query = task_payload.get('query')
        job_title = task_payload.get('job_title')
        result_count = clamp_result_count(
            task_payload.get('result_count'),
            default=10,
            max_value=MAX_HR_SCORECARD_RESULTS,
        )
        user_data = task_payload.get('user_data')

        # Idempotency: if Cloud Tasks retries a task that already finished
        # successfully, do not redo the work (which can clobber the original
        # search_id with a new empty run). Returning 200 OK acknowledges the
        # delivery so the queue drops the retry.
        try:
            db_idem = get_db_manager()
            existing = db_idem.get_task_status(task_id) if (db_idem and task_id) else None
            if existing and existing.get('status') == 'completed':
                progress_existing = existing.get('progress')
                if isinstance(progress_existing, str):
                    try:
                        progress_existing = json.loads(progress_existing)
                    except Exception:
                        progress_existing = {}
                if isinstance(progress_existing, dict) and progress_existing.get('search_id'):
                    print(f"⏩ Task {task_id} already completed with search_id={progress_existing.get('search_id')}, skipping retry")
                    return {"success": True, "skipped": True, "task_id": task_id}
        except Exception as idem_err:
            # Don't fail the task because of an idempotency check failure.
            print(f"⚠️  Idempotency check failed for {task_id}: {idem_err}")

        print(f"\n🏢 HR SCORECARD TASK PROCESSING STARTED")
        print(f"📝 Task ID: {task_id}")
        print(f"📝 Job description: '{query[:100]}...'")
        print(f"💼 Position: {job_title}")
        print(f"📊 Requested results: {result_count}")

        # Update task status to running
        update_task_status(task_id, 'running', {'step': 'initializing'})

        # Initialize token tracking and timing
        token_tracker = TokenTracker()
        start_time = time.time()

        # Run both keyword-extraction LLM calls in parallel (was sequential)
        update_task_status(task_id, 'running', {'step': 'extracting_keywords'})

        jd_keyword_result, search_optimization = await asyncio.gather(
            asyncio.to_thread(extract_standardized_keywords_from_jd, query, token_tracker),
            asyncio.to_thread(extract_keywords_for_single_query, query, token_tracker),
        )
        standardized_keywords = jd_keyword_result.get('all_keywords', [])

        if len(standardized_keywords) == 0:
            error_msg = "No keywords extracted from job description"
            update_task_status(task_id, 'failed', error=error_msg)
            return {"success": False, "error": error_msg}

        optimized_query = search_optimization.get('optimized_query', query)

        # Vector search using company-specific datastore
        update_task_status(task_id, 'running', {'step': 'searching_candidates'})

        client = get_search_client()

        # Reconstruct user object from user_data
        user = user_data

        # Get company code
        company_code = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                company_code = user['company'].get('company_code')
            else:
                company_code = getattr(user['company'], 'company_code', None)

        task_search_criteria = parse_search_criteria(query)
        fetch_count = criteria_fetch_count(
            result_count,
            task_search_criteria,
            max_value=MAX_HR_SCORECARD_FETCH_RESULTS,
        )
        if fetch_count != result_count:
            print(f"🎯 Criteria overfetch enabled: requested={result_count}, pool={fetch_count}")

        request_obj = create_company_search_request(
            query=optimized_query,
            company_code=company_code,
            result_count=fetch_count,
            max_value=MAX_HR_SCORECARD_FETCH_RESULTS,
        )

        response = client.search(request_obj)

        result_list = list(response.results)
        found_count = len(result_list)

        # Fallback: the LLM optimizer occasionally produces a Boolean-style
        # query with parentheses/OR/quoted phrases that Discovery Engine
        # treats literally, returning 0 hits. If we got nothing back, retry
        # with the raw user query before declaring "no candidates".
        if found_count == 0 and optimized_query and optimized_query.strip() != (query or "").strip():
            logging.warning(
                "🔁 Optimized query returned 0 hits; retrying with raw query. "
                "optimized=%r raw=%r", optimized_query[:200], (query or "")[:200]
            )
            try:
                fallback_req = create_company_search_request(
                    query=query,
                    company_code=company_code,
                    result_count=fetch_count,
                    max_value=MAX_HR_SCORECARD_FETCH_RESULTS,
                )
                fallback_resp = client.search(fallback_req)
                fb_results = list(fallback_resp.results)
                if fb_results:
                    logging.info("✅ Raw-query fallback recovered %d results", len(fb_results))
                    result_list = fb_results
                    found_count = len(fb_results)
            except Exception as fb_err:
                logging.error("Raw-query fallback failed: %s", fb_err)

        if found_count == 0:
            update_task_status(task_id, 'completed', {'step': 'completed', 'results': 0})
            return {"success": True, "results": 0, "message": "No candidates found"}

        # Generate HR scorecards for filtered results — in PARALLEL, with a
        # bounded concurrency cap and a per-(file_path, jd_hash) cache lookup
        # so repeat searches against the same JD return without re-LLMing.
        update_task_status(task_id, 'running', {
            'step': 'generating_scorecards',
            'total_candidates': found_count,
            'current_candidate': 0,
            'progress_percent': 0,
        })

        # Stable JD signature: derived ONLY from inputs the user actually
        # provided (raw query + job_title). Earlier versions also mixed in
        # `optimized_query` and `standardized_keywords`, but both come from
        # non-deterministic Gemini calls so identical user-facing searches
        # produced different hashes — the cache effectively never hit. With
        # this stable hash, a repeat search for the same JD returns from
        # cached_resume_analyses without re-LLMing.
        try:
            _q_sig = (query or "").strip().lower()[:500]
            _jt_sig = (job_title or "").strip().lower()[:200]
            jd_hash = hashlib.sha256(f"{_q_sig}|{_jt_sig}".encode("utf-8")).hexdigest()
        except Exception:
            jd_hash = hashlib.sha256((query or "").encode("utf-8")).hexdigest()

        db_cache = None
        try:
            db_cache = get_db_manager()
        except Exception as cache_init_err:
            print(f"⚠️ Cache disabled (no db): {cache_init_err}")

        # Resolve company_id for cache scoping (best-effort)
        cache_company_id = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                cache_company_id = user['company'].get('company_id') or user['company'].get('id')
            else:
                cache_company_id = getattr(user['company'], 'id', None)

        # Concurrency cap — Gemini 2.5 Flash regional quota easily handles
        # 10+ concurrent. Bump via SCORECARD_CONCURRENCY env var.
        max_concurrency = int(os.getenv('SCORECARD_CONCURRENCY', '15'))
        sem = asyncio.Semaphore(max_concurrency)
        progress_lock = asyncio.Lock()
        completed_counter = {'n': 0, 'cache_hits': 0}
        user_bucket, _user_datastore = get_company_resources(user)

        async def _process_one(idx: int, result):
            document_data = {
                "id": result.id,
                "document_name": result.document.name if result.document else "Unknown",
                "relevance_score": getattr(result, 'relevance_score', 0.0),
                "file_path": None,
            }
            if not result.document:
                return document_data

            try:
                file_content = None
                derived_data = None
                if hasattr(result.document, 'derived_struct_data') and result.document.derived_struct_data:
                    derived_data = dict(result.document.derived_struct_data)

                derived_snippet_text = ""
                if derived_data and derived_data.get('snippets'):
                    snippet_texts = []
                    for snip in derived_data['snippets']:
                        try:
                            snip_dict = dict(snip)
                            for key in ('snippet', 'content', 'text', 'extractive_segment'):
                                if snip_dict.get(key):
                                    snippet_texts.append(str(snip_dict[key]))
                                    break
                        except Exception:
                            pass
                    derived_snippet_text = " ".join(snippet_texts)

                # Resolve GCS file_path
                if derived_data and derived_data.get('link', '').startswith('gs://'):
                    gcs_path = derived_data['link'].replace('gs://', '').split('/', 1)
                    if len(gcs_path) == 2 and gcs_path[0] == user_bucket:
                        document_data["file_path"] = gcs_path[1]

                file_path_for_cache = document_data.get("file_path")

                # ── Cache lookup ──────────────────────────────────────
                if db_cache and file_path_for_cache:
                    try:
                        cached = db_cache.get_cached_resume_analysis(
                            file_path_for_cache, jd_hash, company_id=cache_company_id
                        )
                    except Exception as ce:
                        print(f"⚠️ Cache lookup error: {ce}")
                        cached = None
                    # Legacy migration entries (jd_hash='__legacy_no_jd__',
                    # source='migrated_legacy') were scored against an old
                    # query and only contain a partial scorecard. Treating
                    # them as a hit for an unrelated live search produced
                    # candidates with score=100 and "No AI analysis
                    # available". Skip them and force a fresh Gemini scoring.
                    if cached:
                        cached_jd_hash = cached.get('jd_hash')
                        cached_source = (cached.get('source') or '').lower()
                        is_legacy = (
                            cached_jd_hash == '__legacy_no_jd__'
                            or cached_source == 'migrated_legacy'
                        )
                        # Also reject obviously-incomplete cached payloads
                        # (e.g. missing hr_scorecard / candidate_overview)
                        # so a single bad write doesn't poison results.
                        ga_preview = cached.get('gemini_analysis') or {}
                        if isinstance(ga_preview, str):
                            try:
                                ga_preview = json.loads(ga_preview)
                            except Exception:
                                ga_preview = {}
                        hr_sc = ga_preview.get('hr_scorecard') if isinstance(ga_preview, dict) else None
                        has_overview = bool(
                            isinstance(hr_sc, dict)
                            and hr_sc.get('candidate_overview')
                            and (hr_sc.get('candidate_overview', {}).get('experience_years')
                                 or hr_sc.get('candidate_overview', {}).get('position_applied_for')
                                 or hr_sc.get('score_breakdown'))
                        )
                        if is_legacy or not has_overview:
                            print(
                                f"↩️  Skipping cache hit for {file_path_for_cache} "
                                f"(legacy={is_legacy}, complete={has_overview}); will re-score with Gemini"
                            )
                            cached = None
                    if cached and cached.get('gemini_analysis'):
                        ga = cached['gemini_analysis']
                        if isinstance(ga, str):
                            try: ga = json.loads(ga)
                            except Exception: ga = None
                        if ga:
                            ga = dict(ga)
                            # Sanitize identity fields on cached payloads
                            # (older entries may have "Not Provided" names).
                            try:
                                if isinstance(ga.get('hr_scorecard'), dict):
                                    ga['hr_scorecard'] = _sanitize_scorecard_identity(
                                        ga['hr_scorecard'], file_path_for_cache
                                    )
                                    ga['hr_scorecard'] = _enrich_scorecard_evidence(
                                        ga['hr_scorecard'],
                                        derived_snippet_text,
                                        optimized_query or query,
                                    )
                            except Exception as _se:
                                print(f"⚠️ Cache sanitize error: {_se}")
                            ga['from_cache'] = True
                            ga['cache_source'] = cached.get('source', 'cache')
                            document_data["gemini_analysis"] = ga
                            document_data["analysis_type"] = "comprehensive_hr_scorecard"
                            async with progress_lock:
                                completed_counter['n'] += 1
                                completed_counter['cache_hits'] += 1
                                update_task_status(task_id, 'running', {
                                    'step': 'generating_scorecards',
                                    'current_candidate': completed_counter['n'],
                                    'total_candidates': found_count,
                                    'progress_percent': int(completed_counter['n'] / found_count * 100),
                                    'cache_hits': completed_counter['cache_hits'],
                                })
                            return document_data

                # ── Need fresh content / LLM ─────────────────────────
                if file_path_for_cache:
                    file_content, file_bytes = await asyncio.to_thread(get_file_content_and_bytes_from_gcs, file_path_for_cache, user)
                    if not file_content and file_bytes:
                        direct_analysis = await asyncio.to_thread(send_file_to_gemini_directly, file_bytes, file_path_for_cache, optimized_query)
                        if direct_analysis.get('success'):
                            document_data["gemini_analysis"] = direct_analysis
                            document_data["analysis_type"] = "direct_file_basic"
                            file_content = "Direct analysis"

                if not file_content and derived_snippet_text:
                    file_content = derived_snippet_text

                if file_content and 'gemini_analysis' not in document_data:
                    async with sem:
                        hr_analysis = await analyze_resume_with_hr_scorecard(
                            file_content, optimized_query, job_title,
                            standardized_keywords, token_tracker,
                            file_path=file_path_for_cache
                        )
                    document_data["gemini_analysis"] = hr_analysis
                    document_data["analysis_type"] = "comprehensive_hr_scorecard"

                    # Persist to cache on success
                    if (db_cache and file_path_for_cache
                            and hr_analysis.get('success')
                            and hr_analysis.get('hr_scorecard')):
                        try:
                            db_cache.save_cached_resume_analysis(
                                file_path=file_path_for_cache,
                                jd_hash=jd_hash,
                                gemini_analysis=hr_analysis,
                                match_score=hr_analysis.get('match_score'),
                                jd_summary=(optimized_query or query or '')[:500],
                                company_id=cache_company_id,
                                source='live_gemini',
                            )
                        except Exception as save_err:
                            print(f"⚠️ Cache save failed: {save_err}")
                elif 'gemini_analysis' not in document_data:
                    document_data["gemini_analysis"] = {
                        "analysis": "Could not retrieve file content for HR scorecard",
                        "match_score": 0,
                        "success": False,
                        "scorecard_type": "no_content",
                    }
                    document_data["analysis_type"] = "no_content"

            except Exception as e:
                print(f"❌ Error processing candidate {idx}: {e}")
                document_data["gemini_analysis"] = {
                    "analysis": f"HR scorecard generation failed: {e}",
                    "match_score": 0,
                    "success": False,
                    "error": str(e),
                    "scorecard_type": "error",
                }
                document_data["analysis_type"] = "error"

            async with progress_lock:
                completed_counter['n'] += 1
                update_task_status(task_id, 'running', {
                    'step': 'generating_scorecards',
                    'current_candidate': completed_counter['n'],
                    'total_candidates': found_count,
                    'progress_percent': int(completed_counter['n'] / found_count * 100),
                    'cache_hits': completed_counter['cache_hits'],
                })
            return document_data

        # Process candidate analyses in batches. With hard criteria, keep
        # scoring deeper search results only until the requested number of
        # exact matches is reached.
        if task_search_criteria.is_empty():
            results = await asyncio.gather(*[
                _process_one(i, r) for i, r in enumerate(result_list)
            ])
        else:
            results = []
            batch_size = max(result_count, min(max_concurrency, max(1, fetch_count)))
            for batch_start in range(0, len(result_list), batch_size):
                batch = result_list[batch_start:batch_start + batch_size]
                batch_results = await asyncio.gather(*[
                    _process_one(batch_start + i, r) for i, r in enumerate(batch)
                ])
                apply_criteria_to_results(
                    batch_results,
                    task_search_criteria,
                    strict=True,
                    min_kept=0,
                )
                results.extend(batch_results)
                apply_criteria_to_results(
                    results,
                    task_search_criteria,
                    strict=True,
                    min_kept=0,
                )
                results[:] = results[:result_count]
                print(
                    f"🎯 Exact-match fill progress: {len(results)}/{result_count} "
                    f"after scoring {min(batch_start + len(batch), len(result_list))}/{found_count}"
                )
                if len(results) >= result_count:
                    break

        # Build hr_scorecards list from successful analyses
        hr_scorecards = [
            r['gemini_analysis'] for r in results
            if r.get('gemini_analysis', {}).get('success')
            and r.get('gemini_analysis', {}).get('hr_scorecard')
        ]

        # Sort by HR scorecard match score
        hr_scorecards.sort(key=lambda x: x.get('match_score', 0), reverse=True)
        results.sort(key=lambda x: x.get('gemini_analysis', {}).get('match_score', 0), reverse=True)

        # Hard-filter: apply user-specified criteria (location / years / etc.)
        # Cloud Tasks worker path -- mirrors the inline /api/hr-scorecard-search
        # filtering so off-location or over-experienced candidates are dropped
        # before results are persisted.
        try:
            if not task_search_criteria.is_empty():
                before_count = len(results)
                apply_criteria_to_results(results, task_search_criteria, strict=True, min_kept=0)
                results[:] = results[:result_count]
                hr_scorecards = [
                    r['gemini_analysis'] for r in results
                    if r.get('gemini_analysis', {}).get('success')
                    and r.get('gemini_analysis', {}).get('hr_scorecard')
                ]
                print(f"🎯 Hard-criteria filter: {before_count} -> {len(results)} "
                      f"(criteria={task_search_criteria.to_dict()})")
        except Exception as crit_err:
            print(f"⚠️  Hard-criteria filtering failed (non-fatal): {crit_err}")

        # Save results to database
        update_task_status(task_id, 'running', {'step': 'saving_results'})

        try:
            db = get_db_manager()

            # Get company ID from user data
            company_id = None
            if user.get('company'):
                if isinstance(user['company'], dict):
                    company_id = user['company'].get('company_id') or user['company'].get('id')
                else:
                    company_id = getattr(user['company'], 'id', None)

            # Save search history
            search_id = db.save_search_history(
                query=query,
                job_title=job_title,
                result_count=len(results),
                user_id=user['id'],
                company_id=company_id,
                search_method='hr-scorecard-task'
            )

            # Save search results
            if results:
                db.save_search_results(search_id, results, company_id=company_id, user_id=user['id'])

            # Update task status with results
            processing_time = time.time() - start_time
            token_summary = token_tracker.get_summary()

            final_progress = {
                'step': 'completed',
                'total_candidates': found_count,
                'requested_results': result_count,
                'returned_results': len(results),
                'candidate_pool_size': fetch_count,
                'exact_match_shortfall': max(0, result_count - len(results)),
                'search_criteria': task_search_criteria.to_dict(),
                'successful_scorecards': len([r for r in results if r.get('analysis_type') == 'comprehensive_hr_scorecard']),
                'processing_time_seconds': round(processing_time, 2),
                'token_usage': token_summary,
                'search_id': search_id
            }

            update_task_status(task_id, 'completed', final_progress)

            print(f"✅ HR scorecard task completed successfully: {task_id}")
            print(f"📊 Processed {found_count} candidates in {processing_time:.2f} seconds")

            return {
                "success": True,
                "task_id": task_id,
                "results": len(results),
                "requested_result_count": result_count,
                "candidate_pool_size": fetch_count,
                "exact_match_shortfall": max(0, result_count - len(results)),
                "search_id": search_id,
            }

        except Exception as db_error:
            print(f"❌ Database save error: {str(db_error)}")
            update_task_status(task_id, 'completed', {'step': 'completed_with_db_error', 'error': str(db_error)})
            return {"success": True, "task_id": task_id, "results": len(results), "db_error": str(db_error)}

    except Exception as e:
        print(f"❌ HR scorecard task failed: {str(e)}")
        if 'task_id' in locals():
            update_task_status(task_id, 'failed', error=str(e))
        # Return 5xx so Cloud Tasks retries with backoff. Idempotency check
        # at the top of this handler prevents duplicate work on retries that
        # already succeeded.
        raise HTTPException(status_code=500, detail=f"Task failed (will be retried): {str(e)}")

@app.get("/api/hr-scorecard-task-status/{task_id}")
async def get_hr_scorecard_task_status(task_id: str, user: dict = Depends(require_auth)):
    """Get the status of an HR scorecard task"""
    try:
        db = get_db_manager()

        # Get task status from database
        task_status = db.get_task_status(task_id)

        if not task_status:
            raise HTTPException(status_code=404, detail="Task not found")

        # Check if user has permission to view this task
        if task_status.get('user_id') != user['id']:
            # Check if user is super admin or from same company
            if user.get('user_type') != 'super_admin':
                user_company_id = None
                if user.get('company'):
                    if isinstance(user['company'], dict):
                        user_company_id = user['company'].get('company_id') or user['company'].get('id')
                    else:
                        user_company_id = getattr(user['company'], 'id', None)

                if user_company_id != task_status.get('company_id'):
                    raise HTTPException(status_code=403, detail="Access denied")

        return task_status

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error getting task status: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to get task status")

@app.get("/api/hr-scorecard-task-results/{task_id}")
async def get_hr_scorecard_task_results(task_id: str, user: dict = Depends(require_auth)):
    """Get the results of a completed HR scorecard task"""
    try:
        db = get_db_manager()

        # Get task status first
        task_status = db.get_task_status(task_id)

        if not task_status:
            raise HTTPException(status_code=404, detail="Task not found")

        # Check if user has permission to view this task
        if task_status.get('user_id') != user['id']:
            if user.get('user_type') != 'super_admin':
                user_company_id = None
                if user.get('company'):
                    if isinstance(user['company'], dict):
                        user_company_id = user['company'].get('company_id') or user['company'].get('id')
                    else:
                        user_company_id = getattr(user['company'], 'id', None)

                if user_company_id != task_status.get('company_id'):
                    raise HTTPException(status_code=403, detail="Access denied")

        # Check if task is completed
        if task_status.get('status') != 'completed':
            return {
                "task_id": task_id,
                "status": task_status.get('status'),
                "message": "Task not yet completed",
                "progress": task_status.get('progress')
            }

        # Get search results if task is completed
        progress = task_status.get('progress', {})
        if isinstance(progress, str):
            try:
                progress = json.loads(progress)
            except:
                progress = {}

        search_id = progress.get('search_id')
        if search_id:
            search_results = db.get_search_results(search_id)

            # Reshape DB rows into the payload shape the frontend's
            # generateEnhancedResults() expects (matches legacy synchronous
            # /api/hr-scorecard-search response shape).
            normalized_results = []
            hr_scorecards = []
            for row in (search_results or []):
                row_dict = dict(row) if not isinstance(row, dict) else row
                gemini_analysis = row_dict.get('gemini_analysis') or {}
                if isinstance(gemini_analysis, str):
                    try:
                        gemini_analysis = json.loads(gemini_analysis)
                    except Exception:
                        gemini_analysis = {}
                hr_scorecard = row_dict.get('hr_scorecard') or {}
                if isinstance(hr_scorecard, str):
                    try:
                        hr_scorecard = json.loads(hr_scorecard)
                    except Exception:
                        hr_scorecard = {}
                if hr_scorecard and not gemini_analysis.get('hr_scorecard'):
                    gemini_analysis['hr_scorecard'] = hr_scorecard

                normalized_results.append({
                    'document_name': row_dict.get('candidate_name') or 'Candidate',
                    'file_path': row_dict.get('file_path'),
                    'gemini_analysis': gemini_analysis,
                    'analysis_type': 'comprehensive_hr_scorecard',
                    'match_score': row_dict.get('match_score') or gemini_analysis.get('match_score', 0),
                })
                if gemini_analysis.get('success') and gemini_analysis.get('hr_scorecard'):
                    hr_scorecards.append(gemini_analysis)

            successful = len([r for r in normalized_results if (r.get('gemini_analysis') or {}).get('success')])
            total = len(normalized_results)

            return {
                "task_id": task_id,
                "status": "completed",
                "search_id": search_id,
                "task_info": task_status,
                "query": task_status.get('query', ''),
                "job_title": task_status.get('job_title', 'Position'),
                "search_strategy": "hr_scorecard_comprehensive",
                "total_results": total,
                "results": normalized_results,
                "hr_scorecards": hr_scorecards,
                "top_scorecard": hr_scorecards[0] if hr_scorecards else None,
                "standardized_keywords": {
                    "keywords_extracted": 0,
                    "keywords_list": [],
                    "extraction_success": False,
                    "source": "task_results",
                    "fallback_used": False,
                },
                "hr_metrics": {
                    "successful_scorecards": successful,
                    "strong_fits": len([r for r in normalized_results if (r.get('match_score') or 0) >= 75]),
                    "medium_fits": len([r for r in normalized_results if 50 <= (r.get('match_score') or 0) < 75]),
                    "weak_fits": len([r for r in normalized_results if (r.get('match_score') or 0) < 50]),
                    "scorecard_success_rate": round((successful / total * 100), 1) if total else 0,
                },
                "enhanced_search": True,
                "all_analyses": [r.get('gemini_analysis') for r in normalized_results],
                "analyzed_count": total,
                "top_result_analysis": normalized_results[0].get('gemini_analysis') if normalized_results else None,
            }
        else:
            return {
                "task_id": task_id,
                "status": "completed",
                "message": "Task completed but no search results found",
                "task_info": task_status,
                "results": [],
                "total_results": 0,
                "search_strategy": "hr_scorecard_comprehensive",
                "hr_scorecards": [],
            }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error getting task results: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to get task results")

# NOTE: /api/vector-status is defined earlier (~line 6201). Duplicate removed.

def extract_standardized_keywords_from_jd(query: str, token_tracker: TokenTracker = None) -> Dict[str, Any]:
    """Extract a standardized set of keywords from job description for consistent analysis"""
    print(f"🔍 Starting standardized keyword extraction...")
    print(f"📝 Query: '{query[:200]}...'")

    try:
        # First, try to get PROJECT_ID from config
        PROJECT_ID = None
        try:
            config = load_config()
            if config and 'vector_search' in config:
                PROJECT_ID = config['vector_search']['project_id']
                print(f"✅ Found PROJECT_ID in config: {PROJECT_ID}")
            else:
                print(f"❌ PROJECT_ID not found in config")
        except Exception as config_error:
            print(f"❌ Error loading config: {config_error}")

        if not PROJECT_ID:
            raise Exception("❌ PROJECT_ID not found in config - Cannot extract keywords without proper LLM access")

        client = gemini_client()
        print(f"✅ LLM client initialized successfully")

        try:
            print(f"🔍 Formatting prompt with query: '{query[:100]}...'")
            keyword_extraction_prompt = KEYWORD_EXTRACTION_PROMPT.format(
                job_posting=query
            )
            print(f"✅ Prompt formatted successfully: {len(keyword_extraction_prompt)} chars")
        except Exception as prompt_error:
            print(f"❌ Error formatting prompt: {prompt_error}")
            print(f"❌ Error type: {type(prompt_error)}")
            raise Exception(f"Failed to format keyword extraction prompt: {str(prompt_error)}")

        try:
            print(f"🔍 Creating content structure...")
            contents = [
                types.Content(
                    role="user",
                    parts=[types.Part(text=keyword_extraction_prompt)]
                )
            ]
            print(f"✅ Content structure created successfully")
        except Exception as content_error:
            print(f"❌ Error creating content structure: {content_error}")
            raise Exception(f"Failed to create content structure: {str(content_error)}")

        try:
            print(f"🔍 Creating generation config...")
            generate_content_config = types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
                temperature=0.1,  # Very low temperature for consistency
                top_p=0.8,
                max_output_tokens=4096,  # Maximum allowed for Gemini 2.5 Flash
                safety_settings=[
                    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="OFF"),
                    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="OFF"),
                    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="OFF"),
                    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="OFF")
                ],
            )
            print(f"✅ Generation config created successfully")
        except Exception as config_error:
            print(f"❌ Error creating generation config: {config_error}")
            raise Exception(f"Failed to create generation config: {str(config_error)}")

        try:
            print(f"🔍 Calling Gemini API...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=contents,
                config=generate_content_config,
            )
            print(f"✅ API call completed successfully")

            # Track token usage
            if token_tracker:
                token_usage = extract_token_usage(response, "jd_keyword_extraction", "gemini-2.5-flash")
                token_tracker.add_call(
                    operation="jd_keyword_extraction",
                    model="gemini-2.5-flash",
                    input_tokens=token_usage["input_tokens"],
                    output_tokens=token_usage["output_tokens"],
                    success=True
                )
                print(f"📊 Token usage - Input: {token_usage['input_tokens']}, Output: {token_usage['output_tokens']}, Total: {token_usage['total_tokens']}")
        except Exception as api_error:
            print(f"❌ Error calling Gemini API: {api_error}")
            print(f"❌ API Error type: {type(api_error)}")
            raise Exception(f"Failed to call Gemini API: {str(api_error)}")

        # Extract response text with safer debugging
        analysis_text = ""
        try:
            print(f"🔍 Response type: {type(response).__name__}")

            # Try direct text access first
            try:
                if hasattr(response, 'text'):
                    print(f"✅ Response has 'text' attribute")
                    if response.text is not None:
                        analysis_text = str(response.text)
                        print(f"✅ Extracted text from response.text: {len(analysis_text)} chars")
                    else:
                        print(f"⚠️ response.text is None")
                else:
                    print(f"⚠️ Response has no 'text' attribute")
            except Exception as text_error:
                print(f"❌ Error accessing response.text: {text_error}")

            # If no text found, try candidates approach
            if not analysis_text:
                try:
                    if hasattr(response, 'candidates'):
                        print(f"✅ Response has 'candidates' attribute")
                        if response.candidates:
                            print(f"✅ Found {len(response.candidates)} candidates")
                            for i, candidate in enumerate(response.candidates):
                                if candidate and hasattr(candidate, 'content'):
                                    if hasattr(candidate.content, 'parts'):
                                        for part in candidate.content.parts:
                                            if hasattr(part, 'text') and part.text:
                                                analysis_text += str(part.text)
                                                print(f"✅ Added text from candidate {i}")
                        else:
                            print(f"⚠️ candidates list is empty")
                    else:
                        print(f"⚠️ Response has no 'candidates' attribute")
                except Exception as candidates_error:
                    print(f"❌ Error accessing candidates: {candidates_error}")

            # Final check
            if not analysis_text:
                print(f"❌ No text extracted from response")
                raise Exception("Could not extract any text from LLM response")

        except Exception as extract_error:
            print(f"❌ Critical error in response extraction: {extract_error}")
            raise Exception(f"Response extraction failed: {str(extract_error)}")

        if not analysis_text:
            print(f"❌ Empty keyword extraction response from Gemini API")
            raise Exception("LLM returned empty response - Cannot extract keywords")

        # Parse JSON response with improved error handling
        print(f"🔍 Raw LLM response length: {len(analysis_text)} characters")
        print(f"📝 Response length: {len(analysis_text or '')} chars (preview redacted to avoid leaking PII)")

        try:
            clean_text = analysis_text.strip()

            # Remove code block markers more thoroughly
            if clean_text.startswith('```json'):
                clean_text = clean_text[7:].strip()
            elif clean_text.startswith('```'):
                clean_text = clean_text[3:].strip()

            if clean_text.endswith('```'):
                clean_text = clean_text[:-3].strip()

            # Remove any leading/trailing whitespace and newlines
            clean_text = clean_text.strip()

            print(f"🧹 Cleaned text for JSON parsing: '{clean_text[:200]}...'")

            # Check if we have any content to parse
            if not clean_text or len(clean_text) < 10:
                print(f"❌ Cleaned text is too short or empty: '{clean_text}'")
                raise Exception(f"LLM response too short or empty after cleaning: '{clean_text}'")

            keywords_data = json.loads(clean_text)

            # Validate the response structure
            if not isinstance(keywords_data, dict) or 'keywords' not in keywords_data:
                print(f"❌ Invalid JSON structure: {keywords_data}")
                raise Exception(f"LLM returned invalid JSON structure: {keywords_data}")

            # Flatten all keywords into a single list for matching with advanced deduplication
            all_keywords = []
            keywords_by_category = keywords_data.get('keywords', {})

            # Define unwanted words to filter out (expanded list)
            unwanted_words = {
                'about', 'join', 'team', 'role', 'position', 'company', 'work', 'will', 'we', 'our', 'you', 'your',
                'the', 'and', 'for', 'with', 'this', 'that', 'looking', 'seeking', 'hiring', 'candidate', 'ideal',
                'perfect', 'great', 'excellent', 'strong', 'good', 'best', 'top', 'high', 'low', 'new', 'old',
                'big', 'small', 'business', 'growth', 'success', 'opportunity', 'career', 'benefits', 'job',
                'description', 'requirements', 'qualifications', 'responsibilities', 'duties', 'tasks', 'able',
                'ability', 'skills', 'knowledge', 'understanding', 'working', 'experience', 'years', 'level',
                'provide', 'support', 'help', 'assist', 'develop', 'create', 'build', 'design', 'implement',
                'manage', 'lead', 'coordinate', 'collaborate', 'ensure', 'maintain', 'improve', 'enhance',
                # Additional exclusions for better quality
                'senior', 'junior', 'experienced', 'entry', 'mid', 'scalable', 'applications', 'building',
                'development', 'programming', 'software', 'technical', 'proficient', 'expertise', 'knowledge'
            }

            # Global deduplication set (case-insensitive)
            global_keywords_seen = set()

            for category, keyword_list in keywords_by_category.items():
                if isinstance(keyword_list, list):
                    # Clean, validate, filter, and deduplicate each keyword
                    clean_keywords = []
                    for kw in keyword_list:
                        if kw and str(kw).strip():
                            keyword_clean = str(kw).strip()
                            keyword_lower = keyword_clean.lower()

                            # Advanced filtering
                            if (keyword_lower not in unwanted_words and
                                len(keyword_clean) > 2 and
                                keyword_lower not in global_keywords_seen and
                                not keyword_lower.startswith('web ') and  # Avoid "web applications"
                                not keyword_lower.endswith(' experience') and  # Avoid "5+ years experience"
                                not keyword_lower.endswith(' level') and  # Avoid "senior level"
                                not ('scalable' in keyword_lower and len(keyword_lower.split()) > 1)):  # Avoid "scalable systems"

                                clean_keywords.append(keyword_clean)
                                global_keywords_seen.add(keyword_lower)

                    keywords_by_category[category] = clean_keywords
                    all_keywords.extend(clean_keywords)
                    print(f"   📂 {category}: {len(clean_keywords)} keywords")

            print(f"✅ Successfully extracted {len(all_keywords)} standardized keywords")
            print(f"🔑 Sample keywords: {all_keywords[:8]}{'...' if len(all_keywords) > 8 else ''}")

            if len(all_keywords) == 0:
                print(f"⚠️ LLM returned valid JSON but no keywords found")
                raise Exception("LLM returned valid JSON but extracted zero keywords from job description")

            return {
                "success": True,
                "keywords_by_category": keywords_by_category,
                "all_keywords": all_keywords,
                "total_keywords": len(all_keywords),
                "fallback_used": False
            }

        except json.JSONDecodeError as e:
            print(f"❌ Failed to parse keyword JSON: {str(e)}")
            print(f"📝 FULL LLM RESPONSE: '{clean_text}'")
            print(f"📝 Error at position {e.pos}: '{clean_text[max(0, e.pos-20):e.pos+20]}'")
            raise Exception(f"LLM returned invalid JSON format: {str(e)}. Check the response format in logs.")
        except Exception as e:
            print(f"❌ Unexpected error in JSON processing: {str(e)}")
            raise Exception(f"Unexpected error in keyword extraction: {str(e)}")

    except Exception as e:
        print(f"❌ Error in standardized keyword extraction: {str(e)}")
        raise Exception(f"Keyword extraction failed: {str(e)}")

def create_fallback_jd_keywords(query: str) -> Dict[str, Any]:
    """Create fallback keywords using improved text processing"""
    import re

    print(f"🔄 Using fallback keyword extraction")
    print(f"📝 Query preview: '{query[:200]}...'")

    # First, try to extract meaningful phrases (2-4 words) from the query
    query_lower = query.lower()

    # Comprehensive technical keywords to look for (expanded)
    tech_patterns = {
        'programming_languages': [
            'python', 'java', 'javascript', 'typescript', 'sql', r'\br\b', 'scala', 'golang', 'go',
            'rust', r'c\+\+', 'c#', 'php', 'ruby', 'swift', 'kotlin', 'perl', 'matlab', 'julia',
            'visual basic', 'objective-c', 'dart', 'elixir', 'clojure', 'haskell'
        ],
        'cloud_platforms': [
            'aws', 'amazon web services', 'azure', 'microsoft azure', 'gcp', 'google cloud platform',
            'google cloud', 'cloud computing', 'serverless', 'openstack', 'alibaba cloud', 'ibm cloud',
            'oracle cloud', 'digital ocean', 'heroku', 'netlify', 'vercel'
        ],
        'tools_technologies': [
            'docker', 'kubernetes', 'git', 'github', 'gitlab', 'jenkins', 'terraform', 'ansible',
            'tableau', 'power bi', 'excel', 'jira', 'confluence', 'slack', 'grafana', 'prometheus',
            'elasticsearch', 'kibana', 'splunk', 'datadog', 'new relic', 'postman', 'swagger'
        ],
        'frameworks_libraries': [
            'react', 'angular', 'vue', 'vuejs', 'django', 'flask', 'spring', 'spring boot',
            'tensorflow', 'pytorch', 'scikit-learn', 'sklearn', 'pandas', 'numpy', 'keras',
            'spark', 'apache spark', 'hadoop', 'kafka', 'rabbitmq', 'express', 'nodejs', 'node.js',
            'fastapi', '.net', 'rails', 'laravel', 'symfony'
        ],
        'databases': [
            'mysql', 'postgresql', 'postgres', 'mongodb', 'redis', 'elasticsearch', 'cassandra',
            'oracle', 'sql server', 'mssql', 'snowflake', 'bigquery', 'redshift', 'dynamodb',
            'couchdb', 'neo4j', 'influxdb', 'mariadb', 'sqlite', 'firestore', 'cosmos db'
        ],
        'technical_skills': [
            'machine learning', 'ml', 'data science', 'artificial intelligence', 'ai', 'deep learning',
            'neural networks', 'nlp', 'natural language processing', 'computer vision', 'cv',
            'data analysis', 'data analytics', 'statistical modeling', 'predictive analytics',
            'big data', 'data engineering', 'data mining', 'business intelligence', 'bi',
            'web development', 'mobile development', 'full stack', 'backend', 'frontend',
            'devops', 'devsecops', 'cloud architecture', 'microservices', 'api development',
            'rest api', 'graphql', 'software architecture', 'system design'
        ],
        'methodologies': [
            'agile', 'scrum', 'kanban', 'waterfall', 'devops', 'ci/cd', 'continuous integration',
            'continuous deployment', 'mlops', 'dataops', 'gitops', 'data pipeline', 'etl', 'elt',
            'microservices architecture', 'service oriented architecture', 'soa', 'test driven development',
            'tdd', 'behavior driven development', 'bdd', 'pair programming', 'code review'
        ],
        'soft_skills': [
            'leadership', 'communication', 'teamwork', 'problem solving', 'analytical thinking',
            'collaboration', 'project management', 'critical thinking', 'time management',
            'presentation skills', 'mentoring', 'coaching', 'strategic thinking', 'innovation'
        ],
        'experience_levels': [
            'senior', 'lead', 'principal', 'staff', 'manager', 'director', 'architect',
            r'\d+\s*\+?\s*years?', 'entry level', 'junior', 'mid level', 'mid-level', 'experienced'
        ],
        'education': [
            'bachelor', 'master', 'phd', 'doctorate', 'degree', 'computer science', 'cs',
            'data science', 'statistics', 'mathematics', 'engineering', 'information technology',
            'it', 'software engineering', 'electrical engineering', 'physics', 'mba'
        ],
        'certifications': [
            'aws certified', 'azure certified', 'google cloud certified', 'gcp certified',
            'pmp', 'scrum master', 'certified', 'certification', 'cissp', 'ceh', 'comptia',
            'oracle certified', 'microsoft certified', 'cisco', 'ccna', 'ccnp'
        ],
        'industry_terms': [
            'kpi', 'roi', 'analytics', 'insights', 'reporting', 'dashboards', 'metrics',
            'performance', 'optimization', 'automation', 'scalability', 'reliability',
            'high availability', 'disaster recovery', 'sla', 'api', 'sdk', 'saas', 'paas', 'iaas'
        ]
    }

    # Extract matching keywords with regex for better matching
    extracted_keywords = {}
    all_keywords = []
    matched_patterns = set()  # To avoid duplicates

    for category, patterns in tech_patterns.items():
        found_keywords = []
        for pattern in patterns:
            # Use regex for more accurate matching
            try:
                # For simple patterns, add word boundaries
                if '\\' not in pattern and not any(c in pattern for c in ['+', '*', '?', '[', ']', '(', ')']):
                    regex_pattern = r'\b' + pattern + r'\b'
                else:
                    regex_pattern = pattern

                matches = re.findall(regex_pattern, query_lower)
                if matches:
                    # Get the actual matched text
                    for match in matches:
                        clean_match = match.strip()
                        if clean_match and clean_match not in matched_patterns:
                            # Special handling for certain keywords
                            if clean_match == 'r':
                                display_name = 'R'
                            elif clean_match in ['ml', 'ai', 'bi', 'ci/cd', 'api', 'sdk', 'sql', 'nlp', 'cv']:
                                display_name = clean_match.upper()
                            elif '+' in clean_match or '#' in clean_match or '.' in clean_match:
                                display_name = clean_match  # Keep special chars
                            else:
                                display_name = clean_match.title()

                            found_keywords.append(display_name)
                            all_keywords.append(display_name)
                            matched_patterns.add(clean_match)
            except re.error:
                # If regex fails, fall back to simple substring matching
                if pattern in query_lower:
                    display_name = pattern.title()
                    if pattern not in matched_patterns:
                        found_keywords.append(display_name)
                        all_keywords.append(display_name)
                        matched_patterns.add(pattern)

        if found_keywords:
            extracted_keywords[category] = found_keywords

    # More comprehensive stopwords list
    stopwords = {
        'the', 'and', 'for', 'with', 'this', 'that', 'will', 'have', 'from', 'they', 'been', 'were', 'said',
        'each', 'which', 'their', 'time', 'would', 'there', 'could', 'other', 'looking', 'ideal', 'candidate',
        'experience', 'years', 'work', 'working', 'team', 'strong', 'good', 'excellent', 'ability', 'skills',
        'knowledge', 'about', 'join', 'role', 'position', 'company', 'our', 'you', 'your', 'seeking', 'hiring',
        'perfect', 'great', 'best', 'top', 'high', 'low', 'new', 'old', 'big', 'small', 'business', 'growth',
        'success', 'opportunity', 'career', 'benefits', 'job', 'description', 'requirements', 'qualifications',
        'responsibilities', 'duties', 'tasks', 'able', 'understanding', 'provide', 'support', 'help', 'assist',
        'develop', 'create', 'build', 'design', 'implement', 'manage', 'lead', 'coordinate', 'collaborate',
        'ensure', 'maintain', 'improve', 'enhance', 'are', 'who', 'what', 'when', 'where', 'why', 'how',
        'need', 'search', 'store', 'code', 'connect', 'switch', 'respect', 'login', 'data', 'use', 'used',
        'using', 'make', 'made', 'making', 'take', 'took', 'taken', 'give', 'gave', 'given', 'find', 'found',
        'tell', 'told', 'ask', 'asked', 'work', 'worked', 'want', 'wanted', 'look', 'looked', 'become',
        'became', 'leave', 'left', 'feel', 'felt', 'bring', 'brought', 'begin', 'began', 'keep', 'kept'
    }

    # Extract meaningful words from query (4+ characters, not in stopwords)
    words = re.findall(r'\b[a-zA-Z]{4,}\b', query)
    important_words = []

    for word in words:
        word_lower = word.lower()
        if (word_lower not in stopwords and
            word_lower not in matched_patterns and
            len(word) >= 4):  # Minimum 4 characters for meaningful words
            important_words.append(word.title())

    # Only add the most relevant important words (limit to 10)
    for word in important_words[:10]:
        if word not in all_keywords:
            if 'technical_skills' not in extracted_keywords:
                extracted_keywords['technical_skills'] = []
            extracted_keywords['technical_skills'].append(word)
            all_keywords.append(word)

    # Ensure we have at least some meaningful keywords
    if len(all_keywords) < 5:
        print(f"⚠️ Too few keywords found ({len(all_keywords)}), enhancing with basic tech terms")
        # Add some common technical terms that might be relevant
        basic_tech_keywords = [
            'Python', 'Java', 'JavaScript', 'SQL', 'Data Analysis', 'Software Development',
            'Programming', 'Database', 'Web Development', 'API', 'Cloud', 'DevOps'
        ]
        for kw in basic_tech_keywords:
            if kw not in all_keywords:
                all_keywords.append(kw)
                if 'technical_skills' not in extracted_keywords:
                    extracted_keywords['technical_skills'] = []
                extracted_keywords['technical_skills'].append(kw)
                if len(all_keywords) >= 10:
                    break

    print(f"✅ Fallback extracted {len(all_keywords)} keywords: {all_keywords[:15]}{'...' if len(all_keywords) > 15 else ''}")

    # Debug: Show which categories had matches
    for category, keywords in extracted_keywords.items():
        if keywords:
            print(f"   📂 {category}: {len(keywords)} keywords found")

    return {
        "success": True,
        "keywords_by_category": extracted_keywords,
        "all_keywords": all_keywords,
        "total_keywords": len(all_keywords),
        "fallback_used": True
    }

_KEYWORD_BOUNDARY_LEFT = r"(?<![A-Za-z0-9])"
_KEYWORD_BOUNDARY_RIGHT = r"(?![A-Za-z0-9])"
_KEYWORD_MATCH_CACHE: "dict[str, re.Pattern[str]]" = {}


def _compile_keyword_alt(variations: List[str]) -> Optional["re.Pattern[str]"]:
    """Compile an alternation regex that matches any variation with strict
    non-word boundaries on both sides (handles `.net`, `c#`, `node.js`).
    Returns None if no usable variations."""
    cleaned: List[str] = []
    seen: set = set()
    for v in variations:
        if not v:
            continue
        v = v.strip().lower()
        if not v or v in seen:
            continue
        seen.add(v)
        # Reject single-character variations except for known unambiguous tokens.
        # 1-char synonyms like "r" or "c" would false-match everywhere.
        if len(v) < 2:
            continue
        cleaned.append(re.escape(v))
    if not cleaned:
        return None
    # Longest first so the alternation prefers the most specific variant.
    cleaned.sort(key=len, reverse=True)
    cache_key = "|".join(cleaned)
    cached = _KEYWORD_MATCH_CACHE.get(cache_key)
    if cached is not None:
        return cached
    pattern = re.compile(
        _KEYWORD_BOUNDARY_LEFT + r"(?:" + "|".join(cleaned) + r")" + _KEYWORD_BOUNDARY_RIGHT,
        re.IGNORECASE,
    )
    _KEYWORD_MATCH_CACHE[cache_key] = pattern
    return pattern


def enhanced_keyword_matching(resume_content: str, standardized_keywords: List[str]) -> Dict[str, Any]:
    """Word-boundary keyword matching with synonym expansion.

    Uses strict non-word boundaries (`(?<![A-Za-z0-9])` / `(?![A-Za-z0-9])`) so
    short tokens like `ai`, `ml`, `js`, `ts`, `go`, `r` cannot false-match
    inside `available`, `HTML`, `claims`, etc. Handles punctuated tokens
    (`.net`, `c#`, `node.js`) correctly because the boundaries reject only
    adjacent alphanumerics.
    """
    if not resume_content or not standardized_keywords:
        return {
            "matched_keywords": [],
            "missing_keywords": list(standardized_keywords or []),
            "jd_keywords_matched": 0,
            "total_jd_keywords": len(standardized_keywords or []),
        }
    matched_keywords: List[str] = []
    missing_keywords: List[str] = []

    # Comprehensive synonym mapping
    synonym_map = {
        'javascript': ['js', 'java script', 'node.js', 'nodejs', 'node js', 'ecmascript'],
        'typescript': ['ts', 'type script'],
        'python': ['py', 'python3', 'python 3', 'python2', 'python 2'],
        'machine learning': ['ml', 'machinelearning', 'machine-learning', 'ai/ml', 'artificial intelligence'],
        'artificial intelligence': ['ai', 'artificialintelligence', 'machine learning', 'ml'],
        'react': ['reactjs', 'react.js', 'react js', 'react native'],
        'angular': ['angularjs', 'angular.js', 'angular js'],
        'vue': ['vuejs', 'vue.js', 'vue js'],
        'docker': ['containerization', 'containers', 'docker compose'],
        'kubernetes': ['k8s', 'k8', 'container orchestration'],
        'amazon web services': ['aws', 'amazon aws', 'amazon cloud'],
        'google cloud platform': ['gcp', 'google cloud', 'google cloud services'],
        'microsoft azure': ['azure', 'ms azure', 'azure cloud'],
        'postgresql': ['postgres', 'postgre sql', 'postgresql database'],
        'mongodb': ['mongo db', 'mongo', 'mongodb database'],
        'mysql': ['mysql database', 'my sql'],
        'rest': ['rest api', 'restful', 'rest apis', 'restful api'],
        'api': ['apis', 'rest api', 'api development'],
        'agile': ['scrum', 'agile methodology', 'agile development'],
        'scrum': ['agile', 'scrum master', 'scrum methodology'],
        'ci/cd': ['continuous integration', 'continuous deployment', 'cicd'],
        'git': ['github', 'gitlab', 'version control', 'git version control'],
        'github': ['git', 'version control'],
        'sql': ['sql server', 'mysql', 'postgresql', 'database', 'structured query language'],
        'nosql': ['mongodb', 'no sql', 'non-relational database'],
        'redis': ['redis cache', 'redis database'],
        'elasticsearch': ['elastic search', 'elk stack'],
        'kafka': ['apache kafka', 'message queue'],
        'microservices': ['micro services', 'service-oriented architecture', 'soa'],
        'devops': ['dev ops', 'development operations'],
        'terraform': ['infrastructure as code', 'iac'],
        'jenkins': ['ci/cd', 'continuous integration'],
        'ansible': ['configuration management'],
        'css': ['css3', 'cascading style sheets', 'stylesheet'],
        'html': ['html5', 'hypertext markup language'],
        'sass': ['scss', 'css preprocessor'],
        'webpack': ['bundler', 'module bundler'],
        'npm': ['node package manager'],
        'yarn': ['package manager'],
        'express': ['express.js', 'expressjs'],
        'spring': ['spring boot', 'spring framework'],
        'django': ['django framework'],
        'flask': ['flask framework'],
        'laravel': ['laravel framework'],
        'tensorflow': ['tf', 'tensor flow'],
        'pytorch': ['torch', 'py-torch'],
        'scikit-learn': ['sklearn', 'scikit learn', 'sci-kit learn'],
        'pandas': ['data analysis', 'python pandas'],
        'numpy': ['numerical python'],
        'jupyter': ['jupyter notebook', 'ipython'],
        'tableau': ['data visualization'],
        'power bi': ['powerbi', 'microsoft power bi'],
        'excel': ['microsoft excel', 'spreadsheet'],
        'powerpoint': ['microsoft powerpoint', 'ppt'],
        'photoshop': ['adobe photoshop', 'ps'],
        'illustrator': ['adobe illustrator', 'ai'],
        'figma': ['design tool'],
        'sketch': ['design tool'],
        'unity': ['unity3d', 'unity engine'],
        'unreal': ['unreal engine'],
        'blockchain': ['cryptocurrency', 'bitcoin', 'ethereum'],
        'iot': ['internet of things'],
        'ar': ['augmented reality'],
        'vr': ['virtual reality'],
        'ui': ['user interface'],
        'ux': ['user experience'],
        'seo': ['search engine optimization'],
        'crm': ['customer relationship management'],
        'erp': ['enterprise resource planning']
    }

    # Technical context matching - look for related terms that indicate skill presence
    context_indicators = {
        'python': ['developed in python', 'python development', 'python programming', 'python script', 'python application'],
        'javascript': ['javascript development', 'js development', 'frontend development', 'web development'],
        'react': ['react development', 'react component', 'react application'],
        'sql': ['database queries', 'database development', 'wrote sql', 'sql queries'],
        'aws': ['cloud development', 'deployed on aws', 'aws services', 'amazon cloud'],
        'docker': ['containerized applications', 'container deployment', 'deployed containers'],
        'kubernetes': ['orchestrated containers', 'k8s deployment', 'cluster management'],
        'machine learning': ['ml models', 'predictive models', 'data science', 'statistical analysis'],
        'api': ['api development', 'api integration', 'rest services', 'web services']
    }

    for keyword in standardized_keywords:
        keyword_lower = keyword.lower().strip()
        if not keyword_lower:
            missing_keywords.append(keyword)
            continue

        variations = [keyword_lower]

        if keyword_lower in synonym_map:
            variations.extend(synonym_map[keyword_lower])

        if keyword_lower in context_indicators:
            variations.extend(context_indicators[keyword_lower])

        # Punctuation/whitespace variations — only safe forms.
        compact = keyword_lower.replace(' ', '')
        if compact != keyword_lower:
            variations.append(compact)
        if '-' in keyword_lower:
            variations.append(keyword_lower.replace('-', ' '))
            variations.append(keyword_lower.replace('-', ''))
        if '_' in keyword_lower:
            variations.append(keyword_lower.replace('_', ' '))
        if ' ' in keyword_lower:
            variations.append(keyword_lower.replace(' ', '-'))
            variations.append(keyword_lower.replace(' ', '_'))

        # Domain-specific aliases for common punctuated tech tokens.
        if keyword_lower == '.net':
            variations.extend(['dotnet', 'dot net', '.net framework', 'asp.net', 'vb.net'])
        elif keyword_lower == 'c#':
            variations.extend(['csharp', 'c sharp'])
        elif keyword_lower == 'c++':
            variations.extend(['cpp', 'cplusplus'])

        pattern = _compile_keyword_alt(variations)
        if pattern is not None and pattern.search(resume_content):
            matched_keywords.append(keyword)
        else:
            missing_keywords.append(keyword)

    return {
        "matched_keywords": matched_keywords,
        "missing_keywords": missing_keywords,
        "jd_keywords_matched": len(matched_keywords),
        "total_jd_keywords": len(standardized_keywords),
    }


def analyze_candidate_keywords(resume_content: str, standardized_keywords: List[str], token_tracker: TokenTracker = None) -> Dict[str, Any]:
    """Analyze which standardized keywords are matched/missing in candidate resume using LLM-first intelligent approach"""
    print(f"🔍 Starting LLM-first keyword analysis for {len(standardized_keywords)} keywords")

    # If no keywords or resume content, return empty results
    if not standardized_keywords or not resume_content:
        return {
            "matched_keywords": [],
            "missing_keywords": standardized_keywords or [],
            "jd_keywords_matched": 0,
            "total_jd_keywords": len(standardized_keywords) if standardized_keywords else 0
        }

    # 🚀 PRIORITIZE LLM ANALYSIS FOR ACCURACY
    print(f"🤖 Using LLM-based keyword matching for highest accuracy")

    try:
        # Use LLM for intelligent keyword matching
        print(f"🔧 Initializing Gemini client for keyword matching...")
        print(f"   - Project: {PROJECT_ID}")
        print(f"   - Location: europe-west4")

        client = gemini_client()
        print(f"✅ Gemini client initialized successfully")

        keyword_analysis_prompt = KEYWORD_MATCHING_PROMPT.format(
            resume_content=resume_content,
            keywords_list=', '.join(standardized_keywords)
        )

        print(f"🔧 Preparing LLM request for keyword matching...")
        print(f"   - Prompt length: {len(keyword_analysis_prompt)} chars")
        print(f"   - Keywords to analyze: {len(standardized_keywords)}")

        contents = [
            types.Content(
                role="user",
                parts=[types.Part(text=keyword_analysis_prompt)]
            )
        ]

        print(f"🚀 Sending request to Gemini...")
        # Determine appropriate max_output_tokens based on number of keywords
        # More keywords = potentially longer JSON response
        base_output_tokens = 1024
        keyword_factor = min(len(standardized_keywords) * 5, 1024)  # ~5 tokens per keyword, max 1024 extra
        max_output_tokens = base_output_tokens + keyword_factor
        max_output_tokens = min(max_output_tokens, 3072)  # Cap at reasonable limit

        print(f"🔧 Using max_output_tokens: {max_output_tokens} (base: {base_output_tokens}, keywords: {keyword_factor})")

        generate_response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=contents,
            config=types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
                temperature=0.1,
                max_output_tokens=4096,
                safety_settings=[
                    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="OFF"),
                    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="OFF"),
                    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="OFF"),
                    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="OFF")
                ],
            ),
        )
        print(f"✅ Received response from Gemini")

        # Extract response text first with enhanced debugging
        analysis_text = ""
        print(f"🔍 Extracting response text...")

        if hasattr(generate_response, 'text') and generate_response.text:
            analysis_text = generate_response.text
            print(f"✅ Got text directly: {len(analysis_text)} chars")
        elif hasattr(generate_response, 'candidates') and generate_response.candidates:
            print(f"🔍 Processing {len(generate_response.candidates)} candidates...")
            try:
                for i, candidate in enumerate(generate_response.candidates):
                    print(f"   Candidate {i}: {type(candidate)}")
                    print(f"   - Candidate attributes: {[attr for attr in dir(candidate) if not attr.startswith('_')]}")

                    # Check finish reason if available
                    if hasattr(candidate, 'finish_reason'):
                        finish_reason = candidate.finish_reason
                        print(f"   - Finish reason: {finish_reason}")
                        if finish_reason == "MAX_TOKENS":
                            print("   ⚠️ Response cut off due to token limit!")
                            print("   📊 Consider reducing input size or increasing max_output_tokens")
                        elif finish_reason in ["SAFETY", "RECITATION"]:
                            print("   ⚠️ Response blocked by safety filters!")
                        elif finish_reason == "STOP":
                            print("   ✅ Response completed normally")
                        elif finish_reason:
                            print(f"   ⚠️ Unexpected finish reason: {finish_reason}")

                    # Try to get text directly from candidate first
                    if hasattr(candidate, 'text') and candidate.text:
                        candidate_text = candidate.text
                        analysis_text += candidate_text
                        print(f"   - Got text directly from candidate: {len(candidate_text)} chars")
                    elif candidate and hasattr(candidate, 'content') and candidate.content:
                        print(f"   - Has content: {type(candidate.content)}")
                        print(f"   - Content attributes: {[attr for attr in dir(candidate.content) if not attr.startswith('_')]}")

                        # Try to get text directly from content
                        if hasattr(candidate.content, 'text') and candidate.content.text:
                            content_text = candidate.content.text
                            analysis_text += content_text
                            print(f"   - Got text directly from content: {len(content_text)} chars")
                        elif hasattr(candidate.content, 'parts') and candidate.content.parts:
                            print(f"   - Has {len(candidate.content.parts)} parts")
                            for j, part in enumerate(candidate.content.parts):
                                print(f"     Part {j}: {type(part)}")
                                if part and hasattr(part, 'text') and part.text:
                                    part_text = part.text
                                    analysis_text += part_text
                                    print(f"     - Added text: {len(part_text)} chars")
                                else:
                                    print(f"     - No text in part {j}")
                        else:
                            print(f"   - No parts in content or parts is empty, checking deeper...")
                            # Check for other potential text attributes
                            if hasattr(candidate.content, 'role'):
                                print(f"   - Content role: {getattr(candidate.content, 'role', 'unknown')}")

                            # Try different ways to access parts
                            parts_found = False

                            # Method 1: Direct access
                            if hasattr(candidate.content, 'parts'):
                                parts = candidate.content.parts
                                print(f"   - Parts (direct): {type(parts)}")
                                if parts is not None:
                                    try:
                                        parts_list = list(parts) if hasattr(parts, '__iter__') else [parts]
                                        print(f"   - Parts count: {len(parts_list)}")
                                        for k, p in enumerate(parts_list):
                                            print(f"     Part {k}: {type(p)}")
                                            if hasattr(p, 'text') and p.text:
                                                part_text = p.text
                                                analysis_text += part_text
                                                print(f"     - Added text: {len(part_text)} chars")
                                                parts_found = True
                                    except Exception as parts_e:
                                        print(f"     - Error with direct parts access: {parts_e}")

                            # Method 2: Try accessing through dict/model representation
                            if not parts_found:
                                try:
                                    content_dict = candidate.content.model_dump() if hasattr(candidate.content, 'model_dump') else {}
                                    if 'parts' in content_dict:
                                        parts_data = content_dict['parts']
                                        print(f"   - Parts from model_dump: {type(parts_data)}, length: {len(parts_data) if hasattr(parts_data, '__len__') else 'unknown'}")
                                        if isinstance(parts_data, list):
                                            for k, part_data in enumerate(parts_data):
                                                if isinstance(part_data, dict) and 'text' in part_data:
                                                    part_text = part_data['text']
                                                    if part_text:
                                                        analysis_text += part_text
                                                        print(f"     - Added text from dict part {k}: {len(part_text)} chars")
                                                        parts_found = True
                                except Exception as dict_e:
                                    print(f"     - Error with dict access: {dict_e}")

                            # Method 3: Try JSON representation
                            if not parts_found:
                                try:
                                    import json  # Ensure json is available
                                    content_json = candidate.content.model_dump_json() if hasattr(candidate.content, 'model_dump_json') else '{}'
                                    content_data = json.loads(content_json)
                                    if 'parts' in content_data:
                                        parts_data = content_data['parts']
                                        print(f"   - Parts from JSON: {len(parts_data)} items")
                                        for k, part_data in enumerate(parts_data):
                                            if 'text' in part_data and part_data['text']:
                                                part_text = part_data['text']
                                                analysis_text += part_text
                                                print(f"     - Added text from JSON part {k}: {len(part_text)} chars")
                                                parts_found = True
                                except Exception as json_e:
                                    print(f"     - Error with JSON access: {json_e}")

                            if not parts_found:
                                print(f"   - Could not extract text from any parts method")

                            # Try other common attributes
                            for attr in ['message', 'content_text', 'output', 'result']:
                                if hasattr(candidate.content, attr):
                                    attr_value = getattr(candidate.content, attr)
                                    if attr_value and isinstance(attr_value, str):
                                        analysis_text += attr_value
                                        print(f"   - Got text from {attr}: {len(attr_value)} chars")
                                        break
                    else:
                        print(f"   - No content in candidate {i}")
            except Exception as e:
                print(f"❌ Error extracting text from candidates: {e}")
                print(f"🔍 Raw response: {str(generate_response)[:200]}...")
        else:
            print(f"❌ No text or candidates found in response")
            print(f"🔍 Response attributes: {[attr for attr in dir(generate_response) if not attr.startswith('_')]}")

            # Try alternative extraction methods
            try:
                if hasattr(generate_response, 'result') and generate_response.result:
                    analysis_text = str(generate_response.result)
                    print(f"✅ Got text from result: {len(analysis_text)} chars")
                elif hasattr(generate_response, 'content') and generate_response.content:
                    analysis_text = str(generate_response.content)
                    print(f"✅ Got text from content: {len(analysis_text)} chars")
                elif hasattr(generate_response, 'parts') and generate_response.parts:
                    for part in generate_response.parts:
                        if hasattr(part, 'text'):
                            analysis_text += part.text
                    print(f"✅ Got text from parts: {len(analysis_text)} chars")
            except Exception as alt_e:
                print(f"❌ Alternative extraction also failed: {alt_e}")

        print(f"📄 Final extracted text length: {len(analysis_text)} chars")
        if analysis_text:
            print(f"📄 Text length: {len(analysis_text or '')} chars (preview redacted to avoid leaking PII)")

        # Track token usage for keyword matching LLM call (after analysis_text is defined)
        if token_tracker:
            try:
                token_usage = extract_token_usage(generate_response, "keyword_matching", "gemini-2.5-flash")
                token_tracker.add_call(
                    operation="keyword_matching",
                    model="gemini-2.5-flash",
                    input_tokens=token_usage["input_tokens"],
                    output_tokens=token_usage["output_tokens"],
                    success=bool(analysis_text)  # Now analysis_text is properly defined
                )
                print(f"🔢 LLM Keyword Matching Tokens - Input: {token_usage['input_tokens']}, Output: {token_usage['output_tokens']}, Total: {token_usage['total_tokens']}")
            except Exception as token_error:
                print(f"⚠️ Token tracking error: {token_error}")
                # Still try to track the call even if token extraction fails
                token_tracker.add_call(
                    operation="keyword_matching",
                    model="gemini-2.5-flash",
                    input_tokens=0,
                    output_tokens=0,
                    success=bool(analysis_text)
                )

        if not analysis_text:
            print(f"⚠️ Empty LLM response for keyword matching")
            print(f"🔍 Response debug info:")
            print(f"   - Response object: {type(generate_response)}")
            print(f"   - Has text attr: {hasattr(generate_response, 'text')}")
            print(f"   - Has candidates: {hasattr(generate_response, 'candidates')}")
            if hasattr(generate_response, 'candidates') and generate_response.candidates:
                print(f"   - Number of candidates: {len(generate_response.candidates)}")
            print(f"   - Resume content length: {len(resume_content)} chars")
            print(f"   - Keywords count: {len(standardized_keywords)}")

            # Try one more time with a simpler approach
            print("🔄 Trying simple text extraction...")
            try:
                import re  # Ensure re is available
                simple_text = str(generate_response)
                if len(simple_text) > 100:  # If conversion to string gives us meaningful content
                    # Look for JSON patterns in the string representation
                    json_match = re.search(r'\{.*?\}', simple_text, re.DOTALL)
                    if json_match:
                        analysis_text = json_match.group(0)
                        print(f"✅ Extracted JSON from string representation: {len(analysis_text)} chars")
                    else:
                        print("❌ No JSON found in string representation")
                else:
                    print("❌ String representation too short")
            except Exception as str_e:
                print(f"❌ String conversion failed: {str_e}")

            if not analysis_text:
                print(f"🔄 Falling back to enhanced string matching")
                enhanced_result = enhanced_keyword_matching(resume_content, standardized_keywords)
                enhanced_result['analysis_method'] = 'enhanced_string_fallback'
                enhanced_result['fallback_reason'] = 'empty_llm_response'
                return enhanced_result

        # Parse LLM response
        try:
            import json  # Ensure json is available at function scope

            print(f"🔧 Starting JSON parsing...")
            print(f"   - Raw analysis_text length: {len(analysis_text)}")
            print(f"   - Raw analysis_text preview: {repr(analysis_text[:100])}")

            clean_text = analysis_text.strip()
            print(f"   - After strip: {repr(clean_text[:100])}")

            if clean_text.startswith('```json'):
                clean_text = clean_text[7:].strip()
                print(f"   - Removed ```json marker")
            elif clean_text.startswith('```'):
                clean_text = clean_text[3:].strip()
                print(f"   - Removed ``` marker")
            if clean_text.endswith('```'):
                clean_text = clean_text[:-3].strip()
                print(f"   - Removed ending ``` marker")

            print(f"   - Final clean_text preview: {repr(clean_text[:100])}")

            # Try to parse as complete JSON first
            try:
                analysis_data = json.loads(clean_text)
            except json.JSONDecodeError as partial_e:
                # If JSON is incomplete (likely due to MAX_TOKENS), try to repair it
                print(f"⚠️ Attempting to repair partial JSON: {partial_e}")

                # Try to complete common incomplete patterns
                repaired_text = clean_text

                # If it ends abruptly, try to close the JSON
                if not repaired_text.endswith('}'):
                    if repaired_text.endswith(','):
                        repaired_text = repaired_text[:-1]  # Remove trailing comma
                    repaired_text += '}'

                # Try to fix incomplete arrays
                if '"missing_keywords": [' in repaired_text and not repaired_text.count('[') == repaired_text.count(']'):
                    if repaired_text.endswith(','):
                        repaired_text = repaired_text[:-1]
                    repaired_text += ']}'

                try:
                    analysis_data = json.loads(repaired_text)
                    print(f"✅ Successfully repaired partial JSON")
                except json.JSONDecodeError:
                    # If repair failed, re-raise the original error
                    raise partial_e

            matched_keywords = analysis_data.get('matched_keywords', [])
            missing_keywords = analysis_data.get('missing_keywords', [])

            # Validate and clean results
            matched_keywords = [kw for kw in matched_keywords if kw in standardized_keywords]
            missing_keywords = [kw for kw in missing_keywords if kw in standardized_keywords]

            # Ensure all keywords are accounted for
            all_analyzed = set(matched_keywords + missing_keywords)
            for kw in standardized_keywords:
                if kw not in all_analyzed:
                    missing_keywords.append(kw)

            print(f"✅ LLM Analysis: {len(matched_keywords)}/{len(standardized_keywords)} matched")
            print(f"🔍 Matched: {matched_keywords[:5]}{'...' if len(matched_keywords) > 5 else ''}")
            print(f"❌ Missing: {missing_keywords[:5]}{'...' if len(missing_keywords) > 5 else ''}")

            return {
                "matched_keywords": matched_keywords,
                "missing_keywords": missing_keywords,
                "jd_keywords_matched": len(matched_keywords),
                "total_jd_keywords": len(standardized_keywords),
                "analysis_method": "llm_primary",
                "analysis_notes": analysis_data.get('analysis_notes', 'LLM-based accurate matching'),
                "confidence": "high"
            }

        except json.JSONDecodeError as e:
            print(f"❌ Failed to parse LLM response: {e}")
            print(f"📄 Raw LLM response (first 500 chars): {analysis_text[:500]}...")
            print(f"📄 Raw LLM response (last 200 chars): ...{analysis_text[-200:]}")
            print(f"📄 Response length: {len(analysis_text)} chars")
            print(f"📄 Response starts with: {repr(analysis_text[:50])}")
            print(f"📄 Response ends with: {repr(analysis_text[-50:])}")

            # Try to fix common JSON formatting issues
            print("🔧 Attempting JSON repair...")
            try:
                # Remove leading/trailing whitespace and newlines
                fixed_text = analysis_text.strip()

                # If it doesn't start with {, add it
                if not fixed_text.startswith('{'):
                    fixed_text = '{' + fixed_text

                # If it doesn't end with }, add it
                if not fixed_text.endswith('}'):
                    # Remove trailing comma if present
                    if fixed_text.endswith(','):
                        fixed_text = fixed_text[:-1]
                    fixed_text = fixed_text + '}'

                print(f"🔧 Repaired JSON: {fixed_text[:200]}...")
                analysis_data = json.loads(fixed_text)
                print(f"✅ Successfully repaired and parsed JSON!")

                # Continue with normal processing
                matched_keywords = analysis_data.get('matched_keywords', [])
                missing_keywords = analysis_data.get('missing_keywords', [])

                # Validate and clean results
                matched_keywords = [kw for kw in matched_keywords if kw in standardized_keywords]
                missing_keywords = [kw for kw in missing_keywords if kw in standardized_keywords]

                # Ensure all keywords are accounted for
                all_analyzed = set(matched_keywords + missing_keywords)
                for kw in standardized_keywords:
                    if kw not in all_analyzed:
                        missing_keywords.append(kw)

                print(f"✅ Repaired LLM Analysis: {len(matched_keywords)}/{len(standardized_keywords)} matched")
                print(f"🔍 Matched: {matched_keywords[:5]}{'...' if len(matched_keywords) > 5 else ''}")
                print(f"❌ Missing: {missing_keywords[:5]}{'...' if len(missing_keywords) > 5 else ''}")

                return {
                    "matched_keywords": matched_keywords,
                    "missing_keywords": missing_keywords,
                    "jd_keywords_matched": len(matched_keywords),
                    "total_jd_keywords": len(standardized_keywords),
                    "analysis_method": "llm_repaired",
                    "analysis_notes": analysis_data.get('analysis_notes', 'LLM-based matching with JSON repair'),
                    "confidence": "medium"
                }

            except json.JSONDecodeError as repair_e:
                print(f"❌ JSON repair also failed: {repair_e}")
                enhanced_result = enhanced_keyword_matching(resume_content, standardized_keywords)
                enhanced_result['analysis_method'] = 'enhanced_string_fallback'
                enhanced_result['fallback_reason'] = f'json_parse_error: {str(e)}, repair_failed: {str(repair_e)}'
                enhanced_result['confidence'] = 'medium'
                return enhanced_result

    except Exception as e:
        print(f"❌ Error in LLM keyword analysis: {e}")
        print(f"📄 Analysis text length: {len(analysis_text) if 'analysis_text' in locals() else 'undefined'}")
        if 'analysis_text' in locals() and analysis_text:
            print(f"📄 Raw LLM response (first 500 chars): {analysis_text[:500]}...")
            print(f"📄 Raw LLM response (last 200 chars): ...{analysis_text[-200:]}")
            print(f"📄 Response starts with: {repr(analysis_text[:50])}")
            print(f"📄 Response ends with: {repr(analysis_text[-50:])}")
        enhanced_result = enhanced_keyword_matching(resume_content, standardized_keywords)
        enhanced_result['analysis_method'] = 'enhanced_string_fallback'
        enhanced_result['fallback_reason'] = f'llm_error: {str(e)}'
        enhanced_result['confidence'] = 'medium'
        return enhanced_result


def fallback_keyword_matching(resume_content: str, standardized_keywords: List[str]) -> Dict[str, Any]:
    """Lightweight word-boundary keyword matcher used when the LLM matcher fails.

    Thin wrapper around :func:`enhanced_keyword_matching` — the previous
    implementation used substring `in` checks (`'ai' in 'available'`) which
    inflated match counts. Routing through `enhanced_keyword_matching` ensures
    boundary-correct matching everywhere.
    """
    return enhanced_keyword_matching(resume_content, standardized_keywords)

@app.get("/api/search-history")
async def get_search_history(
    limit: int = 25,
    offset: int = 0,
    page: int = None,
    page_size: int = None,
    user: dict = Depends(require_auth)
):
    """Get a paginated slice of search history for the authenticated user.

    Accepts either ``limit``/``offset`` directly, or 1-based ``page``/``page_size``
    for convenience. Both styles return the same response shape.
    """
    try:
        # Resolve effective page_size / offset, clamping to sane bounds.
        effective_page_size = page_size if page_size is not None else limit
        try:
            effective_page_size = int(effective_page_size)
        except (TypeError, ValueError):
            effective_page_size = 25
        effective_page_size = max(1, min(100, effective_page_size))

        if page is not None:
            try:
                page_int = max(1, int(page))
            except (TypeError, ValueError):
                page_int = 1
            effective_offset = (page_int - 1) * effective_page_size
        else:
            try:
                effective_offset = max(0, int(offset))
            except (TypeError, ValueError):
                effective_offset = 0

        db = get_db_manager()

        # For tenant users and tenant admins, only show their own searches
        # For super admins, show all searches
        if user['user_type'] in ['tenant_user', 'tenant_admin']:
            user_filter = {"user_id": user['id']}
        else:  # super_admin
            user_filter = {}

        history = db.get_search_history(
            limit=effective_page_size,
            offset=effective_offset,
            **user_filter,
        )
        total = db.count_search_history(**user_filter)

        # Convert datetime objects to strings for JSON serialization
        for item in history:
            if 'search_timestamp' in item and item['search_timestamp']:
                item['search_timestamp'] = item['search_timestamp'].isoformat()
            # Limit top results to 3 for display
            if 'top_results' in item and item['top_results']:
                item['top_results'] = item['top_results'][:3]

        return {
            "success": True,
            "history": history,
            "total": total,
            "offset": effective_offset,
            "limit": effective_page_size,
            "page": (effective_offset // effective_page_size) + 1,
            "page_size": effective_page_size,
            "has_more": effective_offset + len(history) < total,
        }
    except Exception as e:
        print(f"❌ Error fetching search history: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch search history: {str(e)}")

@app.get("/api/search-results/{search_id}")
async def get_saved_search_results(search_id: int, user: dict = Depends(require_auth)):
    """Get all results for a specific search"""
    try:
        db = get_db_manager()
        company_id = None
        if user.get('user_type') != 'super_admin':
            company_info = user.get('company') or {}
            if isinstance(company_info, dict):
                company_id = company_info.get('company_id') or company_info.get('id')
            else:
                company_id = getattr(company_info, 'company_id', None) or getattr(company_info, 'id', None)

            if not company_id:
                raise HTTPException(status_code=403, detail="Company context required")

        results = db.get_search_results(search_id, company_id=company_id)

        # Convert datetime objects to strings
        for result in results:
            if 'created_at' in result and result['created_at']:
                result['created_at'] = result['created_at'].isoformat()

        # Format results to match the expected structure
        formatted_results = []
        for result in results:
            formatted_result = {
                "id": str(result['id']),
                "document_name": result['candidate_name'],
                "relevance_score": 0.0,
                "file_path": result['file_path'],
                "gemini_analysis": result['gemini_analysis'],
                "analysis_type": "comprehensive_hr_scorecard"
            }
            formatted_results.append(formatted_result)

        return {
            "success": True,
            "search_id": search_id,
            "results": formatted_results,
            "total_results": len(formatted_results)
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error fetching search results: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch search results: {str(e)}")

@app.post("/api/candidate-action")
@limiter.limit("60/minute")
async def save_candidate_action(
    request: Request,
    search_result_id: int = Form(...),
    candidate_name: str = Form(...),
    action_type: str = Form(...),
    action_status: bool = Form(True),
    comments: str = Form(None),
    user: dict = Depends(require_auth)
):
    """Save or update HR action for a candidate"""
    try:
        db = get_db_manager()
        effective_user_id = str(user["id"])
        # Tenant scoping: only super_admin may act on rows outside their company
        is_super = (user.get('user_type') == 'super_admin')
        tenant_company_id = None
        if not is_super:
            tenant_company_id = (user.get('company_info') or {}).get('id') or (
                user.get('company') or {}
            ).get('company_id') if isinstance(user.get('company'), dict) else None
            if tenant_company_id is None and isinstance(user.get('company'), dict):
                tenant_company_id = user['company'].get('id')

        ok = db.save_candidate_action(
            search_result_id=search_result_id,
            candidate_name=candidate_name,
            action_type=action_type.lower(),  # normalize for consistency
            action_status=action_status,
            comments=comments,
            user_id=effective_user_id,
            company_id=tenant_company_id,
        )
        if not ok:
            raise HTTPException(status_code=404, detail="Candidate not found in your workspace")

        return {
            "success": True,
            "message": f"Action '{action_type}' saved for {candidate_name}"
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error saving candidate action: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to save candidate action: {str(e)}")

@app.delete("/api/search-history")
@limiter.limit("5/minute")
async def delete_all_search_history(request: Request, user: dict = Depends(require_auth)):
    """Delete ALL search history rows for the current user."""
    try:
        db = get_db_manager()
        company_id = (user.get("company_info") or {}).get("id")
        deleted = db.delete_all_search_history_for_user(user['id'], company_id)
        return {"success": True, "deleted": deleted}
    except Exception as e:
        print(f"❌ Error clearing search history: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to clear search history: {str(e)}")


@app.delete("/api/search-history/{search_id}")
@limiter.limit("30/minute")
async def delete_search_history(request: Request, search_id: int, user: dict = Depends(require_auth)):
    """Delete a search and all its associated data.

    super_admin can delete any row. Tenant admins/users can only delete rows
    they own within their company.
    """
    try:
        db = get_db_manager()
        is_super = (user.get('user_type') == 'super_admin')
        if is_super:
            # super_admin: no ownership filter
            user_id = None
            tenant_company_id = None
        else:
            user_id = str(user['id'])
            tenant_company_id = (user.get('company_info') or {}).get('id')
            if tenant_company_id is None and isinstance(user.get('company'), dict):
                tenant_company_id = user['company'].get('company_id') or user['company'].get('id')
        success = db.delete_search_history(search_id, user_id, company_id=tenant_company_id)

        if success:
            return {"success": True, "message": "Search history deleted successfully"}
        raise HTTPException(status_code=404, detail="Search not found or unauthorized")
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error deleting search history: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete search history: {str(e)}")

@app.get("/api/download-resume")
async def download_resume(file_path: str, user: dict = Depends(require_auth)):
    """Download resume file from Google Cloud Storage bucket.

    Tenant scoping: super_admin bypasses; everyone else must have an
    upload row OR a search-result row for this file under their company.
    The fetch uses the company-specific bucket (with global-bucket
    fallback for legacy data).
    """
    try:
        print(f"📥 Download request for: {file_path} (user={user.get('email')})")

        if not file_path:
            raise HTTPException(status_code=400, detail="File path is required")

        # Strip any gs:// prefix
        blob_name = file_path
        if blob_name.startswith("gs://"):
            # gs://bucket/blob -> blob
            without_scheme = blob_name[5:]
            if "/" in without_scheme:
                blob_name = without_scheme.split("/", 1)[1]
            else:
                blob_name = without_scheme

        # Reject path traversal in the blob name. Note: GCS object names
        # legitimately contain slashes for "folders" so we allow those but
        # block parent traversal segments.
        if any(seg in ('..', '') for seg in blob_name.split('/')[:-1] if seg in ('..',)) or '\x00' in blob_name:
            raise HTTPException(status_code=400, detail="Invalid file path")

        is_super = (user.get('user_type') == 'super_admin')
        # Tenant ownership check (skipped for super_admin)
        if not is_super:
            tenant_company_id = (user.get('company_info') or {}).get('id')
            if tenant_company_id is None and isinstance(user.get('company'), dict):
                tenant_company_id = user['company'].get('company_id') or user['company'].get('id')
            if tenant_company_id is None:
                raise HTTPException(status_code=403, detail="No company context")
            db = get_db_manager()
            if not db.resume_belongs_to_company(blob_name, tenant_company_id):
                print(f"🚫 Tenant mismatch: user_company={tenant_company_id} file={blob_name}")
                # Return 404 (not 403) to avoid leaking existence
                raise HTTPException(status_code=404, detail="Resume file not found")

        # Resolve bucket: company-specific first, fallback to global
        bucket_name, _ = get_company_resources(user) if not is_super else (None, None)
        if not bucket_name:
            bucket_name = GCS_BUCKET_NAME
        if not bucket_name:
            raise HTTPException(status_code=500, detail="GCS bucket not configured")

        client = _get_storage_client()
        bucket = client.bucket(bucket_name)
        print(f"🔍 Looking for blob: {blob_name} in {bucket_name}")
        blob = bucket.blob(blob_name)
        if not blob.exists():
            # Fallback to legacy global bucket if company bucket misses
            if bucket_name != GCS_BUCKET_NAME and GCS_BUCKET_NAME:
                legacy = client.bucket(GCS_BUCKET_NAME).blob(blob_name)
                if legacy.exists():
                    blob = legacy
                else:
                    raise HTTPException(status_code=404, detail="Resume file not found")
            else:
                raise HTTPException(status_code=404, detail="Resume file not found")

        file_content = blob.download_as_bytes()
        filename = blob_name.split('/')[-1] if '/' in blob_name else blob_name

        content_type = "application/octet-stream"
        lower = filename.lower()
        if lower.endswith('.pdf'):
            content_type = "application/pdf"
        elif lower.endswith('.docx'):
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif lower.endswith('.doc'):
            content_type = "application/msword"
        elif lower.endswith('.txt'):
            content_type = "text/plain"

        print(f"✅ Successfully downloaded: {filename} ({len(file_content)} bytes)")

        # Sanitize filename for Content-Disposition (RFC 5987)
        safe_filename = ''.join(c if c.isalnum() or c in '._- ' else '_' for c in filename) or 'resume'
        from fastapi.responses import Response
        return Response(
            content=file_content,
            media_type=content_type,
            headers={
                "Content-Disposition": f'attachment; filename="{safe_filename}"',
                "Content-Length": str(len(file_content))
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Download failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to download resume: {str(e)}")


@app.get("/api/candidates-by-status")
async def get_candidates_by_status(
    status: str,
    limit: int = 200,
    user: dict = Depends(require_auth),
):
    """List candidates the current user has marked with the given status.

    Status values: ``selected``, ``rejected``, ``shortlisted``, ``interviewed``.
    Used by the dashboard stat-card click-throughs to open a dedicated
    candidate list for each status.
    """
    allowed = {"selected", "rejected", "shortlisted", "interviewed", "hired"}
    s = (status or "").strip().lower()
    if s not in allowed:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid status '{status}'. Must be one of: {sorted(allowed)}",
        )
    try:
        limit_int = max(1, min(1000, int(limit)))
    except (TypeError, ValueError):
        limit_int = 200

    try:
        db = get_db_manager()
        user_id = user["id"]
        company_id = None
        if user.get("company"):
            if isinstance(user["company"], dict):
                company_id = user["company"].get("company_id") or user["company"].get("id")
            else:
                company_id = getattr(user["company"], "id", None)

        rows = db.get_candidates_by_action(user_id, s, company_id, limit=limit_int) or []

        candidates = []
        for row in rows:
            actioned_at = row.get("actioned_at") or row.get("actioned_created_at")
            candidates.append({
                "search_result_id":   row.get("search_result_id"),
                "search_id":          row.get("search_id"),
                "candidate_name":     row.get("candidate_name") or "Unknown",
                "candidate_email":    row.get("candidate_email") or "",
                "candidate_phone":    row.get("candidate_phone") or "",
                "candidate_location": row.get("candidate_location") or "",
                "position_applied":   row.get("position_applied") or "",
                "experience_years":   row.get("experience_years"),
                "match_score":        float(row.get("match_score") or 0),
                "match_status":       row.get("match_status") or "",
                "file_path":          row.get("file_path") or "",
                "comments":           row.get("comments") or "",
                "search_query":       row.get("search_query") or "",
                "actioned_at":        actioned_at.isoformat() if actioned_at else None,
            })

        return {
            "success": True,
            "status": s,
            "count": len(candidates),
            "candidates": candidates,
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error listing candidates by status '{status}': {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to list candidates by status: {e}",
        )


@app.get("/api/dashboard-stats")
async def get_dashboard_stats(user: dict = Depends(require_auth)):
    """Get dashboard statistics for the current user"""
    try:
        db = get_db_manager()
        user_id = user['id']
        company_id = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                company_id = user['company'].get('company_id') or user['company'].get('id')
            else:
                company_id = getattr(user['company'], 'id', None)

        # Get candidate action statistics
        candidate_stats = db.get_candidate_action_stats(user_id, company_id)

        # Get search statistics
        search_stats = db.get_user_search_stats(user_id, company_id)

        # Get trending skills from recent searches
        trending_skills = db.get_trending_skills(company_id)

        # Resource usage (resume storage / search queries / plan) — visible to every
        # authenticated user with a company so the dashboard banner can render.
        resource_usage = None
        if company_id:
            try:
                stats = db.get_company_stats(company_id) or {}
                max_resumes = stats.get('max_resumes') or 1000
                max_searches = stats.get('max_searches') or 10000
                total_resumes = stats.get('total_resumes') or 0
                total_searches = stats.get('total_searches') or 0
                resume_pct = round(min(100.0, (total_resumes / max_resumes) * 100), 1) if max_resumes else 0
                search_pct = round(min(100.0, (total_searches / max_searches) * 100), 1) if max_searches else 0
                resource_usage = {
                    "resumes": {"current": total_resumes, "maximum": max_resumes, "usage_percent": resume_pct},
                    "searches": {"current": total_searches, "maximum": max_searches, "usage_percent": search_pct},
                    "users": {"current": stats.get('total_users') or 0, "maximum": stats.get('max_users') or 10},
                    "subscription_plan": stats.get('subscription_plan') or 'basic',
                }
            except Exception as ru_err:
                print(f"⚠️  Could not load resource usage in dashboard-stats: {ru_err}")
                resource_usage = None

        # Activity stats — always populated so dashboard cards never sit at zero
        # because of missing data. Super admins see system-wide totals; tenant
        # users see totals scoped to their company.
        activity_stats = {
            "total_searches": 0,
            "total_resumes": 0,
            "total_candidates": 0,
            "total_users": 0,
            "total_companies": 0,
        }
        try:
            with db.get_cursor() as (cursor, _conn):
                if user['user_type'] == 'super_admin':
                    cursor.execute("""
                        SELECT
                            (SELECT COUNT(*) FROM search_history)   AS total_searches,
                            (SELECT COUNT(*) FROM resume_uploads)   AS total_resumes,
                            (SELECT COUNT(*) FROM search_results)   AS total_candidates,
                            (SELECT COUNT(*) FROM users)            AS total_users,
                            (SELECT COUNT(*) FROM tenant_companies) AS total_companies
                    """)
                elif company_id:
                    cursor.execute("""
                        SELECT
                            (SELECT COUNT(*) FROM search_history WHERE company_id = %s) AS total_searches,
                            (SELECT COUNT(*) FROM resume_uploads WHERE company_id = %s) AS total_resumes,
                            (SELECT COUNT(*) FROM search_results sr
                                JOIN search_history sh ON sr.search_id = sh.id
                                WHERE sh.company_id = %s)                                AS total_candidates,
                            (SELECT COUNT(*) FROM user_companies WHERE company_id = %s)  AS total_users,
                            1                                                            AS total_companies
                    """, (company_id, company_id, company_id, company_id))
                else:
                    cursor.execute("""
                        SELECT
                            (SELECT COUNT(*) FROM search_history WHERE created_by_user_id = %s) AS total_searches,
                            0 AS total_resumes,
                            (SELECT COUNT(*) FROM search_results sr
                                JOIN search_history sh ON sr.search_id = sh.id
                                WHERE sh.created_by_user_id = %s)                                AS total_candidates,
                            0 AS total_users,
                            0 AS total_companies
                    """, (user_id, user_id))
                row = cursor.fetchone() or {}
                if row:
                    activity_stats.update({k: int(row[k] or 0) for k in activity_stats.keys() if k in row})
        except Exception as as_err:
            print(f"⚠️  Could not load activity stats: {as_err}")

        return {
            "success": True,
            "candidate_stats": candidate_stats,
            "search_stats": search_stats,
            "trending_skills": trending_skills,
            "resource_usage": resource_usage,
            "activity_stats": activity_stats,
        }

    except Exception as e:
        print(f"❌ Error getting dashboard stats: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get dashboard stats: {str(e)}")

@app.get("/api/upcoming-events")
async def get_upcoming_events(user: dict = Depends(require_auth)):
    """Get upcoming events and scheduled activities for the user"""
    try:
        db = get_db_manager()
        user_id = user['id']
        company_id = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                company_id = user['company'].get('company_id') or user['company'].get('id')
            else:
                company_id = getattr(user['company'], 'id', None)

        # Get upcoming events (interviews, follow-ups, etc.)
        events = db.get_upcoming_events(user_id, company_id)

        return {
            "success": True,
            "events": events
        }

    except Exception as e:
        print(f"❌ Error getting upcoming events: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get upcoming events: {str(e)}")

@app.post("/api/migrate-vector-datastore")
async def migrate_existing_files_to_vector_datastore(
    user: dict = Depends(require_super_admin)
):
    """
    Migrate existing files from database to vector datastore with proper company metadata.
    This should be run once to reorganize the vector datastore with company metadata.
    """
    print(f"\n🔄 VECTOR DATASTORE MIGRATION STARTED")
    print(f"👤 Initiated by super admin: {user['email']}")

    try:
        db = get_db_manager()

        # Get all uploaded files from database
        with db.get_cursor() as (cursor, conn):
            cursor.execute("""
                SELECT
                    ru.file_name,
                    ru.file_path,
                    ru.uploaded_by_user_id,
                    ru.company_id,
                    tc.company_code,
                    tc.company_name
                FROM resume_uploads ru
                LEFT JOIN user_companies uc ON ru.company_id = uc.company_id
                LEFT JOIN tenant_companies tc ON uc.company_id = tc.id
                ORDER BY ru.upload_timestamp DESC
            """)
            files_to_migrate = cursor.fetchall()

        print(f"📊 Found {len(files_to_migrate)} files to migrate")

        migration_results = []
        success_count = 0
        error_count = 0

        for file_info in files_to_migrate:
            file_path = file_info['file_path']
            company_id = file_info['company_id']
            company_code = file_info['company_code']
            user_id = str(file_info['uploaded_by_user_id'])

            print(f"\n📄 Migrating: {file_path}")
            print(f"   🏢 Company: {company_code} ({company_id})")
            print(f"   👤 User: {user_id}")

            try:
                # Upload to vector datastore with metadata
                migration_success = upload_to_vector_datastore(
                    file_path=file_path,
                    company_id=company_id,
                    company_code=company_code,
                    user_id=user_id
                )

                migration_results.append({
                    "file_path": file_path,
                    "company_code": company_code,
                    "success": migration_success
                })

                if migration_success:
                    success_count += 1
                    print(f"   ✅ Successfully migrated")
                else:
                    error_count += 1
                    print(f"   ❌ Migration failed")

            except Exception as e:
                error_count += 1
                migration_results.append({
                    "file_path": file_path,
                    "company_code": company_code,
                    "success": False,
                    "error": str(e)
                })
                print(f"   ❌ Migration error: {str(e)}")

        print(f"\n🎉 MIGRATION COMPLETED")
        print(f"📊 Total files: {len(files_to_migrate)}")
        print(f"✅ Successful: {success_count}")
        print(f"❌ Failed: {error_count}")

        return {
            "success": True,
            "total_files": len(files_to_migrate),
            "successful_migrations": success_count,
            "failed_migrations": error_count,
            "migration_results": migration_results
        }

    except Exception as e:
        print(f"❌ Migration failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Migration failed: {str(e)}")

@app.post("/api/schedule-interview")
@limiter.limit("30/minute")
async def schedule_interview(request: Request, user: dict = Depends(require_auth)):
    """Schedule an interview with a candidate.

    Extended fields (all optional, backward compatible):
      - meeting_provider:  "teams" | "meet" | "none"  (default "none")
      - send_invite_email: bool  — send the candidate a templated invite
      - email_template_id: int   — which template to use; otherwise the
                                   default interview_invite is used.
    """
    try:
        try:
            data = await request.json()
        except Exception:
            raise HTTPException(status_code=422, detail="Request body must be valid JSON")
        if not isinstance(data, dict):
            raise HTTPException(status_code=422, detail="Request body must be a JSON object")
        if not data.get("candidate_name"):
            raise HTTPException(status_code=422, detail="Field 'candidate_name' is required")
        db = get_db_manager()

        # Add user and company info
        data['created_by'] = user['id']
        company_id = None
        company_name = None
        if user.get('company'):
            if isinstance(user['company'], dict):
                company_id = user['company'].get('company_id') or user['company'].get('id')
                company_name = user['company'].get('company_name') or user['company'].get('name')
            else:
                company_id = getattr(user['company'], 'id', None)
                company_name = getattr(user['company'], 'name', None)
        data['company_id'] = company_id

        # -- Optional: create Teams or Google Meet meeting -----------------
        provider = (data.get('meeting_provider') or '').strip().lower() or None
        meeting_result = None
        if provider and provider not in ('none', 'in_person', 'phone'):
            try:
                from services.calendar_service import create_meeting
                # Compose ISO datetime from date + time fields.
                start_iso = f"{data.get('interview_date')}T{data.get('interview_time') or '09:00'}:00"
                meeting_result = create_meeting(
                    provider,
                    subject=f"Interview: {data['candidate_name']} — "
                            f"{data.get('position', 'Role')}",
                    start_iso=start_iso,
                    duration_minutes=int(data.get('duration') or 30),
                    attendee_email=data.get('candidate_email') or '',
                    organizer_user_id=user['id'],
                    db=db,
                )
                if meeting_result.get('success'):
                    data['meeting_provider'] = meeting_result.get('provider') or provider
                    data['meeting_join_url'] = meeting_result.get('join_url')
                    data['meeting_event_id'] = meeting_result.get('event_id')
                    # Surface the join URL in `location` too for clients that
                    # only render the location field.
                    if meeting_result.get('join_url') and not data.get('location'):
                        data['location'] = meeting_result['join_url']
                else:
                    # Non-fatal: persist the interview without a link and
                    # report the error back to the caller.
                    data['meeting_provider'] = provider
            except Exception as _meet_err:
                print(f"⚠️ Meeting creation failed: {_meet_err}")
                meeting_result = {"success": False, "error": str(_meet_err)}

        # Save interview to database
        interview_id = db.schedule_interview(data)

        # -- Optional: send the candidate an invite email ------------------
        email_result = None
        if data.get('send_invite_email') and data.get('candidate_email'):
            try:
                from services.email_service import (
                    send_email, render_template, build_default_context,
                )
                tpl = None
                tpl_id = data.get('email_template_id')
                if tpl_id:
                    tpl = db.get_email_template(tpl_id, company_id)
                if not tpl:
                    # Fall back to a default interview_invite (company-scoped first, then global).
                    company_tpls = db.list_email_templates(company_id, kind='interview_invite')
                    if not company_tpls:
                        company_tpls = db.list_email_templates(None, kind='interview_invite')
                    tpl = company_tpls[0] if company_tpls else None
                if tpl:
                    ctx = build_default_context(
                        candidate_name=data['candidate_name'],
                        candidate_email=data.get('candidate_email', ''),
                        position=data.get('position') or 'the role',
                        company=company_name or 'our company',
                        recruiter_name=user.get('full_name') or user.get('email'),
                        recruiter_email=user.get('email'),
                        interview_date=str(data.get('interview_date') or ''),
                        interview_time=str(data.get('interview_time') or ''),
                        interview_location=data.get('location') or '',
                        meeting_join_url=data.get('meeting_join_url') or '',
                        custom_message=data.get('custom_message') or '',
                    )
                    subj, body_html = render_template(tpl['subject'], tpl['body'], ctx)
                    email_result = send_email(
                        to=data['candidate_email'],
                        subject=subj, body_html=body_html,
                        prefer=data.get('email_prefer') or 'auto',
                        organizer_user_id=user['id'], db=db,
                    )
                    db.log_email_sent(
                        company_id=company_id,
                        sent_by_user_id=user['id'],
                        recipient_email=data['candidate_email'],
                        candidate_name=data['candidate_name'],
                        kind='interview_invite',
                        subject=subj,
                        body=body_html,
                        provider=email_result.get('provider'),
                        status='sent' if email_result.get('success') else 'failed',
                        error_message=email_result.get('error'),
                        search_result_id=data.get('search_result_id'),
                        interview_id=interview_id,
                    )
            except Exception as _email_err:
                print(f"⚠️ Invite email send failed: {_email_err}")
                email_result = {"success": False, "error": str(_email_err)}

        return {
            "success": True,
            "interview_id": interview_id,
            "message": f"Interview scheduled with {data['candidate_name']}",
            "meeting": meeting_result,
            "email": email_result,
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error scheduling interview: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to schedule interview: {str(e)}")


# ============================================================================
# Integrations: Email templates, candidate email sending, OAuth (Google/MS)
# ============================================================================
def _get_company_id(user: dict) -> Optional[int]:
    if user.get('company'):
        if isinstance(user['company'], dict):
            return user['company'].get('company_id') or user['company'].get('id')
        return getattr(user['company'], 'id', None)
    return None


@app.get("/api/email-templates")
@limiter.limit("60/minute")
async def list_email_templates_endpoint(request: Request,
                                        kind: Optional[str] = None,
                                        user: dict = Depends(require_auth)):
    db = get_db_manager()
    company_id = _get_company_id(user)
    own = db.list_email_templates(company_id, kind=kind) if company_id else []
    globals_ = db.list_email_templates(None, kind=kind)
    seen = {r['id'] for r in own}
    merged = list(own) + [r for r in globals_ if r['id'] not in seen]
    return {"success": True, "templates": merged}


@app.post("/api/email-templates")
@limiter.limit("30/minute")
async def create_email_template_endpoint(request: Request,
                                         user: dict = Depends(require_auth)):
    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=422, detail="Body must be JSON")
    if not isinstance(data, dict):
        raise HTTPException(status_code=422, detail="Body must be a JSON object")
    for f in ("kind", "name", "subject", "body"):
        if not data.get(f):
            raise HTTPException(status_code=422, detail=f"Field '{f}' is required")
    if data["kind"] not in {"shortlist", "interview_invite", "rejection", "custom"}:
        raise HTTPException(status_code=422,
                            detail="kind must be one of: shortlist, interview_invite, rejection, custom")
    db = get_db_manager()
    tpl_id = db.upsert_email_template(
        company_id=_get_company_id(user),
        kind=data["kind"], name=data["name"][:200],
        subject=data["subject"], body=data["body"],
        is_default=bool(data.get("is_default")),
        created_by=user["id"],
    )
    return {"success": True, "template_id": tpl_id}


@app.put("/api/email-templates/{template_id}")
@limiter.limit("30/minute")
async def update_email_template_endpoint(template_id: int, request: Request,
                                         user: dict = Depends(require_auth)):
    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=422, detail="Body must be JSON")
    db = get_db_manager()
    company_id = _get_company_id(user)
    existing = db.get_email_template(template_id, company_id)
    if not existing:
        raise HTTPException(status_code=404, detail="Template not found")
    if existing.get("company_id") is None:
        raise HTTPException(status_code=403, detail="Cannot modify a global default template")
    res = db.upsert_email_template(
        company_id=company_id,
        kind=existing["kind"],
        name=(data.get("name") or existing["name"])[:200],
        subject=data.get("subject") or existing["subject"],
        body=data.get("body") or existing["body"],
        template_id=template_id,
        is_default=bool(data.get("is_default", existing["is_default"])),
        created_by=user["id"],
    )
    if not res:
        raise HTTPException(status_code=404, detail="Template not found")
    return {"success": True}


@app.delete("/api/email-templates/{template_id}")
@limiter.limit("30/minute")
async def delete_email_template_endpoint(template_id: int, request: Request,
                                         user: dict = Depends(require_auth)):
    db = get_db_manager()
    company_id = _get_company_id(user)
    existing = db.get_email_template(template_id, company_id)
    if not existing or existing.get("company_id") is None:
        raise HTTPException(status_code=404, detail="Template not found")
    ok = db.delete_email_template(template_id, company_id)
    return {"success": bool(ok)}


@app.post("/api/send-candidate-email")
@limiter.limit("60/minute")
async def send_candidate_email_endpoint(request: Request,
                                        user: dict = Depends(require_auth)):
    """Send a custom email to a candidate using a saved template.

    Body: {
      candidate_name (req), candidate_email (req),
      template_id (req unless raw_subject+raw_body provided),
      raw_subject, raw_body,         # ad-hoc send without saved template
      position, custom_message,
      search_result_id,              # for log linkage
      prefer  ("auto"|"smtp"|"gmail"|"microsoft")
    }
    """
    try:
        data = await request.json()
    except Exception:
        raise HTTPException(status_code=422, detail="Body must be JSON")
    if not isinstance(data, dict):
        raise HTTPException(status_code=422, detail="Body must be a JSON object")
    if not data.get("candidate_email"):
        raise HTTPException(status_code=422, detail="candidate_email is required")

    db = get_db_manager()
    company_id = _get_company_id(user)
    company_name = None
    if user.get('company') and isinstance(user['company'], dict):
        company_name = user['company'].get('company_name') or user['company'].get('name')

    subj_tpl = data.get("raw_subject")
    body_tpl = data.get("raw_body")
    kind = data.get("kind") or "custom"
    if not (subj_tpl and body_tpl):
        tpl_id = data.get("template_id")
        if not tpl_id:
            raise HTTPException(status_code=422,
                                detail="Provide template_id, or raw_subject+raw_body")
        tpl = db.get_email_template(tpl_id, company_id)
        if not tpl:
            raise HTTPException(status_code=404, detail="Template not found")
        subj_tpl, body_tpl, kind = tpl["subject"], tpl["body"], tpl["kind"]

    from services.email_service import send_email, render_template, build_default_context
    ctx = build_default_context(
        candidate_name=data.get("candidate_name") or "",
        candidate_email=data["candidate_email"],
        position=data.get("position") or "the role",
        company=company_name or "our company",
        recruiter_name=user.get("full_name") or user.get("email"),
        recruiter_email=user.get("email"),
        custom_message=data.get("custom_message") or "",
        interview_date=str(data.get("interview_date") or ""),
        interview_time=str(data.get("interview_time") or ""),
        interview_location=data.get("interview_location") or "",
        meeting_join_url=data.get("meeting_join_url") or "",
    )
    subject, body_html = render_template(subj_tpl, body_tpl, ctx)
    result = send_email(
        to=data["candidate_email"], subject=subject, body_html=body_html,
        prefer=data.get("prefer") or "auto",
        organizer_user_id=user["id"], db=db,
    )
    db.log_email_sent(
        company_id=company_id,
        sent_by_user_id=user["id"],
        recipient_email=data["candidate_email"],
        candidate_name=data.get("candidate_name") or "",
        kind=kind, subject=subject, body=body_html,
        provider=result.get("provider"),
        status="sent" if result.get("success") else "failed",
        error_message=result.get("error"),
        search_result_id=data.get("search_result_id"),
    )
    if not result.get("success"):
        return JSONResponse(status_code=502, content={
            "success": False,
            "error": result.get("error") or "Email send failed",
            "provider": result.get("provider"),
        })
    return {"success": True, "provider": result.get("provider"),
            "subject": subject}


@app.get("/api/integrations/status")
@limiter.limit("60/minute")
async def integrations_status(request: Request,
                              user: dict = Depends(require_auth)):
    """Report which integrations are configured / connected for this user."""
    from services.calendar_service import calendar_config
    from services.email_service import email_config
    db = get_db_manager()
    cal = calendar_config()
    em = email_config()
    g_row = db.get_oauth_token(user["id"], "google")
    m_row = db.get_oauth_token(user["id"], "microsoft")
    return {
        "success": True,
        "google": {
            "client_configured": bool((cal.get("google") or {}).get("client_id")
                                      and (cal.get("google") or {}).get("client_secret")),
            "connected": bool(g_row),
            "account_email": (g_row or {}).get("email"),
        },
        "microsoft": {
            "client_configured": bool((cal.get("microsoft") or {}).get("client_id")
                                       and (cal.get("microsoft") or {}).get("client_secret")),
            "tenant_app_only": bool((cal.get("microsoft") or {}).get("sender_user_id")
                                     and (cal.get("microsoft") or {}).get("tenant_id")),
            "connected": bool(m_row),
            "account_email": (m_row or {}).get("email"),
        },
        "smtp": {
            "configured": bool((em.get("smtp") or {}).get("host")
                                and (em.get("smtp") or {}).get("from")),
            "host": (em.get("smtp") or {}).get("host"),
            "from": (em.get("smtp") or {}).get("from"),
        },
    }


# ---- OAuth: Google ---------------------------------------------------------
@app.get("/api/oauth/google/start")
@limiter.limit("20/minute")
async def oauth_google_start(request: Request,
                             user: dict = Depends(require_auth)):
    from services.calendar_service import calendar_config, GOOGLE_SCOPES
    cfg = (calendar_config().get("google") or {})
    if not (cfg.get("client_id") and cfg.get("redirect_uri")):
        raise HTTPException(status_code=503,
                            detail="Google OAuth not configured. See docs/INTEGRATIONS_SETUP.md")
    state = secrets.token_urlsafe(32)
    from urllib.parse import urlencode
    params = {
        "client_id": cfg["client_id"],
        "redirect_uri": cfg["redirect_uri"],
        "response_type": "code",
        "scope": " ".join(GOOGLE_SCOPES),
        "access_type": "offline",
        "prompt": "consent",
        "state": state,
    }
    url = "https://accounts.google.com/o/oauth2/v2/auth?" + urlencode(params)
    resp = RedirectResponse(url=url, status_code=302)
    resp.set_cookie("oauth_state_google", f"{state}:{user['id']}",
                    max_age=600, httponly=True, secure=True, samesite="lax", path="/")
    return resp


@app.get("/api/oauth/google/callback")
async def oauth_google_callback(request: Request,
                                code: Optional[str] = None,
                                state: Optional[str] = None,
                                error: Optional[str] = None):
    from urllib.parse import quote as _q
    if error:
        # Bounce to dashboard; banner renders error safely via textContent.
        return RedirectResponse(
            url=f"/?integration=google&status=failed&error={_q(str(error)[:200])}",
            status_code=302,
        )
    if not (code and state):
        return RedirectResponse(
            url="/?integration=google&status=failed&error=missing_code_or_state",
            status_code=302,
        )
    cookie_val = request.cookies.get("oauth_state_google") or ""
    if ":" not in cookie_val:
        raise HTTPException(status_code=400, detail="Missing OAuth state cookie")
    saved_state, _, user_id_str = cookie_val.partition(":")
    if not saved_state or saved_state != state or not user_id_str.isdigit():
        raise HTTPException(status_code=400, detail="Invalid OAuth state")
    user_id = int(user_id_str)

    from services.calendar_service import calendar_config
    cfg = (calendar_config().get("google") or {})
    try:
        r = requests.post(
            "https://oauth2.googleapis.com/token",
            data={
                "code": code,
                "client_id": cfg["client_id"],
                "client_secret": cfg["client_secret"],
                "redirect_uri": cfg["redirect_uri"],
                "grant_type": "authorization_code",
            }, timeout=15,
        )
        r.raise_for_status()
        tok = r.json()
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Token exchange failed: {e}")

    # Try to learn the connected user's email
    account_email = None
    try:
        ui = requests.get(
            "https://openidconnect.googleapis.com/v1/userinfo",
            headers={"Authorization": f"Bearer {tok.get('access_token')}"},
            timeout=10,
        ).json()
        account_email = ui.get("email")
    except Exception:
        pass

    from datetime import datetime as _dt, timezone as _tz, timedelta as _td
    expires_at = _dt.now(_tz.utc) + _td(seconds=int(tok.get("expires_in", 3600)))
    db = get_db_manager()
    db.save_oauth_token(
        user_id, "google",
        access_token=tok.get("access_token"),
        refresh_token=tok.get("refresh_token"),
        expires_at=expires_at,
        scope=tok.get("scope"),
        email=account_email,
    )
    resp = RedirectResponse(url="/?integration=google&status=connected", status_code=302)
    resp.delete_cookie("oauth_state_google", path="/")
    return resp


# ---- OAuth: Microsoft ------------------------------------------------------
@app.get("/api/oauth/microsoft/start")
@limiter.limit("20/minute")
async def oauth_microsoft_start(request: Request,
                                user: dict = Depends(require_auth)):
    from services.calendar_service import calendar_config, MICROSOFT_SCOPES
    cfg = (calendar_config().get("microsoft") or {})
    if not (cfg.get("client_id") and cfg.get("redirect_uri")):
        raise HTTPException(status_code=503,
                            detail="Microsoft OAuth not configured. See docs/INTEGRATIONS_SETUP.md")
    state = secrets.token_urlsafe(32)
    tenant = cfg.get("tenant_id") or "common"
    from urllib.parse import urlencode
    params = {
        "client_id": cfg["client_id"],
        "redirect_uri": cfg["redirect_uri"],
        "response_type": "code",
        "response_mode": "query",
        "scope": " ".join(MICROSOFT_SCOPES),
        "state": state,
        "prompt": "select_account",
    }
    url = f"https://login.microsoftonline.com/{tenant}/oauth2/v2.0/authorize?" + urlencode(params)
    resp = RedirectResponse(url=url, status_code=302)
    resp.set_cookie("oauth_state_microsoft", f"{state}:{user['id']}",
                    max_age=600, httponly=True, secure=True, samesite="lax", path="/")
    return resp


@app.get("/api/oauth/microsoft/callback")
async def oauth_microsoft_callback(request: Request,
                                   code: Optional[str] = None,
                                   state: Optional[str] = None,
                                   error: Optional[str] = None):
    from urllib.parse import quote as _q
    if error:
        return RedirectResponse(
            url=f"/?integration=microsoft&status=failed&error={_q(str(error)[:200])}",
            status_code=302,
        )
    if not (code and state):
        return RedirectResponse(
            url="/?integration=microsoft&status=failed&error=missing_code_or_state",
            status_code=302,
        )
    cookie_val = request.cookies.get("oauth_state_microsoft") or ""
    if ":" not in cookie_val:
        raise HTTPException(status_code=400, detail="Missing OAuth state cookie")
    saved_state, _, user_id_str = cookie_val.partition(":")
    if not saved_state or saved_state != state or not user_id_str.isdigit():
        raise HTTPException(status_code=400, detail="Invalid OAuth state")
    user_id = int(user_id_str)

    from services.calendar_service import calendar_config, MICROSOFT_SCOPES
    cfg = (calendar_config().get("microsoft") or {})
    tenant = cfg.get("tenant_id") or "common"
    try:
        r = requests.post(
            f"https://login.microsoftonline.com/{tenant}/oauth2/v2.0/token",
            data={
                "code": code,
                "client_id": cfg["client_id"],
                "client_secret": cfg["client_secret"],
                "redirect_uri": cfg["redirect_uri"],
                "grant_type": "authorization_code",
                "scope": " ".join(MICROSOFT_SCOPES),
            }, timeout=15,
        )
        r.raise_for_status()
        tok = r.json()
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Token exchange failed: {e}")

    # Discover the user principal name for sendMail/calendar paths.
    account_email = None
    try:
        me = requests.get(
            "https://graph.microsoft.com/v1.0/me",
            headers={"Authorization": f"Bearer {tok.get('access_token')}"},
            timeout=10,
        ).json()
        account_email = me.get("mail") or me.get("userPrincipalName")
    except Exception:
        pass

    from datetime import datetime as _dt, timezone as _tz, timedelta as _td
    expires_at = _dt.now(_tz.utc) + _td(seconds=int(tok.get("expires_in", 3600)))
    db = get_db_manager()
    db.save_oauth_token(
        user_id, "microsoft",
        access_token=tok.get("access_token"),
        refresh_token=tok.get("refresh_token"),
        expires_at=expires_at,
        scope=tok.get("scope"),
        email=account_email,
    )
    resp = RedirectResponse(url="/?integration=microsoft&status=connected", status_code=302)
    resp.delete_cookie("oauth_state_microsoft", path="/")
    return resp


@app.post("/api/oauth/{provider}/disconnect")
@limiter.limit("20/minute")
async def oauth_disconnect(provider: str, request: Request,
                           user: dict = Depends(require_auth)):
    if provider not in ("google", "microsoft"):
        raise HTTPException(status_code=400, detail="Unknown provider")
    db = get_db_manager()
    ok = db.delete_oauth_token(user["id"], provider)
    return {"success": bool(ok)}


# JD Builder API Endpoints

def search_web_for_company_info(company_name: str) -> str:
    """
    Performs a web search to get information about a company using dedicated web search app.
    Falls back to resume search engine if web search app is not available.
    """
    if not company_name:
        return ""

    try:
        print(f"🔍 Searching web for information about: {company_name}")

        # Check for dedicated web search configuration first
        web_search_config = config.get('web_search')
        if web_search_config and web_search_config.get('serving_config'):
            print(f"🌐 Using dedicated web search app")
            return search_with_dedicated_web_app(company_name, web_search_config)

        # Fallback to existing vector search configuration
        print(f"⚠️  No dedicated web search app found, using resume search engine")
        vector_config = config.get('vector_search')
        if not vector_config:
            print(f"⚠️  Vector search not configured.")
            return search_web_fallback_with_gemini(company_name)

        project_id = vector_config.get('project_id')
        location = vector_config.get('location', 'global')
        datastore_id = vector_config.get('datastore_id')

        if not all([project_id, datastore_id]):
            print(f"⚠️  Missing search configuration.")
            return search_web_fallback_with_gemini(company_name)

        # Get search client using existing function
        search_client = get_search_client(location)

        # Create serving config using existing datastore
        serving_config = f"projects/{project_id}/locations/{location}/collections/default_collection/dataStores/{datastore_id}/servingConfigs/default_config"

        # Create search query focused on company information with web search intent
        search_query = f"site:wikipedia.org OR site:linkedin.com OR site:crunchbase.com OR site:*.com {company_name} company about mission values culture business overview"

        # Create search request with web search parameters
        request = discoveryengine_v1.SearchRequest(
            serving_config=serving_config,
            query=search_query,
            page_size=5,  # Get top 5 results for more comprehensive info
            # Enable web search capabilities
            content_search_spec=discoveryengine_v1.SearchRequest.ContentSearchSpec(
                snippet_spec=discoveryengine_v1.SearchRequest.ContentSearchSpec.SnippetSpec(
                    return_snippet=True,
                    max_snippet_count=3
                ),
                summary_spec=discoveryengine_v1.SearchRequest.ContentSearchSpec.SummarySpec(
                    summary_result_count=3,
                    include_citations=True,
                    ignore_adversarial_query=True,
                    ignore_non_summary_seeking_query=True
                ),
                extractive_content_spec=discoveryengine_v1.SearchRequest.ContentSearchSpec.ExtractiveContentSpec(
                    max_extractive_answer_count=3,
                    max_extractive_segment_count=3
                )
            ),
            # Add search parameters that might help with web content
            query_expansion_spec=discoveryengine_v1.SearchRequest.QueryExpansionSpec(
                condition=discoveryengine_v1.SearchRequest.QueryExpansionSpec.Condition.AUTO,
            ),
        )

        print(f"🌐 Using existing search engine for web search: {datastore_id}")

        # Perform the search
        response = search_client.search(request)
        results = list(response.results)

        # Check if we have summary from the search engine
        if hasattr(response, 'summary') and response.summary:
            summary_text = response.summary.summary_text
            if summary_text and len(summary_text) > 50:
                # Check if the summary is actually useful (not just "no results found")
                if not any(phrase in summary_text.lower() for phrase in [
                    "no results could be found",
                    "no results found",
                    "try rephrasing",
                    "no information available",
                    "unable to find"
                ]):
                    print(f"✅ Found useful AI-generated summary for {company_name}")
                    # Clean up and limit length
                    summary_clean = summary_text.replace('\n', ' ').replace('\r', ' ')
                    summary_clean = ' '.join(summary_clean.split())  # Remove extra whitespace

                    if len(summary_clean) > 500:
                        summary_clean = summary_clean[:500] + "..."

                    return summary_clean
                else:
                    print(f"⚠️ AI summary indicates no results found, using Gemini fallback")
                    return search_web_fallback_with_gemini(company_name)

        if not results:
            print(f"⚠️  No search results found for {company_name}")
            return search_web_fallback_with_gemini(company_name)

        # Extract and combine information from search results
        company_info_parts = []

        for result in results[:3]:  # Use top 3 results
            try:
                # Get document data
                doc_data = result.document.derived_struct_data

                # Extract snippet or other relevant text
                if hasattr(doc_data, 'snippet') and doc_data.snippet:
                    snippet = doc_data.snippet.strip()
                    if snippet and len(snippet) > 20:  # Only meaningful snippets
                        company_info_parts.append(snippet)

                # Extract title for context
                if hasattr(doc_data, 'title') and doc_data.title:
                    title = doc_data.title.strip()
                    if title and company_name.lower() in title.lower():
                        company_info_parts.append(f"About {company_name}: {title}")

                # Try to extract more content if available
                if hasattr(doc_data, 'extractive_answers'):
                    for answer in doc_data.extractive_answers[:2]:  # Top 2 answers
                        if hasattr(answer, 'content') and answer.content:
                            content = answer.content.strip()
                            if content and len(content) > 20:
                                company_info_parts.append(content)

            except Exception as e:
                print(f"Error processing search result: {e}")
                continue

        if company_info_parts:
            # Combine and clean up the information
            combined_info = " ".join(company_info_parts)

            # Clean up and limit length
            combined_info = combined_info.replace('\n', ' ').replace('\r', ' ')
            combined_info = ' '.join(combined_info.split())  # Remove extra whitespace

            # Limit to reasonable length (about 500 characters)
            if len(combined_info) > 500:
                combined_info = combined_info[:500] + "..."

            print(f"✅ Found company information for {company_name}")
            return combined_info
        else:
            print(f"⚠️  No useful information extracted for {company_name}")
            return search_web_fallback_with_gemini(company_name)

    except Exception as e:
        print(f"❌ Error during web search for {company_name}: {e}")
        return search_web_fallback_with_gemini(company_name)

def search_with_dedicated_web_app(company_name: str, web_search_config: dict) -> str:
    """
    Perform web search using the dedicated web search app.
    Uses official Google sample code pattern.
    """
    try:
        project_id = web_search_config.get('project_id')
        location = web_search_config.get('location', 'global')
        engine_id = web_search_config.get('engine_id')

        if not all([project_id, location, engine_id]):
            print(f"⚠️  Missing web search configuration")
            return search_web_fallback_with_gemini(company_name)

        # Create search query optimized for company information
        search_query = f"{company_name} company about mission values culture business overview"

        print(f"🌐 Searching with dedicated web search app (engine: {engine_id})")

        # Use official Google pattern for client setup
        from google.api_core.client_options import ClientOptions

        client_options = (
            ClientOptions(api_endpoint=f"{location}-discoveryengine.googleapis.com")
            if location != "global"
            else None
        )

        # Create a client
        client = discoveryengine_v1.SearchServiceClient(client_options=client_options)

        # The full resource name of the search app serving config
        serving_config = f"projects/{project_id}/locations/{location}/collections/default_collection/engines/{engine_id}/servingConfigs/default_config"

        print(f"🔗 Using serving config: {serving_config}")

        # Content search spec for web search
        content_search_spec = discoveryengine_v1.SearchRequest.ContentSearchSpec(
            # For information about snippets
            snippet_spec=discoveryengine_v1.SearchRequest.ContentSearchSpec.SnippetSpec(
                return_snippet=True
            ),
            # For information about search summaries
            summary_spec=discoveryengine_v1.SearchRequest.ContentSearchSpec.SummarySpec(
                summary_result_count=5,
                include_citations=True,
                ignore_adversarial_query=True,
                ignore_non_summary_seeking_query=True,
                model_prompt_spec=discoveryengine_v1.SearchRequest.ContentSearchSpec.SummarySpec.ModelPromptSpec(
                    preamble=f"Provide a comprehensive overview of {company_name} including their business, industry, mission, and key information. Focus on factual, professional details."
                ),
                model_spec=discoveryengine_v1.SearchRequest.ContentSearchSpec.SummarySpec.ModelSpec(
                    version="stable",
                ),
            ),
        )

        # Create search request using official pattern
        request = discoveryengine_v1.SearchRequest(
            serving_config=serving_config,
            query=search_query,
            page_size=10,
            content_search_spec=content_search_spec,
            query_expansion_spec=discoveryengine_v1.SearchRequest.QueryExpansionSpec(
                condition=discoveryengine_v1.SearchRequest.QueryExpansionSpec.Condition.AUTO,
            ),
            spell_correction_spec=discoveryengine_v1.SearchRequest.SpellCorrectionSpec(
                mode=discoveryengine_v1.SearchRequest.SpellCorrectionSpec.Mode.AUTO
            ),
        )

        # Perform the search
        page_result = client.search(request)

        # Process the first response to get summary
        try:
            # Get the first response from the pager
            responses = list(page_result)
            if responses:
                response = responses[0]

                # Check if we have summary from the search engine
                if hasattr(response, 'summary') and response.summary:
                    summary_text = response.summary.summary_text
                    if summary_text and len(summary_text) > 50:
                        # Check if the summary is actually useful
                        if not any(phrase in summary_text.lower() for phrase in [
                            "no results could be found",
                            "no results found",
                            "try rephrasing",
                            "no information available",
                            "unable to find"
                        ]):
                            print(f"✅ Found web search summary for {company_name}")
                            # Clean up and limit length
                            summary_clean = summary_text.replace('\n', ' ').replace('\r', ' ')
                            summary_clean = ' '.join(summary_clean.split())  # Remove extra whitespace

                            if len(summary_clean) > 500:
                                summary_clean = summary_clean[:500] + "..."

                            return summary_clean
                        else:
                            print(f"⚠️ Web search summary indicates no results, using Gemini fallback")
                            return search_web_fallback_with_gemini(company_name)

                # If no summary, try to extract from individual results
                if hasattr(response, 'results') and response.results:
                    company_info_parts = []

                    for result in response.results[:3]:  # Use top 3 results
                        try:
                            doc_data = result.document.derived_struct_data

                            # Extract snippet
                            if hasattr(doc_data, 'snippet') and doc_data.snippet:
                                snippet = doc_data.snippet.strip()
                                if snippet and len(snippet) > 20:
                                    company_info_parts.append(snippet)

                            # Extract title
                            if hasattr(doc_data, 'title') and doc_data.title:
                                title = doc_data.title.strip()
                                if title and company_name.lower() in title.lower():
                                    company_info_parts.append(f"About {company_name}: {title}")

                        except Exception as e:
                            print(f"Error processing web search result: {e}")
                            continue

                    if company_info_parts:
                        # Combine and clean up the information
                        combined_info = " ".join(company_info_parts)
                        combined_info = combined_info.replace('\n', ' ').replace('\r', ' ')
                        combined_info = ' '.join(combined_info.split())

                        if len(combined_info) > 500:
                            combined_info = combined_info[:500] + "..."

                        print(f"✅ Found web search information for {company_name}")
                        return combined_info

        except Exception as e:
            print(f"Error processing search results: {e}")

        # If no useful results from web search, use Gemini fallback
        print(f"⚠️ No useful web search results for {company_name}, using fallback")
        return search_web_fallback_with_gemini(company_name)

    except Exception as e:
        print(f"❌ Error with dedicated web search for {company_name}: {e}")
        print(f"🔍 Error details: {str(e)}")
        return search_web_fallback_with_gemini(company_name)

def search_web_fallback_with_gemini(company_name: str) -> str:
    """
    Fallback web search using basic company information when other methods fail.
    First tries Gemini AI, then falls back to predefined company info.
    """
    # First try Gemini API
    try:
        from google import genai
        from google.genai import types

        print(f"🤖 Using Gemini fallback for {company_name} information")

        client = gemini_client()

        prompt = f"Provide a brief professional overview of {company_name} company including their main business, industry, and key information. Keep it concise and factual."

        contents = [
            types.Content(
                role="user",
                parts=[types.Part(text=prompt)]
            )
        ]

        generate_content_config = types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
            temperature=0.3,
            top_p=0.8,
            max_output_tokens=3000,
        )

        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=contents,
            config=generate_content_config,
        )

        # Try to extract text from response
        if response and hasattr(response, 'text'):
            result_text = getattr(response, 'text', None)
            if result_text and result_text.strip():
                result = result_text.strip()
                print(f"✅ Generated company info for {company_name} using Gemini")
                return result

        print(f"⚠️ Gemini didn't return useful content, using basic company info")

    except Exception as e:
        print(f"⚠️ Gemini fallback error: {e}, using basic company info")

    # Fallback to basic company information database
    return get_basic_company_info(company_name)

def get_basic_company_info(company_name: str) -> str:
    """
    Provide basic company information for well-known companies.
    """
    print(f"📋 Using basic company info database for {company_name}")

    # Basic company information database
    company_info = {
        "google": "Google is a multinational technology company specializing in Internet-related services and products, including search engines, cloud computing, software, and hardware. Founded in 1998, Google is known for innovation in artificial intelligence, machine learning, and digital technologies.",

        "microsoft": "Microsoft is a multinational technology corporation that develops, manufactures, licenses, and supports software products, services, and devices. Best known for Windows operating system and Office productivity suite, Microsoft is also a leader in cloud computing with Azure platform.",

        "apple": "Apple Inc. is a multinational technology company that designs, develops, and sells consumer electronics, computer software, and online services. Known for products like iPhone, iPad, Mac computers, and innovative design philosophy focused on user experience.",

        "amazon": "Amazon is a multinational technology and e-commerce company focusing on online retail, cloud computing (AWS), digital streaming, and artificial intelligence. Started as an online bookstore, Amazon has grown into one of the world's largest technology companies.",

        "tesla": "Tesla is an American electric vehicle and clean energy company. Tesla designs and manufactures electric cars, energy storage systems, and solar panels. The company is known for innovation in sustainable transportation and energy solutions.",

        "meta": "Meta (formerly Facebook) is a multinational technology conglomerate focusing on social media, virtual reality, and metaverse technologies. The company operates Facebook, Instagram, WhatsApp, and is investing heavily in VR/AR technologies.",

        "netflix": "Netflix is a streaming entertainment service with over 200 million paid memberships worldwide. The company offers TV series, documentaries, and feature films across various genres and languages, and has pioneered the streaming industry.",

        "openai": "OpenAI is an artificial intelligence research company focused on developing artificial general intelligence (AGI) that benefits humanity. Known for creating advanced AI models like GPT and DALL-E, OpenAI conducts cutting-edge research in machine learning.",

        "spacex": "SpaceX is a private space exploration company founded by Elon Musk. The company develops spacecraft, satellites, and launch vehicles with the goal of reducing space transportation costs and enabling Mars colonization.",

        "uber": "Uber is a multinational ride-hailing company offering services including ride-hailing, food delivery (Uber Eats), package delivery, and freight transport. Uber has revolutionized urban transportation through its technology platform."
    }

    # Normalize company name for lookup
    company_key = company_name.lower().strip()

    # Check for exact match or partial match
    if company_key in company_info:
        return company_info[company_key]

    # Check for partial matches
    for key, info in company_info.items():
        if company_key in key or key in company_key:
            return info

    # Generic fallback for unknown companies
    return f"{company_name} is a company operating in its respective industry. Specific details about the company's operations, mission, and business model would require additional research to provide accurate information."

def _normalise_jd_html(text: str) -> str:
    """Convert any stray markdown the model leaks into clean HTML.

    The JD prompt asks for pure HTML, but Gemini occasionally still emits:
      - `### Heading` lines
      - `**bold**` runs
      - `- bullet` or `* bullet` lines
      - `---` horizontal rules
      - bare paragraphs separated by blank lines
    The frontend pushes the response straight into innerHTML, so anything
    that isn't already HTML shows up as raw markdown noise. This helper
    fixes those cases without disturbing properly-formed HTML.
    """
    import re as _re
    if not text:
        return ""
    s = text.strip()

    # If output already looks like HTML (starts with a tag), only strip stray
    # ** and ### inside text nodes; otherwise convert markdown wholesale.
    looks_like_html = s.lstrip().startswith("<")

    # Drop horizontal rules.
    s = _re.sub(r"(?m)^\s*---+\s*$", "", s)

    if looks_like_html:
        # Inline cleanup: kill leftover **bold** and ### inside HTML text.
        s = _re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
        s = _re.sub(r"(?m)^\s*#{1,6}\s*", "", s)
        return s.strip()

    # Markdown -> HTML fallback.
    lines = s.split("\n")
    out: list[str] = []
    in_ul = False

    def _close_ul():
        nonlocal in_ul
        if in_ul:
            out.append("</ul>")
            in_ul = False

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            _close_ul()
            continue
        # Headings
        m = _re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            _close_ul()
            level = min(len(m.group(1)), 2) + 0  # h1 stays h1, ## -> h2, ### -> h2
            tag = "h1" if level == 1 else "h2"
            content = _re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", m.group(2).strip())
            out.append(f"<{tag}>{content}</{tag}>")
            continue
        # Bullets
        m = _re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            if not in_ul:
                out.append("<ul>")
                in_ul = True
            content = _re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", m.group(1).strip())
            out.append(f"  <li>{content}</li>")
            continue
        # Paragraph
        _close_ul()
        content = _re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", line.strip())
        out.append(f"<p>{content}</p>")
    _close_ul()
    return "\n".join(out)


def generate_jd_with_gemini(job_details: Dict[str, Any]) -> Dict[str, Any]:
    """Generate job description using Gemini AI"""
    print(f"🤖 Generating JD with Gemini AI...")

    try:
        # Check which SDK to use
        model, sdk_type = get_vertex_ai_model("gemini-2.5-flash")

        if sdk_type == "vertexai" and model:
            # Use vertexai SDK for Cloud Run
            from vertexai.generative_models import GenerationConfig

            print(f"🔧 Using vertexai SDK for generation")
            generation_config = GenerationConfig(
                temperature=0.7,
                top_p=0.9,
                max_output_tokens=5000,
            )
        else:
            # Use google-genai SDK
            print(f"🔧 Using google-genai SDK for generation")
            client = gemini_client()

        # Extract job details
        job_title = job_details.get('job_title', '')
        company_name = job_details.get('company_name', '')
        company_info = job_details.get('company_info', '')
        department = job_details.get('department', '')
        location = job_details.get('location', '')
        experience_level = job_details.get('experience_level', '')
        employment_type = job_details.get('employment_type', '')
        description = job_details.get('description', '')
        skills = job_details.get('skills', [])

        # Build company context section for the prompt
        company_context = ""
        if company_name:
            company_context = f"\nCOMPANY INFORMATION:\n- Company Name: {company_name}"
            if company_info:
                company_context += f"\n- Company Background: {company_info}"

        # Create comprehensive prompt using imported JD_PROMPT
        jd_prompt = JD_PROMPT.format(
            position=job_title,
            company=company_name,
            department=department,
            location=location,
            experience_level=experience_level,
            employment_type=employment_type,
            additional_info=f"{company_context}\n\nADDITIONAL REQUIREMENTS:\n{description}\n\nREQUIRED SKILLS:\n{', '.join(skills) if skills else 'Not specified'}"
        )

        print(f"🔄 Creating request for JD generation...")

        if sdk_type == "vertexai" and model:
            # Use vertexai SDK for Cloud Run
            print(f"🚀 Sending request to Gemini via vertexai SDK...")
            response = model.generate_content(
                jd_prompt,
                generation_config=generation_config
            )
        else:
            # Use google-genai SDK
            contents = [
                types.Content(
                    role="user",
                    parts=[types.Part(text=jd_prompt)]
                )
            ]

            generate_content_config = types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
                temperature=0.7,
                top_p=0.9,
                max_output_tokens=5000,
                safety_settings=[
                    types.SafetySetting(
                        category="HARM_CATEGORY_HATE_SPEECH",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_DANGEROUS_CONTENT",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_SEXUALLY_EXPLICIT",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_HARASSMENT",
                        threshold="OFF"
                    )
                ],
            )

            print(f"🚀 Sending request to Gemini via google-genai SDK...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=contents,
                config=generate_content_config,
            )

        job_description = response.text if hasattr(response, 'text') else str(response)

        # Clean up any markdown code blocks that might wrap the HTML
        if job_description.strip().startswith('```html'):
            job_description = job_description.strip()
            job_description = job_description.replace('```html', '').replace('```', '').strip()
        elif job_description.strip().startswith('```'):
            job_description = job_description.strip()
            # Remove any code block markers
            lines = job_description.split('\n')
            if lines[0].startswith('```'):
                lines = lines[1:]
            if lines and lines[-1].strip() == '```':
                lines = lines[:-1]
            job_description = '\n'.join(lines)

        # Defensive markdown-to-HTML cleanup. Even with strict prompting, Gemini
        # occasionally leaks ** bold markers, ### headings, or --- rules. The
        # frontend renders this via innerHTML, so any literal markdown shows
        # up as raw text. Convert the common patterns to clean HTML.
        job_description = _normalise_jd_html(job_description)

        print(f"✅ JD generated successfully (length: {len(job_description)} chars)")

        return {
            "job_description": job_description,
            "success": True,
            "generated_by": "Gemini 2.5 Flash",
            "word_count": len(job_description.split()),
            "input_details": job_details
        }

    except Exception as e:
        print(f"❌ JD generation failed: {str(e)}")
        return {
            "job_description": f"Failed to generate job description: {str(e)}",
            "success": False,
            "error": str(e),
            "generated_by": "Gemini 2.5 Flash"
        }

def enhance_jd_with_gemini(existing_content: str, skills: List[str] = None) -> Dict[str, Any]:
    """Enhance existing job description using Gemini AI"""
    print(f"🤖 Enhancing JD with Gemini AI...")

    try:
        # Check which SDK to use
        model, sdk_type = get_vertex_ai_model("gemini-2.5-flash")

        if sdk_type == "vertexai" and model:
            # Use vertexai SDK for Cloud Run
            from vertexai.generative_models import GenerationConfig

            print(f"🔧 Using vertexai SDK for enhancement")
            generation_config = GenerationConfig(
                temperature=0.6,
                top_p=0.8,
                max_output_tokens=5000,
            )
        else:
            # Use google-genai SDK
            print(f"🔧 Using google-genai SDK for enhancement")
            client = gemini_client()

        skills_text = f"\nAdditional skills to incorporate: {', '.join(skills)}" if skills else ""

        enhance_prompt = ENHANCE_PROMPT.format(
            job_description=existing_content,
        ) + skills_text

        print(f"🔄 Creating request for JD enhancement...")

        if sdk_type == "vertexai" and model:
            # Use vertexai SDK for Cloud Run
            print(f"🚀 Sending enhancement request to Gemini via vertexai SDK...")
            response = model.generate_content(
                enhance_prompt,
                generation_config=generation_config
            )
        else:
            # Use google-genai SDK
            contents = [
                types.Content(
                    role="user",
                    parts=[types.Part(text=enhance_prompt)]
                )
            ]

            generate_content_config = types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
                temperature=0.6,
                top_p=0.8,
                max_output_tokens=5000,
                safety_settings=[
                    types.SafetySetting(
                        category="HARM_CATEGORY_HATE_SPEECH",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_DANGEROUS_CONTENT",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_SEXUALLY_EXPLICIT",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_HARASSMENT",
                        threshold="OFF"
                    )
                ],
            )

            print(f"🚀 Sending enhancement request to Gemini via google-genai SDK...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=contents,
                config=generate_content_config,
            )

        enhanced_description = response.text if hasattr(response, 'text') else str(response)

        # Clean up any markdown code blocks that might wrap the HTML
        if enhanced_description.strip().startswith('```html'):
            enhanced_description = enhanced_description.strip()
            enhanced_description = enhanced_description.replace('```html', '').replace('```', '').strip()
        elif enhanced_description.strip().startswith('```'):
            enhanced_description = enhanced_description.strip()
            # Remove any code block markers
            lines = enhanced_description.split('\n')
            if lines[0].startswith('```'):
                lines = lines[1:]
            if lines and lines[-1].strip() == '```':
                lines = lines[:-1]
            enhanced_description = '\n'.join(lines)

        print(f"✅ JD enhanced successfully (length: {len(enhanced_description)} chars)")

        return {
            "enhanced_description": enhanced_description,
            "success": True,
            "enhanced_by": "Gemini 2.5 Flash",
            "word_count": len(enhanced_description.split())
        }

    except Exception as e:
        print(f"❌ JD enhancement failed: {str(e)}")
        return {
            "enhanced_description": f"Failed to enhance job description: {str(e)}",
            "success": False,
            "error": str(e),
            "enhanced_by": "Gemini 2.5 Flash"
        }

def extract_keywords_with_gemini(content: str) -> Dict[str, Any]:
    """Extract keywords and skills from job description using Gemini AI"""
    print(f"🤖 Extracting keywords with Gemini AI...")

    try:
        # Check which SDK to use
        model, sdk_type = get_vertex_ai_model("gemini-2.5-flash")

        if sdk_type == "vertexai" and model:
            # Use vertexai SDK for Cloud Run
            from vertexai.generative_models import GenerationConfig

            print(f"🔧 Using vertexai SDK for keyword extraction")
            generation_config = GenerationConfig(
                temperature=0.3,
                top_p=0.7,
                max_output_tokens=8000,
            )
        else:
            # Use google-genai SDK
            print(f"🔧 Using google-genai SDK for keyword extraction")
            client = gemini_client()

        keyword_prompt = JD_KEYWORD_PROMPT.format(
            job_description=content
        )

        print(f"🔄 Creating request for keyword extraction...")

        if sdk_type == "vertexai" and model:
            # Use vertexai SDK for Cloud Run
            print(f"🚀 Sending keyword extraction request to Gemini via vertexai SDK...")
            response = model.generate_content(
                keyword_prompt,
                generation_config=generation_config
            )
        else:
            # Use google-genai SDK
            contents = [
                types.Content(
                    role="user",
                    parts=[types.Part(text=keyword_prompt)]
                )
            ]

            generate_content_config = types.GenerateContentConfig(thinking_config=types.ThinkingConfig(thinking_budget=0),
                temperature=0.3,
                top_p=0.7,
                max_output_tokens=5000,
                safety_settings=[
                    types.SafetySetting(
                        category="HARM_CATEGORY_HATE_SPEECH",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_DANGEROUS_CONTENT",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_SEXUALLY_EXPLICIT",
                        threshold="OFF"
                    ),
                    types.SafetySetting(
                        category="HARM_CATEGORY_HARASSMENT",
                        threshold="OFF"
                    )
                ],
            )

            print(f"🚀 Sending keyword extraction request to Gemini via google-genai SDK...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=contents,
                config=generate_content_config,
            )

        response_text = response.text if hasattr(response, 'text') else str(response)

        # Parse JSON response
        try:
            # Clean the response text
            clean_text = response_text.strip()
            if clean_text.startswith('```json'):
                clean_text = clean_text[7:]
            if clean_text.endswith('```'):
                clean_text = clean_text[:-3]
            clean_text = clean_text.strip()

            keywords_data = json.loads(clean_text)

            print(f"✅ Keywords extracted successfully")
            print(f"   📝 Total keywords: {len(keywords_data.get('keywords', []))}")

            return {
                "keywords": keywords_data.get('keywords', []),
                "technical_skills": keywords_data.get('technical_skills', []),
                "soft_skills": keywords_data.get('soft_skills', []),
                "qualifications": keywords_data.get('qualifications', []),
                "industry_terms": keywords_data.get('industry_terms', []),
                "success": True,
                "extracted_by": "Gemini 2.5 Flash"
            }

        except json.JSONDecodeError as e:
            print(f"⚠️ Failed to parse JSON response: {str(e)}")

            # Fallback - extract keywords from raw text
            import re
            words = re.findall(r'\b[A-Za-z][A-Za-z0-9+#\.]*\b', response_text)
            keywords = list(set([word for word in words if len(word) > 2 and word.lower() not in ['the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'her', 'was', 'one', 'our', 'out', 'day', 'had', 'has', 'his', 'how', 'man', 'new', 'now', 'old', 'see', 'two', 'way', 'who', 'boy', 'did', 'its', 'let', 'put', 'say', 'she', 'too', 'use']]))[:15]

            return {
                "keywords": keywords,
                "technical_skills": [],
                "soft_skills": [],
                "qualifications": [],
                "industry_terms": [],
                "success": True,
                "extracted_by": "Gemini 2.5 Flash (Fallback)",
                "note": "JSON parsing failed, used fallback extraction"
            }

    except Exception as e:
        print(f"❌ Keyword extraction failed: {str(e)}")
        return {
            "keywords": [],
            "technical_skills": [],
            "soft_skills": [],
            "qualifications": [],
            "industry_terms": [],
            "success": False,
            "error": str(e),
            "extracted_by": "Gemini 2.5 Flash"
        }

@app.post("/api/generate-jd")
@limiter.limit("15/minute")
async def generate_job_description(
    request: Request,
    company_name: str = Form(""),
    company_details: str = Form(""),
    job_title: str = Form(...),
    department: str = Form(""),
    location: str = Form(""),
    experience_level: str = Form(""),
    employment_type: str = Form("full-time"),
    description: str = Form(""),
    skills: str = Form("[]"),
    user: dict = Depends(require_auth)
):
    """Generate a job description using AI"""
    try:
        company_name = _sanitize_user_text(company_name, max_len=200)
        company_details = _sanitize_user_text(company_details, max_len=4000)
        job_title = _sanitize_user_text(job_title, max_len=200)
        department = _sanitize_user_text(department, max_len=200)
        location = _sanitize_user_text(location, max_len=200)
        experience_level = _sanitize_user_text(experience_level, max_len=100)
        employment_type = _sanitize_user_text(employment_type, max_len=100) or "full-time"
        description = _sanitize_user_text(description, max_len=8000)
        # Parse skills JSON
        skills_list = json.loads(skills) if skills else []

        # Use LLM-only approach with provided company details
        job_details = {
            "company_name": company_name,
            "company_info": company_details,  # Use provided company details instead of web search
            "job_title": job_title,
            "department": department,
            "location": location,
            "experience_level": experience_level,
            "employment_type": employment_type,
            "description": description,
            "skills": skills_list
        }

        print(f"🏢 Generating JD for: {job_title}")
        if company_name:
            print(f"🏢 Company: {company_name}")
        if company_details:
            print(f"📋 Company Details: {company_details[:100]}{'...' if len(company_details) > 100 else ''}")
        print(f"👤 Requested by: {user['email']}")
        print(f"🤖 Using LLM-only generation with company details")

        result = generate_jd_with_gemini(job_details)

        return result

    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid skills format")
    except Exception as e:
        print(f"❌ Generate JD endpoint error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate job description: {str(e)}")

@app.post("/api/enhance-jd")
@limiter.limit("15/minute")
async def enhance_job_description(
    request: Request,
    existing_content: str = Form(...),
    skills: str = Form("[]"),
    user: dict = Depends(require_auth)
):
    """Enhance an existing job description using AI"""
    try:
        existing_content = _sanitize_user_text(existing_content, max_len=20000)
        # Parse skills JSON
        skills_list = json.loads(skills) if skills else []

        print(f"🔧 Enhancing JD (length: {len(existing_content)} chars)")
        print(f"👤 Requested by: {user['email']}")

        result = enhance_jd_with_gemini(existing_content, skills_list)

        return result

    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid skills format")
    except Exception as e:
        print(f"❌ Enhance JD endpoint error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to enhance job description: {str(e)}")

@app.post("/api/extract-jd-keywords")
@limiter.limit("30/minute")
async def extract_jd_keywords(
    request: Request,
    content: str = Form(...),
    user: dict = Depends(require_auth)
):
    """Extract keywords from job description or requirements"""
    try:
        content = _sanitize_user_text(content, max_len=20000)
        print(f"🔍 Extracting keywords from content (length: {len(content)} chars)")
        print(f"👤 Requested by: {user['email']}")

        result = extract_keywords_with_gemini(content)

        return result

    except Exception as e:
        print(f"❌ Extract keywords endpoint error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to extract keywords: {str(e)}")


@app.post("/api/parse-jd-file")
@limiter.limit("30/minute")
async def parse_jd_file(
    request: Request,
    file: UploadFile = File(...),
    user: dict = Depends(require_auth)
):
    """
    Parse an uploaded job description file (PDF / DOCX / TXT) and return its
    extracted plain text. Used by the search page paperclip button so HR can
    upload a JD instead of pasting it.
    """
    try:
        filename = file.filename or "uploaded.txt"
        ext = os.path.splitext(filename)[1].lower()
        allowed = {".pdf", ".doc", ".docx", ".txt"}
        if ext not in allowed:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type '{ext}'. Use PDF, DOC, DOCX, or TXT.",
            )

        content = await file.read()
        # Guard against empty / oversized uploads (limit mirrors the JS guard).
        if not content:
            raise HTTPException(status_code=400, detail="Uploaded file is empty.")
        if len(content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large (max 10 MB).")

        text = extract_text_from_file(content, filename)
        if not text or not text.strip():
            raise HTTPException(
                status_code=422,
                detail="Could not extract any text from the file. It may be a scanned image PDF.",
            )

        # Cap returned text so we never blow up the chat input field.
        max_chars = 20000
        truncated = False
        if len(text) > max_chars:
            text = text[:max_chars]
            truncated = True

        print(f"📄 Parsed JD file '{filename}' for {user.get('email')}: {len(text)} chars (truncated={truncated})")

        return {
            "success": True,
            "filename": filename,
            "text": text,
            "truncated": truncated,
            "char_count": len(text),
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Parse JD file error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {str(e)}")



if __name__ == "__main__":
    import uvicorn
    # Use PORT environment variable for compatibility with Cloud Run
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
