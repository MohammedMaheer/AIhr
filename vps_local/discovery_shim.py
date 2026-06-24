"""
discovery_shim — drop-in for `google.cloud.discoveryengine_v1`.

Replaces Discovery Engine's vector + keyword search with pgvector queries
against the local `resume_embeddings` table.

Surface implemented (only what main.py touches):
  - SearchServiceClient / DocumentServiceClient / DataStoreServiceClient
  - SearchRequest, SearchRequest.ContentSearchSpec.* (no-op kept for arg compat)
  - SearchRequest.QueryExpansionSpec, SearchRequest.SpellCorrectionSpec
  - Document, DataStore, IndustryVertical, SolutionType, CreateDocumentRequest
  - search() → response.results[*].document.{struct_data, derived_struct_data, name}
  - response.summary  (left empty)
"""
from __future__ import annotations
import os
import mimetypes
import re
import threading
from dataclasses import dataclass, field
from collections import OrderedDict
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional
from urllib.parse import urlparse

# We rely on the same vertexai_shim embedder so query and corpus use the same model.
from . import vertexai_shim
from .profile_extractor import evaluate_profile, profile_from_row, profile_from_text
from search_criteria import (
    LOCATION_COUNTRY_CITIES,
    LOCATION_SYNONYMS,
    parse_search_criteria,
)


ADVANCED_SEARCH_ENABLED = os.environ.get("SMARTHR_ADVANCED_SEARCH", "1").lower() in (
    "1", "true", "yes", "on"
)
ADVANCED_SEARCH_POOL_MIN = int(os.environ.get("SMARTHR_ADVANCED_SEARCH_POOL_MIN", "150"))
ADVANCED_SEARCH_POOL_MULTIPLIER = int(os.environ.get("SMARTHR_ADVANCED_SEARCH_POOL_MULTIPLIER", "12"))
ADVANCED_SEARCH_POOL_MAX = int(os.environ.get("SMARTHR_ADVANCED_SEARCH_POOL_MAX", "700"))
QUERY_EMBED_CACHE_MAX = int(os.environ.get("SMARTHR_QUERY_EMBED_CACHE_MAX", "512"))


# ---------- Constants / enums (kept for arg compatibility only) ----------

class IndustryVertical:
    GENERIC = "GENERIC"


class SolutionType:
    SOLUTION_TYPE_SEARCH = "SOLUTION_TYPE_SEARCH"


# ---------- Document / DataStore / requests ----------

@dataclass
class _ContentInner:
    uri: str = ""
    mime_type: str = ""


@dataclass
class Document:
    id: str = ""
    struct_data: Optional[Dict[str, Any]] = None
    derived_struct_data: Optional[Dict[str, Any]] = None
    content: Optional[_ContentInner] = None
    name: str = ""

    Content = _ContentInner  # nested type used by callers


@dataclass
class DataStore:
    display_name: str = ""
    industry_vertical: str = IndustryVertical.GENERIC
    content_config: str = "CONTENT_REQUIRED"
    solution_types: List[str] = field(default_factory=list)
    name: str = ""

    class ContentConfig:
        CONTENT_REQUIRED = "CONTENT_REQUIRED"


@dataclass
class CreateDocumentRequest:
    parent: str = ""
    document: Optional[Document] = None
    document_id: str = ""


# Nested SearchRequest.* compatibility classes
class SearchRequest:
    def __init__(self, *, serving_config: str = "", query: str = "",
                 page_size: int = 10, content_search_spec=None,
                 query_expansion_spec=None, spell_correction_spec=None,
                 filter: str = "", **kwargs):
        self.serving_config = serving_config
        self.query = query
        self.page_size = page_size
        self.content_search_spec = content_search_spec
        self.query_expansion_spec = query_expansion_spec
        self.spell_correction_spec = spell_correction_spec
        self.filter = filter

    class ContentSearchSpec:
        class SnippetSpec:
            def __init__(self, return_snippet: bool = True, **kwargs):
                self.return_snippet = return_snippet

        class SummarySpec:
            class ModelPromptSpec:
                def __init__(self, preamble: str = "", **kwargs):
                    self.preamble = preamble

            class ModelSpec:
                def __init__(self, version: str = "", **kwargs):
                    self.version = version

            def __init__(self, summary_result_count: int = 3, **kwargs):
                self.summary_result_count = summary_result_count

        class ExtractiveContentSpec:
            def __init__(self, **kwargs):
                pass

        def __init__(self, snippet_spec=None, summary_spec=None,
                     extractive_content_spec=None, **kwargs):
            self.snippet_spec = snippet_spec
            self.summary_spec = summary_spec
            self.extractive_content_spec = extractive_content_spec

    class QueryExpansionSpec:
        class Condition:
            AUTO = "AUTO"
            DISABLED = "DISABLED"

        def __init__(self, condition=None, **kwargs):
            self.condition = condition

    class SpellCorrectionSpec:
        class Mode:
            AUTO = "AUTO"

        def __init__(self, mode=None, **kwargs):
            self.mode = mode


# ---------- Search response shapes ----------

@dataclass
class _Result:
    id: str = ""
    document: Optional[Document] = None
    relevance_score: float = 0.0


@dataclass
class _SearchResponse:
    results: List[_Result] = field(default_factory=list)
    summary: Any = None
    total_size: int = 0

    def __iter__(self):
        return iter(self.results)


# ---------- DB access ----------

_db_pool_singleton = None
_advanced_schema_ensured = False
_query_embedding_cache: OrderedDict[str, List[float]] = OrderedDict()
_query_embedding_cache_lock = threading.Lock()


def _get_pool():
    """Lazy pool, separate from main.py's, so the shim is import-safe."""
    global _db_pool_singleton
    if _db_pool_singleton is not None:
        return _db_pool_singleton
    import psycopg2
    from psycopg2 import pool as pg_pool
    cfg = dict(
        host=os.environ.get("DB_HOST", "smarthr-postgres"),
        port=int(os.environ.get("DB_PORT", "5432")),
        dbname=os.environ.get("DB_NAME", "recruitment"),
        user=os.environ.get("DB_USER", "smarthr"),
        password=os.environ.get("DB_PASSWORD", ""),
    )
    _db_pool_singleton = pg_pool.SimpleConnectionPool(1, 5, **cfg)
    return _db_pool_singleton


def _embed_query(text: str) -> List[float]:
    key = re.sub(r"\s+", " ", (text or "").strip().lower())
    if key and QUERY_EMBED_CACHE_MAX > 0:
        with _query_embedding_cache_lock:
            cached = _query_embedding_cache.get(key)
            if cached is not None:
                _query_embedding_cache.move_to_end(key)
                return cached
    model = vertexai_shim.language_models.TextEmbeddingModel.from_pretrained("all-MiniLM-L6-v2")
    values = model.get_embeddings([text])[0].values
    if key and QUERY_EMBED_CACHE_MAX > 0:
        with _query_embedding_cache_lock:
            _query_embedding_cache[key] = values
            _query_embedding_cache.move_to_end(key)
            while len(_query_embedding_cache) > QUERY_EMBED_CACHE_MAX:
                _query_embedding_cache.popitem(last=False)
    return values


def warm_search_runtime() -> None:
    """Load heavyweight local search dependencies before the first user query."""
    _get_pool()
    _embed_query("startup warmup query")


def _vec_lit(vec: List[float]) -> str:
    return "[" + ",".join(f"{x:.6f}" for x in vec) + "]"


def _ensure_advanced_search_schema(cur) -> None:
    """Ensure optional structured-search columns/indexes exist."""
    global _advanced_schema_ensured
    if _advanced_schema_ensured:
        return
    statements = [
        "ALTER TABLE resume_embeddings ADD COLUMN IF NOT EXISTS candidate_years NUMERIC(5,2)",
        "ALTER TABLE resume_embeddings ADD COLUMN IF NOT EXISTS candidate_location TEXT",
        "ALTER TABLE resume_embeddings ADD COLUMN IF NOT EXISTS candidate_languages TEXT[]",
        "ALTER TABLE resume_embeddings ADD COLUMN IF NOT EXISTS candidate_skills TEXT[]",
        "ALTER TABLE resume_embeddings ADD COLUMN IF NOT EXISTS profile_source TEXT",
        "ALTER TABLE resume_embeddings ADD COLUMN IF NOT EXISTS profile_updated_at TIMESTAMP",
        "CREATE INDEX IF NOT EXISTS idx_resume_emb_candidate_years ON resume_embeddings(candidate_years)",
        "CREATE INDEX IF NOT EXISTS idx_resume_emb_candidate_location ON resume_embeddings(candidate_location)",
        "CREATE INDEX IF NOT EXISTS idx_resume_emb_candidate_languages ON resume_embeddings USING gin(candidate_languages)",
        "CREATE INDEX IF NOT EXISTS idx_resume_emb_candidate_skills ON resume_embeddings USING gin(candidate_skills)",
        "CREATE INDEX IF NOT EXISTS idx_resume_emb_text_fts ON resume_embeddings USING gin(to_tsvector('simple', coalesce(extracted_text, '')))",
    ]
    for stmt in statements:
        cur.execute(stmt)
    _advanced_schema_ensured = True


def _candidate_pool_size(page_size: int, has_criteria: bool) -> int:
    page_size = max(1, int(page_size or 10))
    if not has_criteria:
        return page_size
    return min(
        ADVANCED_SEARCH_POOL_MAX,
        max(page_size, ADVANCED_SEARCH_POOL_MIN, page_size * ADVANCED_SEARCH_POOL_MULTIPLIER),
    )


def _location_sql_terms(location: str) -> list[str]:
    terms: list[str] = []
    wanted = (location or "").strip().lower()
    if not wanted:
        return terms
    terms.append(wanted)
    for aliases in LOCATION_SYNONYMS:
        if wanted in aliases:
            terms.extend(sorted(aliases))
    terms.extend(sorted(LOCATION_COUNTRY_CITIES.get(wanted, set())))
    out: list[str] = []
    for term in terms:
        if term and term not in out:
            out.append(term)
    return out


def _criteria_sql_where(criteria) -> tuple[str, list[Any]]:
    clauses = ["embedding IS NOT NULL"]
    params: list[Any] = []

    if criteria.min_years is not None:
        clauses.append("candidate_years IS NOT NULL AND candidate_years >= %s")
        params.append(criteria.min_years)
    if criteria.max_years is not None:
        clauses.append("candidate_years IS NOT NULL AND candidate_years <= %s")
        params.append(criteria.max_years)

    location_terms: list[str] = []
    for location in criteria.all_locations():
        location_terms.extend(_location_sql_terms(location))
    if location_terms:
        loc_clauses = []
        for term in location_terms:
            loc_clauses.append("candidate_location ILIKE %s")
            params.append(f"%{term}%")
        clauses.append("(" + " OR ".join(loc_clauses) + ")")

    for language in criteria.required_languages:
        clauses.append("%s = ANY(candidate_languages)")
        params.append(language)

    for skill in criteria.required_skills:
        clauses.append(
            "EXISTS (SELECT 1 FROM unnest(candidate_skills) s WHERE s ILIKE %s)"
        )
        params.append(f"%{skill}%")

    return " AND ".join(clauses), params


def _row_tuple_to_dict(r) -> Dict[str, Any]:
    return {
        "id": str(r[0]),
        "filename": r[1],
        "folder": r[2],
        "file_size": r[3],
        "content_type": r[4],
        "text": r[5] or "",
        "text_length": r[6] or 0,
        "score": float(r[7] or 0.0),
        "lexical_score": float(r[8] or 0.0),
        "candidate_years": float(r[9]) if r[9] is not None else None,
        "candidate_location": r[10] or "",
        "candidate_languages": list(r[11] or []),
        "candidate_skills": list(r[12] or []),
        "profile_source": r[13] or "",
    }


def _rank_rows(rows: list[Dict[str, Any]], query: str, page_size: int) -> list[Dict[str, Any]]:
    criteria = parse_search_criteria(query)
    if criteria.is_empty():
        return rows[:page_size]

    ranked = []
    for row in rows:
        profile = profile_from_row(row)
        ev, status = evaluate_profile(profile, criteria)
        row["candidate_profile"] = profile
        row["criteria_match_preview"] = ev
        row["criteria_profile_status"] = status

        vector_score = float(row.get("score") or 0.0)
        lexical_score = min(float(row.get("lexical_score") or 0.0), 1.0)
        profile_bonus = {
            "exact": 1.0,
            "unknown": 0.15,
            "known_fail": -1.0,
            "no_criteria": 0.0,
        }.get(status, 0.0)
        if ev.get("hard_fail"):
            profile_bonus = -1.25

        row["score"] = (0.68 * vector_score) + (0.20 * lexical_score) + (0.12 * profile_bonus)
        row["vector_score"] = vector_score
        row["hybrid_lexical_score"] = lexical_score
        ranked.append(row)

    ranked.sort(
        key=lambda item: (
            1 if item.get("criteria_profile_status") == "exact" else 0,
            item.get("score") or 0.0,
            item.get("vector_score") or 0.0,
        ),
        reverse=True,
    )
    exact = [r for r in ranked if r.get("criteria_profile_status") == "exact"]
    unknown = [r for r in ranked if r.get("criteria_profile_status") == "unknown"]
    if len(exact) >= page_size:
        return exact[:page_size]
    return (exact + unknown)[:page_size]


def _extract_text_from_file(path: Path) -> str:
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    return "\n".join((p.extract_text() or "") for p in pdf.pages)[:80_000]
            except Exception:
                from PyPDF2 import PdfReader
                reader = PdfReader(str(path))
                return "\n".join((p.extract_text() or "") for p in reader.pages)[:80_000]
        if ext in (".docx", ".doc"):
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            return "\n".join(p.text for p in doc.paragraphs)[:80_000]
        if ext in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="ignore")[:80_000]
    except Exception:
        return ""
    return ""


def _paths_for_uri(uri: str) -> tuple[str, list[Path]]:
    parsed = urlparse(uri or "")
    storage_root = Path(os.environ.get("SMARTHR_STORAGE_ROOT", "/app/storage")).resolve()
    candidates: list[Path] = []
    key = ""
    if parsed.scheme == "gs":
        bucket = parsed.netloc
        key = parsed.path.lstrip("/")
        candidates.append(storage_root / bucket / key)
        candidates.append(storage_root / "resumes" / key)
        candidates.append(storage_root / key)
    else:
        raw = Path(uri or "")
        key = raw.name
        candidates.append(raw)
        candidates.append(storage_root / "resumes" / raw)
    return key, candidates


def _upsert_document_embedding(document: Document, document_id: str = "") -> str:
    content = getattr(document, "content", None)
    uri = getattr(content, "uri", "") if content else ""
    key, candidates = _paths_for_uri(uri)
    path = next((p for p in candidates if p.is_file()), None)
    if path is None:
        return "local/operations/create_document_missing_file"

    text = _extract_text_from_file(path)
    filename = Path(key or path.name).name
    folder = Path(key).parts[0] if key and len(Path(key).parts) > 1 else "legacy"
    ctype = (getattr(content, "mime_type", "") if content else "") or mimetypes.guess_type(str(path))[0] or path.suffix.lstrip(".")

    struct = document.struct_data or {}
    company_id = None
    try:
        raw_company_id = struct.get("company_id")
        if raw_company_id and str(raw_company_id).lower() != "system":
            company_id = int(raw_company_id)
    except Exception:
        company_id = None

    profile = profile_from_text(text, source="text_incremental")
    vec = _embed_query(text) if text and len(text) >= 30 else None

    pool = _get_pool()
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            _ensure_advanced_search_schema(cur)
            if vec is not None:
                cur.execute(
                    """
                    INSERT INTO resume_embeddings
                        (filename, folder, company_id, file_size, content_type,
                         extracted_text, text_length, embedding,
                         candidate_years, candidate_location,
                         candidate_languages, candidate_skills, profile_source,
                         profile_updated_at, embedded_at, embed_model)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s::vector,
                            %s, NULLIF(%s, ''), %s, %s, %s, NOW(), NOW(), %s)
                    ON CONFLICT (filename) DO UPDATE SET
                        folder = EXCLUDED.folder,
                        company_id = EXCLUDED.company_id,
                        file_size = EXCLUDED.file_size,
                        content_type = EXCLUDED.content_type,
                        extracted_text = EXCLUDED.extracted_text,
                        text_length = EXCLUDED.text_length,
                        embedding = EXCLUDED.embedding,
                        candidate_years = EXCLUDED.candidate_years,
                        candidate_location = EXCLUDED.candidate_location,
                        candidate_languages = EXCLUDED.candidate_languages,
                        candidate_skills = EXCLUDED.candidate_skills,
                        profile_source = EXCLUDED.profile_source,
                        profile_updated_at = NOW(),
                        embedded_at = NOW(),
                        embed_model = EXCLUDED.embed_model,
                        extraction_error = NULL
                    """,
                    (
                        filename, folder, company_id, path.stat().st_size, ctype,
                        text.replace("\x00", ""), len(text),
                        _vec_lit(vec),
                        profile.get("candidate_years"),
                        profile.get("candidate_location") or "",
                        profile.get("candidate_languages") or [],
                        profile.get("candidate_skills") or [],
                        profile.get("profile_source") or "text_incremental",
                        os.environ.get("SMARTHR_EMBED_MODEL", "all-MiniLM-L6-v2"),
                    ),
                )
            else:
                cur.execute(
                    """
                    INSERT INTO resume_embeddings
                        (filename, folder, company_id, file_size, content_type,
                         extracted_text, text_length, extraction_error,
                         profile_source, profile_updated_at)
                    VALUES (%s, %s, %s, %s, %s, '', 0, 'no_text_extracted',
                            'text_incremental', NOW())
                    ON CONFLICT (filename) DO UPDATE SET
                        folder = EXCLUDED.folder,
                        company_id = EXCLUDED.company_id,
                        file_size = EXCLUDED.file_size,
                        content_type = EXCLUDED.content_type,
                        extraction_error = EXCLUDED.extraction_error,
                        profile_updated_at = NOW()
                    """,
                    (filename, folder, company_id, path.stat().st_size, ctype),
                )
        conn.commit()
    finally:
        pool.putconn(conn)
    return f"local/operations/create_document_{document_id or filename}"


def _vector_search(query: str, page_size: int = 10,
                   filter_str: str = "") -> List[Dict[str, Any]]:
    """Run pgvector ANN search; return list of dicts shaped like Discovery results."""
    if not query.strip():
        return []
    qvec = _embed_query(query)
    criteria = parse_search_criteria(query)
    pool_size = _candidate_pool_size(page_size, not criteria.is_empty())
    pool = _get_pool()
    conn = pool.getconn()
    try:
        with conn.cursor() as cur:
            if not ADVANCED_SEARCH_ENABLED:
                cur.execute(
                    """
                    SELECT id, filename, folder, file_size, content_type,
                           extracted_text, text_length,
                           1 - (embedding <=> %s::vector) AS score
                      FROM resume_embeddings
                     WHERE embedding IS NOT NULL
                     ORDER BY embedding <=> %s::vector
                     LIMIT %s
                    """,
                    (_vec_lit(qvec), _vec_lit(qvec), int(page_size or 10)),
                )
                rows = cur.fetchall()
                return [
                    {
                        "id": str(r[0]),
                        "filename": r[1],
                        "folder": r[2],
                        "file_size": r[3],
                        "content_type": r[4],
                        "text": r[5] or "",
                        "text_length": r[6] or 0,
                        "score": float(r[7] or 0.0),
                    }
                    for r in rows
                ]

            _ensure_advanced_search_schema(cur)
            conn.commit()

            # Cosine distance uses the HNSW vector_cosine_ops index. In
            # advanced mode we pull a deeper semantic pool, add full-text rank
            # and structured-profile signals, then trim in Python.
            cur.execute(
                """
                WITH semantic AS (
                    SELECT id, filename, folder, file_size, content_type,
                           extracted_text, text_length,
                           1 - (embedding <=> %s::vector) AS vector_score,
                           candidate_years, candidate_location,
                           candidate_languages, candidate_skills, profile_source
                      FROM resume_embeddings
                     WHERE embedding IS NOT NULL
                     ORDER BY embedding <=> %s::vector
                     LIMIT %s
                )
                SELECT id, filename, folder, file_size, content_type,
                       extracted_text, text_length,
                       vector_score,
                       ts_rank_cd(
                           to_tsvector('simple', coalesce(extracted_text, '')),
                           websearch_to_tsquery('simple', %s)
                       ) AS lexical_score,
                       candidate_years, candidate_location,
                       candidate_languages, candidate_skills, profile_source
                  FROM semantic
                """,
                (_vec_lit(qvec), _vec_lit(qvec), int(pool_size), query[:1000]),
            )
            rows = cur.fetchall()
            if not criteria.is_empty():
                where_sql, where_params = _criteria_sql_where(criteria)
                supplemental_limit = max(pool_size, page_size * 20)
                cur.execute(
                    f"""
                    SELECT id, filename, folder, file_size, content_type,
                           extracted_text, text_length,
                           1 - (embedding <=> %s::vector) AS vector_score,
                           ts_rank_cd(
                               to_tsvector('simple', coalesce(extracted_text, '')),
                               websearch_to_tsquery('simple', %s)
                           ) AS lexical_score,
                           candidate_years, candidate_location,
                           candidate_languages, candidate_skills, profile_source
                      FROM resume_embeddings
                     WHERE {where_sql}
                     ORDER BY embedding <=> %s::vector
                     LIMIT %s
                    """,
                    (
                        _vec_lit(qvec),
                        query[:1000],
                        *where_params,
                        _vec_lit(qvec),
                        int(supplemental_limit),
                    ),
                )
                seen_ids = {r[0] for r in rows}
                rows = list(rows)
                for row in cur.fetchall():
                    if row[0] not in seen_ids:
                        rows.append(row)
                        seen_ids.add(row[0])
        row_dicts = [_row_tuple_to_dict(r) for r in rows]
        return _rank_rows(row_dicts, query, int(page_size or 10))
    finally:
        pool.putconn(conn)


def _row_to_result(row: Dict[str, Any]) -> _Result:
    file_key = (row.get("folder") or "legacy").rstrip("/") + "/" + (row.get("filename") or "")
    bucket = os.environ.get("GCS_BUCKET_NAME", "smarthr-prod-2026-resume-storage")
    gs_uri = f"gs://{bucket}/{file_key}"
    text = row.get("text") or ""
    snippet = (text[:300] + "…") if len(text) > 300 else text
    derived = {
        "link": gs_uri,
        "snippets": [{"snippet": snippet, "snippet_status": "SUCCESS"}],
        "extractive_segments": [{"content": snippet}] if snippet else [],
    }
    struct = {
        "filename": row.get("filename"),
        "file_path": file_key,
        "gcs_uri": gs_uri,
        "candidate_profile": row.get("candidate_profile") or {},
        "criteria_match_preview": row.get("criteria_match_preview") or {},
        "criteria_profile_status": row.get("criteria_profile_status") or "",
    }
    doc = Document(
        id=row["id"],
        name=f"local/documents/{row['id']}",
        struct_data=struct,
        derived_struct_data=derived,
        content=_ContentInner(uri=gs_uri, mime_type=row.get("content_type") or ""),
    )
    return _Result(id=row["id"], document=doc, relevance_score=row.get("score", 0.0))


# ---------- Clients ----------

class SearchServiceClient:
    def __init__(self, *args, **kwargs):
        pass

    def search(self, request: SearchRequest, **kwargs) -> _SearchResponse:
        rows = _vector_search(query=request.query, page_size=request.page_size,
                              filter_str=getattr(request, "filter", "") or "")
        return _SearchResponse(
            results=[_row_to_result(r) for r in rows],
            total_size=len(rows),
            summary=None,
        )


class DocumentServiceClient:
    def __init__(self, *args, **kwargs):
        pass

    def create_document(self, request: CreateDocumentRequest = None, **kwargs):
        op_name = "local/operations/create_document"
        try:
            if request and request.document:
                op_name = _upsert_document_embedding(
                    request.document,
                    document_id=getattr(request, "document_id", "") or getattr(request.document, "id", ""),
                )
        except Exception as exc:
            op_name = f"local/operations/create_document_failed_{type(exc).__name__}"

        class _Op:
            pass
        _Op.name = op_name
        return _Op()


class DataStoreServiceClient:
    def __init__(self, *args, **kwargs):
        pass

    def list_data_stores(self, parent: str = "", **kwargs):
        # Return a single "local" datastore.
        return [DataStore(name=f"{parent}/dataStores/smarthr-local",
                          display_name="smarthr-local")]

    def create_data_store(self, parent: str = "", data_store: DataStore = None,
                          data_store_id: str = "", **kwargs):
        class _LRO:
            name = f"local/operations/create_datastore_{data_store_id}"

            def result(self, timeout=None):
                return DataStore(name=f"{parent}/dataStores/{data_store_id}",
                                 display_name=data_store_id)
        return _LRO()
